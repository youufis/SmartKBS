"""
学生成长档案 API 路由
聚合展示学生全维度成长数据：考试、积分、点名、任务、对话
"""
from datetime import datetime, timedelta
from collections import defaultdict
from typing import Any

from fastapi import APIRouter, HTTPException, Request, Query

from backend.api.dependencies import get_current_user
from backend.database import execute_query
from backend.question_db import execute_query as q_execute_query
from backend.logger import logger
from backend.prompts import build_ai_role

router = APIRouter()


@router.get("/{username}", summary="获取学生完整成长档案")
async def get_portfolio(username: str, request: Request):
    """获取指定学生的完整成长档案"""
    user = get_current_user(request)
    current_username = user["username"]
    role = user.get("role", 2)

    # 权限：学生只能看自己，教师/管理员可看任何学生
    if role == 2 and current_username != username:
        raise HTTPException(status_code=403, detail="无权查看其他学生的档案")

    # 获取学生基本信息
    user_rows = execute_query(
        "SELECT username, name, class, grade, gender FROM users WHERE username = ?",
        (username,),
    )
    if not user_rows:
        raise HTTPException(status_code=404, detail="学生不存在")

    user_info = {
        "username": user_rows[0][0],
        "name": user_rows[0][1] or user_rows[0][0],
        "class": user_rows[0][2] or "",
        "grade": user_rows[0][3] or "",
        "gender": str(user_rows[0][4] or ""),
    }

    # ── 1. 考试成绩 ──
    exam_results = q_execute_query(
        """SELECT ea.id, e.title, ea.score, ea.total_score, e.pass_score,
                  ea.submitted_at, e.subject
           FROM exam_attempts ea
           JOIN exams e ON ea.exam_id = e.id
           WHERE ea.student_username = ? AND ea.status IN ('submitted', 'graded')
           ORDER BY ea.submitted_at ASC""",
        (username,),
    )

    exam_stats = {}
    if exam_results:
        scores = [r['score'] for r in exam_results]
        total_scores = [r['total_score'] for r in exam_results]
        percentages = [(r['score'] / r['total_score'] * 100) if r['total_score'] > 0 else 0 for r in exam_results]
        passed = sum(1 for r in exam_results if r['score'] >= r['pass_score'])
        exam_stats = {
            "total_exams": len(exam_results),
            "avg_score": round(sum(scores) / len(scores), 1) if scores else 0,
            "avg_percentage": round(sum(percentages) / len(percentages), 1) if percentages else 0,
            "passed_count": passed,
            "failed_count": len(exam_results) - passed,
            "max_score": max(scores) if scores else 0,
            "min_score": min(scores) if scores else 0,
            "trend": percentages,  # 百分比趋势（用于前端画折线图）
            "subjects": list(set(r['subject'] for r in exam_results if r['subject'])),
        }

    # ── 2. 课堂积分 ──
    scores_rows = execute_query(
        """SELECT teacher_username, grade, class_name, score, updated_at
           FROM scores WHERE student_name = ?
           ORDER BY updated_at DESC""",
        (user_info['name'],),
    )

    score_stats = {"total_score": 0, "teacher_count": 0, "class_count": 0, "records": [], "trend": []}
    if scores_rows:
        total_score = sum(r[3] for r in scores_rows)
        teacher_count = len(set(r[0] for r in scores_rows))
        score_stats = {
            "total_score": total_score,
            "teacher_count": teacher_count,
            "class_count": len(set(r[2] for r in scores_rows)),
            "records": [
                {
                    "teacher": r[0],
                    "grade": r[1],
                    "class": r[2],
                    "score": r[3],
                    "updated_at": r[4] or "",
                }
                for r in scores_rows
            ],
        }

        # 积分趋势（按日期分组）
        score_by_date = defaultdict(int)
        for r in scores_rows:
            if r[4]:
                date_key = r[4][:10]
                score_by_date[date_key] += r[3]
        score_stats["trend"] = [
            {"date": k, "score": v}
            for k, v in sorted(score_by_date.items())
        ]

    # ── 2B. 奖励积分（活动自动发放） ──
    reward_points = 0
    reward_history = []
    try:
        reward_row = execute_query(
            "SELECT total_points FROM student_total_points WHERE student_username=?",
            (username,),
        )
        if reward_row and reward_row[0][0]:
            reward_points = reward_row[0][0]

        reward_rows = execute_query(
            """SELECT activity_type, activity_title, reward_type, points, reason, created_at
               FROM activity_rewards
               WHERE student_username=?
               ORDER BY created_at DESC LIMIT 100""",
            (username,),
        )
        reward_history = [
            {
                "activity_type": r[0],
                "activity_title": r[1] or "",
                "reward_type": r[2],
                "points": r[3],
                "reason": r[4] or "",
                "created_at": r[5] or "",
            }
            for r in reward_rows
        ]
    except Exception:
        pass  # student_total_points 表可能不存在

    # ── 3. 点名记录 ──
    rollcall_rows = execute_query(
        """SELECT teacher_username, grade, class_name, result, points, created_at
           FROM rollcall_history WHERE student_name = ?
           ORDER BY created_at DESC""",
        (user_info['name'],),
    )

    rollcall_stats = {"total_calls": 0, "correct_count": 0, "wrong_count": 0, "accuracy": 0, "total_points": 0}
    if rollcall_rows:
        total_calls = len(rollcall_rows)
        correct = sum(1 for r in rollcall_rows if r[3] == "1")
        wrong = sum(1 for r in rollcall_rows if r[3] == "0")
        total_points = sum(r[4] or 0 for r in rollcall_rows)
        rollcall_stats = {
            "total_calls": total_calls,
            "correct_count": correct,
            "wrong_count": wrong,
            "accuracy": round(correct / max(total_calls, 1) * 100, 1),
            "total_points": total_points,
        }

    # ── 4. 任务完成 ──
    task_rows = execute_query(
        """SELECT t.name, t.description, ts.submitted_at, t.created_at
           FROM task_submissions ts
           JOIN tasks t ON ts.task_id = t.id
           WHERE ts.student_username = ?
           ORDER BY ts.submitted_at DESC""",
        (username,),
    )

    task_stats = {}
    if task_rows:
        task_stats = {
            "completed": len(task_rows),
            "tasks": [
                {
                    "name": r[0],
                    "description": r[1] or "",
                    "submitted_at": r[2] or "",
                }
                for r in task_rows
            ],
        }
    else:
        task_stats = {"completed": 0, "tasks": []}

    # ── 5. 对话活跃度 ──
    chat_rows = execute_query(
        """SELECT date, COUNT(*) as cnt
           FROM conversations WHERE username = ?
           GROUP BY date ORDER BY date DESC""",
        (username,),
    )

    chat_stats = {"total_days": 0, "total_chats": 0, "avg_daily": 0, "recent_days": []}
    if chat_rows:
        total_days = len(chat_rows)
        total_chats = sum(r[1] for r in chat_rows)
        chat_stats = {
            "total_days": total_days,
            "total_chats": total_chats,
            "avg_daily": round(total_chats / max(total_days, 1), 1),
            "recent_days": [
                {"date": r[0], "count": r[1]}
                for r in chat_rows[:30]
            ],
        }

    # ── 7. 课程练习（知识点练习） ──
    # 注：ai_practice_results 在 questions.db 中，knowledge_points 在 smartkb.db 中，需分开查询
    cp_rows = q_execute_query(
        """SELECT id, kp_id, score, total_score, accuracy, evaluation, submitted_at
           FROM ai_practice_results
           WHERE student_username = ?
           ORDER BY submitted_at DESC""",
        (username,),
    )

    course_practice_stats = {"total_count": 0, "avg_accuracy": 0, "total_score": 0, "records": []}
    if cp_rows:
        # 批量查询知识点名称
        kp_ids = list(set(r['kp_id'] for r in cp_rows))
        kp_name_map = {}
        if kp_ids:
            placeholders = ",".join("?" for _ in kp_ids)
            kp_rows = execute_query(
                f"SELECT id, name FROM knowledge_points WHERE id IN ({placeholders})",
                tuple(kp_ids),
            )
            for kr in kp_rows:
                kp_name_map[kr[0]] = kr[1]

        total_practices = len(cp_rows)
        avg_accuracy = round(
            sum(r['accuracy'] for r in cp_rows) / total_practices, 1
        )
        total_score_sum = sum(r['score'] for r in cp_rows)
        course_practice_stats = {
            "total_count": total_practices,
            "avg_accuracy": avg_accuracy,
            "total_score": total_score_sum,
            "records": [
                {
                    "id": r['id'],
                    "kp_name": kp_name_map.get(r['kp_id'], f"知识点#{r['kp_id']}"),
                    "score": r['score'],
                    "total_score": r['total_score'],
                    "accuracy": r['accuracy'],
                    "evaluation": r['evaluation'],
                    "submitted_at": r['submitted_at'],
                }
                for r in cp_rows
            ],
        }

    # ── 8. 综合摘要 ──
    summary_parts = []
    if exam_stats:
        summary_parts.append(
            f"参加了 {exam_stats['total_exams']} 场考试，"
            f"平均分 {exam_stats['avg_percentage']}%，"
            f"通过 {exam_stats['passed_count']} 场"
        )
    if score_stats:
        summary_parts.append(
            f"累计获得 {score_stats['total_score']} 课堂积分"
        )
    if reward_points:
        summary_parts.append(
            f"获得 {reward_points} 奖励积分"
        )
    if rollcall_stats:
        summary_parts.append(
            f"被点名 {rollcall_stats['total_calls']} 次，"
            f"正确率 {rollcall_stats['accuracy']}%"
        )
    if task_stats:
        summary_parts.append(
            f"完成 {task_stats['completed']} 个任务"
        )
    if chat_stats:
        summary_parts.append(
            f"AI 对话活跃 {chat_stats['total_days']} 天，共 {chat_stats['total_chats']} 次"
        )
    if course_practice_stats:
        summary_parts.append(
            f"完成 {course_practice_stats['total_count']} 个课程练习，"
            f"平均正确率 {course_practice_stats['avg_accuracy']}%"
        )

    return {
        "user": user_info,
        "summary": " | ".join(summary_parts) if summary_parts else "暂无学习数据",
        "exams": {
            "results": exam_results,
            "stats": exam_stats,
        },
        "scores": score_stats,
        "reward_points": reward_points,
        "reward_history": reward_history,
        "rollcall": rollcall_stats,
        "tasks": task_stats,
        "chats": chat_stats,
        "course_practice": course_practice_stats,
    }


@router.get("/{username}/timeline", summary="获取学生成长时间轴")
async def get_timeline(username: str, request: Request):
    """获取学生成长时间轴事件"""
    user = get_current_user(request)
    current_username = user["username"]
    role = user.get("role", 2)

    if role == 2 and current_username != username:
        raise HTTPException(status_code=403, detail="无权查看")

    user_rows = execute_query(
        "SELECT name FROM users WHERE username = ?",
        (username,),
    )
    student_name = user_rows[0][0] if user_rows and user_rows[0][0] else username

    events = []

    # 考试事件
    exam_events = q_execute_query(
        """SELECT ea.submitted_at, e.title, ea.score, e.total_score, e.pass_score
           FROM exam_attempts ea
           JOIN exams e ON ea.exam_id = e.id
           WHERE ea.student_username = ? AND ea.submitted_at IS NOT NULL
           ORDER BY ea.submitted_at ASC""",
        (username,),
    )
    for ev in exam_events:
        passed = ev['score'] >= ev['pass_score']
        events.append({
            "time": ev['submitted_at'],
            "type": "exam",
            "title": f"完成了考试「{ev['title']}」",
            "detail": f"得分 {ev['score']}/{ev['total_score']} {'✅ 通过' if passed else '❌ 未通过'}",
            "icon": "exam",
        })

    # 积分事件
    score_events = execute_query(
        """SELECT updated_at, score, teacher_username, class_name
           FROM scores WHERE student_name = ?
           ORDER BY updated_at ASC""",
        (student_name,),
    )
    for ev in score_events:
        if ev[0]:
            events.append({
                "time": ev[0],
                "type": "score",
                "title": f"课堂积分变动",
                "detail": f"{'获得' if ev[1] > 0 else '扣除'} {abs(ev[1])} 分 (由 {ev[2]})",
                "icon": "score",
            })

    # 点名事件
    rc_events = execute_query(
        """SELECT created_at, result, points, teacher_username
           FROM rollcall_history WHERE student_name = ?
           ORDER BY created_at ASC""",
        (student_name,),
    )
    for ev in rc_events:
        if ev[0]:
            result_label = "正确" if ev[1] == "1" else ("错误" if ev[1] == "0" else "待定")
            events.append({
                "time": ev[0],
                "type": "rollcall",
                "title": "被点名",
                "detail": f"回答{result_label}，积分变动 {ev[2] or 0}",
                "icon": "rollcall",
            })

    # 任务事件
    task_events = execute_query(
        """SELECT ts.submitted_at, t.name
           FROM task_submissions ts
           JOIN tasks t ON ts.task_id = t.id
           WHERE ts.student_username = ?
           ORDER BY ts.submitted_at ASC""",
        (username,),
    )
    for ev in task_events:
        if ev[0]:
            events.append({
                "time": ev[0],
                "type": "task",
                "title": f"提交了任务「{ev[1]}」",
                "detail": "",
                "icon": "task",
            })

    # 课程练习事件（ai_practice_results 在 questions.db，knowledge_points 在 smartkb.db）
    cp_raw = q_execute_query(
        """SELECT kp_id, submitted_at, score, total_score, accuracy
           FROM ai_practice_results
           WHERE student_username = ? AND submitted_at IS NOT NULL
           ORDER BY submitted_at ASC""",
        (username,),
    )
    if cp_raw:
        cp_kp_ids = list(set(r['kp_id'] for r in cp_raw))
        cp_kp_map = {}
        if cp_kp_ids:
            ph = ",".join("?" for _ in cp_kp_ids)
            kp_rows = execute_query(
                f"SELECT id, name FROM knowledge_points WHERE id IN ({ph})",
                tuple(cp_kp_ids),
            )
            for kr in kp_rows:
                cp_kp_map[kr[0]] = kr[1]
        for ev in cp_raw:
            if ev['submitted_at']:
                kp_name = cp_kp_map.get(ev['kp_id'], f"知识点#{ev['kp_id']}")
                passed = ev['accuracy'] >= 60
                events.append({
                    "time": ev['submitted_at'],
                    "type": "exam",
                    "title": f"完成了课程练习「{kp_name}」",
                    "detail": f"得分 {ev['score']}/{ev['total_score']} · 正确率 {ev['accuracy']}% {'✅' if passed else '🔄'}",
                    "icon": "practice",
                })

    # 按时间排序
    events.sort(key=lambda x: x["time"] or "")
    return events


# ═══════════════════════════════════════════════════════════
# V3.3 新增：AI 学习报告
# ═══════════════════════════════════════════════════════════

@router.get("/{username}/report", summary="AI 生成学生学习报告")
async def get_learning_report(username: str, request: Request):
    """AI 根据学生数据生成个性化学习报告"""
    user = get_current_user(request)
    current_username = user["username"]
    role = user.get("role", 2)

    if role == 2 and current_username != username:
        raise HTTPException(status_code=403, detail="无权查看其他学生的报告")

    # 获取学生信息
    user_rows = execute_query(
        "SELECT username, name, class, grade FROM users WHERE username = ?",
        (username,),
    )
    if not user_rows:
        raise HTTPException(status_code=404, detail="学生不存在")
    student_name = user_rows[0][1] or user_rows[0][0]
    student_grade = user_rows[0][3] or ""
    student_class = user_rows[0][2] or ""

    # 获取报告周期参数
    days = int(request.query_params.get("days", 30))
    period = request.query_params.get("period", f"近{days}天")
    since = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")

    # ── 1. 考试成绩（周期内） ──
    exam_results = q_execute_query(
        """SELECT e.title, ea.score, ea.total_score, e.pass_score, ea.submitted_at, e.subject
           FROM exam_attempts ea
           JOIN exams e ON ea.exam_id = e.id
           WHERE ea.student_username = ? AND ea.status IN ('submitted', 'graded')
           AND ea.submitted_at >= ?
           ORDER BY ea.submitted_at ASC""",
        (username, since),
    )

    exam_text = ""
    if exam_results:
        for r in exam_results:
            pct = round(r['score'] / max(r['total_score'], 1) * 100, 1)
            passed = "通过" if r['score'] >= r['pass_score'] else "未通过"
            exam_text += f"- 《{r['title']}》({r['subject']}): {r['score']}/{r['total_score']}分 ({pct}%) {passed}\n"
    else:
        exam_text = "周期内暂无考试记录"

    # ── 2. 积分 ──
    score_rows = execute_query(
        """SELECT COALESCE(SUM(score),0), COALESCE(AVG(score),0),
                  COUNT(*), COALESCE(MAX(score),0), COALESCE(MIN(score),0)
           FROM scores WHERE student_name = ? AND updated_at >= ?""",
        (student_name, since),
    )
    sr = score_rows[0] if score_rows else (0, 0, 0, 0, 0)
    total_score = sr[0]
    score_count = sr[2]

    # 积分趋势（scores 表用 updated_at）
    score_trend_rows = execute_query(
        """SELECT DATE(updated_at), SUM(score)
           FROM scores WHERE student_name = ? AND updated_at >= ?
           GROUP BY DATE(updated_at) ORDER BY DATE(updated_at)""",
        (student_name, since),
    )
    if score_trend_rows:
        scores_list = [r[1] for r in score_trend_rows]
        if len(scores_list) >= 2:
            if scores_list[-1] > scores_list[0]:
                score_trend = "上升趋势 📈"
            elif scores_list[-1] < scores_list[0]:
                score_trend = "下降趋势 📉"
            else:
                score_trend = "保持稳定 ➡️"
        else:
            score_trend = "数据较少"
    else:
        score_trend = "暂无数据"

    # ── 3. 点名 ──
    rc_rows = execute_query(
        """SELECT COUNT(*), COALESCE(SUM(CASE WHEN result='1' THEN 1 ELSE 0 END), 0)
           FROM rollcall_history WHERE student_name = ? AND created_at >= ?""",
        (student_name, since),
    )
    rc_total = rc_rows[0][0] if rc_rows else 0
    rc_correct = rc_rows[0][1] if rc_rows else 0

    # ── 4. 任务 ──
    task_rows = execute_query(
        "SELECT COUNT(*) FROM task_submissions WHERE student_username = ? AND submitted_at >= ?",
        (username, since),
    )
    task_count = task_rows[0][0] if task_rows else 0

    # ── 5. 对话 ──
    chat_rows = execute_query(
        """SELECT COUNT(DISTINCT date), COUNT(*)
           FROM conversations WHERE username = ? AND date >= ?""",
        (username, since),
    )
    chat_days = chat_rows[0][0] if chat_rows else 0
    chat_total = chat_rows[0][1] if chat_rows else 0

    # ── 构建 Prompt ──
    from backend.prompts.report import LEARNING_REPORT_PROMPT
    from backend.api.chat_router import get_api_keys
    from backend.api.ai_service import call_ai_async

    keys = get_api_keys(username if role == 2 else current_username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    def _safe(s):
        return str(s).replace('{', '{{').replace('}', '}}')

    ai_role = build_ai_role(grade=student_grade)
    prompt = f"{ai_role}" + LEARNING_REPORT_PROMPT.format(
        student_name=_safe(student_name),
        grade=_safe(student_grade),
        cls=_safe(student_class),
        period=_safe(period),
        exam_text=_safe(exam_text),
        total_score=_safe(total_score),
        score_trend=_safe(score_trend),
        rollcall_total=_safe(rc_total),
        rollcall_correct=_safe(rc_correct),
        rollcall_rate=_safe(round(rc_correct / max(rc_total, 1) * 100, 1)),
        task_count=_safe(task_count),
        chat_days=_safe(chat_days),
        chat_total=_safe(chat_total),
    )

    from backend.ai_task_manager import task_manager
    _student_username = username

    async def _do_report() -> dict[str, Any]:
        try:
            result = await call_ai_async(prompt, api_key)
            return {
                "report": result,
                "student": {"username": _student_username, "name": student_name, "grade": student_grade, "class": student_class},
                "period": period,
                "data": {
                    "exams": len(exam_results),
                    "total_score": total_score,
                    "rollcall_rate": round(rc_correct / max(rc_total, 1) * 100, 1),
                    "tasks": task_count,
                    "chat_days": chat_days,
                },
            }
        except Exception as e:
            logger.error(f"AI 学习报告生成失败: {e}")
            return {"error": f"生成学习报告失败: {str(e)}"}

    task_id = await task_manager.create_task(description="AI 学习报告", coro_factory=_do_report)
    return {"task_id": task_id, "message": "AI 学习报告已提交，请稍后查询结果"}


# ═══════════════════════════════════════════════════════════
# V3.3 新增：学习报告导出 Word 文档
# ═══════════════════════════════════════════════════════════

@router.get("/{username}/report/export", summary="导出学习报告为 Word 文档")
async def export_learning_report_docx(username: str, request: Request, token: str = Query("")):
    """导出 AI 学习报告为 Word 文档"""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse
    import urllib.parse

    # 支持 token 参数认证（用于 window.open 下载）
    if token:
        request.state.user = None
        from backend.auth import decode_jwt_token
        payload = decode_jwt_token(token)
        if payload:
            request.state.user = payload

    user = get_current_user(request)
    current_username = user["username"]
    role = user.get("role", 2)

    if role == 2 and current_username != username:
        raise HTTPException(status_code=403, detail="无权查看其他学生的报告")

    # 获取学生信息
    user_rows = execute_query(
        "SELECT username, name, class, grade FROM users WHERE username = ?",
        (username,),
    )
    if not user_rows:
        raise HTTPException(status_code=404, detail="学生不存在")
    student_name = user_rows[0][1] or user_rows[0][0]
    student_grade = user_rows[0][3] or ""
    student_class = user_rows[0][2] or ""

    days = int(request.query_params.get("days", 30))
    period = request.query_params.get("period", f"近{days}天")
    since = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")

    # ── 收集数据 ──
    exam_results = q_execute_query(
        """SELECT e.title, ea.score, ea.total_score, e.pass_score, ea.submitted_at, e.subject
           FROM exam_attempts ea JOIN exams e ON ea.exam_id = e.id
           WHERE ea.student_username = ? AND ea.status IN ('submitted', 'graded')
           AND ea.submitted_at >= ? ORDER BY ea.submitted_at ASC""",
        (username, since),
    )

    score_rows = execute_query(
        """SELECT COALESCE(SUM(score),0), COALESCE(AVG(score),0), COUNT(*), COALESCE(MAX(score),0), COALESCE(MIN(score),0)
           FROM scores WHERE student_name = ? AND updated_at >= ?""",
        (student_name, since),
    )
    sr = score_rows[0] if score_rows else (0, 0, 0, 0, 0)
    total_score = sr[0]

    rc_rows = execute_query(
        """SELECT COUNT(*), COALESCE(SUM(CASE WHEN result='1' THEN 1 ELSE 0 END), 0)
           FROM rollcall_history WHERE student_name = ? AND created_at >= ?""",
        (student_name, since),
    )
    rc_total = rc_rows[0][0] if rc_rows else 0
    rc_correct = rc_rows[0][1] if rc_rows else 0

    task_rows = execute_query(
        "SELECT COUNT(*) FROM task_submissions WHERE student_username = ? AND submitted_at >= ?",
        (username, since),
    )
    task_count = task_rows[0][0] if task_rows else 0

    chat_rows = execute_query(
        """SELECT COUNT(DISTINCT date), COUNT(*) FROM conversations WHERE username = ? AND date >= ?""",
        (username, since),
    )
    chat_days = chat_rows[0][0] if chat_rows else 0
    chat_total = chat_rows[0][1] if chat_rows else 0

    # ── 生成报告文本 ──
    from backend.prompts.report import LEARNING_REPORT_PROMPT
    from backend.api.chat_router import get_api_keys
    from backend.api.ai_service import call_ai_async

    keys = get_api_keys(username if role == 2 else current_username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    def _safe(s):
        return str(s).replace('{', '{{').replace('}', '}}')

    exam_text = ""
    if exam_results:
        for r in exam_results:
            pct = round(r['score'] / max(r['total_score'], 1) * 100, 1)
            passed = "通过" if r['score'] >= r['pass_score'] else "未通过"
            exam_text += f"- 《{r['title']}》({r['subject']}): {r['score']}/{r['total_score']}分 ({pct}%) {passed}\n"
    else:
        exam_text = "周期内暂无考试记录"

    score_trend = "暂无数据"

    ai_role = build_ai_role(grade=student_grade)
    prompt = f"{ai_role}" + LEARNING_REPORT_PROMPT.format(
        student_name=_safe(student_name),
        grade=_safe(student_grade),
        cls=_safe(student_class),
        period=_safe(period),
        exam_text=_safe(exam_text),
        total_score=_safe(total_score),
        score_trend=_safe(score_trend),
        rollcall_total=_safe(rc_total),
        rollcall_correct=_safe(rc_correct),
        rollcall_rate=_safe(round(rc_correct / max(rc_total, 1) * 100, 1)),
        task_count=_safe(task_count),
        chat_days=_safe(chat_days),
        chat_total=_safe(chat_total),
    )

    report_text = await call_ai_async(prompt, api_key)

    # ── 生成 Word 文档 ──
    doc = Document()
    style = doc.styles['Normal']  # type: ignore[union-attr]
    style.font.name = 'Microsoft YaHei'  # type: ignore[attr-defined]
    style.font.size = Pt(11)  # type: ignore[attr-defined]
    style.paragraph_format.line_spacing = 1.5  # type: ignore[attr-defined]

    title = doc.add_heading(f"{student_name} 的学习报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"{student_grade}{student_class}班  |  报告周期：{period}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- **') and '：' in line:
            content = line.lstrip('- ')
            p = doc.add_paragraph()
            bold_end = content.find('**', 2)
            if bold_end > 0:
                run = p.add_run(content[2:bold_end])
                run.bold = True
                p.add_run(content[bold_end + 2:])
            else:
                p.add_run(content)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f'{i}. ') for i in range(1, 10)):
            doc.add_paragraph(line, style='List Number')
        else:
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_filename = urllib.parse.quote(f"{student_name}_学习报告.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )
