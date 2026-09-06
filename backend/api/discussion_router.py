# -*- coding: utf-8 -*-
"""
分组讨论 API 路由
教师创建讨论、学生分组聊天、AI 助教参与、讨论报告
"""
import json
import random
import asyncio
from datetime import datetime
from typing import Any

from fastapi import APIRouter, HTTPException, Request, Query, WebSocket, WebSocketDisconnect
from pydantic import BaseModel

from backend.api.dependencies import get_current_user
from backend.database import execute_query, execute_insert_update, execute_query_dict
from backend.permission_service import (
    filter_activities_by_scope,
    check_activity_visibility,
    is_student_in_teacher_scope,
    parse_legacy_teacher_grade_class,
)
from backend.logger import logger
from backend.ws_manager import manager as ws_manager
from backend.prompts import apply_skills, build_ai_role
from backend.utils import extract_json_from_text

router = APIRouter()


# ═══════════════════════════════════════════════
# 权限助手(S1/S2/S6)
# ═══════════════════════════════════════════════

_DISC_COLS = ["id", "creator_username", "title", "status", "grade", "classes",
              "target_scope", "target_grade", "target_class", "target_users"]


def _disc_dict(disc_id: int) -> dict | None:
    """按列名读取讨论记录

    S11: 旧代码把硬编码列序 zip 到 SELECT * 上, 而 discussions 表的 created_at/updated_at
    位于 target_* 之前, 导致 target_scope/target_grade/... 全部错位,
    学生可见性判定拿到垃圾值后回落到"教师任教范围", 于是管理员建的"仅高一可见"讨论
    对全体年级放行。
    """
    rows = execute_query_dict("SELECT * FROM discussions WHERE id=?", (disc_id,))
    return rows[0] if rows else None


def _disc_brief(disc_id: int) -> dict | None:
    rows = execute_query_dict(
        "SELECT id, creator_username, title, status, grade, classes, target_scope, "
        "target_grade, target_class, target_users FROM discussions WHERE id=?",
        (disc_id,),
    )
    return rows[0] if rows else None


def _group_discussion_id(group_id: int) -> int | None:
    rows = execute_query("SELECT discussion_id FROM discussion_groups WHERE id=?", (group_id,))
    return rows[0][0] if rows else None


def _is_disc_member(username: str, disc_id: int) -> bool:
    return bool(execute_query(
        """SELECT 1 FROM discussion_members dm
           JOIN discussion_groups dg ON dm.group_id = dg.id
           WHERE dg.discussion_id=? AND dm.username=?""",
        (disc_id, username),
    ))


def _teacher_covers_discussion(teacher: str, disc: dict) -> bool:
    """教师可见判定: 创建者本人 / 组内有我任教的学生 / 目标年级在我任教范围内"""
    if not disc:
        return False
    if (disc.get("creator_username") or "") == teacher:
        return True
    members = execute_query(
        """SELECT DISTINCT dm.username FROM discussion_members dm
           JOIN discussion_groups dg ON dg.id = dm.group_id
           WHERE dg.discussion_id=?""",
        (disc.get("id"),),
    ) or []
    for m in members[:60]:
        try:
            if m[0] and is_student_in_teacher_scope(m[0], teacher):
                return True
        except Exception:
            continue
    trow = execute_query("SELECT grade, class FROM users WHERE username=?", (teacher,))
    if not trow:
        return False
    try:
        tmap = parse_legacy_teacher_grade_class(str(trow[0][0] or ""), str(trow[0][1] or ""))
    except Exception:
        tmap = {}
    targets = [g.strip() for g in str(disc.get("target_grade") or disc.get("grade") or "").replace("，", ",").split(",") if g.strip()]
    return bool(targets) and any(g in tmap for g in targets)


def _assert_can_access_discussion(user: dict, disc: dict | None, what: str = "讨论数据") -> None:
    """教师须在该讨论范围内, 学生须是成员, 管理员不限(S1/S6)"""
    role = user.get("role", 2)
    username = user.get("username", "")
    if role == 0:
        return
    if not disc:
        raise HTTPException(status_code=404, detail="讨论不存在")
    if role == 1:
        if _teacher_covers_discussion(username, disc):
            return
        raise HTTPException(status_code=403, detail=f"无权查看该讨论的{what}")
    if _is_disc_member(username, disc["id"]):
        return
    raise HTTPException(status_code=403, detail=f"无权访问该{what}")


def _assert_can_access_group(user: dict, group_id: int, what: str = "小组数据") -> None:
    disc_id = _group_discussion_id(group_id)
    if disc_id is None:
        raise HTTPException(status_code=404, detail="小组不存在")
    _assert_can_access_discussion(user, _disc_brief(disc_id), what)


def _assert_discussion_owner(user: dict, disc: dict | None, action: str = "") -> None:
    """教师仅能管理自己创建的讨论(与 update/delete/start/end 口径一致)"""
    if user.get("role", 2) == 1 and (disc or {}).get("creator_username") != user.get("username", ""):
        raise HTTPException(status_code=403, detail="只能管理自己的讨论")


def _student_scope_ok(user: dict, item: dict) -> bool:
    """学生可见性(S7): 声明按年级/班级共享但目标为空的讨论不应放行"""
    if user.get("role", 2) != 2:
        return True
    scope = (item.get("target_scope") or "teacher_classes")
    if scope in ("grade", "class") and not str(item.get("target_grade") or "").strip():
        return False
    srow = execute_query("SELECT grade, class FROM users WHERE username=?", (user.get("username", ""),))
    grade = str(srow[0][0] or "") if srow else ""
    cls = str(srow[0][1] or "") if srow else ""
    return check_activity_visibility(
        user.get("username", ""), grade, cls,
        item.get("creator_username", ""), scope,
        item.get("target_grade", ""), item.get("target_class", ""), item.get("target_users", ""),
    )


# ── AI 内容审核 + 重复/频率限制（发言合规检查） ──

import time as _dt_time
import hashlib as _dt_hashlib

# 内容审核跟踪：username -> { rejected_hashes: set, rejection_count: int, window_start: float, blocked_until: float }
_discussion_review_tracker: dict[str, dict[str, Any]] = {}
_DT_MAX_REJECTIONS = 3
_DT_WINDOW_SECONDS = 300
_DT_BLOCK_SECONDS = 60


def _check_discussion_rate_limit(username: str, content: str) -> tuple[bool, str]:
    """检查是否触发频率限制或重复提交"""
    now = _dt_time.time()
    tracker = _discussion_review_tracker.setdefault(username, {
        "rejected_hashes": set(),
        "rejection_count": 0,
        "window_start": now,
        "blocked_until": 0,
    })
    if now < tracker["blocked_until"]:
        remain = int(tracker["blocked_until"] - now)
        return False, f"提交过于频繁，请{remain}秒后再试"
    if now - tracker["window_start"] > _DT_WINDOW_SECONDS:
        tracker["rejected_hashes"] = set()
        tracker["rejection_count"] = 0
        tracker["window_start"] = now
    content_hash = _dt_hashlib.md5(content.encode("utf-8")).hexdigest()
    if content_hash in tracker["rejected_hashes"]:
        return False, "该内容与之前被拒绝的内容相同，请修改后重新提交"
    return True, ""


def _record_discussion_rejection(username: str, content: str):
    """记录一次审核拒绝并扣除 2 分"""
    now = _dt_time.time()
    tracker = _discussion_review_tracker.setdefault(username, {
        "rejected_hashes": set(),
        "rejection_count": 0,
        "window_start": now,
        "blocked_until": 0,
    })
    content_hash = _dt_hashlib.md5(content.encode("utf-8")).hexdigest()
    tracker["rejected_hashes"].add(content_hash)
    tracker["rejection_count"] += 1

    # 扣除 2 分
    try:
        from backend.reward_engine import deduct_points
        deduct_points(username, "讨论发言审核不通过", 2)
    except Exception as e:
        logger.warning(f"讨论审核扣分失败 (user={username}): {e}")

    if tracker["rejection_count"] >= _DT_MAX_REJECTIONS:
        tracker["blocked_until"] = now + _DT_BLOCK_SECONDS
        logger.warning(f"用户 {username} 触发审核频率限制，封禁 {_DT_BLOCK_SECONDS} 秒")


def _ai_content_review(content: str, username: str) -> tuple[bool, str]:
    """调用 AI 审核讨论发言是否合规，含重复提交和频率限制，返回 (是否通过, 原因)"""
    # 先检查频率限制和重复提交
    allowed, msg = _check_discussion_rate_limit(username, content)
    if not allowed:
        return False, msg

    prompt = (
        '你是一个课堂内容审核助手。请判断以下学生在分组讨论中的发言是否包含：\n'
        '1. 违反法律法规的内容\n'
        '2. 违反道德规范、社会公序良俗的内容\n'
        '3. 不文明用语、辱骂、攻击性、歧视性言论\n'
        '4. 色情、暴力、恐怖、血腥等内容\n'
        '5. 广告、垃圾信息\n\n'
        '学生发言：' + content + '\n\n'
        '请严格判断。只返回以下JSON格式（不要包含其他文字）：\n'
        '{"safe": true, "reason": ""} 或 {"safe": false, "reason": "简要说明违规原因（10字以内）"}'
    )
    api_key = _get_api_key()
    if not api_key:
        return True, ""

    try:
        import json, re, concurrent.futures
        from backend.api.ai_service import _ai_thread_pool
        future = _ai_thread_pool.submit(_call_review_sync, prompt, api_key)
        try:
            result = future.result(timeout=15)
        except concurrent.futures.TimeoutError:
            logger.warning(f"AI 内容审核超时（15秒），已放行: user={username}")
            return True, ""
        jm = re.search(r'\{[^}]+\}', result)
        if jm:
            data = json.loads(jm.group())
            if not data.get("safe", True):
                _record_discussion_rejection(username, content)
                return False, data.get("reason", "内容不合规")
        return True, ""
    except Exception:
        return True, ""


def _call_review_sync(prompt: str, api_key: str) -> str:
    """内容审核专用同步调用"""
    from backend.api.ai_service import call_ai_sync_direct
    return call_ai_sync_direct(prompt, api_key)


# ── 辅助函数 ──

def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


# ── 请求/响应模型 ──

class DiscussionCreate(BaseModel):
    title: str
    description: str = ""
    subject: str = ""
    group_mode: str = "none"       # auto / manual / random / none
    group_count: int = 0
    members_per_group: int = 4
    ai_role: str = "mixed"         # observer / guide / proactive / judge / mixed
    duration_minutes: int = 30
    grade: str = ""
    classes: str = ""
    require_summary: bool = False
    target_scope: str = "teacher_classes"
    target_grade: str = ""
    target_class: str = ""
    target_users: str = ""


class DiscussionUpdate(BaseModel):
    title: str | None = None
    description: str | None = None
    subject: str | None = None
    duration_minutes: int | None = None
    ai_role: str | None = None
    require_summary: bool | None = None


class MessageSend(BaseModel):
    content: str


class BroadcastMessage(BaseModel):
    content: str


class AiGenerateDiscussion(BaseModel):
    topic: str
    subject: str = ""  # 由前端传递
    group_mode: str = "none"
    ai_role: str = "mixed"
    duration_minutes: int = 30


# ── 创建讨论 ──

@router.post("/discussions", summary="创建讨论活动")
async def create_discussion(req: DiscussionCreate, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可创建讨论")

    if not req.title.strip():
        raise HTTPException(status_code=400, detail="请输入讨论标题")
    if req.duration_minutes < 1 or req.duration_minutes > 120:
        raise HTTPException(status_code=400, detail="讨论时长范围为 1-120 分钟")
    if req.group_count < 0 or req.members_per_group < 1:
        raise HTTPException(status_code=400, detail="分组参数无效")

    now = _now()
    disc_id = execute_insert_update(
        """INSERT INTO discussions
           (creator_username, title, description, subject, group_mode,
            group_count, members_per_group, ai_role, duration_minutes,
            status, grade, classes, require_summary,
            target_scope, target_grade, target_class, target_users,
            created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?, ?, ?,
                   ?, ?, ?, ?, ?, ?)""",
        (user["username"], req.title, req.description, req.subject,
         req.group_mode, req.group_count, req.members_per_group,
         req.ai_role, req.duration_minutes,
         req.grade, req.classes, 1 if req.require_summary else 0,
         req.target_scope, req.target_grade, req.target_class, req.target_users,
         now, now),
    )
    logger.info(f"教师 {user['username']} 创建讨论: {req.title}")
    return {"status": "ok", "discussion_id": disc_id}


# ── AI 生成讨论方案 ──

@router.post("/discussions/ai-generate", summary="AI 自动生成讨论方案")
async def ai_generate_discussion(req: AiGenerateDiscussion, request: Request):
    """AI 根据主题自动生成讨论方案"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可用")

    ai_role_desc = {
        "observer": "旁观者，不主动发言",
        "guide": "适时引导讨论方向的引导者",
        "proactive": "主动参与讨论并提供观点的参与者",
        "judge": "辩论裁判，分析各方论点",
        "mixed": "综合角色：根据讨论情况自动切换角色——需要观察时保持旁观，需要引导时提出启发式问题，需要参与时主动提供观点，需要裁判时分析各方论点",
    }.get(req.ai_role, "引导者")

    from backend.prompts.discussion import DISCUSSION_PLAN_PROMPT
    ai_role = build_ai_role(subject=req.subject)
    prompt = f"{ai_role}\n" + DISCUSSION_PLAN_PROMPT.format(
        subject=req.subject,
        topic=req.topic,
        group_mode=req.group_mode,
        ai_role_desc=ai_role_desc,
        duration_minutes=req.duration_minutes,
    )
    # 注意：不注入技能 — 技能的结构化输出指令与 JSON 格式要求冲突

    import json, re
    from backend.api.chat_router import get_api_keys
    api_key, _ = get_api_keys(user["username"])
    if not api_key:
        return {"status": "error", "content": "AI 功能不可用：请配置 DashScope API Key"}

    from backend.api.ai_service import call_ai_async
    from backend.ai_task_manager import task_manager

    async def _do_generate() -> dict[str, Any]:
        try:
            result = await call_ai_async(prompt, api_key)
            if result:
                data = extract_json_from_text(result)
                if data:
                    return {"status": "ok", "data": data, "raw": result}
                return {"status": "error", "content": result, "raw": result}
            return {"status": "error", "content": "AI 未返回有效结果"}
        except Exception as e:
            logger.warning(f"AI 生成讨论方案失败: {e}")
            return {"status": "error", "content": f"AI 调用出错: {str(e)}"}

    task_id = await task_manager.create_task(description="AI 生成讨论方案", coro_factory=_do_generate)
    return {"task_id": task_id, "message": "AI 生成已提交，请稍后查询结果"}


# ── 获取讨论列表 ──

def _group_col_counts(table: str, ids: list[int]) -> dict[int, int]:
    """按 group_id 批量计数(S7: 取代逐组 COUNT 查询)"""
    if not ids:
        return {}
    ph = ",".join("?" * len(ids))
    rows = execute_query(
        f"SELECT group_id, COUNT(*) FROM {table} WHERE group_id IN ({ph}) GROUP BY group_id",
        tuple(ids),
    )
    return {r[0]: r[1] for r in rows}


def _enrich_discussions(rows: list[dict], username: str, role: int) -> list[dict]:
    """批量富化讨论列表: 分组/消息数/成员数/创建者姓名/报告标记/我的小组"""
    if not rows:
        return []
    disc_ids = [r["id"] for r in rows]
    ph = ",".join("?" * len(disc_ids))

    g_rows = execute_query(
        f"""SELECT id, discussion_id, group_index, name FROM discussion_groups
            WHERE discussion_id IN ({ph}) ORDER BY discussion_id, group_index""",
        tuple(disc_ids),
    )
    group_ids = [g[0] for g in g_rows]
    msg_counts = _group_col_counts("discussion_messages", group_ids)
    mem_counts = _group_col_counts("discussion_members", group_ids)
    groups_by_disc: dict[int, list] = {}
    for gid, did, gidx, gname in g_rows:
        groups_by_disc.setdefault(did, []).append({
            "id": gid, "group_index": gidx, "name": gname,
            "member_count": mem_counts.get(gid, 0),
            "message_count": msg_counts.get(gid, 0),
        })

    creator_names: dict[str, str] = {}
    creators = [r.get("creator_username") for r in rows if r.get("creator_username")]
    if creators:
        cph = ",".join("?" * len(set(creators)))
        for uname, cname in execute_query(
            f"SELECT username, name FROM users WHERE username IN ({cph})", tuple(set(creators))
        ):
            creator_names[uname] = cname or ""

    summary_counts: dict[int, int] = {}
    s_rows = execute_query(
        f"""SELECT discussion_id, COUNT(*) FROM discussion_reports
            WHERE discussion_id IN ({ph}) AND group_id IS NOT NULL GROUP BY discussion_id""",
        tuple(disc_ids),
    )
    summary_counts = {r[0]: r[1] for r in s_rows}

    my_groups: dict[int, tuple] = {}
    if role == 2:
        for did, gid, gidx, gname in execute_query(
            f"""SELECT dg.discussion_id, dg.id, dg.group_index, dg.name
                FROM discussion_members dm JOIN discussion_groups dg ON dm.group_id = dg.id
                WHERE dg.discussion_id IN ({ph}) AND dm.username=?""",
            tuple(disc_ids) + (username,),
        ):
            my_groups.setdefault(did, (gid, gidx, gname))

    results = []
    for item in rows:
        item = dict(item)
        gl = groups_by_disc.get(item["id"], [])
        item["groups"] = gl
        item["total_messages"] = sum(g["message_count"] for g in gl)
        item["total_members"] = sum(g["member_count"] for g in gl)
        item["require_summary"] = bool(item.get("require_summary"))
        creator = item.get("creator_username") or ""
        item["creator_name"] = creator_names.get(creator) or creator
        item["has_summary"] = summary_counts.get(item["id"], 0) > 0
        if role == 2:
            mine = my_groups.get(item["id"])
            item["has_joined"] = mine is not None
            if mine:
                item["my_group"] = {"id": mine[0], "group_index": mine[1], "name": mine[2]}
        results.append(item)
    return results


@router.get("/discussions", summary="获取讨论列表")
async def list_discussions(
    request: Request,
    status: str | None = Query(None, description="筛选状态: pending/active/ended"),
):
    user = get_current_user(request)
    role = user.get("role", 2)
    username = user["username"]

    # ── 先按可见性取出讨论(S7: 学生先过滤再富化, 避免逐条 5 类查询) ──
    conds, params = [], []
    if status:
        conds.append("status=?")
        params.append(status)
    if role == 1:
        # 教师：只查看自己创建的讨论
        conds.append("creator_username=?")
        params.append(username)
    sql = "SELECT * FROM discussions"
    if conds:
        sql += " WHERE " + " AND ".join(conds)
    sql += " ORDER BY created_at DESC"
    rows = execute_query_dict(sql, tuple(params))

    if role == 2:
        # 学生：声明按年级/班级共享但目标为空的讨论不放行, 再走统一可见性过滤
        rows = [r for r in rows if _student_scope_ok(user, r)]
        rows = filter_activities_by_scope(rows, username)

    return _enrich_discussions(rows, username, role)


# ── 获取讨论详情 ──

@router.get("/discussions/{disc_id}", summary="获取讨论详情")
async def get_discussion(disc_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)
    disc = _disc_dict(disc_id)
    if disc is None:
        raise HTTPException(status_code=404, detail="讨论不存在")

    # 教师只能查看自己的讨论，管理员可以全部
    if role == 1 and disc.get("creator_username") != user["username"]:
        raise HTTPException(status_code=403, detail="只能查看自己创建的讨论详情")

    # 学生只能查看自己范围内的讨论(S11: 已加入本讨论的学生始终可查看)
    if role == 2 and not _is_disc_member(user["username"], disc_id):
        if not _student_scope_ok(user, disc):
            raise HTTPException(status_code=403, detail="无权查看该讨论")

    # 获取所有小组
    groups = execute_query(
        "SELECT id, group_index, name FROM discussion_groups WHERE discussion_id=? ORDER BY group_index",
        (disc_id,),
    )
    group_list = []
    for g in groups:
        members = execute_query(
            "SELECT username, role, joined_at FROM discussion_members WHERE group_id=?",
            (g[0],),
        )
        msg_count = execute_query(
            "SELECT COUNT(*) FROM discussion_messages WHERE group_id=?",
            (g[0],),
        )[0][0]
        group_list.append({
            "id": g[0], "group_index": g[1], "name": g[2],
            "members": [{"username": m[0], "role": m[1], "joined_at": m[2]} for m in members],
            "message_count": msg_count,
        })

    # 小组列表此前只有学号, 教师看不出是谁/哪个班: 统一补姓名+年级+班级
    from backend.permission_service import attach_student_info
    for _g in group_list:
        attach_student_info(_g.get("members") or [], key="username", prefix="")
    disc["groups"] = group_list
    disc["require_summary"] = bool(disc["require_summary"])
    # 查询创建者姓名
    creator_info = execute_query(
        "SELECT name FROM users WHERE username=?",
        (disc["creator_username"],),
    )
    disc["creator_name"] = creator_info[0][0] if creator_info and creator_info[0] and creator_info[0][0] else disc["creator_username"]
    return disc


# ── 更新讨论 ──

@router.put("/discussions/{disc_id}", summary="更新讨论")
async def update_discussion(disc_id: int, req: DiscussionUpdate, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)

    _assert_can_access_discussion(user, _disc_brief(disc_id), "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可编辑讨论")

    rows = execute_query("SELECT creator_username FROM discussions WHERE id=?", (disc_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="讨论不存在")
    if rows[0][0] != user["username"] and role != 0:
        raise HTTPException(status_code=403, detail="只能编辑自己的讨论")

    updates = []
    params = []
    for field in ["title", "description", "subject", "duration_minutes", "ai_role"]:
        val = getattr(req, field, None)
        if val is not None:
            updates.append(f"{field}=?")
            params.append(val)
    if req.require_summary is not None:
        updates.append("require_summary=?")
        params.append(1 if req.require_summary else 0)

    if updates:
        updates.append("updated_at=?")
        params.append(_now())
        params.append(disc_id)
        execute_insert_update(
            f"UPDATE discussions SET {', '.join(updates)} WHERE id=?",
            tuple(params),
        )

    return {"status": "ok"}


# ── 开始讨论（自动分组）──

@router.post("/discussions/{disc_id}/start", summary="开始讨论")
async def start_discussion(disc_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可开始讨论")

    disc = _disc_dict(disc_id)
    if disc is None:
        raise HTTPException(status_code=404, detail="讨论不存在")

    # 教师只能管理自己的讨论
    if role == 1 and disc.get("creator_username") != user["username"]:
        raise HTTPException(status_code=403, detail="只能管理自己的讨论")

    if disc["status"] != "pending":
        raise HTTPException(status_code=400, detail="讨论已开始或已结束")

    # 计算分组数量
    group_mode = disc["group_mode"]
    group_count = disc["group_count"]
    members_per_group = disc["members_per_group"]

    if group_mode == "none":
        # 不分组模式：只创建 1 个自由讨论区
        group_count = 1
    elif group_count <= 0 and members_per_group > 0:
        # 根据 members_per_group 估算组数（后续自动分配）
        group_count = members_per_group * 2  # 默认按每组人数*2 估算组数
    elif group_count <= 0:
        group_count = 4  # 默认 4 组

    # 删除旧的临时分组（如果有）
    old_groups = execute_query("SELECT id FROM discussion_groups WHERE discussion_id=?", (disc_id,))
    for og in old_groups:
        execute_insert_update("DELETE FROM discussion_members WHERE group_id=?", (og[0],))
    execute_insert_update("DELETE FROM discussion_groups WHERE discussion_id=?", (disc_id,))

    # 创建空分组，学生后续加入时会自动分配到人数最少的组
    now = _now()
    for idx in range(group_count):
        group_name = "自由讨论区" if group_mode == "none" else f"第{idx + 1}组"
        gid = execute_insert_update(
            "INSERT INTO discussion_groups (discussion_id, group_index, name) VALUES (?, ?, ?)",
            (disc_id, idx + 1, group_name),
        )

    # 更新状态
    execute_insert_update(
        "UPDATE discussions SET status='active', updated_at=? WHERE id=?",
        (_now(), disc_id),
    )

    logger.info(f"教师 {user['username']} 开始了讨论 {disc['title']}, 共 {group_count} 个小组")
    if group_mode == "none":
        return {"status": "ok", "message": "讨论已开始，自由讨论区已创建"}
    return {"status": "ok", "message": f"讨论已开始，共 {group_count} 个小组"}


# ── 结束讨论 ──

@router.post("/discussions/{disc_id}/end", summary="结束讨论")
async def end_discussion(disc_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可结束讨论")

    rows = execute_query("SELECT creator_username FROM discussions WHERE id=?", (disc_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="讨论不存在")
    if role == 1 and rows[0][0] != user["username"]:
        raise HTTPException(status_code=403, detail="只能管理自己的讨论")

    execute_insert_update(
        "UPDATE discussions SET status='ended', updated_at=? WHERE id=?",
        (_now(), disc_id),
    )
    logger.info(f"教师 {user['username']} 结束了讨论 #{disc_id}")

    # 自动生成报告
    try:
        # 在后台异步生成报告
        asyncio.create_task(_auto_generate_report(disc_id))
    except Exception as e:
        logger.warning(f"自动生成报告失败: {e}")

    return {"status": "ok", "message": "讨论已结束，报告生成中"}


async def _auto_generate_report(disc_id: int):
    """异步生成讨论报告（结束讨论时自动调用）"""
    from backend.api.config_router import get_config_value, load_config
    import os

    disc = _disc_dict(disc_id)
    if disc is None:
        return
    # 检查是否要求生成报告
    if not disc.get("require_summary"):
        logger.info(f"讨论 #{disc_id} 未开启总结报告，跳过自动生成")
        return

    groups = execute_query(
        "SELECT id, group_index, name FROM discussion_groups WHERE discussion_id=? ORDER BY group_index",
        (disc_id,),
    )

    now = _now()
    overall_parts = [f"# 讨论报告：{disc['title']}\n"]

    for g in groups:
        members = execute_query(
            "SELECT username FROM discussion_members WHERE group_id=?",
            (g[0],),
        )
        msgs = execute_query(
            "SELECT username, content, msg_type FROM discussion_messages WHERE group_id=? ORDER BY id ASC",
            (g[0],),
        )
        member_names = [m[0] for m in members]
        overall_parts.append(f"\n## {g[2] or f'第{g[1]}组'}")
        overall_parts.append(f"- 成员：{', '.join(member_names) if member_names else '（空）'}")
        overall_parts.append(f"- 消息数：{len(msgs)}")

        ai_texts = [m[1] for m in msgs if m[2] == 'ai_suggest']
        if ai_texts:
            overall_parts.append(f"- AI 介入次数：{len(ai_texts)}")

        # 生成小组 AI 分析摘要（使用新版归纳总结提示词）
        if msgs:
            api_key = _get_api_key()
            if api_key:
                # 组装消息文本
                messages_text = "\n".join(
                    [f"{m[0] or 'AI助教'}: {m[1]}" for m in msgs[-50:]]
                )

                from backend.prompts.discussion import DISCUSSION_AI_SUMMARY_PROMPT
                group_name = g[2] or f"第{g[1]}组"
                ai_role = build_ai_role(subject=disc.get("subject", ""))
                prompt = f"{ai_role}\n" + DISCUSSION_AI_SUMMARY_PROMPT.format(
                    subject=disc.get("subject") or "",
                    title=disc["title"],
                    group_name=group_name,
                    description=disc.get("description") or "",
                    messages_text=messages_text,
                )
                # 注意：不注入技能 — 技能的结构化输出指令与 JSON 格式要求冲突

                try:
                    from backend.api.ai_service import call_ai_async
                    summary = await call_ai_async(prompt, api_key)
                    if summary:
                        # 尝试解析 JSON
                        import json as _json
                        parsed = extract_json_from_text(summary)

                        if parsed:
                            s = parsed.get("summary", "")
                            kp = parsed.get("key_points", [])
                            ac = parsed.get("ai_comment", "")
                            sc = parsed.get("score", "")
                            overall_parts.append(f"\n**AI 分析**：{s}")
                            if kp:
                                overall_parts.append(f"- 关键观点：{'；'.join(kp)}")
                            if ac:
                                overall_parts.append(f"- AI 评价：{ac}")
                            if sc:
                                overall_parts.append(f"- 综合评分：{sc}/10")
                        else:
                            overall_parts.append(f"\n**AI 分析**：{summary[:200]}")

                        # 保存结构化报告
                        report_data = {"raw_content": summary, "parsed": parsed, "generated_at": now}
                        report_json = _json.dumps(report_data, ensure_ascii=False)
                        execute_insert_update(
                            "INSERT INTO discussion_reports (discussion_id, group_id, report_content, generated_at) VALUES (?, ?, ?, ?)",
                            (disc_id, g[0], report_json, now),
                        )
                except Exception:
                    pass

    overall_report = "\n".join(overall_parts)
    execute_insert_update(
        "INSERT INTO discussion_reports (discussion_id, group_id, report_content, generated_at) VALUES (?, NULL, ?, ?)",
        (disc_id, overall_report, now),
    )
    logger.info(f"讨论 #{disc_id} 报告自动生成完成")


# ── 重新开始讨论 ──

@router.post("/discussions/{disc_id}/restart", summary="重新开始讨论")
async def restart_discussion(disc_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)

    _assert_can_access_discussion(user, _disc_brief(disc_id), "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可操作")

    rows = execute_query("SELECT status, creator_username FROM discussions WHERE id=?", (disc_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="讨论不存在")
    if role == 1 and rows[0][1] != user["username"]:
        raise HTTPException(status_code=403, detail="只能管理自己的讨论")
    if rows[0][0] != "ended":
        raise HTTPException(status_code=400, detail="仅已结束的讨论可重新开始")

    execute_insert_update(
        "UPDATE discussions SET status='active', updated_at=? WHERE id=?",
        (_now(), disc_id),
    )
    logger.info(f"教师 {user['username']} 重新开始了讨论 #{disc_id}")
    return {"status": "ok", "message": "讨论已重新开始"}


# ── 删除讨论 ──

@router.delete("/discussions/{disc_id}", summary="删除讨论")
async def delete_discussion(disc_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)

    _assert_can_access_discussion(user, _disc_brief(disc_id), "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可删除讨论")

    rows = execute_query("SELECT creator_username FROM discussions WHERE id=?", (disc_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="讨论不存在")
    if rows[0][0] != user["username"] and role != 0:
        raise HTTPException(status_code=403, detail="只能删除自己的讨论")

    # 删除关联数据
    groups = execute_query("SELECT id FROM discussion_groups WHERE discussion_id=?", (disc_id,))
    for g in groups:
        execute_insert_update("DELETE FROM discussion_messages WHERE group_id=?", (g[0],))
        execute_insert_update("DELETE FROM discussion_members WHERE group_id=?", (g[0],))
    execute_insert_update("DELETE FROM discussion_groups WHERE discussion_id=?", (disc_id,))
    execute_insert_update("DELETE FROM discussion_reports WHERE discussion_id=?", (disc_id,))
    execute_insert_update("DELETE FROM activity_rewards WHERE activity_type='discussion' AND activity_id=?", (str(disc_id),))
    execute_insert_update("DELETE FROM notifications WHERE source_type='discussion' AND source_id=?", (str(disc_id),))
    execute_insert_update("DELETE FROM discussions WHERE id=?", (disc_id,))

    logger.info(f"教师 {user['username']} 删除了讨论 #{disc_id}")
    return {"status": "ok", "message": "讨论已删除"}


# ── 学生加入讨论 ──

@router.post("/discussions/{disc_id}/join", summary="学生加入讨论")
async def join_discussion(disc_id: int, request: Request):
    user = get_current_user(request)
    username = user["username"]

    disc = _disc_brief(disc_id)
    if disc is None:
        raise HTTPException(status_code=404, detail="讨论不存在")
    status = disc["status"]

    # S10: 学生只能加入对自己可见的讨论 —— 旧实现任何学生按 ID 即可加入任意讨论,
    #      拿到成员身份后就能读写该组消息、查看 AI 报告
    if user.get("role", 2) == 2:
        if not _student_scope_ok(user, disc) or not filter_activities_by_scope([disc], username):
            raise HTTPException(status_code=403, detail="该讨论不在你的可见范围内")

    # 检查是否已在某个小组中
    existing = execute_query(
        """SELECT 1 FROM discussion_members dm
           JOIN discussion_groups dg ON dm.group_id = dg.id
           WHERE dg.discussion_id = ? AND dm.username = ?""",
        (disc_id, username),
    )
    if existing:
        # 已在组中，返回其小组 ID
        group_row = execute_query(
            """SELECT dg.id, dg.group_index FROM discussion_groups dg
               JOIN discussion_members dm ON dm.group_id = dg.id
               WHERE dg.discussion_id = ? AND dm.username = ?""",
            (disc_id, username),
        )
        if not group_row:
            raise HTTPException(status_code=500, detail="内部错误：无法获取分组信息")
        return {
            "status": "already_joined",
            "group_id": group_row[0][0],
            "group_index": group_row[0][1],
        }

    if status != "active":
        raise HTTPException(status_code=400, detail="讨论未开始或已结束，暂无法加入")

    # 如果还没有任何分组，先创建一个临时分组（待教师开始后重新分组）
    groups = execute_query(
        "SELECT id, group_index FROM discussion_groups WHERE discussion_id=? ORDER BY group_index",
        (disc_id,),
    )
    if not groups:
        gid = execute_insert_update(
            "INSERT INTO discussion_groups (discussion_id, group_index, name) VALUES (?, 1, '待分组')",
            (disc_id,),
        )
        group_id = gid
    else:
        # 找人数最少的小组加入
        min_count = float("inf")
        group_id = groups[0][0]
        for g in groups:
            cnt = execute_query(
                "SELECT COUNT(*) FROM discussion_members WHERE group_id=?",
                (g[0],),
            )[0][0]
            if cnt < min_count:
                min_count = cnt
                group_id = g[0]

    now = _now()
    execute_insert_update(
        "INSERT OR IGNORE INTO discussion_members (group_id, username, role, joined_at) VALUES (?, ?, 'member', ?)",
        (group_id, username, now),
    )

    # ── 积分奖励 ──
    try:
        from backend.reward_engine import award_participation
        disc_title = execute_query("SELECT title FROM discussions WHERE id=?", (disc_id,))
        title = disc_title[0][0] if disc_title else f"讨论#{disc_id}"
        award_participation(username, "discussion", str(disc_id), title)
    except Exception as e:
        logger.warning(f"讨论积分发放失败 (user={username}, disc_id={disc_id}): {e}")

    return {"status": "joined", "group_id": group_id}


# ── 教师广播 ──

@router.post("/discussions/{disc_id}/broadcast", summary="教师广播消息")
async def broadcast_message(disc_id: int, req: BroadcastMessage, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)

    disc = _disc_brief(disc_id)
    _assert_can_access_discussion(user, disc, "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可广播")
    _assert_discussion_owner(user, disc)

    groups = execute_query("SELECT id FROM discussion_groups WHERE discussion_id=?", (disc_id,))
    now = _now()
    for g in groups:
        execute_insert_update(
            "INSERT INTO discussion_messages (group_id, username, content, msg_type, created_at) VALUES (?, ?, ?, 'broadcast', ?)",
            (g[0], user["username"], req.content, now),
        )
        # WebSocket 广播
        msg_data = {
            "type": "new_message",
            "username": user["username"] or "教师",
            "content": req.content,
            "msg_type": "broadcast",
            "created_at": now,
        }
        asyncio.create_task(ws_manager.broadcast(g[0], msg_data))

    return {"status": "ok", "message": "广播已发送"}


# ── 发送消息 ──

@router.post("/groups/{group_id}/messages", summary="发送消息")
async def send_message(group_id: int, req: MessageSend, request: Request):
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)

    _assert_can_access_group(user, group_id, "数据")

    # 验证是该组成员（管理员/教师例外，可发到任意组）
    if role == 2:
        is_member = execute_query(
            "SELECT 1 FROM discussion_members WHERE group_id=? AND username=?",
            (group_id, username),
        )
        if not is_member:
            raise HTTPException(status_code=403, detail="你不在该小组中")

        # ── AI 内容审核（仅对学生发言审核） ──
        safe, reason = _ai_content_review(req.content, username)
        if not safe:
            logger.info(f"学生 {username} 讨论发言内容不合规已拦截: {reason}")
            raise HTTPException(
                status_code=400,
                detail=f"⚠️ 发言审核未通过：{reason}。请修改后重新发送。"
            )

    now = _now()
    msg_id = execute_insert_update(
        "INSERT INTO discussion_messages (group_id, username, content, msg_type, created_at) VALUES (?, ?, ?, 'text', ?)",
        (group_id, username, req.content, now),
    )
    # 广播新消息到 WebSocket（带上真实 ID，供前端去重）
    msg_data = {
        "type": "new_message",
        "id": msg_id,
        "username": username,
        "content": req.content,
        "msg_type": "text",
        "created_at": now,
    }
    asyncio.create_task(ws_manager.broadcast(group_id, msg_data))
    return {"status": "ok", "id": msg_id}


# ── 获取消息（轮询）──

@router.get("/groups/{group_id}/messages", summary="获取消息")
async def get_messages(
    group_id: int,
    request: Request,
    after_id: int = Query(0, description="只返回大于此 ID 的新消息"),
    limit: int = Query(200, ge=1, le=1000, description="单次最多返回条数"),
):
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)

    _assert_can_access_group(user, group_id, "数据")

    # 校验权限：管理员/教师可看任意组，学生只能看自己所在组
    if role == 2:
        is_member = execute_query(
            "SELECT 1 FROM discussion_members WHERE group_id=? AND username=?",
            (group_id, username),
        )
        if not is_member:
            raise HTTPException(status_code=403, detail="你不在该小组中")

    rows = execute_query(
        """SELECT id, username, content, msg_type, created_at
           FROM discussion_messages
           WHERE group_id=? AND id>?
           ORDER BY id ASC LIMIT ?""",
        (group_id, after_id, limit),
    )
    return [
        {
            "id": r[0],
            "username": r[1] or "AI助教",
            "content": r[2],
            "msg_type": r[3],
            "created_at": r[4],
        }
        for r in rows
    ]


# ── 获取我的讨论（学生视角）──

@router.get("/my-discussions", summary="获取我的讨论")
async def get_my_discussions(request: Request):
    user = get_current_user(request)
    username = user["username"]

    rows = execute_query_dict(
        """SELECT DISTINCT d.* FROM discussions d
           JOIN discussion_groups dg ON dg.discussion_id = d.id
           JOIN discussion_members dm ON dm.group_id = dg.id
           WHERE dm.username = ?
           ORDER BY d.created_at DESC""",
        (username,),
    )
    results = []
    for item in rows:
        item = dict(item)
        item["require_summary"] = bool(item["require_summary"])
        # 找出该学生所在的小组
        my_group = execute_query(
            """SELECT dg.id, dg.group_index, dg.name FROM discussion_groups dg
               JOIN discussion_members dm ON dm.group_id = dg.id
               WHERE dg.discussion_id=? AND dm.username=?""",
            (item["id"], username),
        )
        if my_group:
            item["my_group"] = {
                "id": my_group[0][0],
                "group_index": my_group[0][1],
                "name": my_group[0][2],
            }
        else:
            item["my_group"] = None
        results.append(item)

    return results


# ── AI 助教建议（手动触发）──

@router.post("/groups/{group_id}/ai-suggest", summary="AI 助教建议")
async def ai_suggest(group_id: int, request: Request):
    """手动触发 AI 助教生成讨论引导（仅教师/管理员可用）"""
    user = get_current_user(request)
    role = user.get("role", 2)

    _assert_can_access_group(user, group_id, "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可触发 AI 助教")

    # 获取讨论信息
    disc_row = execute_query(
        """SELECT d.title, d.description, d.ai_role FROM discussions d
           JOIN discussion_groups dg ON dg.discussion_id = d.id
           WHERE dg.id=?""",
        (group_id,),
    )
    if not disc_row:
        raise HTTPException(status_code=404, detail="小组不存在")

    title = disc_row[0][0]
    description = disc_row[0][1]
    ai_role = disc_row[0][2]

    # 获取最近消息
    msgs = execute_query(
        """SELECT username, content FROM discussion_messages
           WHERE group_id=? ORDER BY id DESC LIMIT 10""",
        (group_id,),
    )
    messages_text = "\n".join(
        [f"{m[0] or 'AI助教'}: {m[1]}" for m in reversed(msgs)]
    )

    role_desc = {
        "observer": "观察讨论，不做干预",
        "guide": "适时引导讨论方向，提出启发式问题",
        "proactive": "主动参与讨论，提供观点和论据",
        "judge": "作为辩论裁判，分析各方论点",
        "mixed": "综合角色，根据讨论情况自动切换：观察、引导、主动参与、辩论裁判四种角色",
    }.get(ai_role, "适时引导讨论")

    ai_role_text = build_ai_role()
    prompt = f"""{ai_role_text}你是一位课堂讨论的AI助教，角色是：{role_desc}

讨论主题：{title}
讨论说明：{description}

当前讨论内容：
{messages_text or "（讨论尚未开始）"}

请根据讨论情况给出简短的引导或总结（50-100字）："""
    prompt = apply_skills(prompt, "discussion")

    # 调用 AI
    from backend.api.chat_router import get_api_keys
    api_key, _ = get_api_keys(user["username"])
    if not api_key:
        return {"status": "error", "content": "AI 功能不可用：请配置 API Key"}

    from backend.api.ai_service import call_ai_async

    try:
        content = await call_ai_async(prompt, api_key)
        if content:
            # 将 AI 回复作为消息存入
            now_str = _now()
            msg_id = execute_insert_update(
                "INSERT INTO discussion_messages (group_id, username, content, msg_type, created_at) VALUES (?, NULL, ?, 'ai_suggest', ?)",
                (group_id, content, now_str),
            )
            # WebSocket 广播，实时推送给所有小组成员（带上真实 ID）
            asyncio.create_task(ws_manager.broadcast(group_id, {
                "type": "new_message",
                "id": msg_id,
                "username": None,
                "content": content,
                "msg_type": "ai_suggest",
                "created_at": now_str,
            }))
            return {"status": "ok", "content": content}
        return {"status": "error", "content": "AI 未返回有效结果"}
    except Exception as e:
        logger.warning(f"AI 助教调用失败: {e}")
        return {"status": "error", "content": f"AI 调用出错: {str(e)}"}


# ── AI 归纳总结（互动讨论后的 AI 总结功能）──

def _get_api_key(username: str = "") -> str:
    """获取 API Key 的辅助函数（委托共享的 get_api_keys）"""
    from backend.api.chat_router import get_api_keys
    key, _ = get_api_keys(username) if username else get_api_keys("")
    return key


@router.post("/groups/{group_id}/ai-summary", summary="AI 生成小组讨论归纳总结")
async def generate_group_ai_summary(group_id: int, request: Request):
    """为指定小组生成 AI 讨论归纳总结，包含核心观点、评价与评分（仅教师和管理员可用）"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)

    _assert_can_access_group(user, group_id, "数据")

    # 仅教师和管理员可用
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可使用 AI 总结功能")

    # 获取讨论信息
    disc_row = execute_query(
        """SELECT d.id, d.title, d.description, d.subject, d.status, dg.name
           FROM discussions d
           JOIN discussion_groups dg ON dg.discussion_id = d.id
           WHERE dg.id=?""",
        (group_id,),
    )
    if not disc_row:
        raise HTTPException(status_code=404, detail="小组不存在")

    disc_id = disc_row[0][0]
    title = disc_row[0][1]
    description = disc_row[0][2] or ""
    subject = disc_row[0][3] or ""
    group_name = disc_row[0][5] or f"第{disc_row[0][4]}组"

    # 获取该小组所有消息
    msgs = execute_query(
        """SELECT username, content, msg_type, created_at
           FROM discussion_messages
           WHERE group_id=? ORDER BY id ASC""",
        (group_id,),
    )
    if not msgs:
        return {"status": "error", "content": "该小组暂无讨论消息，无法生成总结"}

    # 组装消息文本
    messages_text = "\n".join(
        [f"{m[0] or 'AI助教'} ({m[3] if m[3] else ''}): {m[1]}"
         for m in msgs]
    )

    # 如果消息太多，只取最近的 80 条（防止 token 超限）
    msg_lines = messages_text.split("\n")
    if len(msg_lines) > 80:
        msg_lines = msg_lines[-80:]
        msg_lines.insert(0, "（以下为最近的部分讨论内容）")
    messages_text = "\n".join(msg_lines)

    api_key = _get_api_key()
    if not api_key:
        return {"status": "error", "content": "AI 功能不可用：请配置 API Key"}

    from backend.prompts.discussion import DISCUSSION_AI_SUMMARY_PROMPT
    ai_role = build_ai_role(subject=subject)
    prompt = f"{ai_role}\n" + DISCUSSION_AI_SUMMARY_PROMPT.format(
        subject=subject,
        title=title,
        group_name=group_name,
        description=description,
        messages_text=messages_text,
    )
    # 注意：不注入技能 — 技能的结构化输出指令与 JSON 格式要求冲突

    from backend.api.ai_service import call_ai_async

    try:
        content = await call_ai_async(prompt, api_key)
        if not content:
            return {"status": "error", "content": "AI 未返回有效结果"}

        # 尝试解析 JSON
        import json as _json
        parsed = extract_json_from_text(content)

        # 构建结构化报告内容
        now_str = _now()
        summary_data = {
            "raw_content": content,
            "parsed": parsed,
            "generated_at": now_str,
        }
        report_json = _json.dumps(summary_data, ensure_ascii=False)

        # 保存到 discussion_reports 表（覆盖已有的同组总结）
        existing = execute_query(
            "SELECT id FROM discussion_reports WHERE discussion_id=? AND group_id=?",
            (disc_id, group_id),
        )
        if existing:
            execute_insert_update(
                "UPDATE discussion_reports SET report_content=?, generated_at=? WHERE id=?",
                (report_json, now_str, existing[0][0]),
            )
        else:
            execute_insert_update(
                "INSERT INTO discussion_reports (discussion_id, group_id, report_content, generated_at) VALUES (?, ?, ?, ?)",
                (disc_id, group_id, report_json, now_str),
            )

        return {
            "status": "ok",
            "content": content,
            "parsed": parsed,
            "generated_at": now_str,
        }
    except Exception as e:
        logger.warning(f"AI 归纳总结生成失败: {e}")
        return {"status": "error", "content": f"AI 调用出错: {str(e)}"}


@router.get("/groups/{group_id}/summary", summary="获取小组讨论归纳总结")
async def get_group_summary(group_id: int, request: Request):
    """获取指定小组已生成的 AI 归纳总结（仅教师和管理员可用）"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)

    _assert_can_access_group(user, group_id, "数据")

    # 仅教师和管理员可用
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可查看 AI 总结")

    # 获取讨论 ID
    disc_row = execute_query(
        "SELECT discussion_id FROM discussion_groups WHERE id=?",
        (group_id,),
    )
    if not disc_row:
        raise HTTPException(status_code=404, detail="小组不存在")
    disc_id = disc_row[0][0]

    rows = execute_query(
        "SELECT id, report_content, generated_at FROM discussion_reports WHERE discussion_id=? AND group_id=? ORDER BY id DESC LIMIT 1",
        (disc_id, group_id),
    )
    if not rows:
        return {"status": "ok", "has_summary": False, "content": None}

    import json as _json
    content = rows[0][1]
    try:
        parsed = _json.loads(content)
    except _json.JSONDecodeError:
        parsed = {"raw_content": content}

    return {
        "status": "ok",
        "has_summary": True,
        "content": parsed,
        "generated_at": rows[0][2],
    }


@router.get("/groups/{group_id}/summary/export", summary="导出小组讨论总结为 Word 文档")
def export_group_summary_docx(
    group_id: int,
    request: Request,
):
    """导出 AI 小组讨论归纳总结为 Word 文档"""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    user = get_current_user(request)
    role = user.get("role", 2)

    _assert_can_access_group(user, group_id, "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可导出")

    # 获取小组信息
    disc_row = execute_query(
        """SELECT d.id, d.title, d.description, d.subject, dg.name
           FROM discussions d
           JOIN discussion_groups dg ON dg.discussion_id = d.id
           WHERE dg.id=?""",
        (group_id,),
    )
    if not disc_row:
        raise HTTPException(status_code=404, detail="小组不存在")

    disc_id = disc_row[0][0]
    disc_title = disc_row[0][1]
    subject = disc_row[0][3] or ""
    group_name = disc_row[0][4] or f"第{disc_row[0][4]}组"

    # 获取已有总结
    rows = execute_query(
        "SELECT id, report_content, generated_at FROM discussion_reports WHERE discussion_id=? AND group_id=? ORDER BY id DESC LIMIT 1",
        (disc_id, group_id),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="该小组暂无 AI 总结，请先生成总结")

    import json as _json
    content = rows[0][1]
    try:
        parsed = _json.loads(content)
    except _json.JSONDecodeError:
        parsed = {"raw_content": content}

    raw_content = parsed.get("raw_content", content)
    parsed_data = parsed.get("parsed")
    generated_at = rows[0][2] or ""

    # ── 生成 Word 文档 ──
    doc = Document()
    style: Any = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_heading(f"讨论总结报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"主题：{disc_title}  小组：{group_name}  学科：{subject}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if generated_at:
        time_info = doc.add_paragraph()
        time_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = time_info.add_run(f"生成时间：{generated_at}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    if parsed_data:
        # 总体归纳
        if parsed_data.get("summary"):
            doc.add_heading("总体归纳", level=2)
            doc.add_paragraph(str(parsed_data["summary"]))

        # 关键观点
        if parsed_data.get("key_points"):
            doc.add_heading("关键观点", level=2)
            for i, point in enumerate(parsed_data["key_points"]):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"观点{i+1}：")
                run.bold = True
                p.add_run(str(point))

        # AI 评价
        if parsed_data.get("ai_comment"):
            doc.add_heading("AI 评价与建议", level=2)
            doc.add_paragraph(str(parsed_data["ai_comment"]))

        # 评分
        if parsed_data.get("score"):
            doc.add_heading("综合评分", level=2)
            doc.add_paragraph(f"得分：{parsed_data['score']}/10")

        # 原始内容
        doc.add_paragraph()
        doc.add_heading("原始 AI 回复", level=2)
        _markdown_to_docx(doc, raw_content)
    else:
        # 无结构化数据，直接写入原始内容
        _markdown_to_docx(doc, raw_content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import urllib.parse
    safe_filename = urllib.parse.quote(f"讨论总结_{group_name}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )


# ── AI 自动生成报告 ──

@router.post("/discussions/{disc_id}/generate-report", summary="AI 生成讨论报告")
async def generate_report(disc_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)

    disc = _disc_dict(disc_id)
    if disc is None:
        raise HTTPException(status_code=404, detail="讨论不存在")
    _assert_can_access_discussion(user, disc, "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可生成报告")
    _assert_discussion_owner(user, disc)

    # 获取所有小组的消息
    groups = execute_query(
        "SELECT id, group_index, name FROM discussion_groups WHERE discussion_id=? ORDER BY group_index",
        (disc_id,),
    )

    reports = []
    now = _now()

    # 生成总体报告
    overall_parts = [f"# 讨论报告：{disc['title']}\n"]
    overall_parts.append(f"\n## 总体概览\n")
    overall_parts.append(f"- 讨论主题：{disc['title']}")
    overall_parts.append(f"- 讨论说明：{disc['description']}")
    overall_parts.append(f"- 小组数量：{len(groups)}")

    group_summaries = []
    for g in groups:
        members = execute_query(
            "SELECT username FROM discussion_members WHERE group_id=?",
            (g[0],),
        )
        msgs = execute_query(
            "SELECT username, content, msg_type FROM discussion_messages WHERE group_id=? ORDER BY id ASC",
            (g[0],),
        )
        member_names = [m[0] for m in members]
        msg_count = len(msgs)
        group_summaries.append(f"\n### {g[2] or f'第{g[1]}组'}")
        group_summaries.append(f"- 成员：{', '.join(member_names)}")
        group_summaries.append(f"- 消息数：{msg_count}")
        # 取最后 5 条消息作为摘要
        recent = msgs[-5:] if msgs else []
        for rm in recent:
            sender = rm[0] or "AI助教"
            content = rm[1][:100]
            group_summaries.append(f"  - {sender}: {content}")

        # 为该小组生成 AI 分析摘要
        if msgs:
            messages_text = "\n".join(
                [f"{m[0] or 'AI助教'}: {m[1][:200]}" for m in msgs[-20:]]
            )
            prompt = f"""请对以下讨论内容进行简要分析（50字以内），总结该小组的关键观点和讨论质量：

讨论主题：{disc['title']}
讨论内容：
{messages_text}

简要分析："""
            prompt = apply_skills(prompt, "discussion")

            ai_summary = ""
            try:
                from backend.api.chat_router import get_api_keys
                api_key, _ = get_api_keys(user["username"])
                if api_key:
                    from backend.api.ai_service import call_ai_async
                    ai_summary = await call_ai_async(prompt, api_key)
            except Exception:
                pass

            if ai_summary:
                group_summaries.append(f"\n**AI 分析**：{ai_summary}")

            # 保存小组报告
            report_content = f"# 小组报告：{g[2] or f'第{g[1]}组'}\n\n"
            report_content += f"## 基本信息\n- 成员：{', '.join(member_names)}\n- 消息数：{msg_count}\n\n"
            report_content += f"## AI 分析\n{ai_summary}\n\n" if ai_summary else ""
            report_content += f"## 讨论内容\n```\n{messages_text}\n```"

            execute_insert_update(
                "INSERT INTO discussion_reports (discussion_id, group_id, report_content, generated_at) VALUES (?, ?, ?, ?)",
                (disc_id, g[0], report_content, now),
            )

    overall_parts.extend(group_summaries)
    overall_report = "\n".join(overall_parts)

    execute_insert_update(
        "INSERT INTO discussion_reports (discussion_id, group_id, report_content, generated_at) VALUES (?, NULL, ?, ?)",
        (disc_id, overall_report, now),
    )

    return {"status": "ok", "message": "报告生成完成"}


# ── 获取报告 ──

@router.get("/discussions/{disc_id}/reports", summary="获取讨论报告")
async def get_reports(disc_id: int, request: Request):
    # S1: 旧实现无任何鉴权, 未登录即可拉到 AI 小组/整体报告全文
    user = get_current_user(request)
    _assert_can_access_discussion(user, _disc_brief(disc_id), "报告")
    rows = execute_query(
        "SELECT dr.id, dr.group_id, dr.report_content, dr.generated_at, "
        "dg.group_index, dg.name "
        "FROM discussion_reports dr "
        "LEFT JOIN discussion_groups dg ON dr.group_id = dg.id "
        "WHERE dr.discussion_id=? "
        "ORDER BY dr.group_id IS NULL DESC, dg.group_index ASC",
        (disc_id,),
    )
    return [
        {
            "id": r[0],
            "group_id": r[1],
            "content": r[2],
            "generated_at": r[3],
            "group_index": r[4],
            "group_name": r[5],
            "is_overall": r[1] is None,
        }
        for r in rows
    ]


# ── WebSocket 实时消息 ──

@router.websocket("/ws/{group_id}")
async def websocket_endpoint(websocket: WebSocket, group_id: int):
    """WebSocket 端点，客户端连接后实时接收新消息

    S2: 握手必须携带有效 token, 且必须是该组成员(学生)/任教范围内教师/管理员。
    旧实现匿名也能 accept 并回 pong, 任意 group_id 均可挂接。
    """
    from backend.auth import authenticate_payload

    token = websocket.query_params.get("token", "")
    payload = authenticate_payload(token)
    if not payload:
        logger.warning(f"[讨论WS] 拒绝未认证连接 group={group_id}")
        await websocket.close(code=4401, reason="未认证")
        return
    username = payload.get("username", "")
    if not username:
        await websocket.close(code=4401, reason="未认证")
        return
    try:
        _assert_can_access_group(
            {"username": username, "role": payload.get("role", 2)}, group_id, "消息",
        )
    except HTTPException:
        logger.warning(f"[讨论WS] 拒绝越权连接 user={username} group={group_id}")
        await websocket.close(code=4403, reason="无权访问该小组")
        return
    except Exception:
        await websocket.close(code=4404, reason="小组不存在")
        return

    await ws_manager.connect(group_id, websocket)
    try:
        while True:
            # 保持连接，等待客户端发来 ping 或断开
            data = await websocket.receive_text()
            if data == "ping":
                await websocket.send_json({"type": "pong"})
    except WebSocketDisconnect:
        ws_manager.disconnect(group_id, websocket)
    except Exception:
        ws_manager.disconnect(group_id, websocket)


# ── 讨论监控（教师/管理员）──

@router.get("/discussions/{disc_id}/monitor", summary="讨论监控面板数据")
async def monitor_discussion(disc_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)

    _assert_can_access_discussion(user, _disc_brief(disc_id), "数据")
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可查看监控")

    rows = execute_query("SELECT status, title FROM discussions WHERE id=?", (disc_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="讨论不存在")
    status, title = rows[0]

    # 获取所有小组状态
    groups = execute_query(
        "SELECT id, group_index, name FROM discussion_groups WHERE discussion_id=? ORDER BY group_index",
        (disc_id,),
    )
    now = datetime.now()
    group_list = []
    total_messages = 0
    total_members = 0
    cold_groups = 0

    for g in groups:
        members = execute_query(
            "SELECT COUNT(*) FROM discussion_members WHERE group_id=?",
            (g[0],),
        )[0][0]

        msg_count = execute_query(
            "SELECT COUNT(*) FROM discussion_messages WHERE group_id=?",
            (g[0],),
        )[0][0]

        last_msg = execute_query(
            "SELECT content, created_at FROM discussion_messages WHERE group_id=? ORDER BY id DESC LIMIT 1",
            (g[0],),
        )

        last_active = ""
        is_cold = False
        last_preview = ""
        if last_msg:
            last_preview = last_msg[0][0][:60]
            last_active = last_msg[0][1]
            try:
                last_time = datetime.strptime(last_active, "%Y-%m-%d %H:%M:%S")
                idle_seconds = (now - last_time).total_seconds()
                is_cold = idle_seconds > 60 and status == "active"
            except ValueError:
                pass

        if is_cold:
            cold_groups += 1
        total_messages += msg_count
        total_members += members

        group_list.append({
            "id": g[0],
            "group_index": g[1],
            "name": g[2] or f"第{g[1]}组",
            "member_count": members,
            "message_count": msg_count,
            "last_active": last_active,
            "last_preview": last_preview,
            "is_cold": is_cold,
        })

    return {
        "discussion_id": disc_id,
        "title": title,
        "status": status,
        "total_groups": len(groups),
        "total_members": total_members,
        "total_messages": total_messages,
        "cold_groups": cold_groups,
        "online_count": sum(ws_manager.get_group_connections(g["id"]) for g in group_list),
        "groups": group_list,
    }


# ── AI 自动触发 ──

@router.post("/discussions/{disc_id}/auto-trigger", summary="触发 AI 自动引导（冷场检测）")
async def auto_trigger_ai(disc_id: int, request: Request):
    """检测所有小组是否需要 AI 介入（冷场超过 60 秒），自动发送引导"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可用")

    rows = execute_query("SELECT status FROM discussions WHERE id=?", (disc_id,))
    if not rows or rows[0][0] != "active":
        return {"status": "ok", "triggered": 0, "message": "讨论未激活"}

    groups = execute_query(
        "SELECT id, group_index FROM discussion_groups WHERE discussion_id=?",
        (disc_id,),
    )
    now = datetime.now()
    triggered = 0
    from backend.api.chat_router import get_api_keys
    api_key, _ = get_api_keys(user["username"])
    if not api_key:
        return {"status": "error", "triggered": 0, "message": "API Key 未配置"}

    from backend.api.ai_service import call_ai_async

    for g in groups:
        last_msg = execute_query(
            "SELECT content FROM discussion_messages WHERE group_id=? AND username IS NOT NULL ORDER BY id DESC LIMIT 1",
            (g[0],),
        )
        if not last_msg:
            continue

        last_content = last_msg[0][0]

        # 获取最近 5 条消息
        recent = execute_query(
            "SELECT username, content FROM discussion_messages WHERE group_id=? ORDER BY id DESC LIMIT 5",
            (g[0],),
        )
        messages_text = "\n".join(
            [f"{m[0] or 'AI助教'}: {m[1][:200]}" for m in reversed(recent)]
        )

        ai_role_text = build_ai_role()
        prompt = f"""{ai_role_text}你是一位课堂讨论的AI助教。
请根据以下讨论内容给出一个简短的引导问题或总结（30-50字），目的是推动讨论继续深入：

讨论内容：
{messages_text}

简短引导："""
        prompt = apply_skills(prompt, "discussion")

        try:
            content = await call_ai_async(prompt, api_key)
            if content:
                now_str = _now()
                execute_insert_update(
                    "INSERT INTO discussion_messages (group_id, username, content, msg_type, created_at) VALUES (?, NULL, ?, 'ai_suggest', ?)",
                    (g[0], content, now_str),
                )
                # WebSocket 广播
                asyncio.create_task(ws_manager.broadcast(g[0], {
                    "type": "new_message",
                    "username": None,
                    "content": content,
                    "msg_type": "ai_suggest",
                    "created_at": now_str,
                }))
                triggered += 1
        except Exception as e:
            logger.warning(f"AI 自动触发失败 小组#{g[0]}: {e}")

    return {"status": "ok", "triggered": triggered}


# ═══════════════════════════════════════════════════════════
# Word 导出工具
# ═══════════════════════════════════════════════════════════

def _markdown_to_docx(doc, text: str):
    """将 Markdown 文本写入 docx 文档"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- **') and '：' in line:
            content = line.lstrip('- ')
            p = doc.add_paragraph()
            bold_end = content.find('**', 2)
            if bold_end > 0:
                run = p.add_run(content[2:bold_end])
                run.bold = True
                p.add_run(content[bold_end + 2:])
            else:
                p.add_run(content)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f'{i}. ') for i in range(1, 10)):
            doc.add_paragraph(line, style='List Number')
        else:
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            else:
                doc.add_paragraph(line)
