"""
AI 智能学情分析 API 路由
利用 DashScope AI 对学习数据进行分析，生成自然语言学情报告
"""
import json
import os
from datetime import datetime

from fastapi import APIRouter, HTTPException, Request, Query

from backend.api.dependencies import get_current_user
from backend.database import execute_query
from backend.question_db import execute_query as q_execute_query
from backend.logger import logger

router = APIRouter()


def _get_dashscope_api_key() -> str:
    """获取 DashScope API Key"""
    key = os.environ.get("DASHSCOPE_API_KEY", "")
    if not key:
        try:
            from backend.api.config_router import load_config
            cfg = load_config()
            key = cfg.get("dashscope_api_key", "")
        except Exception:
            pass
    return key


def _call_ai(prompt: str) -> str:
    """调用 AI 分析（同步，保留兼容）"""
    api_key = _get_dashscope_api_key()
    if not api_key:
        return "⚠️ AI 分析功能不可用：请管理员在「系统配置」中填写 DashScope API Key"

    from backend.api.ai_service import call_ai_sync
    try:
        return call_ai_sync(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 学情分析调用失败: {e}")
        return f"AI 分析出错：{str(e)}"


async def _call_ai_task(description: str, prompt: str) -> str:
    """提交 AI 分析后台任务，返回 task_id"""
    api_key = _get_dashscope_api_key()
    if not api_key:
        # 无 API Key 时直接返回错误信息（不创建任务）
        task_id = __import__('uuid').uuid4().hex[:12]
        from backend.ai_task_manager import task_manager, AITask, TaskStatus
        task = AITask(task_id, description)
        task.status = TaskStatus.FAILED
        task.error = "⚠️ AI 分析功能不可用：请管理员在「系统配置」中填写 DashScope API Key"
        task.completed_at = __import__('time').time()
        task_manager._tasks[task_id] = task
        return task_id

    from backend.api.ai_service import call_ai_async

    async def _do_analysis() -> dict:
        try:
            result = await call_ai_async(prompt, api_key)
            return {"result": result}
        except Exception as e:
            logger.error(f"AI 学情分析调用失败: {e}")
            return {"error": f"AI 分析出错：{str(e)}"}

    from backend.ai_task_manager import task_manager
    return await task_manager.create_task(description=description, coro_factory=_do_analysis)


def _safe_int(val) -> int:
    try:
        return int(val)
    except (TypeError, ValueError):
        return 0


def _safe_float(val) -> float:
    try:
        return float(val)
    except (TypeError, ValueError):
        return 0.0


def _extract_class_num(cls_val: str) -> str:
    """从 '高一1班' 格式中提取班级数字 '1'"""
    import re
    nums = re.findall(r'\d+', cls_val)
    return nums[0] if nums else cls_val


@router.get("/class-overview", summary="班级学情总览（AI 生成）")
async def class_overview(
    request: Request,
    grade: str = Query("", description="年级"),
    cls: str = Query("", description="班级"),
    teacher: str = Query("", description="教师用户名"),
):
    """AI 生成班级学情总览报告"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可查看")

    query_teacher = teacher or username

    # 教师只能查看自己班级的数据
    if role == 1:
        teacher_info = execute_query(
            "SELECT grade, class FROM users WHERE username=?", (username,)
        )
        if teacher_info:
            t_grade = (teacher_info[0][0] or "").strip()
            t_class = (teacher_info[0][1] or "").strip()
            allowed_grades = [g.strip() for g in t_grade.split("|") if g.strip()]
            # 如果教师配置了年级，检查请求的年级是否在允许范围内
            if allowed_grades and grade not in allowed_grades:
                raise HTTPException(status_code=403, detail="无权查看其他年级的数据")

    # 收集班级数据
    # 班级号格式处理：users.class 存数字(1)，下拉框传"高一1班"
    class_num = _extract_class_num(cls)

    # 1. 学生人数
    student_count = execute_query(
        "SELECT COUNT(*) FROM users WHERE role = 2 AND grade = ? AND class = ?",
        (grade, class_num),
    )
    total_students = student_count[0][0] if student_count else 0
    if total_students == 0:
        return {"report": "暂无该班级的学生数据", "data": {}}

    # 2. 积分统计（scores.class_name 存 "高一1班" 格式）
    score_stats = execute_query(
        """SELECT COUNT(*), COALESCE(SUM(score),0), COALESCE(AVG(score),0),
                  COALESCE(MAX(score),0), COALESCE(MIN(score),0)
           FROM scores WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls),
    )

    # 3. 点名统计（rollcall_history.result 存 correct/incorrect/skip）
    rc_stats = execute_query(
        """SELECT COUNT(*),
                  SUM(CASE WHEN result='correct' THEN 1 ELSE 0 END),
                  SUM(CASE WHEN result='incorrect' THEN 1 ELSE 0 END)
           FROM rollcall_history WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls),
    )

    # 4. 考试统计（该班级最近的考试）
    exam_stats = q_execute_query(
        """SELECT e.title, e.subject, COUNT(ea.id) as attempt_count,
                  COALESCE(AVG(ea.score),0) as avg_score,
                  e.pass_score, e.total_score
           FROM exams e
           LEFT JOIN exam_attempts ea ON ea.exam_id = e.id
           WHERE e.status IN ('published', 'ended')
           GROUP BY e.id
           ORDER BY e.created_at DESC LIMIT 5""",
    )

    # 5. 任务完成率
    task_stats = execute_query(
        "SELECT COUNT(*) FROM tasks WHERE creator_username = ? AND status = 'active'",
        (query_teacher,),
    )
    task_submissions = execute_query(
        """SELECT COUNT(DISTINCT ts.student_username)
           FROM task_submissions ts
           JOIN tasks t ON ts.task_id = t.id
           WHERE t.creator_username = ?""",
        (query_teacher,),
    )

    # ── 构造 AI Prompt ──
    sc = score_stats[0] if score_stats else (0, 0, 0, 0, 0)
    rc = rc_stats[0] if rc_stats else (0, 0, 0)

    data_summary = {
        "grade": grade,
        "class": cls,
        "total_students": total_students,
        "score_count": _safe_int(sc[0]),
        "score_total": _safe_int(sc[1]),
        "score_avg": round(_safe_float(sc[2]), 1),
        "score_max": _safe_int(sc[3]),
        "score_min": _safe_int(sc[4]),
        "rollcall_total": _safe_int(rc[0]),
        "rollcall_correct": _safe_int(rc[1]),
        "rollcall_wrong": _safe_int(rc[2]),
        "exams": [
            {
                "title": e['title'],
                "subject": e['subject'],
                "attempts": _safe_int(e['attempt_count']),
                "avg_score": round(_safe_float(e['avg_score']), 1),
                "pass_score": _safe_float(e['pass_score']),
                "total_score": _safe_float(e['total_score']),
            }
            for e in exam_stats
        ],
        "active_tasks": _safe_int(task_stats[0][0] if task_stats else 0),
        "submitted_students": _safe_int(task_submissions[0][0] if task_submissions else 0),
    }

    prompt = f"""你是一位经验丰富的高中信息科技/通用技术教师。请根据以下班级数据，生成一份专业的学情分析报告。

班级：{grade}{cls}班
学生人数：{total_students}

【课堂积分】
有积分记录的学生数：{data_summary['score_count']}
总积分：{data_summary['score_total']}
平均积分：{data_summary['score_avg']}
最高积分：{data_summary['score_max']}

【点名情况】
总点名次数：{data_summary['rollcall_total']}
回答正确次数：{data_summary['rollcall_correct']}
回答错误次数：{data_summary['rollcall_wrong']}

【考试情况】
{chr(10).join(f"- 《{e['title']}》({e['subject']}): 参考{e['attempts']}人, 平均分{e['avg_score']}/{e['total_score']}" for e in data_summary['exams'])}

【任务完成】
活跃任务数：{data_summary['active_tasks']}
已提交学生数：{data_summary['submitted_students']}

请生成包含以下内容的分析报告（以 Markdown 格式输出）：
1. 📊 **班级整体情况**：对该班级的学习状态进行总体评价
2. 📈 **学习亮点**：指出表现突出的方面
3. ⚠️ **需要关注的问题**：指出薄弱环节
4. 💡 **教学建议**：给出具体的改进建议

请使用自然、亲切的语气，直接以分析内容开头，不要出现"根据提供的数据"等冗余表述。"""

    task_id = await _call_ai_task("班级学情分析", prompt)

    return {
        "task_id": task_id,
        "message": "AI 分析已提交，请稍后查询结果",
        "data": data_summary,
    }


@router.get("/student/{target_username}", summary="学生个体 AI 学情分析")
async def student_analytics(target_username: str, request: Request):
    """AI 生成学生个体学情分析"""
    user = get_current_user(request)
    role = user.get("role", 2)

    if role == 2 and user["username"] != target_username:
        raise HTTPException(status_code=403, detail="无权查看")

    # 获取学生信息
    user_rows = execute_query(
        "SELECT name, grade, class FROM users WHERE username = ?",
        (target_username,),
    )
    if not user_rows:
        raise HTTPException(status_code=404, detail="学生不存在")
    student_name = user_rows[0][0] or target_username
    student_grade = user_rows[0][1] or ""
    student_class = user_rows[0][2] or ""

    # 考试成绩
    exam_results = q_execute_query(
        """SELECT e.title, e.subject, ea.score, ea.total_score, e.pass_score, ea.submitted_at
           FROM exam_attempts ea
           JOIN exams e ON ea.exam_id = e.id
           WHERE ea.student_username = ? AND ea.status IN ('submitted', 'graded')
           ORDER BY ea.submitted_at DESC""",
        (target_username,),
    )

    # 积分
    score_rows = execute_query(
        "SELECT COALESCE(SUM(score),0) FROM scores WHERE student_name = ?",
        (student_name,),
    )
    total_score = score_rows[0][0] if score_rows else 0

    # 点名
    rc_rows = execute_query(
        """SELECT COUNT(*), SUM(CASE WHEN result='1' THEN 1 ELSE 0 END)
           FROM rollcall_history WHERE student_name = ?""",
        (student_name,),
    )
    rc_total = rc_rows[0][0] if rc_rows else 0
    rc_correct = rc_rows[0][1] if rc_rows else 0

    # 任务
    task_rows = execute_query(
        "SELECT COUNT(*) FROM task_submissions WHERE student_username = ?",
        (target_username,),
    )
    task_count = task_rows[0][0] if task_rows else 0

    # 对话
    chat_rows = execute_query(
        "SELECT COUNT(*) FROM conversations WHERE username = ?",
        (target_username,),
    )
    chat_count = chat_rows[0][0] if chat_rows else 0

    # 构建数据
    exam_text = ""
    for e in exam_results:
        exam_text += f"- 《{e['title']}》({e['subject']}): {e['score']}/{e['total_score']}分 ({'通过' if e['score'] >= e['pass_score'] else '未通过'})\n"

    prompt = f"""你是一位高中信息科技教师。请根据以下学生数据，生成一份个性化的学情分析报告。

学生：{student_name}
班级：{student_grade}{student_class}

【考试成绩】
{exam_text if exam_text else '暂无考试记录'}

【课堂积分】累计 {total_score} 分
【点名情况】共被点名 {rc_total} 次，回答正确 {rc_correct} 次
【任务完成】已提交 {task_count} 个任务
【AI 对话】共 {chat_count} 次

请生成（Markdown 格式）：
1. 📊 **学习概况**
2. 💪 **优势与进步**
3. 🔧 **改进建议**
4. 🎯 **下一阶段学习建议**

语气亲切、鼓励为主。"""

    task_id = await _call_ai_task("学生个体学情分析", prompt)

    return {
        "task_id": task_id,
        "message": "AI 分析已提交，请稍后查询结果",
        "student": {"username": target_username, "name": student_name, "grade": student_grade, "class": student_class},
    }


@router.get("/exam/{exam_id}/report", summary="单次考试 AI 分析报告")
async def exam_analytics(exam_id: int, request: Request):
    """AI 生成单次考试的分析报告"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可查看")

    exam = q_execute_query("SELECT * FROM exams WHERE id = ?", (exam_id,))
    if not exam:
        raise HTTPException(status_code=404, detail="考试不存在")
    exam = exam[0]

    attempts = q_execute_query(
        "SELECT * FROM exam_attempts WHERE exam_id = ? AND status IN ('submitted', 'graded') AND auto_graded = 1 AND answers IS NOT NULL AND answers != '' ORDER BY score DESC",
        (exam_id,),
    )

    questions = q_execute_query(
        """SELECT qb.id, qb.type, qb.question_text, qb.correct_answer, qb.difficulty, qb.knowledge_points,
                  eq.score as question_score
           FROM exam_questions eq
           JOIN question_bank qb ON eq.question_id = qb.id
           WHERE eq.exam_id = ?
           ORDER BY eq.sort_order""",
        (exam_id,),
    )

    scores = [a['score'] for a in attempts]
    total_count = len(scores)
    avg_score = round(sum(scores) / max(total_count, 1), 1)
    pass_count = sum(1 for s in scores if s >= exam['pass_score'])
    max_score = max(scores) if scores else 0
    min_score = min(scores) if scores else 0

    # 每题正确率
    q_accuracy = []
    for q in questions:
        correct = 0
        for a in attempts:
            if a.get('answers'):
                try:
                    answers = json.loads(a['answers']) if isinstance(a['answers'], str) else a['answers']
                except (json.JSONDecodeError, TypeError):
                    answers = {}
            else:
                answers = {}
            q_key = str(q['id'])
            # 注意: answers 存储的是批改结果对象，如 {student_answer, correct_answer, is_correct}
            ans_entry = answers.get(q_key, {})
            if isinstance(ans_entry, dict):
                if ans_entry.get('is_correct'):
                    correct += 1
            elif ans_entry == q['correct_answer']:
                # 兼容旧格式：直接存答案文本
                correct += 1
        rate = round(correct / max(total_count, 1) * 100, 1)
        q_accuracy.append({
            "id": q['id'],
            "type": q['type'],
            "text": q['question_text'][:80],
            "correct_rate": rate,
            "difficulty": q.get('difficulty', 'medium'),
            "knowledge_points": q.get('knowledge_points', ''),
        })

    # 找出最薄弱的知识点
    weak_points = [q for q in q_accuracy if q['correct_rate'] < 60]

    prompt = f"""你是一位高中{exam['subject']}教师。请根据以下考试数据，生成专业的考试分析报告。

考试名称：{exam['title']}
科目：{exam['subject']}
参考人数：{total_count}
平均分：{avg_score}/{exam['total_score']}
最高分：{max_score}
最低分：{min_score}
及格人数：{pass_count}/{total_count}（及格线{exam['pass_score']}分）

各题正确率（正确率 < 60% 的为薄弱题）：
{chr(10).join(f'- 第{i+1}题 ({q["type"]}, {q["difficulty"]}): 正确率{q["correct_rate"]}%' + (f' [知识点: {q["knowledge_points"]}]' if q['knowledge_points'] else '') for i, q in enumerate(q_accuracy))}

薄弱知识点：
{chr(10).join(('- ' + (q["knowledge_points"] or (f"第{q['id']}题"))) for q in weak_points) if weak_points else '无显著薄弱点'}

请生成（Markdown 格式）：
1. 📊 **总体情况**
2. 📈 **分数分布分析**
3. ⚠️ **薄弱知识点**
4. 💡 **教学改进建议"""

    task_id = await _call_ai_task("考试分析报告", prompt)

    return {
        "task_id": task_id,
        "message": "AI 分析已提交，请稍后查询结果",
        "exam": {"id": exam['id'], "title": exam['title'], "subject": exam['subject']},
        "statistics": {
            "total_students": total_count,
            "avg_score": avg_score,
            "max_score": max_score,
            "min_score": min_score,
            "pass_count": pass_count,
            "pass_rate": round(pass_count / max(total_count, 1) * 100, 1),
            "total_score": exam['total_score'],
        },
        "question_accuracy": q_accuracy,
    }


# ═══════════════════════════════════════════════════════════
# V3.2 新增：AI 教学建议
# ═══════════════════════════════════════════════════════════

@router.get("/teaching-suggestions", summary="AI 教学建议")
async def teaching_suggestions(
    request: Request,
    grade: str = Query(...),
    cls: str = Query(...),
    teacher_username: str = Query(""),
):
    """AI 生成具体可操作的教学建议"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)

    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可查看")

    # 确定查询的教师（管理员可指定，教师只能看自己）
    query_teacher = teacher_username if (role == 0 and teacher_username) else username

    # 班级号格式处理：users.class 存数字(1)，下拉框传"高一1班"
    class_num = _extract_class_num(cls)
    cls_display = class_num
    cls_name = f"{grade}{cls_display}班"

    # ── 收集数据 ──
    # 1. 学生人数
    student_count = execute_query(
        "SELECT COUNT(*) FROM users WHERE role = 2 AND grade = ? AND class = ?",
        (grade, cls_display),
    )
    total_students = student_count[0][0] if student_count else 0
    if total_students == 0:
        return {"suggestions": "暂无该班级的数据", "data": {}}

    # 2. 积分统计
    score_stats = execute_query(
        """SELECT COUNT(*), COALESCE(SUM(score),0), COALESCE(AVG(score),0),
                  COALESCE(MAX(score),0), COALESCE(MIN(score),0)
           FROM scores WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls_name),
    )

    # 3. 点名统计
    rc_stats = execute_query(
        """SELECT COUNT(*),
                  SUM(CASE WHEN result='correct' THEN 1 ELSE 0 END),
                  SUM(CASE WHEN result='incorrect' THEN 1 ELSE 0 END)
           FROM rollcall_history WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls_name),
    )

    # 4. 考试统计
    exam_stats = q_execute_query(
        """SELECT e.title, e.subject, COUNT(ea.id) as attempt_count,
                  COALESCE(AVG(ea.score),0) as avg_score,
                  e.pass_score, e.total_score
           FROM exams e
           LEFT JOIN exam_attempts ea ON ea.exam_id = e.id
           WHERE e.creator_username = ? AND e.status IN ('published', 'ended')
           GROUP BY e.id
           ORDER BY e.created_at DESC LIMIT 5""",
        (query_teacher,),
    )

    # 5. 任务统计
    task_stats = execute_query(
        "SELECT COUNT(*) FROM tasks WHERE creator_username = ? AND status = 'active'",
        (query_teacher,),
    )
    task_submissions = execute_query(
        """SELECT COUNT(DISTINCT ts.student_username)
           FROM task_submissions ts
           JOIN tasks t ON ts.task_id = t.id
           WHERE t.creator_username = ?""",
        (query_teacher,),
    )

    sc = score_stats[0] if score_stats else (0, 0, 0, 0, 0)
    rc = rc_stats[0] if rc_stats else (0, 0, 0)
    rc_total = _safe_int(rc[0])
    rc_correct = _safe_int(rc[1])

    exam_text = ""
    for e in exam_stats:
        exam_text += f"- 《{e['title']}》({e['subject']}): 参考{e['attempt_count']}人, 平均分{round(_safe_float(e['avg_score']),1)}/{e['total_score']}\n"
    if not exam_text:
        exam_text = "暂无考试数据"

    active_tasks = _safe_int(task_stats[0][0] if task_stats else 0)
    submitted = _safe_int(task_submissions[0][0] if task_submissions else 0)

    # ── 构建 Prompt ──
    from backend.prompts.analytics import TEACHING_SUGGESTIONS_PROMPT

    prompt = TEACHING_SUGGESTIONS_PROMPT.format(
        grade=grade,
        cls=cls_display,
        total_students=total_students,
        score_count=_safe_int(sc[0]),
        score_total=_safe_int(sc[1]),
        score_avg=round(_safe_float(sc[2]), 1),
        score_max=_safe_int(sc[3]),
        score_min=_safe_int(sc[4]),
        rollcall_total=rc_total,
        rollcall_correct=rc_correct,
        rollcall_wrong=rc_total - rc_correct,
        rollcall_rate=round(rc_correct / max(rc_total, 1) * 100, 1),
        exam_text=exam_text,
        active_tasks=active_tasks,
        submitted_students=submitted,
        task_rate=round(submitted / max(total_students, 1) * 100, 1),
    )

    task_id = await _call_ai_task("AI 教学建议", prompt)

    data_summary = {
        "total_students": total_students,
        "score_avg": round(_safe_float(sc[2]), 1),
        "rollcall_rate": round(rc_correct / max(rc_total, 1) * 100, 1),
        "exam_count": len(exam_stats),
        "task_rate": round(submitted / max(total_students, 1) * 100, 1),
    }

    return {
        "task_id": task_id,
        "message": "AI 分析已提交，请稍后查询结果",
        "data": data_summary,
    }


# ═══════════════════════════════════════════════════════════
# 导出 WORD 文档
# ═══════════════════════════════════════════════════════════

def _markdown_to_docx(doc, text: str):
    """将 Markdown 文本写入 docx 文档"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- **') and '：' in line:
            content = line.lstrip('- ')
            p = doc.add_paragraph()
            bold_end = content.find('**', 2)
            if bold_end > 0:
                run = p.add_run(content[2:bold_end])
                run.bold = True
                p.add_run(content[bold_end + 2:])
            else:
                p.add_run(content)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f'{i}. ') for i in range(1, 10)):
            doc.add_paragraph(line, style='List Number')
        else:
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            else:
                doc.add_paragraph(line)


@router.get("/class-overview/export", summary="导出班级学情分析为 Word 文档")
async def export_class_overview_docx(
    request: Request,
    grade: str = Query(...),
    cls: str = Query(...),
    teacher: str = Query(""),
    token: str = Query(""),
):
    """导出 AI 班级学情分析报告为 Word 文档"""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    if token:
        request.state.user = None
        from backend.auth import decode_jwt_token
        payload = decode_jwt_token(token)
        if payload:
            request.state.user = payload

    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可导出")

    query_teacher = teacher or username
    class_num = _extract_class_num(cls)

    student_count = execute_query(
        "SELECT COUNT(*) FROM users WHERE role = 2 AND grade = ? AND class = ?",
        (grade, class_num),
    )
    total_students = student_count[0][0] if student_count else 0
    if total_students == 0:
        raise HTTPException(status_code=400, detail="暂无该班级的学生数据")

    score_stats = execute_query(
        """SELECT COUNT(*), COALESCE(SUM(score),0), COALESCE(AVG(score),0),
                  COALESCE(MAX(score),0), COALESCE(MIN(score),0)
           FROM scores WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls),
    )
    rc_stats = execute_query(
        """SELECT COUNT(*),
                  SUM(CASE WHEN result='correct' THEN 1 ELSE 0 END),
                  SUM(CASE WHEN result='incorrect' THEN 1 ELSE 0 END)
           FROM rollcall_history WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls),
    )
    exam_stats = q_execute_query(
        """SELECT e.title, e.subject, COUNT(ea.id) as attempt_count,
                  COALESCE(AVG(ea.score),0) as avg_score, e.pass_score, e.total_score
           FROM exams e
           LEFT JOIN exam_attempts ea ON ea.exam_id = e.id
           WHERE e.status IN ('published', 'ended')
           GROUP BY e.id ORDER BY e.created_at DESC LIMIT 5""",
    )
    task_stats = execute_query(
        "SELECT COUNT(*) FROM tasks WHERE creator_username = ? AND status = 'active'",
        (query_teacher,),
    )
    task_submissions = execute_query(
        """SELECT COUNT(DISTINCT ts.student_username)
           FROM task_submissions ts JOIN tasks t ON ts.task_id = t.id
           WHERE t.creator_username = ?""",
        (query_teacher,),
    )

    sc = score_stats[0] if score_stats else (0, 0, 0, 0, 0)
    rc = rc_stats[0] if rc_stats else (0, 0, 0)

    data_summary = {
        "total_students": total_students,
        "score_count": _safe_int(sc[0]), "score_total": _safe_int(sc[1]),
        "score_avg": round(_safe_float(sc[2]), 1),
        "score_max": _safe_int(sc[3]), "score_min": _safe_int(sc[4]),
        "rollcall_total": _safe_int(rc[0]), "rollcall_correct": _safe_int(rc[1]),
        "rollcall_wrong": _safe_int(rc[2]),
        "exams": [{"title": e['title'], "subject": e['subject'],
                    "attempts": _safe_int(e['attempt_count']),
                    "avg_score": round(_safe_float(e['avg_score']), 1),
                    "total_score": _safe_float(e['total_score'])}
                   for e in exam_stats],
        "active_tasks": _safe_int(task_stats[0][0] if task_stats else 0),
        "submitted_students": _safe_int(task_submissions[0][0] if task_submissions else 0),
    }

    prompt = f"""你是一位经验丰富的高中信息科技/通用技术教师。请根据以下班级数据，生成一份专业的学情分析报告。

班级：{grade}{cls}班
学生人数：{total_students}

【课堂积分】
有积分记录的学生数：{data_summary['score_count']}
总积分：{data_summary['score_total']}
平均积分：{data_summary['score_avg']}
最高积分：{data_summary['score_max']}

【点名情况】
总点名次数：{data_summary['rollcall_total']}
回答正确次数：{data_summary['rollcall_correct']}
回答错误次数：{data_summary['rollcall_wrong']}

【考试情况】
{chr(10).join(f"- 《{e['title']}》({e['subject']}): 参考{e['attempts']}人, 平均分{e['avg_score']}/{e['total_score']}" for e in data_summary['exams'])}

【任务完成】
活跃任务数：{data_summary['active_tasks']}
已提交学生数：{data_summary['submitted_students']}

请生成包含以下内容的分析报告（以 Markdown 格式输出）：
1. 📊 **班级整体情况**：对该班级的学习状态进行总体评价
2. 📈 **学习亮点**：指出表现突出的方面
3. ⚠️ **需要关注的问题**：指出薄弱环节
4. 💡 **教学建议**：给出具体的改进建议

请使用自然、亲切的语气，直接以分析内容开头，不要出现"根据提供的数据"等冗余表述。"""

    api_key = _get_dashscope_api_key()
    if not api_key:
        raise HTTPException(status_code=400, detail="AI 功能不可用：请配置 API Key")

    from backend.api.ai_service import call_ai_async
    try:
        report_text = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 学情分析调用失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 分析出错: {str(e)}")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_heading(f"{grade}{cls}班 学情分析报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"年级：{grade}  班级：{cls}班  学生人数：{total_students}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    _markdown_to_docx(doc, report_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import urllib.parse
    safe_filename = urllib.parse.quote(f"学情分析_{grade}{cls}班.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )


@router.get("/teaching-suggestions/export", summary="导出 AI 教学建议为 Word 文档")
async def export_teaching_suggestions_docx(
    request: Request,
    grade: str = Query(...),
    cls: str = Query(...),
    teacher_username: str = Query(""),
    token: str = Query(""),
):
    """导出 AI 教学建议为 Word 文档"""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    if token:
        request.state.user = None
        from backend.auth import decode_jwt_token
        payload = decode_jwt_token(token)
        if payload:
            request.state.user = payload

    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可导出")

    query_teacher = teacher_username if (role == 0 and teacher_username) else username
    class_num = _extract_class_num(cls)
    cls_name = f"{grade}{class_num}班"

    student_count = execute_query(
        "SELECT COUNT(*) FROM users WHERE role = 2 AND grade = ? AND class = ?",
        (grade, class_num),
    )
    total_students = student_count[0][0] if student_count else 0
    if total_students == 0:
        raise HTTPException(status_code=400, detail="暂无该班级的数据")

    score_stats = execute_query(
        """SELECT COUNT(*), COALESCE(SUM(score),0), COALESCE(AVG(score),0),
                  COALESCE(MAX(score),0), COALESCE(MIN(score),0)
           FROM scores WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls_name),
    )
    rc_stats = execute_query(
        """SELECT COUNT(*),
                  SUM(CASE WHEN result='correct' THEN 1 ELSE 0 END),
                  SUM(CASE WHEN result='incorrect' THEN 1 ELSE 0 END)
           FROM rollcall_history WHERE teacher_username = ? AND grade = ? AND class_name = ?""",
        (query_teacher, grade, cls_name),
    )
    exam_stats = q_execute_query(
        """SELECT e.title, e.subject, COUNT(ea.id) as attempt_count,
                  COALESCE(AVG(ea.score),0) as avg_score, e.total_score
           FROM exams e LEFT JOIN exam_attempts ea ON ea.exam_id = e.id
           WHERE e.status IN ('published', 'ended')
           GROUP BY e.id ORDER BY e.created_at DESC LIMIT 5""",
    )
    task_stats = execute_query(
        "SELECT COUNT(*) FROM tasks WHERE creator_username = ? AND status = 'active'",
        (query_teacher,),
    )
    task_submissions = execute_query(
        """SELECT COUNT(DISTINCT ts.student_username)
           FROM task_submissions ts JOIN tasks t ON ts.task_id = t.id
           WHERE t.creator_username = ?""",
        (query_teacher,),
    )

    sc = score_stats[0] if score_stats else (0, 0, 0, 0, 0)
    rc = rc_stats[0] if rc_stats else (0, 0, 0)
    rc_total = _safe_int(rc[0])
    rc_correct = _safe_int(rc[1])
    active_tasks = _safe_int(task_stats[0][0] if task_stats else 0)
    submitted = _safe_int(task_submissions[0][0] if task_submissions else 0)

    exam_text = ""
    for e in exam_stats:
        exam_text += f"- 《{e['title']}》({e['subject']}): 参考{e['attempt_count']}人, 平均分{round(_safe_float(e['avg_score']), 1)}/{e['total_score']}\n"

    prompt = f"""你是一位高中信息科技教师。请根据以下教学数据，为{grade}{class_num}班生成具体的教学建议。

【班级概况】
年级：{grade}
班级：{class_num}班
学生人数：{total_students}

【课堂积分】
有积分记录学生数：{_safe_int(sc[0])}
总积分：{_safe_int(sc[1])}
平均积分：{round(_safe_float(sc[2]), 1)}

【点名情况】
总点名次数：{rc_total}
回答正确：{rc_correct}
正确率：{round(rc_correct / max(rc_total, 1) * 100, 1)}%

【考试情况】
{exam_text}

【任务情况】
活跃任务数：{active_tasks}
已提交学生数：{submitted}
提交率：{round(submitted / max(total_students, 1) * 100, 1)}%

请生成以下内容（Markdown 格式）：
1. 🎯 **教学重点建议**
2. 📝 **课堂互动建议**
3. ⏰ **任务安排建议**
4. 💪 **分层教学建议**
5. 📊 **预期改进目标**

请直接以建议内容开头，语气亲切专业。"""

    api_key = _get_dashscope_api_key()
    if not api_key:
        raise HTTPException(status_code=400, detail="AI 功能不可用：请配置 API Key")

    from backend.api.ai_service import call_ai_async
    try:
        suggestions_text = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 教学建议调用失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 分析出错: {str(e)}")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_heading(f"{grade}{class_num}班 AI 教学建议", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"年级：{grade}  班级：{class_num}班  学生人数：{total_students}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    _markdown_to_docx(doc, suggestions_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import urllib.parse
    safe_filename = urllib.parse.quote(f"教学建议_{grade}{class_num}班.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )


@router.get("/exam/{exam_id}/report/export", summary="导出考试 AI 分析报告为 Word 文档")
async def export_exam_report_docx(
    exam_id: int,
    request: Request,
    token: str = Query(""),
):
    """导出 AI 考试分析报告为 Word 文档"""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    if token:
        request.state.user = None
        from backend.auth import decode_jwt_token
        payload = decode_jwt_token(token)
        if payload:
            request.state.user = payload

    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可导出")

    exam = q_execute_query("SELECT * FROM exams WHERE id = ?", (exam_id,))
    if not exam:
        raise HTTPException(status_code=404, detail="考试不存在")
    exam = exam[0]

    attempts = q_execute_query(
        "SELECT * FROM exam_attempts WHERE exam_id = ? AND status IN ('submitted', 'graded') AND auto_graded = 1 AND answers IS NOT NULL AND answers != '' ORDER BY score DESC",
        (exam_id,),
    )
    questions = q_execute_query(
        """SELECT qb.id, qb.type, qb.question_text, qb.correct_answer, qb.difficulty, qb.knowledge_points, eq.score as question_score
           FROM exam_questions eq JOIN question_bank qb ON eq.question_id = qb.id
           WHERE eq.exam_id = ? ORDER BY eq.sort_order""",
        (exam_id,),
    )

    scores = [a['score'] for a in attempts]
    total_count = len(scores)
    avg_score = round(sum(scores) / max(total_count, 1), 1)
    pass_count = sum(1 for s in scores if s >= exam['pass_score'])
    max_score = max(scores) if scores else 0
    min_score = min(scores) if scores else 0

    q_accuracy = []
    for q in questions:
        correct = 0
        for a in attempts:
            if a.get('answers'):
                try:
                    answers = json.loads(a['answers']) if isinstance(a['answers'], str) else a['answers']
                except (json.JSONDecodeError, TypeError):
                    answers = {}
            else:
                answers = {}
            q_key = str(q['id'])
            ans_entry = answers.get(q_key, {})
            if isinstance(ans_entry, dict):
                if ans_entry.get('is_correct'):
                    correct += 1
            elif ans_entry == q['correct_answer']:
                correct += 1
        rate = round(correct / max(total_count, 1) * 100, 1)
        q_accuracy.append({
            "id": q['id'], "type": q['type'], "text": q['question_text'][:80],
            "correct_rate": rate, "difficulty": q.get('difficulty', 'medium'),
            "knowledge_points": q.get('knowledge_points', ''),
        })

    weak_points = [q for q in q_accuracy if q['correct_rate'] < 60]

    prompt = f"""你是一位高中{exam['subject']}教师。请根据以下考试数据，生成专业的考试分析报告。

考试名称：{exam['title']}
科目：{exam['subject']}
参考人数：{total_count}
平均分：{avg_score}/{exam['total_score']}
最高分：{max_score}
最低分：{min_score}
及格人数：{pass_count}/{total_count}（及格线{exam['pass_score']}分）

各题正确率（正确率 < 60% 的为薄弱题）：
{chr(10).join(f'- 第{i+1}题 ({q["type"]}, {q["difficulty"]}): 正确率{q["correct_rate"]}%' + (f' [知识点: {q["knowledge_points"]}]' if q['knowledge_points'] else '') for i, q in enumerate(q_accuracy))}

薄弱知识点：
{chr(10).join(('- ' + (q["knowledge_points"] or (f"第{q['id']}题"))) for q in weak_points) if weak_points else '无显著薄弱点'}

请生成（Markdown 格式）：
1. 📊 **总体情况**
2. 📈 **分数分布分析**
3. ⚠️ **薄弱知识点**
4. 💡 **教学改进建议"""

    api_key = _get_dashscope_api_key()
    if not api_key:
        raise HTTPException(status_code=400, detail="AI 功能不可用：请配置 API Key")

    from backend.api.ai_service import call_ai_async
    try:
        report_text = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 考试分析调用失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 分析出错: {str(e)}")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_heading(f"考试分析报告 - {exam['title']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"科目：{exam['subject']}  参考人数：{total_count}  平均分：{avg_score}/{exam['total_score']}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    _markdown_to_docx(doc, report_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import urllib.parse
    safe_filename = urllib.parse.quote(f"考试分析_{exam['title']}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )

