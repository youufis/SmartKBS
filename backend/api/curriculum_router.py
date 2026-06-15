"""
课程大纲 API 路由
课程 → 章/节 → 知识点 四级树形结构管理
支持资源绑定、学习进度追踪
"""
import json
import os
from datetime import datetime
from typing import Any, Optional

from fastapi import APIRouter, HTTPException, Request, Query
from pydantic import BaseModel

from backend.database import execute_query_dict as execute_query, execute_insert_update, get_connection
from backend.question_db import execute_query as q_execute_query
from backend.api.dependencies import get_current_user
from backend.auth import is_admin, is_teacher
from backend.logger import logger
from backend.api.chat_router import get_api_keys

router = APIRouter()


# ═══════════════════════════════════════════════════════════
# Pydantic 请求/响应模型
# ═══════════════════════════════════════════════════════════

class CourseCreate(BaseModel):
    name: str
    code: str = ""
    description: str = ""
    grade: str = ""
    cover_image: str = ""
    sort_order: int = 0
    subject: str = ""

class CourseUpdate(BaseModel):
    name: str | None = None
    code: str | None = None
    description: str | None = None
    grade: str | None = None
    cover_image: str | None = None
    sort_order: int | None = None
    status: str | None = None
    subject: str | None = None

class ChapterCreate(BaseModel):
    course_id: int
    parent_id: int | None = None
    name: str
    description: str = ""
    sort_order: int = 0

class ChapterUpdate(BaseModel):
    name: str | None = None
    description: str | None = None
    parent_id: int | None = None
    sort_order: int | None = None
    status: str | None = None

class KnowledgePointCreate(BaseModel):
    chapter_id: int
    name: str
    description: str = ""
    learning_objectives: str = ""
    difficulty: str = "medium"
    estimated_minutes: int = 0
    sort_order: int = 0

class KnowledgePointUpdate(BaseModel):
    name: str | None = None
    description: str | None = None
    learning_objectives: str | None = None
    difficulty: str | None = None
    estimated_minutes: int | None = None
    sort_order: int | None = None
    status: str | None = None

class BindingCreate(BaseModel):
    knowledge_point_id: int
    resource_type: str
    resource_id: int
    sort_order: int = 0

class ProgressUpdate(BaseModel):
    status: str = "in_progress"  # not_started | in_progress | completed
    score: float | None = None


# ═══════════════════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════════════════

def _can_manage(user: dict[str, Any]) -> bool:
    """检查是否有管理权限（教师/管理员）"""
    role = user.get("role", 2)
    return role in (0, 1)


def _check_resource_ownership(resource_type: str, resource_id: int, username: str) -> bool:
    """校验指定资源是否属于该用户"""
    try:
        if resource_type in ("html", "download"):
            row = execute_query_one(
                "SELECT 1 FROM shared_resources WHERE id=? AND owner_username=?",
                (resource_id, username),
            )
            return row is not None
        elif resource_type == "exam":
            row = q_execute_query(
                "SELECT 1 FROM exams WHERE id=? AND creator_username=?",
                (resource_id, username),
            )
            return len(row) > 0
        elif resource_type == "discussion":
            row = execute_query_one(
                "SELECT 1 FROM discussions WHERE id=? AND creator_username=?",
                (resource_id, username),
            )
            return row is not None
        elif resource_type == "interaction_quiz":
            row = execute_query_one(
                "SELECT 1 FROM interaction_quizzes WHERE id=? AND creator_username=?",
                (resource_id, username),
            )
            return row is not None
        elif resource_type == "task":
            row = execute_query_one(
                "SELECT 1 FROM tasks WHERE id=? AND creator_username=?",
                (str(resource_id), username),
            )
            return row is not None
    except Exception:
        pass
    return False


def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


# 课程名称 → 所属大类的推断映射（用于已有课程回填）
_SUBJECT_KEYWORDS: list[tuple[list[str], str]] = [
    (["技术与设计"], "通用技术"),
    (["数据与计算", "信息系统与社会"], "信息科技"),
    (["人工智能"], "人工智能"),
]


def _infer_subject(course_name: str) -> str:
    """根据课程名称推断所属大类，无法推断时返回空字符串"""
    for keywords, subject in _SUBJECT_KEYWORDS:
        for kw in keywords:
            if kw in course_name:
                return subject
    return ""


def _ensure_subject(course_dict: dict[str, Any]) -> dict[str, Any]:
    """确保课程有 subject 字段，为空时尝试从名称推断"""
    subject = course_dict.get("subject", "") or ""
    if not subject:
        subject = _infer_subject(course_dict.get("name", ""))
        course_dict["subject"] = subject
    return course_dict


def _build_course_tree(course_id: int) -> list[dict[str, Any]]:
    """构建课程的完整章节-知识点树（批量查询优化版）

    优化说明：将 N+1 次递归查询合并为 2 次批量查询，
    在内存中用字典组装树结构，大幅减少 SQL 查询次数。
    """
    # 1) 一次性查询该课程所有章节
    all_chapters = execute_query(
        """SELECT * FROM chapters
           WHERE course_id=? AND status='active'
           ORDER BY sort_order, id""",
        (course_id,),
    )
    # 2) 一次性查询该课程所有知识点
    all_kps = execute_query(
        """SELECT kp.* FROM knowledge_points kp
           JOIN chapters ch ON ch.id = kp.chapter_id
           WHERE ch.course_id=? AND kp.status='active'
           ORDER BY kp.sort_order, kp.id""",
        (course_id,),
    )

    # 2.5) 批量查询所有知识点的资源绑定数量
    kp_ids = [kp["id"] for kp in all_kps]
    resource_count_map: dict[int, int] = {}
    if kp_ids:
        placeholders = ",".join("?" for _ in kp_ids)
        counts = execute_query(
            f"SELECT knowledge_point_id, COUNT(*) as cnt FROM curriculum_bindings WHERE knowledge_point_id IN ({placeholders}) GROUP BY knowledge_point_id",
            tuple(kp_ids),
        )
        for row in counts:
            resource_count_map[row["knowledge_point_id"]] = row["cnt"]

    # 3) 构建 parent_id → 章节列表 的映射
    children_map: dict[int, list[dict[str, Any]]] = {}
    for ch in all_chapters:
        pid = ch["parent_id"] or 0  # 顶层用 0 表示
        children_map.setdefault(pid, []).append(ch)

    # 4) 构建 chapter_id → 知识点列表 的映射
    kp_map: dict[int, list[dict[str, Any]]] = {}
    for kp in all_kps:
        kp_map.setdefault(kp["chapter_id"], []).append(kp)

    # 5) 递归组装树
    def _build_node(ch: dict[str, Any]) -> dict[str, Any]:
        node = dict(ch)
        # 子章节
        node["children"] = [_build_node(c) for c in children_map.get(ch["id"], [])]
        # 知识点
        kps = kp_map.get(ch["id"], [])
        for kp in kps:
            kp["resource_count"] = resource_count_map.get(kp["id"], 0)
        node["knowledge_points"] = kps
        return node

    return [_build_node(ch) for ch in children_map.get(0, [])]


def _inject_progress(kps: list[dict[str, Any]], username: str):
    """为学生注入学习进度状态"""
    if not kps:
        return
    kp_ids = [kp["id"] for kp in kps]
    placeholders = ",".join("?" for _ in kp_ids)
    rows = execute_query(
        f"SELECT knowledge_point_id, status, score FROM learning_progress WHERE student_username=? AND knowledge_point_id IN ({placeholders})",
        (username, *kp_ids),
    )
    progress_map = {r["knowledge_point_id"]: r for r in rows}
    for kp in kps:
        p = progress_map.get(kp["id"])
        if p:
            kp["progress_status"] = p["status"]
            kp["progress_score"] = p["score"]
        else:
            kp["progress_status"] = "not_started"
            kp["progress_score"] = 0


def _get_resource_info(resource_type: str, resource_id: int) -> dict[str, Any]:
    """根据资源类型和 ID 获取资源名称和访问路径"""
    result = {"name": "", "url": ""}
    try:
        if resource_type in ("html", "download"):
            rows = execute_query(
                "SELECT file_name, file_path FROM shared_resources WHERE id=? AND resource_type=?",
                (resource_id, resource_type),
            )
            if rows:
                result["name"] = rows[0]["file_name"]
                result["url"] = "/api/files/" + rows[0]["file_path"].lstrip("/")
        elif resource_type == "question":
            rows = q_execute_query(
                "SELECT question_text FROM question_bank WHERE id=? AND status='active'",
                (resource_id,),
            )
            if rows:
                text = rows[0]["question_text"]
                result["name"] = text[:60] + ("..." if len(text) > 60 else "")
                result["url"] = f"/questions?highlight={resource_id}"
        elif resource_type == "exam":
            rows = q_execute_query(
                "SELECT title FROM exams WHERE id=?",
                (resource_id,),
            )
            if rows:
                result["name"] = rows[0]["title"]
                result["url"] = f"/exam?highlight={resource_id}"
        elif resource_type == "discussion":
            rows = execute_query(
                "SELECT title FROM discussions WHERE id=?",
                (resource_id,),
            )
            if rows:
                result["name"] = rows[0]["title"]
                result["url"] = f"/discussion?highlight={resource_id}"
        elif resource_type == "interaction_quiz":
            rows = execute_query(
                "SELECT title FROM interaction_quizzes WHERE id=?",
                (resource_id,),
            )
            if rows:
                result["name"] = rows[0]["title"]
                result["url"] = f"/interaction?highlight={resource_id}"
        elif resource_type == "task":
            rows = execute_query(
                "SELECT name FROM tasks WHERE id=?",
                (str(resource_id),),
            )
            if rows:
                result["name"] = rows[0]["name"]
                result["url"] = f"/tasks?highlight={resource_id}"
    except Exception:
        pass
    if not result["name"]:
        result["name"] = f"[{resource_type}:{resource_id}]"
    return result


# ═══════════════════════════════════════════════════════════
# AI 辅助生成课程结构
# ═══════════════════════════════════════════════════════════

class AIGenerateRequest(BaseModel):
    """AI 生成课程请求"""
    content: str = ""           # 文本内容
    subject: str = ""  # 科目（由前端传递）
    grade: str = ""    # 年级（由前端传递）
    course_name: str = ""       # 课程名称（留空由 AI 推断）
    auto_save: bool = False     # 是否自动保存到数据库

@router.post("/ai-generate", summary="AI 辅助生成课程结构")
async def ai_generate_curriculum(req: AIGenerateRequest, request: Request):
    """上传教学内容文本，AI 自动提取课程→章节→知识点结构"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    if not req.content.strip():
        raise HTTPException(status_code=400, detail="请输入教学内容文本")

    # 获取 API Key（复用 chat_router 的缓存逻辑）
    api_key, _ = get_api_keys(user["username"])
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key，请在系统配置中设置")

    # 构造 Prompt
    course_hint = f"课程名称：{req.course_name}" if req.course_name else "请根据内容推断课程名称"
    prompt = f"""你是教学大纲设计师。将以下教学内容提取为 JSON 大纲。

【三条铁律】
1. 结构: 章→节→知识点，知识点只放在节下
2. ⚠️ 知识点名称绝不能与节名相同
3. 每个节至少拆出 1~2 个知识点

科目：{req.subject}，年级：{req.grade}，{course_hint}

✅ 正确: 节="机器学习基础" → 知识点=["监督学习算法","模型评估方法"]
❌ 禁止: 节="机器学习基础" → 知识点=["机器学习基础"]  ← 完全重复！

输出格式：
{{"course_name":"...","chapters":[{{"name":"章","children":[
  {{"name":"节","knowledge_points":[{{"name":"知识点","difficulty":"easy","estimated_minutes":20}}]}}
]}}]}}

内容：
{req.content[:8000]}"""

    try:
        from backend.api.ai_service import call_ai_async
        ai_response = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 生成课程失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 调用失败: {str(e)}")

    # 解析 AI 返回的 JSON
    result = _parse_ai_json(ai_response)
    if not result:
        raise HTTPException(status_code=500, detail="AI 返回格式异常，无法解析为课程结构")

    # 后处理：修正 AI 输出的不合格知识点（同名、缺失等）
    _fix_kp_names(result)

    # 自动保存
    if req.auto_save:
        saved = _save_ai_result(result, req.subject, req.grade, user["username"])
        result["saved"] = saved

    return result


@router.post("/ai-generate-from-file", summary="上传文档 AI 生成课程结构")
async def ai_generate_from_file(request: Request):
    """上传文档文件（txt/md/pdf/docx），AI 自动提取课程→章节→知识点结构"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    # 获取 API Key
    api_key, _ = get_api_keys(user["username"])
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key，请在系统配置中设置")

    # 解析 multipart 表单
    from starlette.datastructures import UploadFile
    form = await request.form()
    file_raw = form.get("file")
    subject_val = form.get("subject", "")
    grade_val = form.get("grade", "")
    course_name = form.get("course_name", "")
    auto_save = form.get("auto_save", "false") == "true"

    if not file_raw or not isinstance(file_raw, UploadFile) or not file_raw.filename:
        raise HTTPException(status_code=400, detail="请上传文件")
    file: UploadFile = file_raw
    subject: str = str(subject_val) if subject_val else ""
    grade: str = str(grade_val) if grade_val else ""

    # 读取文件内容
    content_bytes = await file.read()
    filename = (file.filename or "").lower()

    # 提取文本（按扩展名处理）
    text_content = ""
    if filename.endswith(".txt") or filename.endswith(".md"):
        text_content = content_bytes.decode("utf-8", errors="replace")
    elif filename.endswith(".pdf"):
        try:
            import io
            import importlib
            PyPDF2 = importlib.import_module('PyPDF2')
            reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            text_content = "\n".join(page.extract_text() for page in reader.pages)
        except ImportError:
            # 无 PyPDF2 时尝试 pdfminer
            try:
                import io
                from pdfminer.high_level import extract_text as pdf_extract  # type: ignore
                text_content = pdf_extract(io.BytesIO(content_bytes))
            except ImportError:
                raise HTTPException(status_code=400, detail="缺少 PDF 解析库，请安装 PyPDF2 或 pdfminer")
    elif filename.endswith(".docx"):
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(content_bytes))
            text_content = "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            raise HTTPException(status_code=400, detail="缺少 DOCX 解析库，请安装 python-docx")
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {filename}，支持 txt/md/pdf/docx")

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="文件内容为空或无法提取文本")

    if len(text_content) > 20000:
        text_content = text_content[:20000] + "\n\n[内容已截断，仅处理前 20000 字符]"

    # 构造 Prompt
    course_hint = f"课程名称：{course_name}" if course_name else "请根据内容推断课程名称"
    prompt = f"""你是教学大纲设计师。将文件「{filename}」中的教学内容提取为 JSON 大纲。

【三条铁律】
1. 结构: 章→节→知识点，知识点只放在节下
2. ⚠️ 知识点名称绝不能与节名相同
3. 每个节至少拆出 1~2 个知识点

科目：{subject}，年级：{grade}，{course_hint}

✅ 正确: 节="机器学习基础" → 知识点=["监督学习算法","模型评估方法"]
❌ 禁止: 节="机器学习基础" → 知识点=["机器学习基础"]  ← 完全重复！

输出格式：
{{"course_name":"...","chapters":[{{"name":"章","children":[
  {{"name":"节","knowledge_points":[{{"name":"知识点","difficulty":"easy","estimated_minutes":20}}]}}
]}}]}}

内容：
{text_content}"""

    try:
        from backend.api.ai_service import call_ai_async
        ai_response = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 文件生成课程失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 调用失败: {str(e)}")

    # 解析 AI 返回的 JSON
    result = _parse_ai_json(ai_response)
    if not result:
        raise HTTPException(status_code=500, detail="AI 返回格式异常，无法解析为课程结构")

    # 后处理：修正 AI 输出的不合格知识点（同名、缺失等）
    _fix_kp_names(result)

    # 自动保存
    if auto_save:
        saved = _save_ai_result(result, subject, grade, user["username"])
        result["saved"] = saved

    result["source_file"] = filename
    return result


def _parse_ai_json(text: str) -> dict[str, Any] | None:
    """从 AI 返回文本中提取 JSON"""
    # 尝试直接解析
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # 尝试从 ```json ``` 代码块中提取
    import re
    match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    # 尝试从 { 到 } 提取最外层 JSON
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            pass

    return None


def _save_ai_result(result: dict[str, Any], subject: str, grade: str, username: str) -> dict[str, Any]:
    """将 AI 生成的结构保存到数据库"""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    saved = {"course_id": None, "chapters": 0, "knowledge_points": 0}

    # 1. 创建课程
    course_name = result.get("course_name", f"{subject}课程")
    course_code = result.get("course_code", "")
    course_desc = result.get("course_description", f"AI 自动生成的{subject}课程大纲")

    course_id = execute_insert_update(
        """INSERT INTO courses (name, code, description, grade, subject, sort_order, status, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, 0, 'active', ?, ?)""",
        (course_name, course_code, course_desc, grade, subject, now, now),
    )
    saved["course_id"] = course_id

    # 2. 创建章节和知识点
    chapters = result.get("chapters", [])
    for ch_idx, ch in enumerate(chapters):
        ch_id = execute_insert_update(
            """INSERT INTO chapters (course_id, parent_id, name, description, sort_order, status, created_at, updated_at)
               VALUES (?, NULL, ?, ?, ?, 'active', ?, ?)""",
            (course_id, ch.get("name", ""), ch.get("description", ""), ch_idx, now, now),
        )
        saved["chapters"] += 1

        # 子章节（节）
        children = ch.get("children", [])
        top_kps = ch.get("knowledge_points", [])

        # 后处理：如果章既有子节又有章级知识点，将章级知识点合并到子节中
        # 确保知识点不会与节平级（用户期望的结构）
        if top_kps and children:
            for kp in top_kps:
                kp_name = kp.get("name", "")
                # 查找同名节
                matched = [s for s in children if s.get("name", "").strip() == kp_name.strip()]
                if matched:
                    # 追加到同名节的 knowledge_points 中
                    matched[0].setdefault("knowledge_points", []).append(kp)
                else:
                    # 以知识点名称创建新节，将知识点放入其中
                    children.append({
                        "name": kp_name,
                        "description": kp.get("description", ""),
                        "knowledge_points": [kp],
                    })
            ch["knowledge_points"] = []  # 清空章级知识点，避免重复

        for sec_idx, sec in enumerate(children):
            sec_id = execute_insert_update(
                """INSERT INTO chapters (course_id, parent_id, name, description, sort_order, status, created_at, updated_at)
                   VALUES (?, ?, ?, ?, ?, 'active', ?, ?)""",
                (course_id, ch_id, sec.get("name", ""), sec.get("description", ""), sec_idx, now, now),
            )
            saved["chapters"] += 1

            # 该节下专属的知识点
            if sec_id is not None:
                for kp_idx, kp in enumerate(sec.get("knowledge_points", [])):
                    _insert_kp(sec_id, kp, kp_idx, now)
                    saved["knowledge_points"] += 1

        # 只有章完全没有子节时，知识点才挂在章下
        if top_kps and not children and ch_id is not None:
            for kp_idx, kp in enumerate(top_kps):
                _insert_kp(ch_id, kp, kp_idx, now)
                saved["knowledge_points"] += 1

    logger.info(f"AI 生成课程已保存: {course_name} (id={course_id}), {saved}")
    return saved


def _insert_kp(chapter_id: int, kp: dict[str, Any], sort_order: int, now: str):
    """插入单个知识点"""
    execute_insert_update(
        """INSERT INTO knowledge_points (chapter_id, name, description, learning_objectives, difficulty, estimated_minutes, sort_order, status, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, 'active', ?, ?)""",
        (
            chapter_id,
            kp.get("name", ""),
            kp.get("description", ""),
            kp.get("learning_objectives", ""),
            kp.get("difficulty", "medium"),
            kp.get("estimated_minutes", 30),
            sort_order,
            now,
            now,
        ),
    )


# ═══════════════════════════════════════════════════════════
# AI 结果后处理：修正不合格的知识点
# ═══════════════════════════════════════════════════════════

def _fix_kp_names(result: dict[str, Any]):
    """修正 AI 返回结果中知识点名称不合格的情况"""
    for ch in result.get("chapters", []):
        for sec in ch.get("children", []):
            kps = sec.get("knowledge_points", [])
            sec_name = sec.get("name", "").strip()

            # 修正1: 知识点名与节名相同时，自动生成更具体的名称
            # 修正2: 知识点数量太少时，补充默认知识点
            new_kps = []
            for kp in kps:
                kp_name = kp.get("name", "").strip()
                if kp_name == sec_name:
                    # 根据节名自动拆解知识点
                    generated = _split_section_to_kps(sec_name, kp)
                    new_kps.extend(generated)
                else:
                    new_kps.append(kp)

            # 如果处理后仍然没有知识点，生成默认知识点
            if not new_kps and sec_name:
                new_kps.append({
                    "name": f"{sec_name}概述",
                    "description": sec.get("description", ""),
                    "learning_objectives": f"掌握{sec_name}的基本概念",
                    "difficulty": "medium",
                    "estimated_minutes": 25,
                })

            sec["knowledge_points"] = new_kps


def _split_section_to_kps(sec_name: str, original_kp: dict[str, Any]) -> list[dict[str, Any]]:
    """将节名拆解为多个具体的知识点名称"""
    difficulty = original_kp.get("difficulty", "medium")
    minutes = original_kp.get("estimated_minutes", 30)
    desc = original_kp.get("description", "")
    obj = original_kp.get("learning_objectives", "")

    # 尝试从节名中提取子知识点
    # 1) 包含"与"的：拆分为左右两部分
    if "与" in sec_name:
        parts = sec_name.split("与", 1)
        return [
            {
                "name": parts[0].strip(),
                "description": desc,
                "learning_objectives": obj or f"理解{parts[0].strip()}",
                "difficulty": difficulty,
                "estimated_minutes": max(minutes // 2, 15),
            },
            {
                "name": parts[1].strip(),
                "description": desc,
                "learning_objectives": obj or f"掌握{parts[1].strip()}",
                "difficulty": difficulty,
                "estimated_minutes": max(minutes - minutes // 2, 15),
            },
        ]

    # 2) 其他：生成"概述/核心概念/应用实践"等通用知识点
    return [
        {
            "name": f"{sec_name}基本概念",
            "description": desc,
            "learning_objectives": obj or f"理解{sec_name}的基本概念",
            "difficulty": "easy",
            "estimated_minutes": max(minutes // 2, 15),
        },
        {
            "name": f"{sec_name}核心原理",
            "description": desc,
            "learning_objectives": obj or f"掌握{sec_name}的核心原理",
            "difficulty": difficulty,
            "estimated_minutes": max(minutes - minutes // 2, 15),
        },
    ]


# ═══════════════════════════════════════════════════════════
# 完整树查询（核心端点）
# ═══════════════════════════════════════════════════════════

@router.get("/tree", summary="获取全部课程完整树形结构")
async def get_curriculum_tree(request: Request):
    """返回课程→章→节→知识点的完整嵌套结构，学生视图会注入学习进度"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    user_grade = user.get("grade", "")

    # 查询课程列表（学生按年级匹配）
    if role == 2 and user_grade:
        grades = [g.strip() for g in user_grade.split("|")]
        conditions = []
        params = []
        for g in grades:
            conditions.append("grade LIKE ?")
            params.append(f"%{g}%")
        where = " OR ".join(conditions)
        courses = execute_query(
            f"SELECT * FROM courses WHERE status='active' AND ({where}) ORDER BY sort_order, id",
            tuple(params),
        )
    else:
        courses = execute_query(
            "SELECT * FROM courses WHERE status='active' ORDER BY sort_order, id",
        )

    result = []
    for course in courses:
        course_dict = dict(course)
        course_dict = _ensure_subject(course_dict)
        course_dict["chapters"] = _build_course_tree(course["id"])

        # 为所有知识点注入进度（学生视图）
        if role == 2:
            all_kps = []
            def _collect_kps(nodes):
                for node in nodes:
                    _collect_kps(node.get("children", []))
                    all_kps.extend(node.get("knowledge_points", []))
            _collect_kps(course_dict["chapters"])
            _inject_progress(all_kps, username)

            # 计算进度统计
            total = len(all_kps)
            completed = sum(1 for kp in all_kps if kp.get("progress_status") == "completed")
            course_dict["progress"] = {"total": total, "completed": completed}

        result.append(course_dict)

    return result


# ═══════════════════════════════════════════════════════════
# 课程 CRUD
# ═══════════════════════════════════════════════════════════

@router.get("/courses", summary="获取课程列表")
async def list_courses(
    request: Request,
    grade: str = Query(None, description="筛选年级"),
    status: str = Query(None, description="筛选状态"),
):
    """获取课程列表，支持按年级和状态筛选"""
    user = get_current_user(request)
    conditions = []
    params = []

    if grade:
        conditions.append("grade LIKE ?")
        params.append(f"%{grade}%")
    if status:
        conditions.append("status = ?")
        params.append(status)

    where = " AND ".join(conditions) if conditions else "1=1"
    rows = execute_query(
        f"SELECT * FROM courses WHERE {where} ORDER BY sort_order, id",
        tuple(params),
    )
    for row in rows:
        _ensure_subject(row)
    return {"courses": rows, "total": len(rows)}


@router.get("/courses/{course_id}", summary="获取课程详情（含树结构）")
async def get_course(course_id: int, request: Request):
    """获取课程信息及完整的章节-知识点树"""
    user = get_current_user(request)
    course = execute_query_one("SELECT * FROM courses WHERE id=?", (course_id,))
    if not course:
        raise HTTPException(status_code=404, detail="课程不存在")

    course_dict = dict(course)
    course_dict = _ensure_subject(course_dict)
    course_dict["chapters"] = _build_course_tree(course_id)

    # 学生注入进度
    if user.get("role") == 2:
        all_kps = []
        def _collect_kps(nodes):
            for node in nodes:
                _collect_kps(node.get("children", []))
                all_kps.extend(node.get("knowledge_points", []))
        _collect_kps(course_dict["chapters"])
        _inject_progress(all_kps, user["username"])
        total = len(all_kps)
        completed = sum(1 for kp in all_kps if kp.get("progress_status") == "completed")
        course_dict["progress"] = {"total": total, "completed": completed}

    return course_dict


@router.post("/courses", summary="创建课程")
async def create_course(req: CourseCreate, request: Request):
    """创建新课程（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    if not req.name.strip():
        raise HTTPException(status_code=400, detail="请输入课程名称")

    now = _now()
    course_id = execute_insert_update(
        """INSERT INTO courses (name, code, description, grade, cover_image, sort_order, subject, status, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, 'active', ?, ?)""",
        (req.name, req.code, req.description, req.grade, req.cover_image, req.sort_order, req.subject, now, now),
    )
    logger.info(f"用户 {user['username']} 创建课程: {req.name} (id={course_id})")
    return {"message": f"课程「{req.name}」创建成功", "course_id": course_id}


@router.put("/courses/{course_id}", summary="更新课程")
async def update_course(course_id: int, req: CourseUpdate, request: Request):
    """更新课程信息（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    course = execute_query_one("SELECT * FROM courses WHERE id=?", (course_id,))
    if not course:
        raise HTTPException(status_code=404, detail="课程不存在")

    updates = {}
    for field in ["name", "code", "description", "grade", "cover_image", "sort_order", "status", "subject"]:
        val = getattr(req, field, None)
        if val is not None:
            updates[field] = val

    if not updates:
        return {"message": "无更新内容"}

    now = _now()
    updates["updated_at"] = now
    set_clause = ", ".join(f"{k}=?" for k in updates)
    params = list(updates.values()) + [course_id]
    execute_insert_update(
        f"UPDATE courses SET {set_clause} WHERE id=?", tuple(params),
    )
    logger.info(f"用户 {user['username']} 更新课程 id={course_id}")
    return {"message": "课程更新成功"}


@router.delete("/courses/{course_id}", summary="删除课程（硬删除）")
async def delete_course(course_id: int, request: Request):
    """删除课程（管理员）—— 硬删除课程及所有关联数据"""
    user = get_current_user(request)
    if not is_admin(user.get("username", "")):
        raise HTTPException(status_code=403, detail="权限不足：需要管理员权限")

    course_rows = execute_query("SELECT * FROM courses WHERE id=?", (course_id,))
    if not course_rows:
        raise HTTPException(status_code=404, detail="课程不存在")
    course = course_rows[0]

    # 获取所有级联的章节ID和知识点ID
    chapter_rows = execute_query(
        "SELECT id FROM chapters WHERE course_id=?", (course_id,),
    )
    all_chapter_ids = [ch["id"] for ch in chapter_rows]
    # 也获取子章节
    if all_chapter_ids:
        placeholders = ",".join("?" for _ in all_chapter_ids)
        child_rows = execute_query(
            f"SELECT id FROM chapters WHERE parent_id IN ({placeholders})",
            tuple(all_chapter_ids),
        )
        all_chapter_ids.extend(ch["id"] for ch in child_rows if ch["id"] not in all_chapter_ids)

    if all_chapter_ids:
        placeholders = ",".join("?" for _ in all_chapter_ids)
        kp_rows = execute_query(
            f"SELECT id FROM knowledge_points WHERE chapter_id IN ({placeholders})",
            tuple(all_chapter_ids),
        )
        kp_ids = [kp["id"] for kp in kp_rows]
        if kp_ids:
            kp_placeholders = ",".join("?" for _ in kp_ids)
            kp_tuple = tuple(kp_ids)
            execute_insert_update(
                f"DELETE FROM learning_progress WHERE knowledge_point_id IN ({kp_placeholders})", kp_tuple,
            )
            execute_insert_update(
                f"DELETE FROM curriculum_bindings WHERE knowledge_point_id IN ({kp_placeholders})", kp_tuple,
            )
            execute_insert_update(
                f"DELETE FROM knowledge_prerequisites WHERE knowledge_point_id IN ({kp_placeholders})", kp_tuple,
            )
            execute_insert_update(
                f"DELETE FROM knowledge_prerequisites WHERE prerequisite_id IN ({kp_placeholders})", kp_tuple,
            )
            # 清理 AI 练习独立成绩
            from backend.question_db import execute_insert as q_exec2
            q_exec2(f"DELETE FROM ai_practice_results WHERE kp_id IN ({kp_placeholders})", kp_tuple)
        execute_insert_update(
            f"DELETE FROM knowledge_points WHERE chapter_id IN ({placeholders})", tuple(all_chapter_ids),
        )
        execute_insert_update(
            f"DELETE FROM chapters WHERE id IN ({placeholders})", tuple(all_chapter_ids),
        )
        execute_insert_update(
            f"DELETE FROM chapters WHERE parent_id IN ({placeholders})", tuple(all_chapter_ids),
        )

    execute_insert_update("DELETE FROM courses WHERE id=?", (course_id,))

    logger.info(f"管理员 {user['username']} 硬删除课程 id={course_id}")
    return {"message": f"课程「{course['name']}」已永久删除"}


# ═══════════════════════════════════════════════════════════
# 章节 CRUD
# ═══════════════════════════════════════════════════════════

def _build_chapter_node(ch: dict[str, Any]) -> dict[str, Any]:
    """构建单个章节的节点（含子章节和知识点）"""
    node = dict(ch)
    # 子章节
    children = execute_query(
        "SELECT * FROM chapters WHERE parent_id=? AND status='active' ORDER BY sort_order, id",
        (ch["id"],),
    )
    node["children"] = [_build_chapter_node(c) for c in children]
    # 知识点
    kps = execute_query(
        "SELECT * FROM knowledge_points WHERE chapter_id=? AND status='active' ORDER BY sort_order, id",
        (ch["id"],),
    )
    node["knowledge_points"] = kps
    return node


@router.get("/chapters/{chapter_id}", summary="获取章节详情")
async def get_chapter(chapter_id: int, request: Request):
    """获取章节信息（含子章节和知识点）"""
    get_current_user(request)
    ch = execute_query_one("SELECT * FROM chapters WHERE id=?", (chapter_id,))
    if not ch:
        raise HTTPException(status_code=404, detail="章节不存在")
    return _build_chapter_node(ch)


@router.post("/chapters", summary="创建章节")
async def create_chapter(req: ChapterCreate, request: Request):
    """创建章节（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")
    if not req.name.strip():
        raise HTTPException(status_code=400, detail="请输入章节名称")

    # 校验课程存在
    course = execute_query_one("SELECT id FROM courses WHERE id=?", (req.course_id,))
    if not course:
        raise HTTPException(status_code=404, detail="课程不存在")

    # 校验父章节存在（如果有）
    if req.parent_id:
        parent = execute_query_one("SELECT id FROM chapters WHERE id=?", (req.parent_id,))
        if not parent:
            raise HTTPException(status_code=404, detail="父章节不存在")

    now = _now()
    chapter_id = execute_insert_update(
        """INSERT INTO chapters (course_id, parent_id, name, description, sort_order, status, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, 'active', ?, ?)""",
        (req.course_id, req.parent_id, req.name, req.description, req.sort_order, now, now),
    )
    logger.info(f"用户 {user['username']} 创建章节: {req.name} (id={chapter_id})")
    return {"message": f"章节「{req.name}」创建成功", "chapter_id": chapter_id}


@router.put("/chapters/{chapter_id}", summary="更新章节")
async def update_chapter(chapter_id: int, req: ChapterUpdate, request: Request):
    """更新章节信息（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    ch = execute_query_one("SELECT * FROM chapters WHERE id=?", (chapter_id,))
    if not ch:
        raise HTTPException(status_code=404, detail="章节不存在")

    updates = {}
    for field in ["name", "description", "parent_id", "sort_order", "status"]:
        val = getattr(req, field, None)
        if val is not None:
            updates[field] = val

    if not updates:
        return {"message": "无更新内容"}

    now = _now()
    updates["updated_at"] = now
    set_clause = ", ".join(f"{k}=?" for k in updates)
    params = list(updates.values()) + [chapter_id]
    execute_insert_update(
        f"UPDATE chapters SET {set_clause} WHERE id=?", tuple(params),
    )
    logger.info(f"用户 {user['username']} 更新章节 id={chapter_id}")
    return {"message": "章节更新成功"}


@router.delete("/chapters/{chapter_id}", summary="删除章节（硬删除）")
async def delete_chapter(chapter_id: int, request: Request):
    """删除章节（教师/管理员）—— 硬删除章节及所有关联数据"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    ch = execute_query("SELECT * FROM chapters WHERE id=?", (chapter_id,))
    if not ch:
        raise HTTPException(status_code=404, detail="章节不存在")

    # 收集所有子章节ID
    children = execute_query(
        "SELECT id FROM chapters WHERE parent_id=?", (chapter_id,),
    )
    all_chapter_ids = [chapter_id] + [c["id"] for c in children]

    placeholders = ",".join("?" for _ in all_chapter_ids)
    ch_tuple = tuple(all_chapter_ids)

    # 获取所有知识点ID
    kp_rows = execute_query(
        f"SELECT id FROM knowledge_points WHERE chapter_id IN ({placeholders})", ch_tuple,
    )
    kp_ids = [kp["id"] for kp in kp_rows]
    if kp_ids:
        kp_placeholders = ",".join("?" for _ in kp_ids)
        kp_tuple = tuple(kp_ids)
        execute_insert_update(
            f"DELETE FROM learning_progress WHERE knowledge_point_id IN ({kp_placeholders})", kp_tuple,
        )
        execute_insert_update(
            f"DELETE FROM curriculum_bindings WHERE knowledge_point_id IN ({kp_placeholders})", kp_tuple,
        )
    execute_insert_update(
        f"DELETE FROM knowledge_points WHERE chapter_id IN ({placeholders})", ch_tuple,
    )
    execute_insert_update(
        f"DELETE FROM chapters WHERE parent_id IN ({placeholders})", ch_tuple,
    )
    execute_insert_update(
        f"DELETE FROM chapters WHERE id IN ({placeholders})", ch_tuple,
    )

    logger.info(f"用户 {user['username']} 硬删除章节 id={chapter_id}")
    return {"message": "章节已永久删除"}


# ═══════════════════════════════════════════════════════════
# 知识点 CRUD
# ═══════════════════════════════════════════════════════════

@router.get("/knowledge-points/{kp_id}", summary="获取知识点详情")
async def get_knowledge_point(kp_id: int, request: Request):
    """获取知识点详情（含绑定的资源列表）"""
    user = get_current_user(request)
    kp = execute_query_one("SELECT * FROM knowledge_points WHERE id=?", (kp_id,))
    if not kp:
        raise HTTPException(status_code=404, detail="知识点不存在")

    kp_dict = dict(kp)

    # 获取绑定的资源
    bindings = execute_query(
        """SELECT cb.* FROM curriculum_bindings cb
           WHERE cb.knowledge_point_id=? ORDER BY cb.sort_order, cb.id""",
        (kp_id,),
    )
    resources = []
    for b in bindings:
        info = _get_resource_info(b["resource_type"], b["resource_id"])
        resources.append({
            "binding_id": b["id"],
            "resource_type": b["resource_type"],
            "resource_id": b["resource_id"],
            "resource_name": info["name"],
            "resource_url": info["url"],
            "sort_order": b["sort_order"],
            "created_at": b["created_at"],
        })
    kp_dict["resources"] = resources

    # 学生注入进度
    if user.get("role") == 2:
        progress = execute_query_one(
            "SELECT status, score FROM learning_progress WHERE student_username=? AND knowledge_point_id=?",
            (user["username"], kp_id),
        )
        if progress:
            kp_dict["progress_status"] = progress["status"]
            kp_dict["progress_score"] = progress["score"]
        else:
            kp_dict["progress_status"] = "not_started"
            kp_dict["progress_score"] = 0

    return kp_dict


@router.post("/knowledge-points", summary="创建知识点")
async def create_knowledge_point(req: KnowledgePointCreate, request: Request):
    """创建知识点（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")
    if not req.name.strip():
        raise HTTPException(status_code=400, detail="请输入知识点名称")

    # 校验章节存在
    ch = execute_query_one("SELECT id FROM chapters WHERE id=?", (req.chapter_id,))
    if not ch:
        raise HTTPException(status_code=404, detail="章节不存在")

    now = _now()
    kp_id = execute_insert_update(
        """INSERT INTO knowledge_points (chapter_id, name, description, learning_objectives, difficulty, estimated_minutes, sort_order, status, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, 'active', ?, ?)""",
        (req.chapter_id, req.name, req.description, req.learning_objectives,
         req.difficulty, req.estimated_minutes, req.sort_order, now, now),
    )
    logger.info(f"用户 {user['username']} 创建知识点: {req.name} (id={kp_id})")
    return {"message": f"知识点「{req.name}」创建成功", "kp_id": kp_id}


@router.put("/knowledge-points/{kp_id}", summary="更新知识点")
async def update_knowledge_point(kp_id: int, req: KnowledgePointUpdate, request: Request):
    """更新知识点信息（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    kp = execute_query_one("SELECT * FROM knowledge_points WHERE id=?", (kp_id,))
    if not kp:
        raise HTTPException(status_code=404, detail="知识点不存在")

    updates = {}
    for field in ["name", "description", "learning_objectives", "difficulty", "estimated_minutes", "sort_order", "status"]:
        val = getattr(req, field, None)
        if val is not None:
            updates[field] = val

    if not updates:
        return {"message": "无更新内容"}

    now = _now()
    updates["updated_at"] = now
    set_clause = ", ".join(f"{k}=?" for k in updates)
    params = list(updates.values()) + [kp_id]
    execute_insert_update(
        f"UPDATE knowledge_points SET {set_clause} WHERE id=?", tuple(params),
    )
    logger.info(f"用户 {user['username']} 更新知识点 id={kp_id}")
    return {"message": "知识点更新成功"}


@router.delete("/knowledge-points/{kp_id}", summary="删除知识点（硬删除）")
async def delete_knowledge_point(kp_id: int, request: Request):
    """删除知识点（教师/管理员）—— 硬删除并清理关联数据"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    kp = execute_query("SELECT * FROM knowledge_points WHERE id=?", (kp_id,))
    if not kp:
        raise HTTPException(status_code=404, detail="知识点不存在")

    execute_insert_update("DELETE FROM learning_progress WHERE knowledge_point_id=?", (kp_id,))
    execute_insert_update("DELETE FROM curriculum_bindings WHERE knowledge_point_id=?", (kp_id,))
    # 清理 AI 练习独立成绩
    from backend.question_db import execute_insert as q_exec
    q_exec("DELETE FROM ai_practice_results WHERE kp_id=?", (kp_id,))
    execute_insert_update("DELETE FROM knowledge_points WHERE id=?", (kp_id,))

    logger.info(f"用户 {user['username']} 硬删除知识点 id={kp_id}")
    return {"message": "知识点已永久删除"}


# ═══════════════════════════════════════════════════════════
# 资源绑定
# ═══════════════════════════════════════════════════════════

@router.get("/knowledge-points/{kp_id}/resources", summary="获取知识点绑定的资源列表")
async def get_kp_resources(kp_id: int, request: Request):
    """获取指定知识点下绑定的所有资源"""
    get_current_user(request)
    kp = execute_query_one("SELECT id FROM knowledge_points WHERE id=?", (kp_id,))
    if not kp:
        raise HTTPException(status_code=404, detail="知识点不存在")

    bindings = execute_query(
        """SELECT cb.* FROM curriculum_bindings cb
           WHERE cb.knowledge_point_id=? ORDER BY cb.sort_order, cb.id""",
        (kp_id,),
    )
    resources = []
    for b in bindings:
        info = _get_resource_info(b["resource_type"], b["resource_id"])
        resources.append({
            "binding_id": b["id"],
            "knowledge_point_id": b["knowledge_point_id"],
            "resource_type": b["resource_type"],
            "resource_id": b["resource_id"],
            "resource_name": info["name"],
            "resource_url": info["url"],
            "sort_order": b["sort_order"],
            "created_at": b["created_at"],
        })
    return {"resources": resources, "total": len(resources)}


@router.post("/bindings", summary="绑定资源到知识点")
async def create_binding(req: BindingCreate, request: Request):
    """将已有资源绑定到知识点（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    # 校验知识点存在
    kp = execute_query_one("SELECT id FROM knowledge_points WHERE id=?", (req.knowledge_point_id,))
    if not kp:
        raise HTTPException(status_code=404, detail="知识点不存在")

    valid_types = {"html", "download", "exam", "discussion", "interaction_quiz", "task"}
    if req.resource_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"无效的资源类型，可选: {', '.join(sorted(valid_types))}")

    # 校验资源所有权：教师只能绑定自己的资源
    username = user["username"]
    ownership_ok = _check_resource_ownership(req.resource_type, req.resource_id, username)
    if not ownership_ok:
        logger.warning(f"用户 {username} 尝试绑定非自己的资源 {req.resource_type}:{req.resource_id}")
        raise HTTPException(status_code=403, detail="只能绑定自己的资源")

    # 检查是否已绑定
    existing = execute_query_one(
        "SELECT id FROM curriculum_bindings WHERE knowledge_point_id=? AND resource_type=? AND resource_id=?",
        (req.knowledge_point_id, req.resource_type, req.resource_id),
    )
    if existing:
        raise HTTPException(status_code=409, detail="该资源已绑定到此知识点")

    now = _now()
    binding_id = execute_insert_update(
        """INSERT INTO curriculum_bindings (knowledge_point_id, resource_type, resource_id, sort_order, created_at)
           VALUES (?, ?, ?, ?, ?)""",
        (req.knowledge_point_id, req.resource_type, req.resource_id, req.sort_order, now),
    )
    # 重新绑定练习资源时清空学生作答记录
    if req.resource_type == "html":
        try:
            from backend.question_db import execute_insert as q_clear
            q_clear("DELETE FROM ai_practice_results WHERE kp_id=?", (req.knowledge_point_id,))
        except Exception:
            pass
    logger.info(f"用户 {user['username']} 绑定资源 {req.resource_type}:{req.resource_id} 到知识点 {req.knowledge_point_id}")
    return {"message": "资源绑定成功", "binding_id": binding_id}


@router.delete("/bindings/{binding_id}", summary="解绑资源")
async def delete_binding(binding_id: int, request: Request):
    """从知识点解绑资源（教师/管理员）"""
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    binding = execute_query_one("SELECT * FROM curriculum_bindings WHERE id=?", (binding_id,))
    if not binding:
        raise HTTPException(status_code=404, detail="绑定记录不存在")

    # 解绑练习资源时清空学生作答记录
    if binding["resource_type"] == "html":
        try:
            from backend.question_db import execute_insert as q_clear
            q_clear("DELETE FROM ai_practice_results WHERE kp_id=?", (binding["knowledge_point_id"],))
        except Exception:
            pass

    execute_insert_update("DELETE FROM curriculum_bindings WHERE id=?", (binding_id,))
    logger.info(f"用户 {user['username']} 解绑资源 binding_id={binding_id}")
    return {"message": "资源已解绑"}


@router.get("/bindings/available", summary="获取可绑定的候选资源")
async def get_available_resources(
    request: Request,
    resource_type: str = Query(..., description="资源类型"),
    keyword: str = Query("", description="搜索关键词"),
    kp_id: int = Query(None, description="知识点 ID（排除已绑定的）"),
):
    """获取指定类型的可选资源列表（用于绑定界面选择，仅返回当前教师自己的资源）"""
    user = get_current_user(request)
    username = user["username"]

    valid_types = {"html", "download", "exam", "discussion", "interaction_quiz", "task"}
    if resource_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"无效的资源类型")

    results = []

    try:
        if resource_type == "html":
            sql = "SELECT id, file_name as name FROM shared_resources WHERE resource_type='html' AND owner_username=?"
            params = [username]
            if keyword:
                sql += " AND file_name LIKE ?"
                params.append(f"%{keyword}%")
            sql += " ORDER BY id DESC LIMIT 200"
            results = execute_query(sql, tuple(params))

        elif resource_type == "download":
            sql = "SELECT id, file_name as name FROM shared_resources WHERE resource_type='download' AND owner_username=?"
            params = [username]
            if keyword:
                sql += " AND file_name LIKE ?"
                params.append(f"%{keyword}%")
            sql += " ORDER BY id DESC LIMIT 200"
            results = execute_query(sql, tuple(params))

        elif resource_type == "exam":
            sql = "SELECT id, title as name FROM exams WHERE status='published' AND creator_username=?"
            params = [username]
            if keyword:
                sql += " AND title LIKE ?"
                params.append(f"%{keyword}%")
            sql += " ORDER BY id DESC LIMIT 200"
            results = q_execute_query(sql, tuple(params))

        elif resource_type == "discussion":
            sql = "SELECT id, title as name FROM discussions WHERE (status='pending' OR status='active') AND creator_username=?"
            params = [username]
            if keyword:
                sql += " AND title LIKE ?"
                params.append(f"%{keyword}%")
            sql += " ORDER BY id DESC LIMIT 200"
            results = execute_query(sql, tuple(params))

        elif resource_type == "interaction_quiz":
            sql = "SELECT id, title as name FROM interaction_quizzes WHERE status='active' AND creator_username=?"
            params = [username]
            if keyword:
                sql += " AND title LIKE ?"
                params.append(f"%{keyword}%")
            sql += " ORDER BY id DESC LIMIT 200"
            results = execute_query(sql, tuple(params))

        elif resource_type == "task":
            sql = "SELECT id, name FROM tasks WHERE status='active' AND creator_username=?"
            params = [username]
            if keyword:
                sql += " AND name LIKE ?"
                params.append(f"%{keyword}%")
            sql += " ORDER BY id DESC LIMIT 200"
            results = execute_query(sql, tuple(params))
    except Exception as e:
        logger.warning(f"查询候选资源失败 (type={resource_type}): {e}")

    # 排除已绑定到该知识点的资源
    if kp_id and results:
        bound = execute_query(
            "SELECT resource_id FROM curriculum_bindings WHERE knowledge_point_id=? AND resource_type=?",
            (kp_id, resource_type),
        )
        bound_ids = set(r["resource_id"] for r in bound)
        results = [r for r in results if r["id"] not in bound_ids]

    return {"resources": results, "total": len(results)}


# ═══════════════════════════════════════════════════════════
# 节点排序（拖动）
# ═══════════════════════════════════════════════════════════

class ReorderItem(BaseModel):
    """排序项"""
    type: str  # "chapter" | "knowledge_point"
    id: int
    sort_order: int
    parent_id: int | None = None   # 章节新父级ID（None=顶层）
    chapter_id: int | None = None  # 知识点新所属章节ID

class ReorderRequest(BaseModel):
    """拖动排序请求"""
    items: list[ReorderItem]

@router.put("/reorder", summary="拖动排序章节/知识点")
async def reorder_nodes(req: ReorderRequest, request: Request):
    """拖动树节点后批量更新 sort_order（教师/管理员）
    支持同级排序和跨层级拖动（改变 parent_id / chapter_id）
    """
    user = get_current_user(request)
    if not _can_manage(user):
        raise HTTPException(status_code=403, detail="权限不足")

    if not req.items:
        raise HTTPException(status_code=400, detail="排序列表为空")

    now = _now()
    updated = {"chapters": 0, "knowledge_points": 0}

    for item in req.items:
        if item.type == "chapter":
            if item.parent_id is not None:
                execute_insert_update(
                    "UPDATE chapters SET sort_order=?, parent_id=?, updated_at=? WHERE id=?",
                    (item.sort_order, item.parent_id if item.parent_id > 0 else None, now, item.id),
                )
            else:
                execute_insert_update(
                    "UPDATE chapters SET sort_order=?, updated_at=? WHERE id=?",
                    (item.sort_order, now, item.id),
                )
            updated["chapters"] += 1
        elif item.type == "knowledge_point":
            if item.chapter_id is not None:
                execute_insert_update(
                    "UPDATE knowledge_points SET sort_order=?, chapter_id=?, updated_at=? WHERE id=?",
                    (item.sort_order, item.chapter_id, now, item.id),
                )
            else:
                execute_insert_update(
                    "UPDATE knowledge_points SET sort_order=?, updated_at=? WHERE id=?",
                    (item.sort_order, now, item.id),
                )
            updated["knowledge_points"] += 1

    logger.info(f"用户 {user['username']} 拖动排序: {updated}")
    return {"message": "排序已更新", "updated": updated}


# ═══════════════════════════════════════════════════════════
# 学习进度
# ═══════════════════════════════════════════════════════════

@router.get("/progress", summary="学生查看自己的学习进度")
async def get_my_progress(request: Request):
    """获取当前学生的全部学习进度"""
    user = get_current_user(request)
    if user.get("role") != 2:
        raise HTTPException(status_code=403, detail="仅学生可查看学习进度")

    rows = execute_query(
        """SELECT lp.*, kp.name as knowledge_point_name, kp.chapter_id,
                  ch.name as chapter_name, ch.course_id,
                  c.name as course_name
           FROM learning_progress lp
           JOIN knowledge_points kp ON kp.id = lp.knowledge_point_id
           JOIN chapters ch ON ch.id = kp.chapter_id
           JOIN courses c ON c.id = ch.course_id
           WHERE lp.student_username=?
           ORDER BY lp.updated_at DESC""",
        (user["username"],),
    )
    return {"progress": rows, "total": len(rows)}


@router.put("/progress/{kp_id}", summary="更新知识点学习状态")
async def update_progress(kp_id: int, req: ProgressUpdate, request: Request):
    """学生更新单个知识点的学习状态"""
    user = get_current_user(request)
    if user.get("role") != 2:
        raise HTTPException(status_code=403, detail="仅学生可更新学习进度")

    if req.status not in ("not_started", "in_progress", "completed"):
        raise HTTPException(status_code=400, detail="无效的状态值")

    kp = execute_query_one("SELECT id FROM knowledge_points WHERE id=?", (kp_id,))
    if not kp:
        raise HTTPException(status_code=404, detail="知识点不存在")

    now = _now()
    existing = execute_query_one(
        "SELECT id FROM learning_progress WHERE student_username=? AND knowledge_point_id=?",
        (user["username"], kp_id),
    )

    if existing:
        completed_at = now if req.status == "completed" else None
        if req.score is not None:
            execute_insert_update(
                "UPDATE learning_progress SET status=?, score=?, completed_at=?, updated_at=? WHERE id=?",
                (req.status, req.score, completed_at, now, existing["id"]),
            )
        else:
            execute_insert_update(
                "UPDATE learning_progress SET status=?, completed_at=?, updated_at=? WHERE id=?",
                (req.status, completed_at, now, existing["id"]),
            )
    else:
        completed_at = now if req.status == "completed" else None
        score = req.score or 0
        execute_insert_update(
            """INSERT INTO learning_progress (student_username, knowledge_point_id, status, score, completed_at, updated_at)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (user["username"], kp_id, req.status, score, completed_at, now),
        )

    # ── 完成知识点积分奖励（仅学生，仅首次完成） ──
    if req.status == "completed" and user.get("role") == 2:
        try:
            # 检查是否之前已完成过（幂等判断：仅首次完成给奖励）
            old_row = execute_query_one(
                "SELECT status FROM learning_progress WHERE student_username=? AND knowledge_point_id=?",
                (user["username"], kp_id),
            )
            if not old_row or old_row["status"] != "completed":
                from backend.reward_engine import award_participation
                kp_name = execute_query_one("SELECT name FROM knowledge_points WHERE id=?", (kp_id,))
                title = kp_name["name"] if kp_name else f"知识点#{kp_id}"
                award_participation(user["username"], "learning", str(kp_id), title)
        except Exception:
            pass

    return {"message": "学习进度已更新"}


@router.get("/progress/stats", summary="课程维度进度统计")
async def get_progress_stats(request: Request, course_id: int = Query(None)):
    """获取当前学生在各课程上的进度统计"""
    user = get_current_user(request)
    if user.get("role") != 2:
        raise HTTPException(status_code=403, detail="仅学生可查看")

    if course_id:
        # 统计指定课程
        kp_rows = execute_query(
            """SELECT kp.id FROM knowledge_points kp
               JOIN chapters ch ON ch.id = kp.chapter_id
               WHERE ch.course_id=? AND kp.status='active'""",
            (course_id,),
        )
        total = len(kp_rows)
        kp_ids = [r["id"] for r in kp_rows]
        if kp_ids:
            placeholders = ",".join("?" for _ in kp_ids)
            completed = execute_query(
                f"SELECT COUNT(*) as cnt FROM learning_progress WHERE student_username=? AND knowledge_point_id IN ({placeholders}) AND status='completed'",
                (user["username"], *kp_ids),
            )
            done = completed[0]["cnt"] if completed else 0
        else:
            done = 0
        return {"course_id": course_id, "total": total, "completed": done, "rate": round(done / total * 100, 1) if total else 0}
    else:
        # 统计所有课程
        courses = execute_query("SELECT id, name FROM courses WHERE status='active'")
        stats = []
        for c in courses:
            kp_rows = execute_query(
                """SELECT kp.id FROM knowledge_points kp
                   JOIN chapters ch ON ch.id = kp.chapter_id
                   WHERE ch.course_id=? AND kp.status='active'""",
                (c["id"],),
            )
            total = len(kp_rows)
            if total == 0:
                continue
            kp_ids = [r["id"] for r in kp_rows]
            placeholders = ",".join("?" for _ in kp_ids)
            completed = execute_query(
                f"SELECT COUNT(*) as cnt FROM learning_progress WHERE student_username=? AND knowledge_point_id IN ({placeholders}) AND status='completed'",
                (user["username"], *kp_ids),
            )
            done = completed[0]["cnt"] if completed else 0
            stats.append({
                "course_id": c["id"],
                "course_name": c["name"],
                "total": total,
                "completed": done,
                "rate": round(done / total * 100, 1),
            })
        return {"stats": stats}


@router.get("/progress/overview", summary="班级进度总览（教师用）")
async def get_class_progress_overview(
    request: Request,
    course_id: int = Query(None, description="课程 ID"),
    grade: str = Query(None, description="年级"),
    class_name: str = Query(None, description="班级"),
):
    """教师查看指定班级的课程完成进度概览"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="权限不足")

    username = user["username"]

    # 构建学生查询条件
    conditions = ["role=2"]
    params = []

    # 教师只能查看自己班级的学生
    if role == 1:
        from backend.permission_service import get_students_in_scope, get_grade_by_name
        # 使用统一权限服务获取教师管辖学生
        scoped = get_students_in_scope(username)
        if scoped:
            scoped_usernames = [s["username"] for s in scoped]
            placeholders = ",".join("?" * len(scoped_usernames))
            conditions.append(f"username IN ({placeholders})")
            params.extend(scoped_usernames)
        else:
            # 没有任教信息，不返回任何学生
            return {"students": [], "total": 0}

    # 用户选择的额外筛选条件
    if grade:
        from backend.permission_service import get_grade_by_name
        ginfo = get_grade_by_name(grade)
        if ginfo:
            conditions.append("grade_id=?")
            params.append(ginfo["id"])
        else:
            conditions.append("grade=?")
            params.append(grade)
    if class_name:
        import re
        cn = re.sub(r'[^\d]', '', str(class_name))
        gid = None
        if grade:
            from backend.permission_service import get_grade_by_name
            gi = get_grade_by_name(grade)
            if gi:
                gid = gi["id"]
        if gid:
            conditions.append("class_id=(SELECT id FROM classes WHERE grade_id=? AND (name=? OR name=?))")
            params.extend([gid, f"{cn}班", cn])
        else:
            conditions.append("(class=? OR class=?)")
            params.extend([cn, f"{cn}班"])

    where = " AND ".join(conditions)
    students = execute_query(
        f"SELECT username, name, grade, class FROM users WHERE {where} ORDER BY grade, class, name",
        tuple(params),
    )

    if not students:
        return {"students": [], "total": 0}

    # 确定课程
    if course_id:
        courses_list = [{"id": course_id}]
    else:
        courses_list = execute_query("SELECT id, name FROM courses WHERE status='active'")

    # 构建知识点 ID 列表
    kp_map = {}  # course_id -> [(kp_id, kp_name)]
    for c in courses_list:
        kps = execute_query(
            """SELECT kp.id, kp.name FROM knowledge_points kp
               JOIN chapters ch ON ch.id = kp.chapter_id
               WHERE ch.course_id=? AND kp.status='active'
               ORDER BY kp.sort_order, kp.id""",
            (c["id"],),
        )
        if kps:
            kp_map[c["id"]] = kps

    # 构建每位学生的进度矩阵
    result = []
    for stu in students:
        stu_progress = {
            "username": stu["username"],
            "name": stu["name"],
            "grade": stu["grade"],
            "class": stu["class"],
            "courses": [],
        }
        for c in courses_list:
            kps = kp_map.get(c["id"], [])
            if not kps:
                continue
            kp_ids = [k["id"] for k in kps]
            placeholders = ",".join("?" for _ in kp_ids)
            progress_rows = execute_query(
                f"SELECT knowledge_point_id, status FROM learning_progress WHERE student_username=? AND knowledge_point_id IN ({placeholders})",
                (stu["username"], *kp_ids),
            )
            p_map = {r["knowledge_point_id"]: r["status"] for r in progress_rows}
            completed = sum(1 for kp_id in kp_ids if p_map.get(kp_id) == "completed")
            stu_progress["courses"].append({
                "course_id": c["id"],
                "total_kps": len(kps),
                "completed_kps": completed,
                "rate": round(completed / len(kps) * 100, 1) if kps else 0,
                "details": [
                    {"kp_id": k["id"], "kp_name": k["name"], "status": p_map.get(k["id"], "not_started")}
                    for k in kps
                ],
            })
        result.append(stu_progress)

    return {"students": result, "total": len(result)}


# ═══════════════════════════════════════════════════════════
# 辅助：获取单个行（复用 execute_query_one 不存在的情况）
# ═══════════════════════════════════════════════════════════

def execute_query_one(sql: str, params: tuple[Any, ...] = ()):
    """执行查询并返回单条结果"""
    rows = execute_query(sql, params)
    return rows[0] if rows else None


# ═══════════════════════════════════════════════════════════
# V3.1 新增：AI 备课助手
# ═══════════════════════════════════════════════════════════

@router.get("/ai-lesson-plan")
async def ai_lesson_plan(
    request: Request,
    knowledge_point_id: int = Query(..., description="知识点 ID"),
):
    """AI 备课助手：根据知识点生成完整教案"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)

    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可使用备课助手")

    # 获取知识点信息
    kp_rows = execute_query(
        """SELECT kp.id, kp.name, kp.chapter_id, c.name as chapter_name,
                  co.name as course_name, co.grade
           FROM knowledge_points kp
           JOIN chapters c ON c.id = kp.chapter_id
           JOIN courses co ON co.id = c.course_id
           WHERE kp.id = ?""",
        (knowledge_point_id,),
    )
    if not kp_rows:
        raise HTTPException(status_code=404, detail="知识点不存在")
    kp = kp_rows[0]

    from backend.prompts.teaching import LESSON_PLAN_PROMPT
    from backend.api.chat_router import get_api_keys
    from backend.api.ai_service import call_ai_async

    keys = get_api_keys(username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key，请在系统配置中设置")

    def _safe(s):
        return str(s).replace('{', '{{').replace('}', '}}')

    prompt = LESSON_PLAN_PROMPT.format(
        course_name=_safe(kp["course_name"]),
        chapter_name=_safe(kp["chapter_name"]),
        knowledge_point=_safe(kp["name"]),
        grade=_safe(kp.get("grade", "")),
    )

    from backend.ai_task_manager import task_manager

    async def _do_plan() -> dict[str, Any]:
        try:
            result = await call_ai_async(prompt, api_key)
            return {
                "knowledge_point": kp["name"],
                "chapter_name": kp["chapter_name"],
                "course_name": kp["course_name"],
                "lesson_plan": result,
            }
        except Exception as e:
            logger.error(f"AI 备课助手生成失败: {e}")
            return {"error": f"教案生成失败: {str(e)}"}

    task_id = await task_manager.create_task(description="AI 备课", coro_factory=_do_plan)
    return {"task_id": task_id, "message": "AI 备课已提交，请稍后查询结果"}


# ═══════════════════════════════════════════════════════════
# V3.1 新增：导出教案为 Word 文档
# ═══════════════════════════════════════════════════════════

@router.get("/ai-lesson-plan/{kp_id}/export")
async def export_lesson_plan_docx(kp_id: int, request: Request, token: str = Query("")):
    """导出 AI 教案为 Word 文档"""
    import io
    import traceback
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    # 支持 token 参数认证（用于 window.open 下载）
    if token:
        request.state.user = None
        from backend.auth import decode_jwt_token
        payload = decode_jwt_token(token)
        if payload:
            request.state.user = payload

    try:
        user = get_current_user(request)
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"认证失败: {str(e)}")

    username = user["username"]
    role = user.get("role", 2)

    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可导出教案")

    try:
        kp_rows = execute_query(
            """SELECT kp.id, kp.name, kp.chapter_id, c.name as chapter_name,
                      co.name as course_name, co.grade
               FROM knowledge_points kp
               JOIN chapters c ON c.id = kp.chapter_id
               JOIN courses co ON co.id = c.course_id
               WHERE kp.id = ?""",
            (kp_id,),
        )
        if not kp_rows:
            raise HTTPException(status_code=404, detail="知识点不存在")
        kp = kp_rows[0]
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"查询知识点失败: {e}")
        raise HTTPException(status_code=500, detail=f"查询知识点失败: {str(e)}")

    from backend.prompts.teaching import LESSON_PLAN_PROMPT
    from backend.api.chat_router import get_api_keys
    from backend.api.ai_service import call_ai_async

    keys = get_api_keys(username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    def _safe(s):
        return str(s).replace('{', '{{').replace('}', '}}')

    prompt = LESSON_PLAN_PROMPT.format(
        course_name=_safe(kp["course_name"]),
        chapter_name=_safe(kp["chapter_name"]),
        knowledge_point=_safe(kp["name"]),
        grade=_safe(kp.get("grade", "")),
    )

    try:
        lesson_plan_text = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 备课助手生成失败: {e}")
        raise HTTPException(status_code=500, detail=f"教案生成失败: {str(e)}")

    doc = Document()
    style = doc.styles['Normal']  # type: ignore[union-attr]
    style.font.name = 'Microsoft YaHei'  # type: ignore[union-attr]
    style.font.size = Pt(11)  # type: ignore[union-attr]
    style.paragraph_format.line_spacing = 1.5  # type: ignore[union-attr]

    title = doc.add_heading(kp["name"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"课程：{kp['course_name']}  章节：{kp['chapter_name']}  年级：{kp.get('grade', '')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for line in lesson_plan_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- **') and '：' in line:
            content = line.lstrip('- ')
            p = doc.add_paragraph()
            bold_end = content.find('**', 2)
            if bold_end > 0:
                run = p.add_run(content[2:bold_end])
                run.bold = True
                p.add_run(content[bold_end + 2:])
            else:
                p.add_run(content)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f'{i}. ') for i in range(1, 10)):
            doc.add_paragraph(line, style='List Number')
        else:
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import urllib.parse
    safe_filename = urllib.parse.quote(f"{kp['name']}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )


# ═══════════════════════════════════════════════════════════
# AI 课件生成
# ═══════════════════════════════════════════════════════════

@router.post("/ai-courseware/{kp_id}")
async def ai_generate_courseware(kp_id: int, request: Request):
    """[教师] AI 根据知识点生成 HTML 课件"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可使用")

    keys = get_api_keys(username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    # 获取知识点信息
    kp_rows = execute_query(
        """SELECT kp.id, kp.name, kp.description, kp.learning_objectives,
                  c.name as chapter_name, co.name as course_name
           FROM knowledge_points kp
           JOIN chapters c ON c.id = kp.chapter_id
           JOIN courses co ON co.id = c.course_id
           WHERE kp.id = ?""",
        (kp_id,),
    )
    if not kp_rows:
        raise HTTPException(status_code=404, detail="知识点不存在")
    kp = kp_rows[0]

    from backend.prompts.teaching import COURSEWARE_GENERATE_PROMPT
    from backend.api.ai_service import call_ai_async
    from backend.utils import get_account_html_dir

    subject = kp["course_name"] or "信息科技"

    def _safe(s):
        return str(s or "").replace('{', '{{').replace('}', '}}')

    prompt = COURSEWARE_GENERATE_PROMPT.format(
        subject=_safe(subject),
        course_name=_safe(kp["course_name"]),
        chapter_name=_safe(kp["chapter_name"]),
        kp_name=_safe(kp["name"]),
        kp_description=_safe(kp.get("description", "")),
        learning_objectives=_safe(kp.get("learning_objectives", "")),
    )

    from backend.ai_task_manager import task_manager

    async def _generate() -> dict[str, Any]:
        try:
            logger.info(f"AI 课件开始生成: kp_id={kp_id}, kp_name={kp['name']}")
            result = await call_ai_async(prompt, api_key)
            logger.info(f"AI 响应已收到，长度={len(result)}")

            # 清理 AI 可能添加的 markdown 代码块标记
            html_content = result.strip()
            if html_content.startswith("```html"):
                html_content = html_content[7:]
            elif html_content.startswith("```"):
                html_content = html_content[3:]
            if html_content.endswith("```"):
                html_content = html_content[:-3]
            html_content = html_content.strip()

            if not html_content:
                logger.error("AI 返回的 HTML 内容为空")
                return {"error": "AI 返回的 HTML 内容为空"}

            # 保存到用户的 html 目录
            html_dir = get_account_html_dir(username)
            os.makedirs(html_dir, exist_ok=True)
            safe_name = kp["name"].replace(" ", "_").replace("/", "_").replace("\\", "_")
            filename = f"{kp_id}_{safe_name}_课件.html"
            filepath = os.path.join(html_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(html_content)

            # 生成访问 URL
            from backend.config import BASE_DIR
            rel_path = os.path.relpath(filepath, str(BASE_DIR)).replace("\\", "/")
            file_url = f"/api/files/{rel_path}"

            logger.info(f"AI 课件已保存: {filepath}")
            return {
                "kp_name": kp["name"],
                "file_url": file_url,
                "filename": filename,
                "filepath": filepath,
            }
        except Exception as e:
            logger.error(f"AI 课件生成失败: {e}", exc_info=True)
            return {"error": f"课件生成失败: {str(e)}"}

    task_id = await task_manager.create_task(description="AI 课件生成", coro_factory=_generate)
    return {"task_id": task_id, "message": "AI 课件生成已开始，请稍候..."}


@router.get("/ai-courseware/{kp_id}/preview")
async def preview_courseware(kp_id: int, request: Request):
    """预览已生成的 AI 课件 HTML 内容"""
    user = get_current_user(request)
    username = user["username"]

    from backend.utils import get_account_html_dir
    html_dir = get_account_html_dir(username)
    import glob
    pattern = os.path.join(html_dir, f"{kp_id}_*_课件.html")
    files = glob.glob(pattern)
    if not files:
        raise HTTPException(status_code=404, detail="尚未生成课件，请先使用 AI 生成")
    # 取最新的文件
    latest = max(files, key=os.path.getmtime)
    with open(latest, "r", encoding="utf-8") as f:
        content = f.read()

    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=content)


# ═══════════════════════════════════════════════════════════
# AI 练习生成
# ═══════════════════════════════════════════════════════════

def _parse_ai_questions(text: str) -> list[dict[str, Any]] | None:
    """从 AI 返回文本中解析 JSON 题目数组"""
    import re
    # 尝试直接解析
    try:
        data = json.loads(text)
        if isinstance(data, list):
            return data
    except json.JSONDecodeError:
        pass
    # 尝试从 ```json ``` 代码块提取
    match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if match:
        try:
            data = json.loads(match.group(1))
            if isinstance(data, list):
                return data
        except json.JSONDecodeError:
            pass
    # 尝试从 [ 到 ] 提取最外层数组
    start = text.find('[')
    end = text.rfind(']')
    if start != -1 and end != -1 and end > start:
        try:
            data = json.loads(text[start:end + 1])
            if isinstance(data, list):
                return data
        except json.JSONDecodeError:
            pass
    return None


@router.post("/ai-practice/{kp_id}")
async def ai_generate_practice(kp_id: int, request: Request):
    """[教师] AI 根据知识点生成10道单选题 + 创建练习任务 + 生成HTML答题页面"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可使用")

    keys = get_api_keys(username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    # 获取知识点信息
    kp_rows = execute_query(
        """SELECT kp.id, kp.name, kp.description, kp.learning_objectives, kp.difficulty,
                  c.name as chapter_name, co.name as course_name, co.subject
           FROM knowledge_points kp
           JOIN chapters c ON c.id = kp.chapter_id
           JOIN courses co ON co.id = c.course_id
           WHERE kp.id = ?""",
        (kp_id,),
    )
    if not kp_rows:
        raise HTTPException(status_code=404, detail="知识点不存在")
    kp = kp_rows[0]

    from backend.prompts.teaching import PRACTICE_SINGLE_CHOICE_PROMPT
    from backend.api.ai_service import call_ai_async
    from backend.utils import get_account_html_dir
    from backend.question_db import execute_insert as q_insert, execute_update as q_update
    from backend.config import BASE_DIR

    subject = kp.get("subject") or kp["course_name"] or "信息科技"
    difficulty_map = {"easy": "简单", "medium": "中等", "hard": "困难"}
    difficulty_desc = difficulty_map.get(kp.get("difficulty", "medium"), "中等")

    def _safe(s):
        return str(s or "").replace('{', '{{').replace('}', '}}')

    prompt = PRACTICE_SINGLE_CHOICE_PROMPT.format(
        subject=_safe(subject),
        course_name=_safe(kp["course_name"]),
        chapter_name=_safe(kp["chapter_name"]),
        knowledge_point=_safe(kp["name"]),
        kp_description=_safe(kp.get("description", "")),
        difficulty_desc=_safe(difficulty_desc),
    )

    from backend.ai_task_manager import task_manager

    async def _generate() -> dict[str, Any]:
        try:
            logger.info(f"AI 练习开始生成: kp_id={kp_id}, kp_name={kp['name']}")
            result_text = await call_ai_async(prompt, api_key)
            logger.info(f"AI 响应已收到，长度={len(result_text)}")

            questions = _parse_ai_questions(result_text)
            if not questions:
                logger.error("AI 返回格式异常，未能解析出题目")
                return {"error": "AI 返回格式异常，未能解析出题目"}

            # 限制最多10道，且必须全部为 single 类型
            questions = questions[:10]
            for q in questions:
                q["type"] = "single"

            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            question_ids = []

            # 逐题入库
            for q in questions:
                opts = json.dumps(q.get("options", {}), ensure_ascii=False) if q.get("options") else ""
                svg_code = q.get("svg_code") or ""
                has_svg = 1 if svg_code.strip() else 0
                media_placeholders = json.dumps(q.get("media_placeholders") or [], ensure_ascii=False)
                qid = q_insert(
                    """INSERT INTO question_bank (type,question_text,options,correct_answer,explanation,
                        knowledge_points,subject,difficulty,creator_username,source,status,created_at,updated_at,
                        svg_content,has_svg,media_placeholders)
                       VALUES (?,?,?,?,?,?,?,?,?,'ai','active',?,?,?,?,?)""",
                    ("single", q.get("question", ""), opts,
                     q.get("answer", ""), q.get("explanation", ""),
                     kp["name"], subject,
                     q.get("difficulty", "medium"), username, now, now,
                     svg_code, has_svg, media_placeholders),
                )
                q["id"] = qid
                q["index"] = qid
                if "svg_code" in q and "svg_content" not in q:
                    q["svg_content"] = q["svg_code"]
                if "has_svg" not in q:
                    q["has_svg"] = 1 if q.get("svg_code") or q.get("svg_content") else 0

                # 自动生图（有占位符时）
                placeholders = q.get("media_placeholders") or []
                media_files = []
                if placeholders:
                    try:
                        from backend.api.image_gen_service import generate_and_save_image
                        from backend.prompts.chat import IMAGE_GEN_PROMPT_TEMPLATE
                        from pathlib import Path as PPath
                        media_dir = BASE_DIR / "question_media" / str(qid)
                        for ph in placeholders:
                            ph_prompt = IMAGE_GEN_PROMPT_TEMPLATE.format(
                                subject=subject,
                                purpose=ph.get("purpose", "示意图"),
                                description=ph["description"],
                            )
                            local_path = await generate_and_save_image(ph_prompt, media_dir)
                            if local_path:
                                ph["status"] = "generated"
                                media_files.append({
                                    "key": ph["key"],
                                    "type": "image",
                                    "url": f"/api/files/question_media/{qid}/{PPath(local_path).name}",
                                    "alt": ph["description"],
                                    "created_at": now,
                                })
                            else:
                                ph["status"] = "failed"
                        q_update(
                            "UPDATE question_bank SET media_placeholders=?, media_files=? WHERE id=?",
                            (json.dumps(placeholders, ensure_ascii=False),
                             json.dumps(media_files, ensure_ascii=False), qid)
                        )
                    except Exception as img_err:
                        logger.warning(f"自动生图失败: {img_err}")
                q["media_files"] = media_files
                question_ids.append(qid)

            # 创建练习任务（用于成绩记录和积分发放）
            title = f"{kp['name']} 练习"
            session_id = q_insert(
                """INSERT INTO practice_sessions
                   (title, knowledge_points, creator_username, subject, question_count,
                    total_score, target_grade, target_class, target_students, source, status, created_at, updated_at)
                   VALUES (?,?,?,?,?,?,?,?,?,'ai','active',?,?)""",
                (title, kp["name"], username, subject,
                 len(question_ids), len(question_ids) * 10,
                 "", "", "", now, now),
            )

            # 建立题目关联
            for i, qid in enumerate(question_ids):
                from backend.question_db import execute_insert as q_insert2
                q_insert2(
                    "INSERT INTO practice_session_questions (session_id, question_id, sort_order, score) VALUES (?,?,?,?)",
                    (session_id, qid, i, 10),
                )

            # 更新总分
            q_update("UPDATE practice_sessions SET total_score=? WHERE id=?", (len(question_ids) * 10, session_id))

            # 生成 HTML 答题页面
            if session_id is None:
                return {"error": "创建练习任务失败"}
            html_content = _generate_practice_html(kp, questions, session_id, subject, kp_id)
            html_dir = get_account_html_dir(username)
            os.makedirs(html_dir, exist_ok=True)
            safe_name = kp["name"].replace(" ", "_").replace("/", "_").replace("\\", "_")
            filename = f"{kp_id}_{safe_name}_练习.html"
            filepath = os.path.join(html_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(html_content)

            rel_path = os.path.relpath(filepath, str(BASE_DIR)).replace("\\", "/")
            file_url = f"/api/files/{rel_path}"

            logger.info(f"AI 练习已生成: session_id={session_id}, file={filepath}")
            return {
                "session_id": session_id,
                "file_url": file_url,
                "filename": filename,
                "questions": questions,
                "total": len(questions),
                "kp_name": kp["name"],
            }
        except Exception as e:
            logger.error(f"AI 练习生成失败: {e}", exc_info=True)
            return {"error": f"练习生成失败: {str(e)}"}

    task_id = await task_manager.create_task(description="AI 练习生成", coro_factory=_generate)
    return {"task_id": task_id, "message": "AI 练习生成已开始，请稍候..."}


def _generate_practice_html(kp: dict[str, Any], questions: list[dict[str, Any]], session_id: int, subject: str, kp_id: int = 0) -> str:
    """生成自包含的 HTML 答题页面"""
    import html as html_mod

    kp_name = html_mod.escape(kp.get("name", ""))
    chapter_name = html_mod.escape(kp.get("chapter_name", ""))
    course_name = html_mod.escape(kp.get("course_name", ""))
    subject_esc = html_mod.escape(subject)

    questions_json = json.dumps(questions, ensure_ascii=False)

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{kp_name} - AI练习</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
<script src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"></script>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    min-height: 100vh;
    padding: 20px;
}}
.container {{
    max-width: 900px;
    margin: 0 auto;
}}
.header {{
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    border-radius: 16px;
    padding: 32px 40px;
    color: white;
    margin-bottom: 24px;
    box-shadow: 0 8px 32px rgba(102, 126, 234, 0.3);
}}
.header h1 {{ font-size: 24px; margin-bottom: 8px; }}
.header .meta {{ font-size: 14px; opacity: 0.9; }}
.header .meta span {{ margin-right: 16px; }}
.progress-bar {{
    display: flex; align-items: center; gap: 12px; margin-top: 16px;
    background: rgba(255,255,255,0.2); border-radius: 8px; padding: 12px 16px;
}}
.progress-track {{
    flex: 1; height: 6px; background: rgba(255,255,255,0.3); border-radius: 3px; overflow: hidden;
}}
.progress-fill {{
    height: 100%; background: #fff; border-radius: 3px; transition: width 0.3s ease; width: 0%;
}}
.progress-text {{ font-size: 13px; white-space: nowrap; }}
.question-card {{
    background: #fff; border-radius: 12px; padding: 24px 28px; margin-bottom: 16px;
    box-shadow: 0 2px 12px rgba(0,0,0,0.06); transition: box-shadow 0.2s;
    border-left: 4px solid #667eea;
}}
.question-card:hover {{ box-shadow: 0 4px 20px rgba(0,0,0,0.1); }}
.question-card.correct {{ border-left-color: #52c41a; }}
.question-card.wrong {{ border-left-color: #ff4d4f; }}
.q-number {{
    display: inline-flex; align-items: center; justify-content: center;
    width: 28px; height: 28px; border-radius: 50%; background: #667eea; color: #fff;
    font-size: 14px; font-weight: 600; margin-right: 10px; flex-shrink: 0;
}}
.q-header {{ display: flex; align-items: flex-start; margin-bottom: 12px; }}
.q-text {{ font-size: 16px; line-height: 1.6; flex: 1; }}
.media-area {{ margin: 12px 0; text-align: center; }}
.media-area svg {{ max-width: 100%; height: auto; border-radius: 8px; background: #fafafa; padding: 8px; }}
.media-area img {{ max-width: 100%; max-height: 200px; border-radius: 8px; object-fit: contain; }}
.options {{ margin: 12px 0 4px; }}
.option-item {{
    display: flex; align-items: flex-start; padding: 10px 14px; margin-bottom: 6px;
    border: 2px solid #e8e8e8; border-radius: 10px; cursor: pointer;
    transition: all 0.2s; font-size: 15px; line-height: 1.5;
}}
.option-item:hover {{ border-color: #667eea; background: #f8f9ff; }}
.option-item.selected {{ border-color: #667eea; background: #eef0ff; }}
.option-item.correct-answer {{ border-color: #52c41a; background: #f6ffed; }}
.option-item.wrong-answer {{ border-color: #ff4d4f; background: #fff2f0; }}
.option-label {{
    display: inline-flex; align-items: center; justify-content: center;
    width: 26px; height: 26px; border-radius: 50%; background: #f0f0f0;
    font-size: 13px; font-weight: 600; margin-right: 10px; flex-shrink: 0;
}}
.option-item.selected .option-label {{ background: #667eea; color: #fff; }}
.option-item.correct-answer .option-label {{ background: #52c41a; color: #fff; }}
.option-item.wrong-answer .option-label {{ background: #ff4d4f; color: #fff; }}
.option-content {{ flex: 1; }}
.explanation-box {{
    margin-top: 12px; padding: 14px 16px; background: #f9f9f9; border-radius: 8px;
    border-left: 3px solid #faad14; display: none;
}}
.explanation-box.show {{ display: block; }}
.explanation-box .label {{ font-weight: 600; color: #faad14; margin-bottom: 4px; }}
.explanation-box .text {{ font-size: 14px; line-height: 1.6; color: #555; }}
.submit-area {{ text-align: center; margin: 32px 0; }}
.btn-submit {{
    padding: 14px 48px; font-size: 18px; border: none; border-radius: 12px;
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: #fff; cursor: pointer; font-weight: 600;
    box-shadow: 0 4px 16px rgba(102, 126, 234, 0.4);
    transition: all 0.2s;
}}
.btn-submit:hover {{ transform: translateY(-2px); box-shadow: 0 6px 24px rgba(102, 126, 234, 0.5); }}
.btn-submit:disabled {{ opacity: 0.5; cursor: not-allowed; transform: none; }}
.result-area {{
    display: none; background: #fff; border-radius: 16px; padding: 32px;
    margin-bottom: 24px; box-shadow: 0 4px 24px rgba(0,0,0,0.08); text-align: center;
}}
.result-area.show {{ display: block; }}
.result-score {{ font-size: 48px; font-weight: 700; margin: 16px 0; }}
.result-grade {{ font-size: 20px; margin-bottom: 8px; }}
.result-stats {{ display: flex; justify-content: center; gap: 32px; margin: 16px 0; }}
.result-stat {{ text-align: center; }}
.result-stat .num {{ font-size: 24px; font-weight: 600; }}
.result-stat .label {{ font-size: 13px; color: #888; }}
.grade-excellent {{ color: #52c41a; }}
.grade-good {{ color: #1677ff; }}
.grade-medium {{ color: #faad14; }}
.grade-poor {{ color: #ff4d4f; }}
.reattempt-banner {{
    display: none;
    background: #fff3cd;
    border: 1px solid #ffc107;
    border-radius: 10px;
    padding: 12px 18px;
    margin-bottom: 16px;
    color: #856404;
    font-size: 15px;
    line-height: 1.6;
}}
.reattempt-banner .label {{ font-weight: 700; margin-right: 8px; }}
.reattempt-banner .score {{ font-weight: 700; font-size: 18px; color: #e67e22; }}
.reattempt-banner .hint {{ font-size: 13px; opacity: 0.8; margin-top: 4px; }}
.footer {{ text-align: center; color: #999; font-size: 13px; padding: 20px 0; }}
.error-msg {{ color: #ff4d4f; text-align: center; padding: 10px; }}
.loading {{ text-align: center; padding: 60px 0; color: #666; }}
@media (max-width: 640px) {{
    .header {{ padding: 20px; }}
    .question-card {{ padding: 16px; }}
    .option-item {{ padding: 8px 12px; }}
    .btn-submit {{ width: 100%; }}
}}
</style>
</head>
<body>
<div class="container">
    <div class="header">
        <h1>📝 {kp_name}</h1>
        <div class="meta">
            <span>📚 {course_name}</span>
            <span>📖 {chapter_name}</span>
            <span>🏷️ {subject_esc}</span>
            <span>📋 共 10 题 · 每题 10 分</span>
        </div>
        <div class="progress-bar" id="progress-bar">
            <span style="font-size:14px;">⏳ 进度</span>
            <div class="progress-track"><div class="progress-fill" id="progressFill"></div></div>
            <span class="progress-text" id="progressText">0/10</span>
        </div>
    </div>

    <div class="reattempt-banner" id="reattemptBanner">
        <div><span class="label">📋 已答过</span>
        上次成绩：<span class="score" id="prevScore">0</span> 分
        <span id="prevAccuracy" style="margin-left:8px;font-size:14px;color:#856404;"></span></div>
        <div class="hint" id="prevSubmittedAt"></div>
    </div>

    <div id="questionsContainer"></div>

    <div id="existingResults" style="display:none;"></div>

    <div class="submit-area" id="submitArea">
        <button class="btn-submit" id="btnSubmit" onclick="submitPractice()">📤 提交答案</button>
    </div>

    <div class="result-area" id="resultArea">
        <div style="font-size: 48px;" id="resultEmoji">🎉</div>
        <div class="result-grade" id="resultGrade"></div>
        <div class="result-score" id="resultScore"></div>
        <div class="result-stats">
            <div class="result-stat"><div class="num" id="statCorrect">0</div><div class="label">正确</div></div>
            <div class="result-stat"><div class="num" id="statWrong">0</div><div class="label">错误</div></div>
            <div class="result-stat"><div class="num" id="statAccuracy">0%</div><div class="label">正确率</div></div>
        </div>
        <div id="resultDetails" style="margin-top: 16px; font-size: 14px; color: #666;"></div>
        <div id="resultNote" style="margin-top: 8px; font-size: 13px; color: #1677ff; display:none;"></div>
        <div id="resultTime" style="margin-top: 4px; font-size: 12px; color: #999; display:none;"></div>
        <button class="btn-submit" style="margin-top: 20px; padding: 10px 32px; font-size: 15px;" onclick="location.reload()">🔄 重新答题</button>
    </div>

    <div class="footer">AI 智能练习 · 系统自动批改</div>
</div>

<script>
const questions = {questions_json};
const sessionId = {session_id};
const kpId = {kp_id};
const userAnswers = {{}};

function renderQuestions() {{
    const container = document.getElementById('questionsContainer');
    container.innerHTML = '';
    questions.forEach((q, i) => {{
        const card = document.createElement('div');
        card.className = 'question-card';
        card.id = 'qcard_' + i;
        card.dataset.index = i;

        let mediaHtml = '';
        if (q.svg_code && q.svg_code.trim()) {{
            mediaHtml += '<div class="media-area">' + q.svg_code + '</div>';
        }}
        if (q.media_files && q.media_files.length > 0) {{
            mediaHtml += '<div class="media-area">';
            q.media_files.forEach(f => {{
                mediaHtml += '<img src="' + f.url + '" alt="' + (f.alt || '') + '" loading="lazy">';
            }});
            mediaHtml += '</div>';
        }}

        let optionsHtml = '<div class="options">';
        const labels = ['A', 'B', 'C', 'D'];
        if (q.options) {{
            labels.forEach(k => {{
                if (q.options[k] !== undefined) {{
                    optionsHtml += '<div class="option-item" id="opt_' + i + '_' + k + '" onclick="selectOption(' + i + ',\\'' + k + '\\')">';
                    optionsHtml += '<span class="option-label">' + k + '</span>';
                    optionsHtml += '<span class="option-content">' + q.options[k] + '</span>';
                    optionsHtml += '</div>';
                }}
            }});
        }}
        optionsHtml += '</div>';

        const explanationHtml = '<div class="explanation-box" id="expl_' + i + '">' +
            '<div class="label">💡 解析</div>' +
            '<div class="text">' + (q.explanation || '') + '</div>' +
        '</div>';

        card.innerHTML = '<div class="q-header">' +
            '<span class="q-number">' + (i + 1) + '</span>' +
            '<div class="q-text">' + q.question + '</div>' +
        '</div>' + mediaHtml + optionsHtml + explanationHtml;

        container.appendChild(card);
    }});
    updateProgress();
}}

function selectOption(qIdx, key) {{
    if (document.getElementById('resultArea').classList.contains('show')) return;
    userAnswers[qIdx] = key;
    const card = document.getElementById('qcard_' + qIdx);
    const opts = card.querySelectorAll('.option-item');
    opts.forEach(o => o.classList.remove('selected'));
    document.getElementById('opt_' + qIdx + '_' + key).classList.add('selected');
    updateProgress();
}}

function updateProgress() {{
    const answered = Object.keys(userAnswers).length;
    document.getElementById('progressFill').style.width = (answered / questions.length * 100) + '%';
    document.getElementById('progressText').textContent = answered + '/' + questions.length;
}}

// 检查是否已作答过（独立 API，不依赖 practice_sessions）
function checkPreviousAttempt() {{
    var token = localStorage.getItem('smartkb_token');
    if (!token || !kpId) return;
    fetch('/api/curriculum/ai-practice/' + kpId + '/my-result', {{
        headers: {{ 'Authorization': 'Bearer ' + token }}
    }})
    .then(function(r) {{
        if (r.status === 404) return null;
        if (!r.ok) throw new Error('HTTP ' + r.status);
        return r.json();
    }})
    .then(function(data) {{
        if (data && data.result) {{
            var r = data.result;
            // 已作答过 → 完全阻止再次答题
            // 1. 隐藏进度条和提交区域
            var pb = document.getElementById('progress-bar');
            if (pb) pb.style.display = 'none';
            document.getElementById('submitArea').style.display = 'none';
            // 2. 显示已作答横幅（含评价）
            var banner = document.getElementById('reattemptBanner');
            if (banner) {{
                var gradeHtml = '<div><span class="label">📋 已作答</span>' +
                    '成绩：<span class="score">' + r.score + '</span> / ' + r.total_score + ' 分';
                if (r.accuracy > 0) {{
                    gradeHtml += ' <span style="font-size:14px;color:#856404;">（正确率 ' + r.accuracy + '%）</span>';
                }}
                gradeHtml += '</div>';
                if (r.evaluation) {{
                    gradeHtml += '<div style="margin-top:6px;font-size:15px;">' + r.evaluation + '</div>';
                }}
                if (r.reward_points > 0) {{
                    gradeHtml += '<div style="margin-top:4px;font-size:13px;color:#e67e22;">🎁 获得 ' + r.reward_points + ' 积分奖励</div>';
                }}
                gradeHtml += '<div class="hint">提交时间：' + (r.submitted_at || '') + '</div>';
                banner.innerHTML = gradeHtml;
                banner.style.display = 'block';
            }}
            // 3. 如果有结果明细，直接展示
            if (data.allResults && data.allResults.length > 0) {{
                renderPreviousResults(r, data.allResults);
            }}
            // 4. 禁用所有选项点击
            document.querySelectorAll('.option-item').forEach(function(el) {{
                el.style.cursor = 'default';
                el.onclick = null;
            }});
        }}
    }})
    .catch(function(err) {{
        console.error('检查历史作答失败:', err);
    }});
}}

function renderPreviousResults(r, allResults) {{
    var results = allResults || [];
    var earned = r.score;
    var totalScore = r.total_score;
    var accuracy = totalScore > 0 ? Math.round(earned / totalScore * 100) : 0;
    showResults(accuracy, earned, totalScore, null, results);
}}

// 页面加载时直接渲染题目 + 检查历史
var _prevResults = null;
document.addEventListener('DOMContentLoaded', function() {{
    renderQuestions();
    renderMath();
    checkPreviousAttempt();
}});

function renderMath() {{
    if (typeof renderMathInElement === 'function') {{
        setTimeout(function() {{
            renderMathInElement(document.body, {{
                delimiters: [
                    {{left: '$$', right: '$$', display: true}},
                    {{left: '$', right: '$', display: false}}
                ]
            }});
        }}, 500);
    }}
}}

// 本地批改：直接对比答案，即时显示结果
function submitPractice() {{
    const total = questions.length;
    const answered = Object.keys(userAnswers).length;
    if (answered < total) {{
        if (!confirm('还有 ' + (total - answered) + ' 道题未作答，确定提交吗？')) return;
    }}

    // 1. 本地批改
    let earned = 0;
    const totalScore = total * 10;
    const results = {{}};
    questions.forEach((q, i) => {{
        const studentAns = userAnswers[i] || '';
        const correctAns = q.answer || '';
        const isCorrect = studentAns.toUpperCase() === correctAns.toUpperCase();
        const s = isCorrect ? 10 : 0;
        earned += s;
        results[q.id] = {{
            student_answer: studentAns,
            correct_answer: correctAns,
            score: s,
            max_score: 10,
            is_correct: isCorrect
        }};
    }});
    const accuracy = Math.round(earned / totalScore * 100);

    // 2. 即时显示结果
    showResults(accuracy, earned, totalScore, results);

    // 3. 提交到后端（独立 API，不依赖 practice_sessions）
    const token = localStorage.getItem('smartkb_token');
    if (!token) {{
        var errMsg = document.getElementById('resultNote');
        if (errMsg) {{ errMsg.textContent = '⚠️ 未登录，成绩无法保存'; errMsg.style.display = 'block'; }}
        return;
    }}
    const submitBtn = document.getElementById('btnSubmit');
    if (submitBtn) {{ submitBtn.disabled = true; submitBtn.textContent = '⏳ 提交中...'; }}
    fetch('/api/curriculum/ai-practice/' + kpId + '/save-result', {{
        method: 'POST',
        headers: {{
            'Authorization': 'Bearer ' + token,
            'Content-Type': 'application/json'
        }},
        body: JSON.stringify({{
            score: earned,
            total_score: totalScore,
            answers: results
        }})
    }})
    .then(function(r) {{
        if (submitBtn) {{ submitBtn.textContent = '📤 提交答案'; }}
        if (!r.ok) {{
            return r.json().then(function(e) {{ throw new Error(e.detail || '提交失败'); }});
        }}
        return r.json();
    }})
    .then(function(data) {{
        var noteEl = document.getElementById('resultNote');
        if (noteEl) {{
            var txt = '✅ 成绩已记录';
            if (data.evaluation) txt = data.evaluation;
            if (data.reward_points > 0) txt += ' 🎁 +' + data.reward_points + '积分';
            noteEl.textContent = txt;
            noteEl.style.display = 'block';
        }}
    }})
    .catch(function(err) {{
        var noteEl = document.getElementById('resultNote');
        if (noteEl) {{
            noteEl.textContent = '⚠️ ' + (err.message || '提交失败，请重试');
            noteEl.style.display = 'block';
        }}
    }});
}}

function showResults(accuracy, score, totalScore, results, prevResults) {{
    document.getElementById('submitArea').style.display = 'none';
    const area = document.getElementById('resultArea');
    area.classList.add('show');

    document.getElementById('resultScore').textContent = score + ' / ' + totalScore + ' 分';

    let grade, emoji, gradeClass;
    if (accuracy >= 90) {{ grade = '🏆 优秀！'; emoji = '🎉'; gradeClass = 'grade-excellent'; }}
    else if (accuracy >= 80) {{ grade = '🌟 良好！'; emoji = '😊'; gradeClass = 'grade-good'; }}
    else if (accuracy >= 60) {{ grade = '📖 及格'; emoji = '🤔'; gradeClass = 'grade-medium'; }}
    else {{ grade = '💪 继续努力'; emoji = '📚'; gradeClass = 'grade-poor'; }}

    document.getElementById('resultEmoji').textContent = emoji;
    const gradeEl = document.getElementById('resultGrade');
    gradeEl.textContent = grade;
    gradeEl.className = 'result-grade ' + gradeClass;

    // 使用 prevResults（来自API的结构化结果）或本地 results（来自 submitPractice）
    var resultData = prevResults || results;
    var usePrev = !!prevResults;

    let correct = 0, wrong = 0;
    questions.forEach((q, i) => {{
        var qid = q.id;
        // prevResults 是数组，按索引查找；results 是对象，按 q.id 查找
        var res = null;
        if (usePrev) {{
            res = prevResults[i] || null;
        }} else {{
            res = results && results[qid] ? results[qid] : null;
        }}
        var isCorrect = res && res.is_correct;
        // 从 prevResults 中取学生答案
        var studentAns = '';
        var correctAns = q.answer || '';
        if (usePrev && res) {{
            studentAns = res.student_answer || '';
            correctAns = res.correct_answer || correctAns;
        }} else {{
            studentAns = userAnswers[i] || '';
        }}
        if (isCorrect) correct++; else wrong++;

        const card = document.getElementById('qcard_' + i);
        card.classList.add(isCorrect ? 'correct' : 'wrong');

        if (studentAns) {{
            const el = document.getElementById('opt_' + i + '_' + studentAns);
            if (el) el.classList.add(isCorrect ? 'correct-answer' : 'wrong-answer');
        }}
        if (correctAns && correctAns !== studentAns) {{
            const el = document.getElementById('opt_' + i + '_' + correctAns);
            if (el) el.classList.add('correct-answer');
        }}

        const expl = document.getElementById('expl_' + i);
        if (expl) expl.classList.add('show');
    }});

    document.getElementById('statCorrect').textContent = correct;
    document.getElementById('statWrong').textContent = wrong;
    document.getElementById('statAccuracy').textContent = accuracy + '%';

    let details = '';
    if (accuracy >= 80) details = '掌握情况良好，继续保持！💪';
    else if (accuracy >= 60) details = '基础尚可，建议复习错题巩固。📖';
    else details = '需要加强练习，建议回顾知识点后重试。📚';
    document.getElementById('resultDetails').textContent = details;

    // 渲染公式
    if (typeof renderMathInElement === 'function') {{
        renderMathInElement(document.body, {{
            delimiters: [
                {{left: '$$', right: '$$', display: true}},
                {{left: '$', right: '$', display: false}}
            ]
        }});
    }}
}}

//（旧的DOMContentLoaded已迁移到上方统一处理）
</script>
</body>
</html>"""


@router.post("/ai-practice/{kp_id}/from-bank")
async def ai_practice_from_bank(kp_id: int, request: Request):
    """[教师] 从题库选取已有题目生成练习"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可使用")

    body = await request.json()
    question_ids = body.get("question_ids", [])

    if len(question_ids) < 1 or len(question_ids) > 10:
        raise HTTPException(status_code=400, detail="请选择1-10道题")

    # 获取知识点信息
    kp_rows = execute_query(
        """SELECT kp.id, kp.name, kp.description, kp.difficulty,
                  c.name as chapter_name, co.name as course_name, co.subject
           FROM knowledge_points kp
           JOIN chapters c ON c.id = kp.chapter_id
           JOIN courses co ON co.id = c.course_id
           WHERE kp.id = ?""",
        (kp_id,),
    )
    if not kp_rows:
        raise HTTPException(status_code=404, detail="知识点不存在")
    kp = kp_rows[0]

    # 从题库获取题目
    from backend.question_db import execute_query as q_exec, execute_insert as q_insert, execute_update as q_update

    placeholders = ",".join("?" for _ in question_ids)
    questions = q_exec(
        f"""SELECT * FROM question_bank
            WHERE id IN ({placeholders}) AND type='single' AND status='active'
            ORDER BY CASE id {' '.join(f'WHEN ? THEN {i}' for i, _ in enumerate(question_ids))} END""",
        tuple(question_ids) + tuple(question_ids),
    )
    if not questions:
        raise HTTPException(status_code=404, detail="未找到有效题目")

    # 解析 options/media_files JSON
    for q in questions:
        for field in ["options", "media_placeholders", "media_files"]:
            val = q.get(field)
            if val and isinstance(val, str):
                try:
                    q[field] = json.loads(val)
                except (json.JSONDecodeError, TypeError):
                    q[field] = {} if field == "options" else []
        if q.get("options") is None:
            q["options"] = {}

    # 创建练习任务
    from datetime import datetime
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    subject = kp.get("subject") or kp["course_name"] or "信息科技"
    title = f"{kp['name']} 练习"
    session_id = q_insert(
        """INSERT INTO practice_sessions
           (title, knowledge_points, creator_username, subject, question_count,
            total_score, target_grade, target_class, target_students, source, status, created_at, updated_at)
           VALUES (?,?,?,?,?,?,?,?,?,'bank','active',?,?)""",
        (title, kp["name"], username, subject,
         len(questions), len(questions) * 10,
         "", "", "", now, now),
    )

    # 建立题目关联
    for i, q in enumerate(questions):
        q_insert(
            "INSERT INTO practice_session_questions (session_id, question_id, sort_order, score) VALUES (?,?,?,?)",
            (session_id, q["id"], i, 10),
        )
    q_update("UPDATE practice_sessions SET total_score=? WHERE id=?", (len(questions) * 10, session_id))

    # 生成 HTML 答题页面
    from backend.utils import get_account_html_dir
    from backend.config import BASE_DIR

    # 统一字段名（AI生成和题库字段名可能不同）
    for q in questions:
        if "question_text" in q and "question" not in q:
            q["question"] = q["question_text"]
        if "correct_answer" in q and "answer" not in q:
            q["answer"] = q["correct_answer"]
        if "svg_content" in q and "svg_code" not in q:
            q["svg_code"] = q.get("svg_content") or ""

    if session_id is None:
        raise HTTPException(status_code=500, detail="创建练习任务失败")

    html_content = _generate_practice_html(kp, questions, session_id, subject, kp_id)
    html_dir = get_account_html_dir(username)
    os.makedirs(html_dir, exist_ok=True)
    safe_name = kp["name"].replace(" ", "_").replace("/", "_").replace("\\", "_")
    filename = f"{kp_id}_{safe_name}_练习.html"
    filepath = os.path.join(html_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)

    rel_path = os.path.relpath(filepath, str(BASE_DIR)).replace("\\", "/")
    file_url = f"/api/files/{rel_path}"

    logger.info(f"教师 {username} 从题库生成练习: session_id={session_id}, file={filepath}")
    return {
        "session_id": session_id,
        "file_url": file_url,
        "filename": filename,
        "questions": questions,
        "total": len(questions),
        "kp_name": kp["name"],
        "message": f"已从题库选取 {len(questions)} 道题生成练习",
    }


@router.get("/ai-practice/{kp_id}/preview")
async def preview_ai_practice(kp_id: int, request: Request):
    """预览已生成的AI练习HTML页面"""
    user = get_current_user(request)
    username = user["username"]

    from backend.utils import get_account_html_dir
    html_dir = get_account_html_dir(username)
    import glob
    pattern = os.path.join(html_dir, f"{kp_id}_*_练习.html")
    files = glob.glob(pattern)
    if not files:
        raise HTTPException(status_code=404, detail="尚未生成练习，请先使用 AI 生成")
    latest = max(files, key=os.path.getmtime)
    with open(latest, "r", encoding="utf-8") as f:
        content = f.read()

    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=content)


# ═══════════════════════════════════════════════════════════
# AI 练习独立成绩记录（不依赖 practice_sessions）
# ═══════════════════════════════════════════════════════════

class SavePracticeResultRequest(BaseModel):
    """保存练习结果请求"""
    score: int = 0
    total_score: int = 0
    answers: dict[str, Any] = {}


@router.post("/ai-practice/{kp_id}/save-result", summary="保存AI练习成绩（独立存储，仅一次）")
async def save_ai_practice_result(kp_id: int, req: SavePracticeResultRequest, request: Request):
    """保存AI练习的作答成绩（仅一次，含积分奖励）"""
    user = get_current_user(request)
    username = user["username"]

    from backend.question_db import execute_insert, execute_query_one
    import json
    from datetime import datetime

    # ── 防重复：已有记录则拒绝 ──
    existing = execute_query_one(
        "SELECT id, score FROM ai_practice_results WHERE kp_id=? AND student_username=?",
        (kp_id, username),
    )
    if existing:
        raise HTTPException(status_code=409, detail="你已作答过此练习，不能重复提交")

    # ── 生成评价 ──
    accuracy = round(req.score / max(req.total_score, 1) * 100, 1)
    if accuracy >= 90:
        evaluation = "🏆 优秀！掌握情况非常好！"
    elif accuracy >= 80:
        evaluation = "🌟 良好！继续保持！"
    elif accuracy >= 60:
        evaluation = "📖 及格，建议复习错题巩固"
    else:
        evaluation = "💪 需要加强练习，建议回顾知识点后重试"

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    answers_json = json.dumps(req.answers, ensure_ascii=False)

    # ── 发放积分奖励 ──
    from backend.reward_engine import award_participation, award_grade

    kp_name_row = execute_query("SELECT name FROM knowledge_points WHERE id=?", (kp_id,))
    kp_title = kp_name_row[0]["name"] if kp_name_row else f"知识点#{kp_id}"

    total_reward = 0
    try:
        total_reward += award_participation(
            username, "practice", str(kp_id), kp_title, user.get("username", ""),
        )
        total_reward += award_grade(
            username, "practice", str(kp_id),
            req.score, req.total_score, kp_title, user.get("username", ""),
        )
    except Exception as e:
        logger.warning(f"积分发放失败 (kp_id={kp_id}): {e}")

    # ── 保存成绩 ──
    execute_insert(
        """INSERT INTO ai_practice_results
           (kp_id, student_username, score, total_score, accuracy, evaluation,
            reward_points, answers, submitted_at)
           VALUES (?,?,?,?,?,?,?,?,?)""",
        (kp_id, username, req.score, req.total_score, accuracy, evaluation,
         total_reward, answers_json, now),
    )

    logger.info(f"AI 练习成绩已保存: kp_id={kp_id}, username={username}, score={req.score}/{req.total_score}, reward={total_reward}")
    return {
        "message": "成绩已记录",
        "score": req.score,
        "total_score": req.total_score,
        "evaluation": evaluation,
        "reward_points": total_reward,
    }


@router.get("/ai-practice/{kp_id}/my-result", summary="获取我的AI练习历史成绩")
async def get_my_ai_practice_result(kp_id: int, request: Request):
    """获取当前用户在指定知识点上的AI练习历史成绩（独立存储）"""
    user = get_current_user(request)
    username = user["username"]

    from backend.question_db import execute_query_one
    import json

    row = execute_query_one(
        "SELECT * FROM ai_practice_results WHERE kp_id=? AND student_username=?",
        (kp_id, username),
    )
    if not row:
        raise HTTPException(status_code=404, detail="暂无作答记录")

    answers_data = row.get("answers", "{}")
    if isinstance(answers_data, str):
        try:
            answers_data = json.loads(answers_data)
        except (json.JSONDecodeError, TypeError):
            answers_data = {}

    all_results = []
    for qid_str, ans in answers_data.items():
        if isinstance(ans, dict):
            all_results.append({
                "question_id": int(qid_str) if qid_str.isdigit() else qid_str,
                "student_answer": ans.get("student_answer", ""),
                "correct_answer": ans.get("correct_answer", ""),
                "score": ans.get("score", 0),
                "max_score": ans.get("max_score", 10),
                "is_correct": ans.get("is_correct", False),
            })

    return {
        "result": {
            "score": row["score"],
            "total_score": row["total_score"],
            "accuracy": row["accuracy"],
            "evaluation": row.get("evaluation", ""),
            "reward_points": row.get("reward_points", 0),
            "submitted_at": row["submitted_at"],
            "questions": all_results,
        },
        "allResults": all_results,
    }
