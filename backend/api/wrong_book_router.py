"""
错题本 API 路由
自动归集学生错题，AI 生成复习计划
"""
import json
from datetime import datetime
from typing import Any

from fastapi import APIRouter, HTTPException, Request, Query
from pydantic import BaseModel

from backend.api.dependencies import get_current_user
from backend.question_db import execute_query, execute_query_one
from backend.database import execute_query as user_query
from backend.logger import logger
from backend.api.chat_router import get_api_keys
from backend.api.ai_service import call_ai_async
from backend.prompts import apply_skills, build_ai_role
from backend.database import (
    execute_query as db_exec,
    execute_query_dict as db_dict,
    execute_insert_update as db_insert,
)
from backend.permission_service import (
    parse_legacy_teacher_grade_class,
    is_student_in_teacher_scope,
)
from backend.async_utils import spawn_bg

router = APIRouter()

def _check_teacher_can_view_student(teacher_username: str, student_username: str):
    """检查教师是否有权限查看该学生的错题（使用统一权限服务）"""
    from backend.auth import is_admin
    if is_admin(teacher_username):
        return
    if not is_student_in_teacher_scope(student_username, teacher_username):
        raise HTTPException(status_code=403, detail="无权查看其他班级学生的错题")


# ═══════════════════════════════════════════════
# 访问守卫与错题读模型（W1/W2/W4/W5/W14）
# ═══════════════════════════════════════════════

# 已确认"无错题可回填"的学生, 避免每次请求重复全量扫描考试记录
_BACKFILL_SCANNED: set[str] = set()


def _assert_can_access_student(user: dict, requested: str) -> str:
    """统一的学生数据访问守卫(W1/W2)。

    - 学生: 一律强制为本人(忽略请求参数)
    - 教师: 指定他人时必须在其任教范围内
    - 管理员(root): 不限
    """
    username = user.get("username", "")
    role = user.get("role", 2)
    requested = (requested or "").strip()
    if role == 2:
        return username
    if role == 0:
        return requested or username
    if not requested or requested == username:
        return requested or username
    _check_teacher_can_view_student(username, requested)
    return requested


def _db_write(sql: str, params: tuple = ()) -> int:
    """主库写操作并返回受影响行数。

    database.execute_query 不提交(execute_query 里的 DML 会被静默回滚),
    execute_insert_update 返回 lastrowid 而非 rowcount, 故 UPDATE 计数需自己取。
    """
    from backend.database import get_connection
    with get_connection() as conn:
        cur = conn.execute(sql, params)
        conn.commit()
        return cur.rowcount


def _to_int(v: Any) -> int | None:
    try:
        return int(str(v).strip())
    except (TypeError, ValueError):
        return None


def _question_meta_map(qids: list) -> dict[int, dict[str, Any]]:
    """批量取题目详情。

    题库在 questions.db, wrong_book 在 smartkb.db, 不能跨库 JOIN(主库里的
    question_bank 是历史遗留的空壳表, JOIN 会静默返回全 NULL 的题干与答案)。
    注意不过滤 status: 题被软删后错题记录仍应可见(W14)。
    """
    ids = [i for i in (_to_int(q) for q in qids) if i is not None]
    if not ids:
        return {}
    ph = ",".join("?" for _ in ids)
    rows = execute_query(
        f"""SELECT id, type, question_text, options, correct_answer, explanation,
                   knowledge_points, subject, svg_content, has_svg, media_files, media_placeholders, status
            FROM question_bank WHERE id IN ({ph})""",
        tuple(ids),
    )
    return {r["id"]: r for r in rows}


def _activity_meta_map(student: str, source: str, ids: list[int]) -> dict[int, dict[str, Any]]:
    """批量取来源活动(考试/练习/随堂测验)的标题与成绩信息"""
    ids = [i for i in ids if i]
    if not ids:
        return {}
    ph = ",".join("?" for _ in ids)
    try:
        if source == "exam":
            rows = execute_query(
                f"""SELECT ea.exam_id AS rid, e.title, e.subject, ea.score, ea.total_score,
                          ea.submitted_at, ea.id AS attempt_id
                    FROM exam_attempts ea JOIN exams e ON e.id = ea.exam_id
                    WHERE ea.student_username = ? AND ea.status = 'submitted' AND ea.exam_id IN ({ph})
                    ORDER BY ea.submitted_at DESC""",
                (student, *ids),
            )
        elif source == "practice":
            rows = execute_query(
                f"""SELECT pa.session_id AS rid, ps.title, ps.subject, pa.score, pa.total_score,
                          pa.submitted_at, pa.id AS attempt_id
                    FROM practice_attempts pa JOIN practice_sessions ps ON ps.id = pa.session_id
                    WHERE pa.student_username = ? AND pa.status = 'submitted' AND pa.session_id IN ({ph})
                    ORDER BY pa.submitted_at DESC""",
                (student, *ids),
            )
        else:
            return {}
    except Exception as e:
        logger.warning(f"读取错题来源活动失败 ({source}): {e}")
        return {}
    out: dict[int, dict[str, Any]] = {}
    for r in rows:              # 已按时间倒序, 只保留每题最近一次
        out.setdefault(r["rid"], r)
    return out


def _backfill_wrong_book(student: str) -> int:
    """把只存在于 exam_attempts.answers 里的历史错题回填进 wrong_book(W5)。

    每人只跑一次: 表内已有记录或已确认无错题则跳过。
    """
    if not student or student in _BACKFILL_SCANNED:
        return 0
    try:
        has = db_exec("SELECT 1 FROM wrong_book WHERE student_username=? LIMIT 1", (student,))
        if has:
            _BACKFILL_SCANNED.add(student)
            return 0
        attempts = execute_query(
            """SELECT exam_id, answers, submitted_at FROM exam_attempts
               WHERE student_username = ? AND status = 'submitted'
               ORDER BY submitted_at DESC""",
            (student,),
        )
    except Exception as e:
        logger.warning(f"错题回填失败 ({student}): {e}")
        return 0

    wrong: dict[int, tuple[str, str, int]] = {}
    for a in attempts:
        data = a.get("answers")
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except (json.JSONDecodeError, TypeError):
                continue
        if not isinstance(data, dict):
            continue
        for qid_raw, ans in data.items():
            if not isinstance(ans, dict) or ans.get("is_correct", False):
                continue
            qid = _to_int(qid_raw)
            if qid is None:
                continue
            wrong.setdefault(qid, (
                str(ans.get("student_answer", ""))[:500],
                a.get("submitted_at") or "",
                a.get("exam_id") or 0,
            ))
    if not wrong:
        _BACKFILL_SCANNED.add(student)
        return 0

    meta = _question_meta_map(list(wrong.keys()))
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    inserted = 0
    for qid, (wrong_answer, when, exam_id) in wrong.items():
        m = meta.get(qid, {})
        db_insert(
            """INSERT OR IGNORE INTO wrong_book
               (student_username, question_id, source, source_id, knowledge_points, question_type,
                status, wrong_answer, created_at, wrong_count, last_wrong_at)
               VALUES (?, ?, 'exam', ?, ?, ?, 'pending', ?, ?, 1, ?)""",
            (student, qid, exam_id, m.get("knowledge_points", ""), m.get("type", ""),
             wrong_answer, when or now, when or now),
        )
        inserted += 1
    _BACKFILL_SCANNED.add(student)
    if inserted:
        logger.info(f"错题本回填: {student} 共 {inserted} 道历史错题")
    return inserted


def _collect_wrong_for_plan(student: str) -> tuple[list[dict], set, set, str]:
    """给复习计划用的错题汇总(W10: 读表 + 一次批量查题, 不再逐考试解析 JSON)"""
    _backfill_wrong_book(student)
    rows = db_dict(
        """SELECT question_id, wrong_answer, knowledge_points, question_type
           FROM wrong_book WHERE student_username = ? AND status = 'pending'
           ORDER BY COALESCE(NULLIF(last_wrong_at, ''), created_at) DESC, id DESC""",
        (student,),
    ) or []
    if not rows:
        return [], set(), set(), ""
    meta = _question_meta_map([r["question_id"] for r in rows])
    all_wrong: list[dict] = []
    kp_set: set = set()
    type_set: set = set()
    for r in rows:
        q = meta.get(r["question_id"], {})
        kp = (q.get("knowledge_points") or r.get("knowledge_points") or "未知")
        kp_set.add(kp)
        qtype = q.get("type") or r.get("question_type") or ""
        if qtype:
            type_set.add(qtype)
        all_wrong.append({
            "question": q.get("question_text") or f"(题目ID:{r['question_id']}，已不在题库)",
            "type": qtype,
            "correct": q.get("correct_answer", "") or "",
            "your_answer": r.get("wrong_answer", "") or "",
            "knowledge": kp,
        })
    subjects = {(m.get("subject") or "") for m in meta.values()} - {""}
    subject = "、".join(sorted(subjects)) if subjects else "通用"
    return all_wrong, kp_set, type_set, subject


_SOURCE_FALLBACK_TITLES = {
    "exam": "考试错题",
    "practice": "智能练习错题",
    "quiz": "随堂测验错题",
    "auto": "系统自动归集",
}


@router.get("/list", summary="获取错题本(以 wrong_book 表为唯一事实源)")
async def get_wrong_questions(
    request: Request,
    status: str = Query("pending", description="pending=待掌握 mastered=已掌握 all=全部"),
):
    """错题列表: 读 wrong_book 表 + 批量补题目详情, 按来源活动分组返回。

    旧实现是每次现场解析该生全部 exam_attempts JSON, 并叠加两处"知识点连坐"过滤
    (_is_kp_mastered / _is_question_mastered), 既与掌握状态表互不相干,
    又会因逐题查询在错题多时明显变慢。
    """
    user = get_current_user(request)
    target_username = _assert_can_access_student(
        user, request.query_params.get("student_username", "")
    )
    if not target_username:
        return {"total_wrong": 0, "exams": [], "mastered_total": 0, "student_username": ""}

    _backfill_wrong_book(target_username)

    rows = db_dict(
        """SELECT id, question_id, source, source_id, knowledge_points, question_type,
                  status, wrong_answer, created_at, mastered_at, wrong_count, last_wrong_at
           FROM wrong_book
           WHERE student_username = ? AND (? = 'all' OR status = ?)""",
        (target_username, status, status),
    ) or []
    mastered_rows = db_exec(
        "SELECT COUNT(*) FROM wrong_book WHERE student_username=? AND status='mastered'",
        (target_username,),
    )
    mastered_total = mastered_rows[0][0] if mastered_rows else 0

    qmeta = _question_meta_map([r["question_id"] for r in rows])
    grouped: dict[tuple[str, int], list[dict]] = {}
    for r in rows:
        grouped.setdefault((r["source"] or "exam", r["source_id"] or 0), []).append(r)
    ameta = {}
    for src in {k[0] for k in grouped}:
        ameta[src] = _activity_meta_map(target_username, src, [k[1] for k in grouped if k[0] == src])

    groups: list[dict] = []
    total_wrong = 0
    for (src, sid), items in grouped.items():
        info = ameta.get(src, {}).get(sid, {})
        wrong_questions = []
        for it in sorted(items, key=lambda x: str(x.get("last_wrong_at") or x.get("created_at") or ""), reverse=True):
            q = qmeta.get(it["question_id"], {})
            options_raw = q.get("options") or ""
            options = {}
            if options_raw:
                try:
                    opts = json.loads(options_raw) if isinstance(options_raw, str) else options_raw
                    if isinstance(opts, dict):
                        options = opts
                except (json.JSONDecodeError, TypeError):
                    options = {}
            wrong_questions.append({
                "question_id": str(it["question_id"]),
                "wrong_book_id": it["id"],
                "question_text": q.get("question_text") or f"(题目ID:{it['question_id']}，已不在题库)",
                "question_type": it.get("question_type") or q.get("type") or "",
                "options": options,
                "correct_answer": q.get("correct_answer", "") or "",
                "student_answer": it.get("wrong_answer", "") or "",
                "score": 0,
                "max_score": 0,
                "knowledge_points": it.get("knowledge_points") or q.get("knowledge_points") or "",
                "svg_content": q.get("svg_content", "") or "",
                "has_svg": q.get("has_svg", 0) or 0,
                "media_files": q.get("media_files", "") or "",
                "media_placeholders": q.get("media_placeholders", "") or "",
                "wrong_count": it.get("wrong_count") or 1,
                "wb_status": it.get("status"),
                "mastered_at": it.get("mastered_at"),
                "question_status": q.get("status") or "missing",
            })
        total_wrong += len(wrong_questions)
        groups.append({
            "exam_id": sid,
            "exam_title": info.get("title") or _SOURCE_FALLBACK_TITLES.get(src, "错题"),
            "exam_subject": info.get("subject") or "",
            "submitted_at": info.get("submitted_at") or (items[0].get("last_wrong_at") or items[0].get("created_at")),
            "score": info.get("score"),
            "total_score": info.get("total_score"),
            "source": src,
            "wrong_count": len(wrong_questions),
            "wrong_questions": wrong_questions,
        })

    groups.sort(key=lambda g: str(g.get("submitted_at") or ""), reverse=True)
    return {
        "total_wrong": total_wrong,
        "exams": groups,
        "mastered_total": mastered_total,
        "student_username": target_username,
    }


@router.get("/students", summary="获取有错题记录的学生列表")
async def get_students_with_wrong(
    request: Request,
    grade: str = Query("", description="年级筛选"),
    class_name: str = Query("", description="班级筛选"),
):
    """获取有错题记录的学生列表（教师只返回自己所教班级的学生）"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可用")

    # 先查有错题记录的学生用户名
    rows = execute_query(
        """SELECT DISTINCT ea.student_username FROM exam_attempts ea
           WHERE ea.status = 'submitted'"""
    )
    usernames = [r["student_username"] for r in rows]
    if not usernames:
        return {"students": []}

    # 构建查询条件
    conditions = ["username IN ({})".format(",".join("?" * len(usernames)))]
    params = list(usernames)

    if grade:
        conditions.append("grade = ?")
        params.append(grade)
    if class_name:
        conditions.append("class = ?")
        params.append(class_name)

    # 教师只能看自己所教班级的学生
    if role == 1:
        teacher_rows = user_query(
            "SELECT grade, class FROM users WHERE username=?", (username,)
        )
        teacher_grade = (teacher_rows[0][0] or "").strip() if teacher_rows else ""
        teacher_class = str(teacher_rows[0][1] or "").strip() if teacher_rows else ""
        grade_class_map = parse_legacy_teacher_grade_class(teacher_grade, teacher_class)

        if not grade:
            # 没有指定年级，限制在所有所教年级和班级
            grade_conditions = []
            for g, classes in grade_class_map.items():
                if classes:
                    cls_placeholders = ",".join("?" * len(classes))
                    grade_conditions.append("(grade = ? AND class IN ({}))".format(cls_placeholders))
                    params.extend([g] + classes)
                else:
                    grade_conditions.append("grade = ?")
                    params.append(g)
            if grade_conditions:
                conditions.append("(" + " OR ".join(grade_conditions) + ")")
        else:
            # 指定了年级，只限制该年级下的班级
            allowed_classes = grade_class_map.get(grade, [])
            if allowed_classes:
                cls_placeholders = ",".join("?" * len(allowed_classes))
                conditions.append("class IN ({})".format(cls_placeholders))
                params.extend(allowed_classes)

    where = " AND ".join(conditions)
    student_rows = user_query(
        f"""SELECT username, name, grade, class FROM users
           WHERE {where}
           ORDER BY grade, class""",
        tuple(params),
    )

    students = []
    for r in student_rows:
        students.append({
            "username": r[0],
            "name": r[1] or r[0],
            "grade": r[2] or "",
            "class": str(r[3] or ""),
        })
    return {"students": students}


@router.get("/grades", summary="获取有错题记录的年级列表")
async def get_grades_with_wrong(request: Request):
    """获取有错题记录的年级列表（教师只返回自己所教年级）"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可用")

    # 先查有错题记录的学生用户名
    rows = execute_query(
        """SELECT DISTINCT ea.student_username FROM exam_attempts ea
           WHERE ea.status = 'submitted'"""
    )
    usernames = [r["student_username"] for r in rows]
    if not usernames:
        return {"grades": []}

    # 管理员看到所有年级
    if role == 0:
        placeholders = ",".join("?" * len(usernames))
        grade_rows = user_query(
            f"""SELECT DISTINCT grade FROM users
               WHERE username IN ({placeholders}) AND grade IS NOT NULL AND grade != ''
               ORDER BY grade""",
            tuple(usernames),
        )
        return {"grades": [r[0] for r in grade_rows]}

    # 教师只返回自己所教年级
    teacher_rows = user_query(
        "SELECT grade, class FROM users WHERE username=?", (username,)
    )
    teacher_grade = (teacher_rows[0][0] or "").strip() if teacher_rows else ""
    teacher_class = str(teacher_rows[0][1] or "").strip() if teacher_rows else ""
    grade_class_map = parse_legacy_teacher_grade_class(teacher_grade, teacher_class)
    teacher_grades = list(grade_class_map.keys())

    if not teacher_grades:
        return {"grades": []}

    placeholders = ",".join("?" * len(usernames))
    grade_rows = user_query(
        f"""SELECT DISTINCT grade FROM users
           WHERE username IN ({placeholders}) AND grade IS NOT NULL AND grade != ''
           AND grade IN ({','.join('?' * len(teacher_grades))})
           ORDER BY grade""",
        tuple(usernames + teacher_grades),
    )
    return {"grades": [r[0] for r in grade_rows]}


@router.get("/classes", summary="获取指定年级下有错题记录的班级列表")
async def get_classes_with_wrong(
    request: Request,
    grade: str = Query("", description="年级"),
):
    """获取指定年级下有错题记录的班级列表（教师只返回自己所教班级）"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可用")

    if not grade:
        return {"classes": []}

    # 先查有错题记录的学生用户名
    rows = execute_query(
        """SELECT DISTINCT ea.student_username FROM exam_attempts ea
           WHERE ea.status = 'submitted'"""
    )
    usernames = [r["student_username"] for r in rows]
    if not usernames:
        return {"classes": []}

    # 管理员看到该年级所有班级
    placeholders = ",".join("?" * len(usernames))
    params = [grade] + usernames

    if role == 0:
        class_rows = user_query(
            f"""SELECT DISTINCT class FROM users
               WHERE grade = ? AND username IN ({placeholders})
               AND class IS NOT NULL AND class != ''
               ORDER BY class""",
            tuple(params),
        )
        return {"classes": [str(r[0]) for r in class_rows]}

    # 教师只返回自己所教班级
    teacher_rows = user_query(
        "SELECT grade, class FROM users WHERE username=?", (username,)
    )
    teacher_grade = (teacher_rows[0][0] or "").strip() if teacher_rows else ""
    teacher_class = str(teacher_rows[0][1] or "").strip() if teacher_rows else ""
    grade_class_map = parse_legacy_teacher_grade_class(teacher_grade, teacher_class)
    allowed_classes = grade_class_map.get(grade, [])

    if not allowed_classes:
        return {"classes": []}

    class_placeholders = ",".join("?" * len(allowed_classes))
    class_rows = user_query(
        f"""SELECT DISTINCT class FROM users
           WHERE grade = ? AND username IN ({placeholders})
           AND class IS NOT NULL AND class != ''
           AND class IN ({class_placeholders})
           ORDER BY class""",
        tuple(params + allowed_classes),
    )
    return {"classes": [str(r[0]) for r in class_rows]}


@router.get("/review-plan", summary="AI 生成错题复习计划")
async def get_review_plan(request: Request):
    """AI 根据错题本生成个性化复习计划(异步任务)"""
    user = get_current_user(request)
    username = user["username"]
    # W2: 教师指定他人时必须在其任教范围内
    target_username = _assert_can_access_student(
        user, request.query_params.get("student_username", "")
    )

    user_rows = user_query(
        "SELECT name, grade, class FROM users WHERE username=?", (target_username,)
    )
    student_name = user_rows[0][0] if user_rows and user_rows[0][0] else target_username
    student_grade = user_rows[0][1] if user_rows else ""
    student_class = str(user_rows[0][2]) if user_rows and user_rows[0][2] is not None else ""

    all_wrong, kp_set, type_set, subject = _collect_wrong_for_plan(target_username)
    if not all_wrong:
        return {"plan": "暂无错题，继续保持！", "total_wrong": 0, "knowledge_points": [], "weak_types": []}

    max_detail = 30
    wrong_text = ""
    for i, w in enumerate(all_wrong[:max_detail], 1):
        wrong_text += f"{i}. [{w['type']}] {w['question']}\n"
        wrong_text += f"   你的答案：{w['your_answer']} | 正确答案：{w['correct']}\n"
        wrong_text += f"   知识点：{w['knowledge']}\n\n"
    if len(all_wrong) > max_detail:
        wrong_text += f"...（共 {len(all_wrong)} 道错题，以上为最近 {max_detail} 道）\n"

    type_labels = {"single": "单选题", "multiple": "多选题", "true_false": "判断题", "short": "简答题",
                   "fill": "填空题", "essay": "作文", "subjective": "主观题"}
    weak_types = "、".join(str(type_labels.get(t, t)) for t in type_set)

    from backend.prompts.wrong_book import WRONG_BOOK_REVIEW_PROMPT

    ai_role = build_ai_role(grade=student_grade)
    prompt = f"{ai_role}\n" + WRONG_BOOK_REVIEW_PROMPT.format(
        student_name=student_name,
        grade=student_grade,
        cls=student_class,
        total_wrong=len(all_wrong),
        knowledge_points="、".join(sorted(kp_set)[:10]),
        weak_types=weak_types,
        wrong_questions=wrong_text,
        subject=subject or "通用",
    )
    prompt = apply_skills(prompt, "wrong-book")

    keys = get_api_keys(username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    from backend.ai_task_manager import task_manager

    async def _do_plan() -> dict[str, Any]:
        try:
            result = await call_ai_async(prompt, api_key)
            return {"plan": result}
        except Exception as e:
            logger.error(f"AI 复习计划生成失败: {e}")
            return {"error": f"生成复习计划失败: {str(e)}"}

    task_id = await task_manager.create_task(
        description=f"错题复习计划({target_username})", coro_factory=_do_plan
    )

    return {
        "task_id": task_id,
        "message": "AI 复习计划已提交，请稍后查询结果",
        "total_wrong": len(all_wrong),
        "knowledge_points": list(kp_set),
        "weak_types": list(type_set),
        "student_username": target_username,
    }


@router.get("/review-plan/export", summary="导出 AI 复习计划为 Word 文档")
async def export_review_plan_docx(
    request: Request,
    student_username: str = Query("", description="学生用户名"),
):
    """导出 AI 复习计划为 Word 文档"""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)

    # W2: 教师导出别班学生的复习计划必须先过任教范围校验
    target_username = _assert_can_access_student(user, student_username)

    user_rows = user_query(
        "SELECT name, grade, class FROM users WHERE username=?",
        (target_username,),
    )
    student_name = user_rows[0][0] if user_rows and user_rows[0][0] else target_username
    student_grade = user_rows[0][1] if user_rows else ""
    student_class = str(user_rows[0][2]) if user_rows and user_rows[0][2] is not None else ""

    all_wrong, kp_set, type_set, plan_subject = _collect_wrong_for_plan(target_username)

    if not all_wrong:
        raise HTTPException(status_code=400, detail="暂无错题，无需生成复习计划")

    wrong_text = ""
    for i, w in enumerate(all_wrong[:15], 1):
        wrong_text += f"{i}. [{w['type']}] {w['question']}\n"
        wrong_text += f"   你的答案：{w['your_answer']} | 正确答案：{w['correct']}\n"
        wrong_text += f"   知识点：{w['knowledge']}\n\n"

    type_labels = {"single": "单选题", "multiple": "多选题", "true_false": "判断题", "short": "简答题",
                   "fill": "填空题", "essay": "作文", "subjective": "主观题"}
    weak_types = "、".join(str(type_labels.get(t, t)) for t in type_set)

    from backend.prompts.wrong_book import WRONG_BOOK_REVIEW_PROMPT

    subject = plan_subject or "通用"

    ai_role = build_ai_role(grade=student_grade)
    prompt = f"{ai_role}\n" + WRONG_BOOK_REVIEW_PROMPT.format(
        student_name=student_name,
        grade=student_grade,
        cls=student_class,
        total_wrong=len(all_wrong),
        knowledge_points="、".join(sorted(kp_set)[:10]),
        weak_types=weak_types,
        wrong_questions=wrong_text,
        subject=subject,
    )

    keys = get_api_keys(username)
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    try:
        plan_text = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 复习计划生成失败: {e}")
        raise HTTPException(status_code=500, detail=f"生成复习计划失败: {str(e)}")

    # 生成 Word 文档
    doc = Document()
    style = doc.styles['Normal']  # type: ignore[union-attr]
    style.font.name = 'Microsoft YaHei'  # type: ignore[attr-defined]
    style.font.size = Pt(11)  # type: ignore[attr-defined]
    style.paragraph_format.line_spacing = 1.5  # type: ignore[attr-defined]

    title = doc.add_heading(f"{student_name} 的 AI 复习计划", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"学生：{student_name}  年级：{student_grade}  班级：{student_class}班  错题数：{len(all_wrong)} 道")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for line in plan_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- **') and '：' in line:
            content = line.lstrip('- ')
            p = doc.add_paragraph()
            bold_end = content.find('**', 2)
            if bold_end > 0:
                run = p.add_run(content[2:bold_end])
                run.bold = True
                p.add_run(content[bold_end + 2:])
            else:
                p.add_run(content)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f'{i}. ') for i in range(1, 10)):
            doc.add_paragraph(line, style='List Number')
        else:
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import urllib.parse
    safe_filename = urllib.parse.quote(f"{student_name}_复习计划.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )


# ═══════════════════════════════════════════════
# 错题记账（W6/W7/W8/W9）
# ═══════════════════════════════════════════════

def _upsert_wrong(student: str, qid: int, source: str, source_id: int,
                  kp: str, qtype: str, wrong_answer: str, when: str) -> int:
    """写入一条错题。

    唯一索引 ux_wb_student_question 保证"一人一题一行":
    - 新题: 插入 pending
    - 已 pending: 累加 wrong_count 并更新最近错误答案/时间
    - 曾 mastered 又答错: 重开为 pending(错题本不能只减不加)
    """
    existing = db_exec(
        "SELECT id, status FROM wrong_book WHERE student_username=? AND question_id=?",
        (student, qid),
    )
    if existing:
        return _db_write(
            """UPDATE wrong_book
               SET status='pending', mastered_at=NULL,
                   wrong_count=COALESCE(wrong_count, 1) + 1,
                   wrong_answer=?, last_wrong_at=?, source=?, source_id=?,
                   knowledge_points=CASE WHEN ? <> '' THEN ? ELSE knowledge_points END,
                   question_type=CASE WHEN ? <> '' THEN ? ELSE question_type END
               WHERE id=?""",
            (wrong_answer[:500], when, source, source_id, kp, kp, qtype, qtype, existing[0][0]),
        )
    return db_insert(
        """INSERT OR IGNORE INTO wrong_book
           (student_username, question_id, source, source_id, knowledge_points, question_type,
            status, wrong_answer, created_at, wrong_count, last_wrong_at)
           VALUES (?, ?, ?, ?, ?, ?, 'pending', ?, ?, 1, ?)""",
        (student, qid, source, source_id, kp, qtype, wrong_answer[:500], when, when),
    )


def _iter_graded(items: dict, want_correct: bool):
    """安全遍历判分结果: 跳过非数字题目 ID(W8), 返回 (qid, student_answer)"""
    out = []
    skipped = 0
    for qid_raw, ans in (items or {}).items():
        if not isinstance(ans, dict):
            skipped += 1
            continue
        if bool(ans.get("is_correct", False)) != want_correct:
            continue
        qid = _to_int(qid_raw)
        if qid is None:
            skipped += 1
            continue
        out.append((qid, str(ans.get("student_answer", "") or "")))
    return out, skipped


def record_wrong_answers(student_username: str, source_id: int, graded_answers: dict[str, Any],
                         source: str = "exam") -> int:
    """提交后把答错的题写入 wrong_book(考试/练习/随堂测验共用)"""
    pairs, skipped = _iter_graded(graded_answers, want_correct=False)
    if skipped:
        logger.warning(f"错题入库跳过 {skipped} 条无效判分记录 (student={student_username})")
    if not pairs:
        return 0
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    meta = _question_meta_map([qid for qid, _ in pairs])
    written = 0
    for qid, wrong_answer in pairs:
        m = meta.get(qid, {})
        _upsert_wrong(
            student_username, qid, source, source_id or 0,
            m.get("knowledge_points", "") or "", m.get("type", "") or "",
            wrong_answer, now,
        )
        written += 1
    logger.info(f"错题入库: {student_username} +{written} 条 (source={source}, id={source_id})")
    return written


def record_single_wrong(student_username: str, question_id: Any, wrong_answer: str,
                        source: str = "quiz", source_id: int = 0) -> int:
    """单条错题入库(随堂测验等零散场景用; 题已不在题库时只记 ID)"""
    qid = _to_int(question_id)
    if qid is None or not student_username:
        return 0
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    meta = _question_meta_map([qid]).get(qid, {})
    return _upsert_wrong(
        student_username, qid, source, source_id or 0,
        meta.get("knowledge_points", "") or "", meta.get("type", "") or "",
        str(wrong_answer or "")[:500], now,
    )


def mark_wrong_mastered(student_username: str, correct_answers: dict[str, Any]) -> int:
    """答对后把对应错题标为已掌握。

    W6: 旧实现还会按 knowledge_points LIKE 把同知识点的其它错题一并标掌握,
    实测"只答对 1 题"就能清空整片错题, 现改为只按题目 ID 精确标记。
    """
    pairs, _skipped = _iter_graded(correct_answers, want_correct=True)
    if not pairs:
        return 0
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    marked = 0
    for qid, _ans in pairs:
        marked += _db_write(
            """UPDATE wrong_book SET status='mastered', mastered_at=?
               WHERE student_username=? AND question_id=? AND status='pending'""",
            (now, student_username, qid),
        ) or 0
    if marked:
        logger.info(f"错题标记已掌握: {student_username} 共 {marked} 条")
    return marked


class PracticeGenerateRequest(BaseModel):
    count: int = 5
    subjects: list[str] = ["信息科技"]
    knowledge_points: str = ""
    student_username: str = ""  # 教师端指定学生

@router.post("/practice/generate", summary="错题练习—生成题目")
async def generate_wrong_practice(req: PracticeGenerateRequest, request: Request):
    """从错题本取出该生待掌握的原题作为练习(教师布置前预览用)。

    W1: 教师指定其他学生时必须在任教范围内;
    W3: 学生自己调用时不下发 correct_answer / explanation, 避免未答先知。
    """
    user = get_current_user(request)
    role = user.get("role", 2)
    target = _assert_can_access_student(user, req.student_username)
    if req.count < 1 or req.count > 20:
        req.count = 5

    _backfill_wrong_book(target)
    rows = db_dict(
        """SELECT question_id FROM wrong_book
           WHERE student_username = ? AND status = 'pending'
           ORDER BY COALESCE(NULLIF(last_wrong_at, ''), created_at) DESC, id DESC""",
        (target,),
    ) or []
    qids = [r["question_id"] for r in rows]
    if not qids:
        return {"questions": [], "total": 0, "message": "暂无待巩固的错题"}

    import random as _random
    picked = list(qids)
    _random.shuffle(picked)
    meta = _question_meta_map(picked)

    show_answer = role != 2
    questions = []
    for qid in picked:
        if len(questions) >= req.count:
            break
        q = meta.get(qid)
        if not q or q.get("status") != "active":
            continue        # 题已不在题库: 不出练习, 但错题本仍保留记录(W14)
        try:
            opts = json.loads(q.get("options") or "{}")
            if not isinstance(opts, dict):
                opts = {}
        except (json.JSONDecodeError, TypeError):
            opts = {}
        item = {
            "id": qid, "question_id": qid,
            "question": q.get("question_text", ""), "question_text": q.get("question_text", ""),
            "type": q.get("type", ""), "options": opts,
            "knowledge_points": q.get("knowledge_points", ""),
            "difficulty": q.get("difficulty", ""),
            "svg_content": q.get("svg_content", ""), "has_svg": q.get("has_svg", 0),
            "media_files": q.get("media_files", ""),
        }
        if show_answer:
            item["answer"] = q.get("correct_answer", "")
            item["correct_answer"] = q.get("correct_answer", "")
            item["explanation"] = q.get("explanation", "")
        else:
            item["answer"] = ""
            item["correct_answer"] = ""
            item["explanation"] = ""
        questions.append(item)

    return {
        "questions": questions,
        "total": len(questions),
        "pending_total": len(qids),
        "student_username": target,
    }


def check_and_auto_generate_wrong_practice(student_username: str, threshold: int = 30) -> int:
    """错题数超过阈值时自动生成"错题巩固练习"。

    W10: 只查 wrong_book(表已由回填保证有数据), 不再每次全量解析历史答卷 JSON;
    本函数是同步的, 调用方应通过 backend.async_utils.spawn_bg 放到后台线程执行。
    返回新建的 session_id, 0 表示未生成。
    """
    import random as _random
    from backend.question_db import execute_insert as q_ins, execute_update as q_upd

    if not student_username:
        return 0
    _backfill_wrong_book(student_username)
    rows = db_dict(
        "SELECT question_id FROM wrong_book WHERE student_username=? AND status='pending'",
        (student_username,),
    ) or []
    pending_ids = [r["question_id"] for r in rows if r.get("question_id") is not None]
    if len(pending_ids) <= threshold:
        return 0

    # 已有未作答的错题巩固练习就不再重复生成(一次 JOIN 判断)
    outstanding = execute_query(
        """SELECT ps.id FROM practice_sessions ps
           LEFT JOIN practice_attempts pa
                  ON pa.session_id = ps.id AND pa.student_username = ?
           WHERE ps.source = 'wrong_book' AND ps.status = 'active'
             AND ps.target_students LIKE ? AND pa.id IS NULL
           LIMIT 1""",
        (student_username, f'%"{student_username}"%'),
    )
    if outstanding:
        return 0

    meta = _question_meta_map(pending_ids)
    usable = [qid for qid in pending_ids if (meta.get(qid) or {}).get("status") == "active"]
    if not usable:
        return 0
    _random.shuffle(usable)
    selected = usable[:5]

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    title = f"错题巩固练习（{now[:10]}）"
    target_students_str = json.dumps([student_username], ensure_ascii=False)
    session_id = q_ins(
        """INSERT INTO practice_sessions
           (title, knowledge_points, creator_username, subject, question_count,
            total_score, target_grade, target_class, target_students, source, status, created_at, updated_at)
           VALUES (?,?,?,?,?,?,?,?,?,'wrong_book','active',?,?)""",
        (title, "错题巩固", "system", "", len(selected), 0, "", "", target_students_str, now, now),
    )
    total_score = 0
    for i, qid in enumerate(selected):
        q_ins(
            "INSERT INTO practice_session_questions (session_id, question_id, sort_order, score) VALUES (?,?,?,?)",
            (session_id, int(qid), i, 10),
        )
        total_score += 10
    q_upd("UPDATE practice_sessions SET total_score=? WHERE id=?", (total_score, session_id))
    logger.info(
        f"自动生成错题巩固练习 session={session_id} 学生={student_username} "
        f"抽题={len(selected)} 待掌握错题={len(pending_ids)}"
    )
    return session_id or 0


def _students_with_pending_wrong(limit: int = 200) -> list[str]:
    rows = db_dict(
        "SELECT DISTINCT student_username FROM wrong_book WHERE status='pending' LIMIT ?",
        (limit,),
    ) or []
    return [str(r["student_username"]) for r in rows if r.get("student_username")]


@router.get("/practice/check-auto", summary="学生端检查是否需要错题巩固练习")
async def check_auto_wrong_practice(request: Request):
    """学生登录后触发: 超过阈值则后台生成巩固练习(W10: 不再在请求里做重活)"""
    user = get_current_user(request)
    username = user["username"]
    if user.get("role", 2) != 2:
        return {"message": "仅学生触发", "checked": False}

    pending_rows = db_exec(
        "SELECT COUNT(*) FROM wrong_book WHERE student_username=? AND status='pending'",
        (username,),
    )
    pending_total = pending_rows[0][0] if pending_rows else 0
    spawned = spawn_bg(check_and_auto_generate_wrong_practice, username, name=f"错题巩固练习检查({username})")
    return {
        "message": "错题巩固练习检查已提交后台" if spawned else "已同步完成检查",
        "checked": True,
        "pending_total": pending_total,
    }


@router.post("/practice/auto-trigger", summary="手动触发自动生成错题巩固练习")
async def trigger_auto_wrong_practice(request: Request):
    """教师/管理员触发错题巩固练习生成(W10: 限定任教范围 + 全部放后台执行)"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可触发")

    try:
        body = await request.json()
    except Exception:
        body = {}
    target = (body.get("student_username") or request.query_params.get("student_username") or "").strip()

    if target:
        if role != 0:
            _check_teacher_can_view_student(username, target)
        spawn_bg(check_and_auto_generate_wrong_practice, target, name=f"错题巩固练习({target})")
        return {"message": f"已为 {target} 提交后台检查", "dispatched": [target]}

    candidates = _students_with_pending_wrong()
    if role != 0:
        allowed = [
            stu for stu in candidates
            if is_student_in_teacher_scope(stu, username)
        ]
    else:
        allowed = candidates
    for stu in allowed:
        spawn_bg(check_and_auto_generate_wrong_practice, stu, name=f"错题巩固练习({stu})")
    return {
        "message": f"已提交 {len(allowed)} 名学生的错题巩固练习检查(后台执行)",
        "dispatched": allowed[:50],
        "dispatched_total": len(allowed),
    }
