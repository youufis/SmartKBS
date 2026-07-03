"""
协作白板 API 路由
房间管理、页面管理、WebSocket 通信、AI 辅助
"""
import json
from datetime import datetime, timedelta
from typing import Any, Optional

from fastapi import APIRouter, HTTPException, Query, Request, WebSocket, WebSocketDisconnect
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from backend.api.dependencies import get_current_user
from backend.api.chat_router import get_api_keys
from backend.api.config_router import get_config_value
from backend.database import execute_query, execute_insert_update, get_connection
from backend.whiteboard_ws import whiteboard_manager
from backend.logger import logger

router = APIRouter()


def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _get_user_role(user: dict) -> int: # pyright: ignore[reportMissingTypeArgument]
    return user.get("role", 2)


def _is_teacher_or_admin(user: dict) -> bool: # type: ignore
    return _get_user_role(user) in (0, 1)


def _is_admin(user: dict) -> bool: # type: ignore
    return _get_user_role(user) == 0


def _get_ai_timeout() -> int:
    """获取 AI 请求超时配置（秒）"""
    return int(get_config_value("AI_REQUEST_TIMEOUT", 300))


# ── 请求模型 ──

class CreateRoomRequest(BaseModel):
    title: str
    room_type: str = "classroom"        # classroom / course / temporary
    mode: str = "demo"                  # demo / interactive / self_study
    course_kp_id: int | None = None
    grade: str = ""
    class_name: str = ""
    max_pages: int = 20


class UpdateRoomRequest(BaseModel):
    title: str | None = None
    mode: str | None = None
    allow_student_draw: bool | None = None


class JoinByCodeRequest(BaseModel):
    room_code: str


class SavePageRequest(BaseModel):
    snapshot_data: str                 # TLDraw 快照 JSON 字符串
    thumbnail: str = ""               # base64 缩略图
    title: str = ""


class GrantControlRequest(BaseModel):
    username: str


class SaveToResourceRequest(BaseModel):
    kp_id: int


# ═══════════════════════════════════════════════════════════
# 房间管理 API
# ═══════════════════════════════════════════════════════════

@router.post("/rooms", summary="创建白板房间")
async def create_room(req: CreateRoomRequest, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师和管理员可创建白板房间")

    # 如果教师未指定年级，自动从用户信息中补全
    grade = req.grade
    class_name = req.class_name
    if _get_user_role(user) == 1 and not grade:
        from backend.permission_service import get_teacher_grades
        grades = get_teacher_grades(user["username"])
        if grades:
            grade = grades[0]["name"]

    room_code = whiteboard_manager.generate_room_code()
    now = _now()
    room_id = execute_insert_update(
        """INSERT INTO whiteboard_rooms
           (room_code, title, room_type, mode, creator_username, course_kp_id,
            grade, class_name, max_pages, created_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (room_code, req.title, req.room_type, req.mode, user["username"],
         req.course_kp_id, grade, class_name, req.max_pages, now),
    )
    # 创建默认第1页
    execute_insert_update(
        """INSERT INTO whiteboard_pages (room_id, page_number, title, is_current)
           VALUES (?, 1, '第1页', 1)""",
        (room_id,),
    )
    return {"id": room_id, "room_code": room_code, "status": "ok"}


@router.get("/rooms", summary="获取白板房间列表")
async def list_rooms(
    request: Request,
    page: int = Query(1, ge=1),
    size: int = Query(20, ge=1, le=100),
):
    user = get_current_user(request)
    if _is_admin(user):
        # 管理员：查看所有房间
        rows = execute_query(
            """SELECT r.id, r.room_code, r.title, r.room_type, r.mode,
                      r.creator_username, COALESCE(NULLIF(u.name, ''), r.creator_username),
                      r.status, r.student_count, r.created_at
               FROM whiteboard_rooms r
               LEFT JOIN users u ON u.username = r.creator_username
               ORDER BY r.created_at DESC
               LIMIT ? OFFSET ?""",
            (size, (page - 1) * size),
        )
    elif _get_user_role(user) == 1:
        # 教师：只看自己的房间
        rows = execute_query(
            """SELECT r.id, r.room_code, r.title, r.room_type, r.mode,
                      r.creator_username, COALESCE(NULLIF(u.name, ''), r.creator_username),
                      r.status, r.student_count, r.created_at
               FROM whiteboard_rooms r
               LEFT JOIN users u ON u.username = r.creator_username
               WHERE r.creator_username=?
               ORDER BY r.created_at DESC
               LIMIT ? OFFSET ?""",
            (user["username"], size, (page - 1) * size),
        )
    else:
        # 学生：和考试发布一样的权限逻辑
        student_rows = execute_query(
            "SELECT grade, class FROM users WHERE username=?", (user["username"],),
        )
        student_grade = (student_rows[0][0] or "").strip() if student_rows else ""
        student_class = str(student_rows[0][1] or "").strip() if student_rows else ""

        # 获取所有活跃房间
        all_rows = execute_query(
            """SELECT r.id, r.room_code, r.title, r.room_type, r.mode,
                      r.creator_username, COALESCE(NULLIF(u.name, ''), r.creator_username),
                      r.status, r.student_count, r.created_at
               FROM whiteboard_rooms r
               LEFT JOIN users u ON u.username = r.creator_username
               WHERE r.status='active'
               ORDER BY r.created_at DESC""",
        )

        # 管理员列表
        admin_rows = execute_query("SELECT username FROM users WHERE role=0")
        admin_set = set(row[0] for row in admin_rows) if admin_rows else set()

        from backend.permission_service import parse_legacy_teacher_grade_class

        filtered = []
        for room in all_rows:
            creator = room[5]  # creator_username
            # 管理员创建的房间 → 全部可见
            if creator in admin_set:
                filtered.append(room)
                continue
            # 查询创建者的年级班级信息
            teacher_rows = execute_query(
                "SELECT grade, class FROM users WHERE username=?", (creator,),
            )
            if not teacher_rows:
                filtered.append(room)
                continue
            t_grade = (teacher_rows[0][0] or "").strip()
            t_class = str(teacher_rows[0][1] or "").strip()
            # 教师未设置年级班级 → 对所有学生可见（兼容旧数据）
            if not t_grade and not t_class:
                filtered.append(room)
                continue
            # 用和考试一致的解析函数
            grade_class_map = parse_legacy_teacher_grade_class(t_grade, t_class)
            matched = False
            if student_grade and student_grade in grade_class_map:
                allowed_classes = grade_class_map[student_grade]
                if not allowed_classes:
                    matched = True  # 有年级无班级限制 → 该年级全部可见
                elif student_class in allowed_classes:
                    matched = True
            if matched:
                filtered.append(room)

        # 分页
        total = len(filtered)
        offset = (page - 1) * size
        rows = filtered[offset:offset + size]
    return [
        {
            "id": r[0], "room_code": r[1], "title": r[2], "room_type": r[3],
            "mode": r[4], "creator_username": r[5], "creator_name": r[6],
            "status": r[7], "student_count": r[8], "created_at": r[9],
        }
        for r in rows
    ]


@router.get("/rooms/{room_id}", summary="获取房间详情")
async def get_room(room_id: int, request: Request):
    user = get_current_user(request)
    rows = execute_query(
        "SELECT * FROM whiteboard_rooms WHERE id=?", (room_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="房间不存在")
    r = rows[0]
    cols = ["id", "room_code", "title", "room_type", "mode", "creator_username",
            "course_kp_id", "grade", "class_name", "allow_student_draw", "max_pages",
            "auto_save_interval", "status", "student_count", "created_at", "ended_at"]
    return {cols[i]: r[i] for i in range(len(cols))}


@router.get("/rooms/{room_id}/snapshot", summary="获取当前快照（HTTP 兜底，IIS 环境用）")
async def get_snapshot(room_id: int, request: Request):
    """通过 HTTP 获取当前快照和授权状态，IIS 下 WebSocket 不可用时作为兜底"""
    user = get_current_user(request)
    username = user["username"]
    # 优先从内存取快照
    snap = whiteboard_manager.rooms.get(room_id, {}).get("last_snapshot", "")
    if not snap:
        rows = execute_query(
            "SELECT snapshot_data FROM whiteboard_pages WHERE room_id=? AND is_current=1 ORDER BY updated_at DESC LIMIT 1",
            (room_id,),
        )
        if rows and rows[0][0] and rows[0][0] != "{}":
            snap = rows[0][0]
    # 获取当前模式和学生授权状态
    mode = whiteboard_manager.get_mode(room_id)
    granted = whiteboard_manager.is_granted(room_id, username)
    return {"snapshot": snap, "mode": mode, "granted": granted}


@router.patch("/rooms/{room_id}", summary="更新房间配置")
async def update_room(room_id: int, req: UpdateRoomRequest, request: Request):
    user = get_current_user(request)
    rows = execute_query(
        "SELECT creator_username FROM whiteboard_rooms WHERE id=?", (room_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="房间不存在")
    if rows[0][0] != user["username"] and not _is_admin(user):
        raise HTTPException(status_code=403, detail="仅房主和管理员可修改")

    fields = []
    values = []
    if req.title is not None:
        fields.append("title=?")
        values.append(req.title)
    if req.mode is not None:
        fields.append("mode=?")
        values.append(req.mode)
        whiteboard_manager.set_mode(room_id, req.mode)
    if req.allow_student_draw is not None:
        fields.append("allow_student_draw=?")
        values.append(1 if req.allow_student_draw else 0)
    if fields:
        values.append(room_id)
        execute_insert_update(
            f"UPDATE whiteboard_rooms SET {', '.join(fields)} WHERE id=?",
            tuple(values),
        )
    return {"status": "ok"}


@router.post("/rooms/{room_id}/end", summary="结束房间")
async def end_room(room_id: int, request: Request):
    user = get_current_user(request)
    rows = execute_query(
        "SELECT creator_username FROM whiteboard_rooms WHERE id=? AND status='active'",
        (room_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="房间不存在或已结束")
    if rows[0][0] != user["username"] and not _is_admin(user):
        raise HTTPException(status_code=403, detail="仅房主和管理员可结束")

    now = _now()
    execute_insert_update(
        "UPDATE whiteboard_rooms SET status='ended', ended_at=? WHERE id=?",
        (now, room_id),
    )
    await whiteboard_manager.broadcast(room_id, {"type": "room_ended"})
    whiteboard_manager.rooms.pop(room_id, None)
    return {"status": "ok"}


@router.delete("/rooms/{room_id}", summary="删除房间")
async def delete_room(room_id: int, request: Request):
    user = get_current_user(request)
    rows = execute_query(
        "SELECT creator_username FROM whiteboard_rooms WHERE id=?", (room_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="房间不存在")
    if rows[0][0] != user["username"] and not _is_admin(user):
        raise HTTPException(status_code=403, detail="权限不足")

    execute_insert_update("DELETE FROM whiteboard_operations WHERE room_id=?", (room_id,))
    execute_insert_update("DELETE FROM whiteboard_room_members WHERE room_id=?", (room_id,))
    execute_insert_update("DELETE FROM whiteboard_pages WHERE room_id=?", (room_id,))
    execute_insert_update("DELETE FROM whiteboard_rooms WHERE id=?", (room_id,))
    return {"status": "ok"}


# ═══════════════════════════════════════════════════════════
# 加入/离开
# ═══════════════════════════════════════════════════════════

@router.post("/join-by-code", summary="通过房间码加入")
async def join_by_code(req: JoinByCodeRequest, request: Request):
    user = get_current_user(request)
    rows = execute_query(
        "SELECT id, title, mode, grade, class_name, creator_username FROM whiteboard_rooms WHERE room_code=? AND status='active'",
        (req.room_code.upper().strip(),),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="房间不存在或已结束")
    r = rows[0]

    # 学生权限：和考试发布一致的逻辑
    if _get_user_role(user) == 2:  # student
        # 检查创建者是否为管理员
        creator_rows = execute_query(
            "SELECT role FROM users WHERE username=?", (r[5],),
        )
        is_admin_room = creator_rows and creator_rows[0][0] == 0
        if not is_admin_room:
            student_rows = execute_query(
                "SELECT grade, class FROM users WHERE username=?",
                (user["username"],),
            )
            if student_rows:
                s_grade = (student_rows[0][0] or "").strip()
                s_class = str(student_rows[0][1] or "").strip()
                # 查询创建者的年级班级
                t_rows = execute_query(
                    "SELECT grade, class FROM users WHERE username=?", (r[5],),
                )
                if t_rows:
                    t_grade = (t_rows[0][0] or "").strip()
                    t_class = str(t_rows[0][1] or "").strip()
                    if t_grade or t_class:
                        from backend.permission_service import parse_legacy_teacher_grade_class
                        gcm = parse_legacy_teacher_grade_class(t_grade, t_class)
                        matched = False
                        if s_grade and s_grade in gcm:
                            allowed = gcm[s_grade]
                            if not allowed:
                                matched = True
                            elif s_class in allowed:
                                matched = True
                        if not matched:
                            raise HTTPException(status_code=403, detail="该白板房间不属于你的年级或班级")

    # 插入或更新成员记录
    execute_insert_update(
        """INSERT OR REPLACE INTO whiteboard_room_members
           (room_id, username, role, join_time)
           VALUES (?, ?, ?, ?)""",
        (r[0], user["username"], "teacher" if _is_teacher_or_admin(user) else "student", _now()),
    )
    return {
        "room_id": r[0],
        "title": r[1],
        "mode": r[2],
        "grade": r[3],
        "class_name": r[4],
    }


@router.post("/rooms/{room_id}/leave", summary="离开房间")
async def leave_room_api(room_id: int, request: Request):
    user = get_current_user(request)
    execute_insert_update(
        "UPDATE whiteboard_room_members SET leave_time=? WHERE room_id=? AND username=?",
        (_now(), room_id, user["username"]),
    )
    return {"status": "ok"}


# ═══════════════════════════════════════════════════════════
# 页面管理
# ═══════════════════════════════════════════════════════════

@router.get("/rooms/{room_id}/pages", summary="获取所有页面")
async def list_pages(room_id: int, request: Request):
    get_current_user(request)
    rows = execute_query(
        """SELECT page_number, title, snapshot_data, thumbnail, is_current, duration_seconds
           FROM whiteboard_pages WHERE room_id=?
           ORDER BY page_number""",
        (room_id,),
    )
    return [
        {
            "page_number": r[0], "title": r[1], "snapshot_data": r[2],
            "thumbnail": r[3], "is_current": bool(r[4]),
            "duration_seconds": r[5],
        }
        for r in rows
    ]


@router.get("/rooms/{room_id}/pages/{page_number}", summary="获取指定页面快照")
async def get_page(room_id: int, page_number: int, request: Request):
    get_current_user(request)
    rows = execute_query(
        "SELECT snapshot_data, title FROM whiteboard_pages WHERE room_id=? AND page_number=?",
        (room_id, page_number),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="页面不存在")
    return {"snapshot_data": rows[0][0], "title": rows[0][1]}


@router.put("/rooms/{room_id}/pages/{page_number}", summary="保存页面快照")
async def save_page(room_id: int, page_number: int, req: SavePageRequest, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可保存页面")
    now = _now()
    execute_insert_update(
        """INSERT OR REPLACE INTO whiteboard_pages
           (room_id, page_number, snapshot_data, thumbnail, title, is_current, updated_at, created_at)
           VALUES (?, ?, ?, ?, ?, 1, ?, COALESCE((SELECT created_at FROM whiteboard_pages WHERE room_id=? AND page_number=?), ?))""",
        (room_id, page_number, req.snapshot_data, req.thumbnail, req.title, now,
         room_id, page_number, now),
    )
    return {"status": "ok"}


@router.post("/rooms/{room_id}/pages", summary="新增页面")
async def add_page(room_id: int, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可新增页面")

    # 获取最大页码
    rows = execute_query(
        "SELECT COALESCE(MAX(page_number), 0) FROM whiteboard_pages WHERE room_id=?",
        (room_id,),
    )
    next_num = (rows[0][0] or 0) + 1
    execute_insert_update(
        "INSERT INTO whiteboard_pages (room_id, page_number, title) VALUES (?, ?, ?)",
        (room_id, next_num, f"第{next_num}页"),
    )
    return {"page_number": next_num, "status": "ok"}


@router.delete("/rooms/{room_id}/pages/{page_number}", summary="删除页面")
async def delete_page(room_id: int, page_number: int, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可删除页面")
    execute_insert_update(
        "DELETE FROM whiteboard_pages WHERE room_id=? AND page_number=?",
        (room_id, page_number),
    )
    return {"status": "ok"}


# ═══════════════════════════════════════════════════════════
# 控制权管理（互动模式）
# ═══════════════════════════════════════════════════════════

@router.post("/rooms/{room_id}/control/grant", summary="授权学生操作")
async def grant_control(room_id: int, req: GrantControlRequest, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可授权")
    await whiteboard_manager.grant_control(room_id, req.username, user["username"])
    return {"status": "ok"}


@router.post("/rooms/{room_id}/control/revoke", summary="收回操作权")
async def revoke_control(room_id: int, req: GrantControlRequest, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可收回")
    await whiteboard_manager.revoke_control(room_id, req.username)
    return {"status": "ok"}


# ═══════════════════════════════════════════════════════════
# 自习模式：学生列表
# ═══════════════════════════════════════════════════════════

@router.get("/rooms/{room_id}/students", summary="获取房间内学生列表（教师巡览）")
async def list_students(room_id: int, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可查看")

    # 从 WebSocket 内存中获取当前真实在线的学生
    room_data = whiteboard_manager.rooms.get(room_id)
    connected_students = set()
    if room_data:
        connected_students = {
            username for username, conn in room_data.get("connections", {}).items()
            if conn.get("role") == "student"
        }

    # 清理超过40分钟的旧记录（WS 断开时可能漏掉标记 leave_time）
    expiry = (datetime.now() - timedelta(minutes=40)).strftime("%Y-%m-%d %H:%M:%S")
    execute_insert_update(
        "UPDATE whiteboard_room_members SET leave_time=? "
        "WHERE room_id=? AND role='student' AND leave_time IS NULL "
        "AND join_time < ?",
        (expiry, room_id, expiry),
    )

    # 从数据库获取所有未离开的学生（覆盖 WS 连接和 HTTP 轮询两种场景）
    rows = execute_query(
        """SELECT m.username, m.role, m.self_snapshot, u.name, u.class,
                  m.join_time
           FROM whiteboard_room_members m
           LEFT JOIN users u ON u.username = m.username
           WHERE m.room_id=? AND m.role='student' AND m.leave_time IS NULL
           ORDER BY m.join_time""",
        (room_id,),
    )
    return [
        {
            "username": r[0], "role": r[1], "self_snapshot": r[2],
            "name": r[3] or r[0], "class": r[4] or "",
            "join_time": r[5],
        }
        for r in rows
    ]


@router.post("/rooms/{room_id}/register", summary="学生进入房间时注册")
async def register_room(room_id: int, request: Request):
    """HTTP 方式注册学生进入房间（IIS 下 WS 不通时替代 join_room）"""
    user = get_current_user(request)
    role = "teacher" if _is_teacher_or_admin(user) else "student"
    execute_insert_update(
        """INSERT OR REPLACE INTO whiteboard_room_members
           (room_id, username, role, join_time)
           VALUES (?, ?, ?, ?)""",
        (room_id, user["username"], role, _now()),
    )
    return {"status": "ok"}


@router.post("/rooms/{room_id}/spotlight", summary="投屏学生白板到全班")
async def spotlight_student(room_id: int, request: Request):
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可投屏")
    body = await request.json()
    target = body.get("username", "")
    rows = execute_query(
        "SELECT self_snapshot FROM whiteboard_room_members WHERE room_id=? AND username=?",
        (room_id, target),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="学生不在房间中")
    await whiteboard_manager.broadcast(room_id, {
        "type": "spotlight",
        "username": target,
        "snapshot": rows[0][0] or "{}",
    })
    return {"status": "ok"}


# ═══════════════════════════════════════════════════════════
# AI 辅助功能
# ═══════════════════════════════════════════════════════════

def _get_snapshot_text(room_id: int) -> str:
    """从内存或数据库获取白板当前内容的文字和图形描述"""

    def _extract_text(props: dict) -> str:
        """从形状 props 中提取文字（兼容 richText 和 text）"""
        rt = props.get("richText")
        if rt and isinstance(rt, dict):
            try:
                texts = []
                for node in rt.get("content", []):
                    for child in node.get("content", []):
                        t = child.get("text", "")
                        if t:
                            texts.append(t)
                return "".join(texts)
            except Exception:
                pass
        t = props.get("text", "")
        return t if t else ""

    snap = whiteboard_manager.rooms.get(room_id, {}).get("last_snapshot", "")
    if not snap or snap == "{}":
        rows = execute_query(
            "SELECT snapshot_data FROM whiteboard_pages WHERE room_id=? AND is_current=1 ORDER BY updated_at DESC LIMIT 1",
            (room_id,),
        )
        if rows and rows[0][0] and rows[0][0] != "{}":
            snap = rows[0][0]
    if not snap or snap == "{}":
        logger.info(f"[白板AI] room={room_id} 快照为空")
        return "白板当前为空"

    try:
        parsed = json.loads(snap)
        # TLDraw getSnapshot() 格式: { document: { store: { shapeId: {...}, ... } }, session: {...} }
        doc = parsed.get("document", {})
        store = doc.get("store", parsed.get("store", {}))
        # 筛选 typeName 为 shape 的记录（兼容新旧格式）
        shapes = {k: v for k, v in store.items() if isinstance(v, dict) and v.get("typeName") == "shape"}
        if not shapes:
            shapes = store.get("shapes", {})
        if not shapes:
            return "白板当前为空"

        type_names = {"geo": "几何形状", "arrow": "箭头", "text": "文本", "draw": "手绘", "image": "图片", "line": "线条"}
        geo_names = {"rectangle": "矩形", "ellipse": "椭圆", "diamond": "菱形", "triangle": "三角形",
                     "cloud": "云形", "pentagon": "五边形", "hexagon": "六边形", "trapezoid": "梯形",
                     "rhombus": "菱形", "star": "星形", "arrow-up": "上箭头", "arrow-down": "下箭头",
                     "arrow-left": "左箭头", "arrow-right": "右箭头", "cross": "叉形", "x-box": "X框"}

        lines = []
        texts = []
        arrows = []
        for sid, shape in shapes.items():
            props = shape.get("props", {}) or {}
            stype = shape.get("type", "unknown")
            sname = type_names.get(stype, stype)
            text = _extract_text(props)
            x = int(shape.get("x", 0))
            y = int(shape.get("y", 0))

            if text:
                texts.append(text)

            if stype == "geo":
                geo = props.get("geo", "rectangle")
                gname = geo_names.get(geo, geo)
                color = props.get("color", "black")
                fill = props.get("fill", "none")
                w = int(props.get("w", 0))
                h = int(props.get("h", 0))
                desc = f"位于({x},{y})的{gname}({w}x{h})"
                if color != "black":
                    desc += f"，{color}色边框"
                if fill and fill not in ("none", "null"):
                    desc += f"，{fill}色填充"
                if text:
                    desc += f"，文字「{text[:30]}」"
                lines.append(desc)

            elif stype == "arrow":
                start = props.get("start", {})
                end = props.get("end", {})
                sx = int(start.get("x", 0) + shape.get("x", 0))
                sy = int(start.get("y", 0) + shape.get("y", 0))
                ex = int(end.get("x", 0) + shape.get("x", 0))
                ey = int(end.get("y", 0) + shape.get("y", 0))
                arrows.append(f"从({sx},{sy})指向({ex},{ey})的箭头")

            elif stype == "image":
                src = props.get("src", "")
                desc = f"位于({x},{y})的图片"
                if src and isinstance(src, str):
                    desc += f"，来源: {src[:80]}"
                lines.append(desc)

            elif stype == "draw":
                lines.append(f"位于({x},{y})的手绘笔迹")

            elif text:
                lines.append(f"位于({x},{y})的文字「{text[:30]}」")

        # 构建结构化描述
        parts = []
        if texts:
            parts.append(f"## 白板中的文字内容\n{' | '.join(texts)}")
        if lines:
            parts.append(f"## 白板中的图形\n{chr(10).join('- ' + l for l in lines)}")
        if arrows:
            parts.append(f"## 图形之间的连线\n{chr(10).join('- ' + a for a in arrows)}")

        result = "\n\n".join(parts) if parts else "白板当前有内容，但无法识别具体元素"
        return result

    except Exception as e:
        logger.warning(f"解析白板快照失败: {e}")
        return "白板当前有内容"


def _get_snapshot_images(room_id: int) -> list[str]:
    """从白板快照中提取图片，保存为临时文件，返回文件路径列表"""
    import base64, tempfile, os, re

    snap = whiteboard_manager.rooms.get(room_id, {}).get("last_snapshot", "")
    if not snap or snap == "{}":
        rows = execute_query(
            "SELECT snapshot_data FROM whiteboard_pages WHERE room_id=? AND is_current=1 ORDER BY updated_at DESC LIMIT 1",
            (room_id,),
        )
        if rows and rows[0][0] and rows[0][0] != "{}":
            snap = rows[0][0]
    if not snap or snap == "{}":
        return []

    try:
        parsed = json.loads(snap)
        store = parsed.get("document", {}).get("store", parsed.get("store", {}))
        if not store:
            return []

        # 收集所有 assets
        assets = {}
        for k, v in store.items():
            if isinstance(v, dict) and v.get("typeName") == "asset" and v.get("type") == "image":
                assets[k] = v

        # 收集所有 image shapes
        image_shapes = []
        for k, v in store.items():
            if isinstance(v, dict) and v.get("typeName") == "shape" and v.get("type") == "image":
                image_shapes.append(v)

        if not image_shapes or not assets:
            return []

        saved_paths = []
        for shape in image_shapes:
            # TLDraw v3 image shape 用 assetId 引用 asset，不是 src
            src_ref = (shape.get("props") or {}).get("assetId", "")
            if not src_ref or not isinstance(src_ref, str):
                continue

            # assetId 已经是 asset:xxx 格式
            asset_src = ""
            if src_ref.startswith("asset:") and src_ref in assets:
                asset_src = (assets[src_ref].get("props") or {}).get("src", "")
            elif src_ref.startswith("data:") or src_ref.startswith("http"):
                asset_src = src_ref

            if not asset_src:
                continue

            # 处理 data URL
            if asset_src.startswith("data:"):
                try:
                    m = re.match(r'data:image/(\w+);base64,(.+)', asset_src)
                    if m:
                        ext = m.group(1)
                        raw = base64.b64decode(m.group(2))
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}")
                        tmp.write(raw)
                        tmp.close()
                        saved_paths.append(tmp.name)
                except Exception:
                    continue

            # 处理 http/https URL
            elif asset_src.startswith("http"):
                try:
                    import requests as sync_req
                    r = sync_req.get(asset_src, timeout=10)
                    if r.status_code == 200:
                        ext = "png"
                        ct = r.headers.get("content-type", "")
                        if "jpeg" in ct or "jpg" in ct:
                            ext = "jpg"
                        elif "gif" in ct:
                            ext = "gif"
                        elif "webp" in ct:
                            ext = "webp"
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}")
                        tmp.write(r.content)
                        tmp.close()
                        saved_paths.append(tmp.name)
                except Exception:
                    continue

        return saved_paths
    except Exception as e:
        logger.warning(f"[白板AI] 提取图片失败: {e}")
        return []


# ═══════════════════════════════════════════════════════════
# 白板视觉理解增强
# ═══════════════════════════════════════════════════════════

# TLDraw 颜色名 → hex 映射
_TL_COLOR_MAP = {
    "black": "#1a1a1a", "grey": "#8c8c8c", "light-violet": "#f0e6ff",
    "violet": "#722ed1", "blue": "#1890ff", "light-blue": "#e6f7ff",
    "yellow": "#fadb14", "orange": "#fa8c16", "green": "#52c41a",
    "light-green": "#f6ffed", "light-red": "#fff1f0", "red": "#ff4d4f",
    "white": "#ffffff",
}
_TL_GEO_MAP = {
    "rectangle": "rect", "ellipse": "ellipse", "diamond": "polygon",
    "triangle": "polygon", "cloud": "rect", "pentagon": "polygon",
    "hexagon": "polygon", "trapezoid": "polygon", "rhombus": "polygon",
    "star": "polygon", "arrow-up": "polygon", "arrow-down": "polygon",
    "arrow-left": "polygon", "arrow-right": "polygon", "cross": "path",
    "x-box": "rect",
}


def _snapshot_to_svg(snapshot_json: str, max_width: int = 1200, max_height: int = 800) -> str:
    """将 TLDraw 快照渲染为 SVG 图片（供视觉模型理解白板布局）"""
    try:
        parsed = json.loads(snapshot_json)
    except json.JSONDecodeError:
        return ""

    doc = parsed.get("document", {})
    store = doc.get("store", parsed.get("store", {}))
    shapes = {k: v for k, v in store.items()
              if isinstance(v, dict) and v.get("typeName") == "shape"}
    if not shapes:
        return ""

    # 计算边界
    min_x, min_y, max_right, max_bottom = 0, 0, max_width, max_height
    for s in shapes.values():
        sx = s.get("x", 0) or 0
        sy = s.get("y", 0) or 0
        sw = (s.get("props") or {}).get("w", 100) or 100
        sh = (s.get("props") or {}).get("h", 40) or 40
        min_x = min(min_x, sx)
        min_y = min(min_y, sy)
        max_right = max(max_right, sx + sw)
        max_bottom = max(max_bottom, sy + sh)

    pad = 20
    vw = max_right - min_x + pad * 2
    vh = max_bottom - min_y + pad * 2
    vw = max(vw, 400)
    vh = max(vh, 200)

    elements = []
    # 背景
    elements.append(f'<rect x="0" y="0" width="{vw}" height="{vh}" fill="#f9f9f9" rx="4"/>')

    for shape_id, shape in shapes.items():
        props = shape.get("props", {}) or {}
        stype = shape.get("type", "unknown")
        x = (shape.get("x", 0) or 0) - min_x + pad
        y = (shape.get("y", 0) or 0) - min_y + pad
        w = props.get("w", 100) or 100
        h = props.get("h", 40) or 40
        color = _TL_COLOR_MAP.get(props.get("color", "black"), "#1a1a1a")
        fill = props.get("fill", "none")
        opacity = "0.3" if fill and fill not in ("none", "null") else "0.1"
        text = _extract_text(props)

        if stype == "geo":
            geo = props.get("geo", "rectangle")
            svg_type = _TL_GEO_MAP.get(geo, "rect")

            if svg_type == "rect":
                elements.append(
                    f'<rect x="{x}" y="{y}" width="{w}" height="{h}" '
                    f'fill="{color}" fill-opacity="{opacity}" stroke="{color}" '
                    f'stroke-width="2" rx="4"/>'
                )
            elif svg_type == "ellipse":
                rx, ry = w / 2, h / 2
                elements.append(
                    f'<ellipse cx="{x + rx}" cy="{y + ry}" rx="{rx}" ry="{ry}" '
                    f'fill="{color}" fill-opacity="{opacity}" stroke="{color}" stroke-width="2"/>'
                )
            elif svg_type == "polygon" and geo == "triangle":
                points = f"{x + w / 2},{y} {x + w},{y + h} {x},{y + h}"
                elements.append(
                    f'<polygon points="{points}" fill="{color}" '
                    f'fill-opacity="{opacity}" stroke="{color}" stroke-width="2"/>'
                )
            elif svg_type == "polygon" and geo == "diamond":
                points = f"{x + w / 2},{y} {x + w},{y + h / 2} {x + w / 2},{y + h} {x},{y + h / 2}"
                elements.append(
                    f'<polygon points="{points}" fill="{color}" '
                    f'fill-opacity="{opacity}" stroke="{color}" stroke-width="2"/>'
                )
            else:
                elements.append(
                    f'<rect x="{x}" y="{y}" width="{w}" height="{h}" '
                    f'fill="{color}" fill-opacity="{opacity}" stroke="{color}" stroke-width="2" rx="2"/>'
                )

            # 文字
            if text:
                font_size = 14
                if props.get("size") == "xl":
                    font_size = 22
                elif props.get("size") == "l":
                    font_size = 18
                elif props.get("size") == "s":
                    font_size = 11
                lines = text.split("\n")
                line_h = font_size + 4
                total_text_h = len(lines) * line_h
                text_start_y = y + (h - total_text_h) / 2 + font_size
                for li, line in enumerate(lines):
                    elements.append(
                        f'<text x="{x + w / 2}" y="{text_start_y + li * line_h}" '
                        f'font-family="sans-serif" font-size="{font_size}" '
                        f'fill="{color}" text-anchor="middle" dominant-baseline="auto">'
                        f'{_xml_escape(line[:80])}</text>'
                    )

        elif stype == "text":
            if text:
                elements.append(
                    f'<text x="{x}" y="{y + 16}" font-family="sans-serif" '
                    f'font-size="16" fill="{color}" dominant-baseline="auto">'
                    f'{_xml_escape(text[:100])}</text>'
                )

        elif stype == "arrow":
            start = props.get("start", {})
            end = props.get("end", {})
            sx = (start.get("x", 0) or 0) + (shape.get("x", 0) or 0) - min_x + pad
            sy = (start.get("y", 0) or 0) + (shape.get("y", 0) or 0) - min_y + pad
            ex = (end.get("x", 0) or 0) + (shape.get("x", 0) or 0) - min_x + pad
            ey = (end.get("y", 0) or 0) + (shape.get("y", 0) or 0) - min_y + pad
            elements.append(
                f'<line x1="{sx}" y1="{sy}" x2="{ex}" y2="{ey}" '
                f'stroke="{color}" stroke-width="2" marker-end="url(#arrowhead)"/>'
            )

        elif stype == "draw":
            # 手绘笔迹用虚线框表示
            elements.append(
                f'<rect x="{x}" y="{y}" width="{w}" height="{h}" '
                f'fill="none" stroke="{color}" stroke-width="1" stroke-dasharray="4,3" rx="2"/>'
            )
            elements.append(
                f'<text x="{x + w / 2}" y="{y + h / 2 + 4}" font-family="sans-serif" '
                f'font-size="11" fill="{color}" text-anchor="middle">✏️ 手绘</text>'
            )

        elif stype == "image":
            elements.append(
                f'<rect x="{x}" y="{y}" width="{w}" height="{h}" '
                f'fill="#e6f7ff" stroke="#1890ff" stroke-width="1" stroke-dasharray="4,3" rx="2"/>'
            )
            elements.append(
                f'<text x="{x + w / 2}" y="{y + h / 2 + 4}" font-family="sans-serif" '
                f'font-size="12" fill="#1890ff" text-anchor="middle">🖼️ 图片</text>'
            )

        elif stype == "line":
            # 简单线条
            elements.append(
                f'<line x1="{x}" y1="{y}" x2="{x + w}" y2="{y + h}" '
                f'stroke="{color}" stroke-width="2"/>'
            )

    # 如果有箭头，添加箭头标记定义
    has_arrow = any(s.get("type") == "arrow" for s in shapes.values())
    defs = ""
    if has_arrow:
        defs = (
            '<defs><marker id="arrowhead" markerWidth="10" markerHeight="7" '
            'refX="10" refY="3.5" orient="auto">'
            '<polygon points="0 0, 10 3.5, 0 7" fill="#1a1a1a"/></marker></defs>'
        )

    svg = (
        f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {vw} {vh}" '
        f'width="{vw}" height="{vh}" style="background:#f9f9f9;font-family:sans-serif">'
        f'{defs}{"".join(elements)}</svg>'
    )
    return svg


def _xml_escape(text: str) -> str:
    """转义 XML 特殊字符"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")\
               .replace('"', "&quot;").replace("'", "&apos;")


def _snapshot_to_base64_svg(snapshot_json: str) -> str:
    """将 TLDraw 快照渲染为 base64 SVG data URL"""
    svg = _snapshot_to_svg(snapshot_json)
    if not svg:
        return ""
    import base64
    return "data:image/svg+xml;base64," + base64.b64encode(svg.encode("utf-8")).decode("utf-8")


async def _understand_whiteboard_with_vision(
    room_id: int, api_key: str, snapshot_text: str
) -> str:
    """
    增强白板理解：
    1. 从快照提取精确的布局描述（形状位置、颜色、文字、箭头关系）
    2. 将白板上的真实图片发送给视觉模型识别
    3. 融合文本解析 + 图片视觉理解，返回增强描述
    """
    # 获取快照
    snap = whiteboard_manager.rooms.get(room_id, {}).get("last_snapshot", "")
    if not snap or snap == "{}":
        rows = execute_query(
            "SELECT snapshot_data FROM whiteboard_pages WHERE room_id=? AND is_current=1 ORDER BY updated_at DESC LIMIT 1",
            (room_id,),
        )
        if rows and rows[0][0] and rows[0][0] != "{}":
            snap = rows[0][0]
    if not snap or snap == "{}":
        return snapshot_text

    # ── 生成精确布局描述（替代 SVG 图片，因为视觉模型不支持 SVG 格式）──
    layout_desc = _snapshot_to_layout_text(snap)

    # 获取白板上的真实图片（PNG/JPG — 视觉模型支持这些格式）
    image_paths = _get_snapshot_images(room_id)

    if image_paths:
        import requests as sync_requests
        from backend.api.chat_router import get_config_value as get_cfg
        from backend.utils import encode_image_to_base64, get_image_mime_type

        model_name = get_cfg("MODEL_VL_NAME", "qwen3-vl-plus")
        api_base = get_cfg("QWEN_OPENAI_API_BASE",
                           "https://dashscope.aliyuncs.com/compatible-mode/v1")

        content = []
        # 白板上的真实图片
        for fp in image_paths:
            try:
                b64 = encode_image_to_base64(fp)
                mime = get_image_mime_type(fp)
                content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"},
                })
            except Exception as e:
                logger.warning(f"[白板AI-视觉增强] 图片编码失败 {fp}: {e}")

        if content:
            vision_prompt = (
                "请仔细查看白板上的这张图片内容。描述你看到了什么，包括：\n"
                "1. 图片中有什么物体/文字\n"
                "2. 图片和白板教学内容的关联\n"
                "3. 提取图片中任何可读的文字"
            )
            content.append({"type": "text", "text": vision_prompt})

            payload = {
                "model": model_name,
                "messages": [{"role": "user", "content": content}],
                "stream": False,
            }

            try:
                resp = sync_requests.post(
                    f"{api_base}/chat/completions",
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    json=payload,
                    timeout=60,
                )
                if resp.status_code == 200:
                    data = resp.json()
                    vision_text = data["choices"][0]["message"]["content"]
                    # 融合文本提取 + 布局描述 + 图片视觉识别
                    enhanced = (
                        f"## 白板文字提取（结构化解析）\n{snapshot_text}\n\n"
                        f"## 白板布局描述（精确位置）\n{layout_desc}\n\n"
                        f"## 白板图片视觉识别\n{vision_text}"
                    )
                    logger.info(f"[白板AI-视觉增强] room={room_id} 文本+图片融合完成")
                    return enhanced
            except Exception as e:
                logger.warning(f"[白板AI-视觉增强] 图片视觉分析异常: {e}")

    # 没有图片或图片分析失败时，返回文本 + 精确布局描述
    enhanced = (
        f"## 白板文字提取（结构化解析）\n{snapshot_text}\n\n"
        f"## 白板布局描述（精确位置）\n{layout_desc}"
    )
    # 统一清理临时图片
    for fp in image_paths:
        try:
            import os
            os.unlink(fp)
        except Exception:
            pass

    return enhanced


def _snapshot_to_layout_text(snapshot_json: str) -> str:
    """将 TLDraw 快照转换为精确的布局文字描述（替代 SVG 图片）"""
    try:
        parsed = json.loads(snapshot_json)
    except json.JSONDecodeError:
        return ""

    doc = parsed.get("document", {})
    store = doc.get("store", parsed.get("store", {}))
    shapes = {k: v for k, v in store.items()
              if isinstance(v, dict) and v.get("typeName") == "shape"}
    if not shapes:
        return "白板当前为空"

    type_names = {
        "geo": "几何形状", "arrow": "箭头", "text": "文本",
        "draw": "手绘", "image": "图片", "line": "线条"
    }
    geo_names = {
        "rectangle": "矩形", "ellipse": "椭圆", "diamond": "菱形",
        "triangle": "三角形", "cloud": "云形", "star": "星形",
        "arrow-up": "上箭头", "arrow-down": "下箭头",
        "arrow-left": "左箭头", "arrow-right": "右箭头",
    }
    tl_colors = {
        "black": "黑色", "grey": "灰色", "blue": "蓝色",
        "light-blue": "浅蓝", "green": "绿色", "light-green": "浅绿",
        "orange": "橙色", "yellow": "黄色", "red": "红色",
        "light-red": "浅红", "violet": "紫色", "light-violet": "浅紫",
        "white": "白色",
    }

    # 按位置排序（从上到下，从左到右）
    sorted_shapes = sorted(
        shapes.values(),
        key=lambda s: ((s.get("y", 0) or 0), (s.get("x", 0) or 0))
    )

    items = []
    arrow_count = 0
    for shape in sorted_shapes:
        props = shape.get("props", {}) or {}
        stype = shape.get("type", "unknown")
        x = int(shape.get("x", 0) or 0)
        y = int(shape.get("y", 0) or 0)
        w = int(props.get("w", 0) or 0)
        h = int(props.get("h", 0) or 0)
        color = tl_colors.get(props.get("color", "black"), props.get("color", "black"))
        text = _extract_text(props)

        if stype == "geo":
            geo = props.get("geo", "rectangle")
            gname = geo_names.get(geo, geo)
            fill = props.get("fill", "none")
            fill_info = f"，{color}色填充" if fill and fill not in ("none", "null") else ""
            text_info = f"，内容「{text}」" if text else ""
            items.append(f"[{gname}] 位置({x},{y}) 大小({w}x{h}){fill_info}{text_info}")

        elif stype == "arrow":
            start = props.get("start", {})
            end = props.get("end", {})
            sx = int((start.get("x", 0) or 0) + (shape.get("x", 0) or 0))
            sy = int((start.get("y", 0) or 0) + (shape.get("y", 0) or 0))
            ex = int((end.get("x", 0) or 0) + (shape.get("x", 0) or 0))
            ey = int((end.get("y", 0) or 0) + (shape.get("y", 0) or 0))
            arrow_count += 1
            label = text or ""
            label_info = f"，标注「{label}」" if label else ""
            items.append(f"[箭头#{arrow_count}] 从({sx},{sy})指向({ex},{ey}){label_info}")

        elif stype == "text":
            if text:
                items.append(f"[文本] 位置({x},{y}) 内容「{text}」")

        elif stype == "draw":
            items.append(f"[手绘] 位置({x},{y}) 大小({w}x{h})")

        elif stype == "image":
            items.append(f"[图片] 位置({x},{y}) 大小({w}x{h})")

        elif stype == "line":
            items.append(f"[线条] 从({x},{y})到({x + w},{y + h})")

    if not items:
        return "白板当前有内容，但无法识别具体元素"

    # 添加布局总结
    total = len(items)
    arrows_found = sum(1 for i in items if i.startswith("[箭头"))
    texts_found = sum(1 for i in items if "内容「" in i)

    summary = (
        f"白板共有 {total} 个元素，其中 {texts_found} 个包含文字，"
        f"{arrows_found} 个箭头连接。\n"
        f"元素按从上到下、从左到右排列如下：\n"
    )
    for idx, item in enumerate(items, 1):
        summary += f"  {idx}. {item}\n"

    # 尝试识别箭头连接关系（文字相近的形状之间的箭头）
    if arrows_found > 0:
        summary += "\n结构关系分析："
        for i, item in enumerate(items):
            if item.startswith("[箭头"):
                # 找出该箭头前后的形状
                before = items[i - 1] if i > 0 else None
                after = items[i + 1] if i < len(items) - 1 else None
                if before and after and ("几何" in before or "文本" in before):
                    before_text = before.split("内容「")[-1].rstrip("」") if "内容「" in before else ""
                    after_text = after.split("内容「")[-1].rstrip("」") if "内容「" in after else ""
                    summary += f"\n  - 箭头连接："
                    if before_text:
                        summary += f"「{before_text}」→"
                    if after_text:
                        summary += f"「{after_text}」"

    return summary


@router.post("/ai/chat-stream", summary="AI 白板助手流式对话")
async def ai_chat_stream(request: Request):
    """白板 AI 助手 SSE 流式对话"""
    user = get_current_user(request)
    username = user["username"]
    body = await request.json()
    prompt = body.get("prompt", "")
    room_id = body.get("room_id")

    if not prompt:
        raise HTTPException(status_code=400, detail="请输入问题")

    # 获取 API Key
    dashscope_api_key, _ = get_api_keys(username)
    if not dashscope_api_key:
        return StreamingResponse(
            _error_stream("API Key 未配置，请联系管理员"),
            media_type="text/event-stream",
        )

    # 获取白板上下文
    mode = whiteboard_manager.get_mode(room_id) if room_id else "demo"
    kp_name = body.get("kp_name", "")
    subject = body.get("subject", "通用技术")
    use_vision = body.get("use_vision", False)  # 前端可控制是否使用视觉理解

    # ── 获取白板理解文本（视觉增强版） ──
    snapshot_text = _get_snapshot_text(room_id) if room_id else "无白板内容"
    if room_id and use_vision and dashscope_api_key:
        try:
            enhanced = await _understand_whiteboard_with_vision(room_id, dashscope_api_key, snapshot_text)
            if enhanced != snapshot_text:
                snapshot_text = enhanced
                logger.info(f"[白板AI] room={room_id} 使用视觉增强理解")
        except Exception as e:
            logger.warning(f"[白板AI] 视觉增强失败，降级到文本解析: {e}")

    # 构建系统提示词
    from backend.prompts.whiteboard_ai import WHITEBOARD_TEACHER_ASSISTANT
    system_prompt = WHITEBOARD_TEACHER_ASSISTANT.format(
        mode=mode,
        kp_name=kp_name or "未指定",
        subject=subject,
        snapshot_text=snapshot_text,
    )
    enhanced_prompt = f"{system_prompt}\n\n---\n\n用户提问：{prompt}"

    # 检测白板中是否有图片，有则走视觉流式（传统逻辑）
    image_paths = _get_snapshot_images(room_id) if room_id else []
    if image_paths and not use_vision:
        return StreamingResponse(
            _whiteboard_ai_vision_stream(system_prompt, prompt, image_paths, username, dashscope_api_key),
            media_type="text/event-stream",
        )

    from backend.api.chat_router import _chat_event_generator as _cg
    return StreamingResponse(
        _whiteboard_ai_stream(enhanced_prompt, username, dashscope_api_key),
        media_type="text/event-stream",
    )


def _whiteboard_ai_stream(prompt: str, username: str, api_key: str):
    """白板 AI 流式生成器（后端自行累加，前端只替换不追加）"""
    import asyncio
    from backend.api.chat_router import _agent_chat_stream
    try:
        full = ""
        for chunk in _agent_chat_stream(prompt, None, api_key, username):
            full += chunk["text"]
            yield f"data: {json.dumps({'type': 'delta', 'content': full})}\n\n"
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except GeneratorExit:
        # 客户端断开连接
        logger.info("[白板AI] 流式对话客户端断开")
    except Exception as e:
        logger.error(f"白板 AI 流式生成失败: {e}")
        yield f"data: {json.dumps({'type': 'error', 'content': f'AI 响应失败：{str(e)}'})}\n\n"


def _whiteboard_ai_vision_stream(system_prompt: str, user_prompt: str, image_paths: list[str],
                                 username: str, api_key: str):
    """白板 AI 视觉流式生成：将白板图片+文字一起送入视觉模型"""
    import requests as sync_requests
    from backend.api.chat_router import get_config_value as get_cfg
    from backend.utils import encode_image_to_base64, get_image_mime_type

    model_name = get_cfg("MODEL_VL_NAME", "qwen3-vl-plus")
    api_base = get_cfg("QWEN_OPENAI_API_BASE",
                       "https://dashscope.aliyuncs.com/compatible-mode/v1")

    # 构建多模态 content 数组
    content = []
    # 先加图片
    for fp in image_paths:
        try:
            b64 = encode_image_to_base64(fp)
            mime = get_image_mime_type(fp)
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{b64}"},
            })
        except Exception as e:
            logger.warning(f"[白板AI] 图片编码失败 {fp}: {e}")
    # 再加用户提问
    content.append({"type": "text", "text": f"{system_prompt}\n\n用户提问：{user_prompt}"})

    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": content}],
        "stream": True,
    }

    try:
        resp = sync_requests.post(
            f"{api_base}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=payload,
            stream=True,
            timeout=120,
        )
        if resp.status_code != 200:
            yield f"data: {json.dumps({'type': 'error', 'content': f'视觉模型调用失败: HTTP {resp.status_code}'})}\n\n"
            return

        full_text = ""
        for line in resp.iter_lines():
            if not line:
                continue
            decoded = line.decode("utf-8") if isinstance(line, bytes) else line
            if decoded.startswith("data:"):  # type: ignore[arg-type]
                data_str = decoded[5:]  # type: ignore[union-attr]
                if data_str.strip() == "[DONE]":
                    break
                try:
                    data = json.loads(data_str)
                    if "choices" in data and data["choices"]:
                        delta = data["choices"][0].get("delta", {})
                        content_delta = delta.get("content", "")
                        if content_delta:
                            full_text += content_delta
                            yield f"data: {json.dumps({'type': 'delta', 'content': full_text})}\n\n"
                except json.JSONDecodeError:
                    continue
        yield f"data: {json.dumps({'type': 'done'})}\n\n"

        # 清理临时图片文件
        for fp in image_paths:
            try:
                import os
                os.unlink(fp)
            except Exception:
                pass

    except Exception as e:
        logger.error(f"[白板AI] 视觉流式生成失败: {e}")
        yield f"data: {json.dumps({'type': 'error', 'content': f'AI 响应失败：{str(e)}'})}\n\n"
    except GeneratorExit:
        logger.info("[白板AI] 视觉流式客户端断开")


def _error_stream(msg: str):
    """生成错误流"""
    yield f"data: {json.dumps({'type': 'error', 'content': msg})}\n\n"


@router.post("/ai/generate-diagram", summary="AI 生成图示（SSE 流式进度，SVG 优先）")
async def ai_generate_diagram(request: Request):
    """AI 图示生成：SSE 流式返回进度和结果，防止 IIS 超时"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    body = await request.json()
    description = body.get("description", "")
    subject = body.get("subject", "通用技术")
    if not description:
        raise HTTPException(status_code=400, detail="请输入描述")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    return StreamingResponse(
        _generate_diagram_stream(description, subject, dashscope_api_key),
        media_type="text/event-stream",
    )


async def _generate_diagram_stream(description: str, subject: str, api_key: str):
    """AI 图示生成流式生成器（SSE 事件流）"""
    import asyncio

    def _sse(event: str, data: dict) -> str:
        return f"event: {event}\ndata: {json.dumps(data)}\n\n"

    try:
        # ── 阶段1：AI 分析需求 ──
        yield _sse("progress", {"phase": "analyzing", "message": "AI 正在分析您的需求..."})
        await asyncio.sleep(0.1)

        from backend.prompts.whiteboard_ai import DIAGRAM_GENERATION_PROMPT
        prompt = DIAGRAM_GENERATION_PROMPT.format(description=description, subject=subject)
        timeout = _get_ai_timeout()

        from backend.api.ai_service import call_ai_sync_with_timeout
        result = await call_ai_sync_with_timeout(prompt, api_key, timeout=timeout)

        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        if not jm:
            yield _sse("error", {"message": "AI 返回格式异常"})
            return
        data = json.loads(jm.group())

        mode = data.get("mode", "svg")

        # ── SVG 模式 ──
        if mode == "svg":
            svg_content = data.get("svg", "")
            if not svg_content:
                yield _sse("error", {"message": "AI 未生成 SVG 内容"})
                return
            if "<svg" not in svg_content:
                yield _sse("error", {"message": "AI 生成内容不是有效 SVG"})
                return

            yield _sse("progress", {"phase": "svg", "message": "AI 正在生成 SVG 图示..."})
            await asyncio.sleep(0.1)
            yield _sse("result", {
                "mode": "svg",
                "svg": svg_content,
                "width": data.get("width", 800),
                "height": data.get("height", 600),
                "title": data.get("title", ""),
            })
            return

        # ── 图片模式（调用通义万相） ──
        if mode == "image":
            image_prompt = data.get("prompt", description)
            yield _sse("progress", {"phase": "image_gen", "message": "AI 正在调用图片生成服务（通义万相），预计需要 30-60 秒..."})

            from backend.api.image_gen_service import generate_and_save_image
            import uuid
            from pathlib import Path

            save_dir = Path("question_media") / "whiteboard_ai"
            save_dir.mkdir(parents=True, exist_ok=True)
            filename = f"wb_{uuid.uuid4().hex}"

            # 图片生成也有超时控制
            try:
                local_path = await asyncio.wait_for(
                    generate_and_save_image(
                        prompt=image_prompt,
                        save_dir=save_dir,
                        filename=filename,
                    ),
                    timeout=120,
                )
                if local_path:
                    yield _sse("result", {
                        "mode": "image",
                        "image_url": f"/api/files/question_media/whiteboard_ai/{Path(local_path).name}",
                        "title": data.get("title", ""),
                    })
                    return
                else:
                    yield _sse("error", {"message": "图片生成失败，请稍后重试", "fallback": image_prompt, "mode": "text"})
                    return
            except asyncio.TimeoutError:
                logger.warning(f"通义万相生图超时: {image_prompt[:100]}")
                yield _sse("error", {"message": "图片生成超时（超过120秒），请稍后重试或简化描述", "fallback": image_prompt, "mode": "text"})
                return

        yield _sse("error", {"message": f"未知模式: {mode}"})

    except TimeoutError as e:
        yield _sse("error", {"message": str(e)})
    except asyncio.CancelledError:
        # 客户端断开连接时的正常清理
        logger.info("[白板AI] 图示生成被客户端取消")
    except Exception as e:
        logger.error(f"AI 生成图示失败: {e}")
        yield _sse("error", {"message": f"AI 生成失败: {str(e)}"})


@router.post("/ai/generate-board", summary="AI 根据知识点生成完整板书")
async def ai_generate_board(request: Request):
    """一键板书：根据知识点生成结构化板书形状"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    body = await request.json()
    kp_name = body.get("kp_name", "")
    subject = body.get("subject", "通用技术")
    grade = body.get("grade", "")
    if not kp_name:
        raise HTTPException(status_code=400, detail="请指定知识点")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    from backend.prompts.whiteboard_ai import BOARD_GENERATION_PROMPT
    prompt = BOARD_GENERATION_PROMPT.format(kp_name=kp_name, subject=subject, grade=grade)

    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        if jm:
            data = json.loads(jm.group())
            return {"title": data.get("title", kp_name), "shapes": data.get("shapes", [])}
        return {"title": kp_name, "shapes": [], "raw": result}
    except TimeoutError as e:
        logger.warning(f"AI 生成板书超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        logger.error(f"AI 生成板书失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 生成板书失败: {str(e)}")


@router.post("/ai/beautify-board", summary="AI 美化排版白板内容")
async def ai_beautify_board(request: Request):
    """板书美化+自动排版：读取白板当前内容，AI 重新组织为整洁的结构化板书"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    body = await request.json()
    room_id = body.get("room_id")
    subject = body.get("subject", "通用技术")
    if not room_id:
        raise HTTPException(status_code=400, detail="请提供房间 ID")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    snapshot_text = _get_snapshot_text(room_id)
    if not snapshot_text or snapshot_text == "白板当前为空":
        raise HTTPException(status_code=400, detail="白板当前为空，无内容可美化")

    from backend.prompts.whiteboard_ai import BEAUTIFY_BOARD_PROMPT
    prompt = BEAUTIFY_BOARD_PROMPT.format(
        snapshot_text=snapshot_text,
        subject=subject,
    )

    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        if jm:
            data = json.loads(jm.group())
            return {"title": data.get("title", "美化板书"), "shapes": data.get("shapes", [])}
        return {"title": "美化板书", "shapes": [], "raw": result}
    except TimeoutError as e:
        logger.warning(f"AI 美化排版超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        logger.error(f"AI 美化排版失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 美化排版失败: {str(e)}")


@router.post("/ai/smart-annotation", summary="AI 智能批注选中内容")
async def ai_smart_annotation(request: Request):
    """智能批注：分析白板选中的内容并给出标注建议"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    body = await request.json()
    selection_desc = body.get("selection_desc", "")
    mode = body.get("mode", "demo")
    if not selection_desc:
        raise HTTPException(status_code=400, detail="请先在白板上选中内容")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    from backend.prompts.whiteboard_ai import SMART_LABEL_PROMPT
    prompt = SMART_LABEL_PROMPT.format(selection_desc=selection_desc, mode=mode)

    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        if jm:
            data = json.loads(jm.group())
            return data
        return {"summary": "", "label_type": "comment", "label_text": result, "color": "#ff4d4f"}
    except TimeoutError as e:
        logger.warning(f"AI 智能批注超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        logger.error(f"AI 智能批注失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 批注失败: {str(e)}")


@router.post("/ai/generate-mindmap", summary="AI 根据板书生成思维导图")
async def ai_generate_mindmap(request: Request):
    """思维导图：根据白板内容生成结构化思维导图"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    body = await request.json()
    room_id = body.get("room_id")
    subject = body.get("subject", "通用技术")
    if not room_id:
        raise HTTPException(status_code=400, detail="请提供房间 ID")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    snapshot_text = _get_snapshot_text(room_id)
    if not snapshot_text or snapshot_text == "白板当前为空":
        raise HTTPException(status_code=400, detail="白板当前为空")

    from backend.prompts.whiteboard_ai import MIND_MAP_PROMPT
    prompt = MIND_MAP_PROMPT.format(snapshot_text=snapshot_text, subject=subject)

    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        if jm:
            data = json.loads(jm.group())
            return {"title": data.get("title", "思维导图"), "shapes": data.get("shapes", [])}
        return {"title": "思维导图", "shapes": [], "raw": result}
    except TimeoutError as e:
        logger.warning(f"AI 思维导图生成超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        logger.error(f"AI 思维导图生成失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 思维导图生成失败: {str(e)}")


@router.post("/ai/suggest", summary="AI 根据当前内容推荐下一步")
async def ai_suggest(request: Request):
    user = get_current_user(request)
    body = await request.json()
    current_content = body.get("content", "")
    kp_name = body.get("kp_name", "")
    if not current_content:
        raise HTTPException(status_code=400, detail="请提供当前白板内容")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    prompt = f"""你是一位教学助手。课堂上白板当前内容如下，请给出简短的下一步教学建议（50字以内）。

{current_content}
知识点：{kp_name}

直接回复文字，不要JSON格式。"""
    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        return {"suggestion": result.strip()}
    except TimeoutError as e:
        logger.warning(f"AI 教学建议超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 建议失败: {str(e)}")


@router.get("/ai/export-summary/{room_id}", summary="导出板书总结（Word）")
async def export_board_summary(room_id: int, request: Request):
    """板书总结 + 导出 Word：AI 总结白板内容并生成 docx 文件"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    username = user["username"]

    dashscope_api_key, _ = get_api_keys(username)
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    snapshot_text = _get_snapshot_text(room_id)
    if not snapshot_text or snapshot_text == "白板当前为空":
        raise HTTPException(status_code=400, detail="白板当前为空，无内容可导出")

    # 获取房间信息
    room_rows = execute_query(
        "SELECT title FROM whiteboard_rooms WHERE id=?", (room_id,),
    )
    room_title = room_rows[0][0] if room_rows else "白板"
    subject = "通用技术"

    # AI 生成总结
    from backend.prompts.whiteboard_ai import BOARD_SUMMARY_PROMPT
    prompt = BOARD_SUMMARY_PROMPT.format(snapshot_text=snapshot_text, subject=subject)
    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        data = json.loads(jm.group()) if jm else {}
    except TimeoutError as e:
        logger.warning(f"AI 导出总结超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        logger.error(f"AI 总结失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 总结失败: {str(e)}")

    # 生成 Word 文档
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io, urllib.parse
    from datetime import datetime

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 标题
    title_text = data.get("title", f"{room_title} - 板书总结")
    t = doc.add_heading(title_text, level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"课程：{room_title}  学科：{subject}  导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # 内容概括
    summary = data.get("summary", "")
    if summary:
        doc.add_heading("内容概括", level=2)
        doc.add_paragraph(summary)

    # 核心要点
    key_points = data.get("key_points", [])
    if key_points:
        doc.add_heading("核心要点", level=2)
        for i, point in enumerate(key_points):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"要点{i+1}：")
            run.bold = True
            p.add_run(str(point))

    # 难点
    difficulties = data.get("difficulties", [])
    if difficulties:
        doc.add_heading("重点难点", level=2)
        for d in difficulties:
            doc.add_paragraph(d, style='List Bullet')

    # 课后作业
    homework = data.get("homework", "")
    if homework:
        doc.add_heading("课后作业", level=2)
        doc.add_paragraph(homework)

    # 原始板书内容
    doc.add_paragraph()
    doc.add_heading("板书原始内容", level=2)
    doc.add_paragraph(snapshot_text[:2000])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_fn = urllib.parse.quote(f"板书总结_{room_title}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_fn}"},
    )


@router.post("/ai/generate-quiz", summary="AI 根据板书生成随堂提问")
async def ai_generate_quiz(request: Request):
    """随堂提问：根据白板内容生成一道选择题"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    body = await request.json()
    room_id = body.get("room_id")
    subject = body.get("subject", "通用技术")
    kp_name = body.get("kp_name", "")
    if not room_id:
        raise HTTPException(status_code=400, detail="请提供房间 ID")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    snapshot_text = _get_snapshot_text(room_id)
    if not snapshot_text or snapshot_text == "白板当前为空":
        raise HTTPException(status_code=400, detail="白板当前为空，请先在白板上书写内容")

    from backend.prompts.whiteboard_ai import QUIZ_GENERATION_PROMPT
    prompt = QUIZ_GENERATION_PROMPT.format(
        snapshot_text=snapshot_text,
        kp_name=kp_name or "未指定",
        subject=subject,
    )

    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        if jm:
            data = json.loads(jm.group())
            return data
        return {"error": "AI 返回格式异常", "raw": result}
    except TimeoutError as e:
        logger.warning(f"AI 生成随堂提问超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        logger.error(f"AI 生成随堂提问失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 生成失败: {str(e)}")


@router.post("/ai/generate-bilingual", summary="AI 将板书转换为中英双语")
async def ai_generate_bilingual(request: Request):
    """双语板书：读取当前白板内容，生成中英双语对照文本"""
    user = get_current_user(request)
    if not _is_teacher_or_admin(user):
        raise HTTPException(status_code=403, detail="仅教师可使用")
    body = await request.json()
    room_id = body.get("room_id")
    subject = body.get("subject", "通用技术")
    if not room_id:
        raise HTTPException(status_code=400, detail="请提供房间 ID")

    dashscope_api_key, _ = get_api_keys(user["username"])
    if not dashscope_api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    snapshot_text = _get_snapshot_text(room_id)
    if not snapshot_text or snapshot_text == "白板当前为空":
        raise HTTPException(status_code=400, detail="白板当前为空，请先在白板上书写内容")

    from backend.prompts.whiteboard_ai import BILINGUAL_BOARD_PROMPT
    prompt = BILINGUAL_BOARD_PROMPT.format(snapshot_text=snapshot_text, subject=subject)

    try:
        from backend.api.ai_service import call_ai_sync_with_timeout
        timeout = _get_ai_timeout()
        result = await call_ai_sync_with_timeout(prompt, dashscope_api_key, timeout=timeout)
        import re
        jm = re.search(r'\{[\s\S]*\}', result.strip())
        if jm:
            data = json.loads(jm.group())
            pairs = data.get("pairs", [])
            # 将双语对转为 TLDraw 形状
            shapes = []
            for i, pair in enumerate(pairs):
                base_y = pair.get("y", 80 + i * 80)
                base_x = pair.get("x", 50)
                # 中文矩形
                shapes.append({
                    "type": "geo",
                    "x": base_x,
                    "y": base_y,
                    "props": {
                        "geo": "rectangle",
                        "w": 320,
                        "h": 36,
                        "color": "#1890ff",
                        "fill": "#e6f7ff",
                        "text": pair.get("chinese", ""),
                        "fontSize": 15,
                        "fontWeight": "normal",
                        "textAlign": "middle",
                        "size": "m",
                    },
                })
                # 英文矩形（右侧对齐）
                shapes.append({
                    "type": "geo",
                    "x": base_x + 340,
                    "y": base_y,
                    "props": {
                        "geo": "rectangle",
                        "w": 320,
                        "h": 36,
                        "color": "#52c41a",
                        "fill": "#f6ffed",
                        "text": pair.get("english", ""),
                        "fontSize": 14,
                        "fontWeight": "normal",
                        "textAlign": "middle",
                        "size": "m",
                    },
                })
            return {"shapes": shapes, "title": "中英双语板书"}
        return {"shapes": [], "raw": result}
    except TimeoutError as e:
        logger.warning(f"AI 生成双语板书超时: {e}")
        raise HTTPException(status_code=504, detail=str(e))
    except Exception as e:
        logger.error(f"AI 生成双语板书失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 生成双语板书失败: {str(e)}")


# ═══════════════════════════════════════════════════════════
# WebSocket 端点
# ═══════════════════════════════════════════════════════════

@router.websocket("/ws/{room_id}")
async def whiteboard_websocket(websocket: WebSocket, room_id: int):
    """白板 WebSocket 实时通信"""
    # 1. 认证
    from backend.auth import decode_jwt_token
    token = websocket.query_params.get("token", "")
    if not token:
        await websocket.close(code=4001, reason="缺少认证令牌")
        return

    try:
        payload = decode_jwt_token(token)
        if payload is None:
            await websocket.close(code=4001, reason="认证失败")
            return
        username = payload.get("username", "") or ""
    except Exception:
        await websocket.close(code=4001, reason="认证失败")
        return

    if not username:
        await websocket.close(code=4001, reason="认证失败")
        return

    # 2. 查询用户详情
    rows = execute_query(
        "SELECT username, role, name FROM users WHERE username=?", (username,),
    )
    if not rows:
        await websocket.close(code=4001, reason="用户不存在")
        return

    row = rows[0]
    username = row[0]
    role_num = row[1]
    role = "teacher" if role_num in (0, 1) else "student"

    # 3. 验证房间存在
    room_rows = execute_query(
        "SELECT id, mode, creator_username FROM whiteboard_rooms WHERE id=? AND status='active'",
        (room_id,),
    )
    if not room_rows:
        await websocket.close(code=4003, reason="房间不存在或已结束")
        return

    room_mode = room_rows[0][1]

    # 4. 加入房间
    await whiteboard_manager.join_room(room_id, username, role, websocket)
    # 同步房间模式
    whiteboard_manager.set_mode(room_id, room_mode)

    try:
        while True:
            data = await websocket.receive_json()
            msg_type = data.get("type", "")

            # ── 白板操作 ──
            if msg_type == "op":
                # 使用管理器中的实时模式，而非连接时的缓存值
                live_mode = whiteboard_manager.get_mode(room_id)
                if _can_operate(room_id, username, role, live_mode):
                    await whiteboard_manager.handle_op(room_id, username, data)
                else:
                    logger.warning(f"[白板WS] op denied for {username}({role}), room={room_id}, mode={live_mode}")

            # ── 光标同步 ──
            elif msg_type == "cursor":
                await whiteboard_manager.handle_cursor(room_id, username, data)

            # ── 切换页面（教师）──
            elif msg_type == "switch_page":
                if role == "teacher":
                    page = data.get("page", 1)
                    whiteboard_manager.set_current_page(room_id, page)
                    # 加载新页面快照
                    page_rows = execute_query(
                        "SELECT snapshot_data FROM whiteboard_pages WHERE room_id=? AND page_number=?",
                        (room_id, page),
                    )
                    await whiteboard_manager.broadcast(room_id, {
                        "type": "page_switched",
                        "page": page,
                        "snapshot": page_rows[0][0] if page_rows else "{}",
                    })

            # ── 切换模式（教师）──
            elif msg_type == "mode_change":
                if role == "teacher":
                    new_mode = data.get("mode", "demo")
                    old_mode = whiteboard_manager.get_mode(room_id)
                    whiteboard_manager.set_mode(room_id, new_mode)
                    execute_insert_update(
                        "UPDATE whiteboard_rooms SET mode=? WHERE id=?",
                        (new_mode, room_id),
                    )
                    # 退出互动模式时清除所有学生授权
                    if old_mode == "interactive" and new_mode != "interactive":
                        room = whiteboard_manager.rooms.get(room_id)
                        if room:
                            room.get("granted_users", set()).clear()
                            for conn in room.get("connections", {}).values():
                                conn["granted"] = False
                    await whiteboard_manager.broadcast(room_id, {
                        "type": "mode_changed",
                        "mode": new_mode,
                    })

            # ── 举手 ──
            elif msg_type == "raise_hand":
                await whiteboard_manager.broadcast(room_id, {
                    "type": "hand_raised",
                    "username": username,
                })

            # ── 学生请求同步当前快照 ──
            elif msg_type == "request_sync":
                last_snap = whiteboard_manager.rooms.get(room_id, {}).get("last_snapshot", "")
                if last_snap:
                    await whiteboard_manager.send_to_user(room_id, username, {
                        "type": "op_broadcast",
                        "sender": "system",
                        "data": {"snapshot": last_snap},
                    })

            # ── 自习模式：学生保存自己的白板 ──
            elif msg_type == "self_save":
                if role == "student" and whiteboard_manager.get_mode(room_id) == "self_study":
                    execute_insert_update(
                        "UPDATE whiteboard_room_members SET self_snapshot=? WHERE room_id=? AND username=?",
                        (data.get("snapshot", "{}"), room_id, username),
                    )

            # ── 自习模式：学生提交 ──
            elif msg_type == "self_submit":
                if role == "student" and whiteboard_manager.get_mode(room_id) == "self_study":
                    execute_insert_update(
                        "UPDATE whiteboard_room_members SET self_snapshot=? WHERE room_id=? AND username=?",
                        (data.get("snapshot", "{}"), room_id, username),
                    )
                    await whiteboard_manager.broadcast(room_id, {
                        "type": "student_submitted",
                        "username": username,
                    })

    except WebSocketDisconnect:
        await whiteboard_manager.leave_room(room_id, username)
    except Exception as e:
        logger.warning(f"白板 WS 异常 room#{room_id} user={username}: {e}")
        await whiteboard_manager.leave_room(room_id, username)


def _can_operate(room_id: int, username: str, role: str, mode: str) -> bool:
    """检查操作权限"""
    if role == "teacher":
        return True
    if mode == "interactive":
        return whiteboard_manager.is_granted(room_id, username)
    if mode == "self_study":
        return True  # 自习模式各自可操作
    return False  # demo 模式学生不可操作
