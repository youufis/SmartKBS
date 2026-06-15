# -*- coding: utf-8 -*-
"""
课堂互动 API 路由
随堂测验、快速投票、课堂提问
"""
import asyncio
import json
import re
from datetime import datetime

from fastapi import APIRouter, HTTPException, Request, Query
from pydantic import BaseModel

from backend.api.dependencies import get_current_user
from backend.database import execute_query, execute_insert_update, execute_batch
from backend.logger import logger

router = APIRouter()


def _get_user_grade_class(username: str) -> tuple:
    """查询用户的年级(grade)和班级(class)"""
    rows = execute_query(
        "SELECT grade, class FROM users WHERE username = ?",
        (username,),
    )
    if rows and rows[0]:
        return str(rows[0][0] or ""), str(rows[0][1] or "")
    return "", ""


# ── 请求模型 ──

class QuizCreate(BaseModel):
    title: str
    description: str = ""
    questions: str  # JSON 字符串 [{type, question, options, answer, score}]


class QuizAnswerSubmit(BaseModel):
    answers: str  # JSON 字符串 [{question_index, answer}]


class PollCreate(BaseModel):
    question: str
    options: list[str]
    poll_type: str = "single"  # single / multiple


class QuestionCreate(BaseModel):
    content: str
    is_anonymous: bool = False


class QuestionAnswer(BaseModel):
    answer: str


class AiGenerateQuiz(BaseModel):
    topic: str
    subject: str = ""  # 由前端传递
    count: int = 1
    question_type: str = "single"  # single / true_false / mixed


class AiGeneratePoll(BaseModel):
    topic: str
    option_count: int = 4


# ── AI 辅助函数 ──

def _call_ai(prompt: str) -> str:
    """调用 AI（非流式）- 支持智能体/直接调大模型双模式"""
    import os
    api_key = os.environ.get("DASHSCOPE_API_KEY", "")
    if not api_key:
        try:
            from backend.api.config_router import load_config
            cfg = load_config()
            api_key = cfg.get("dashscope_api_key", "")
        except Exception:
            pass
    if not api_key:
        return "⚠️ AI 功能不可用：请配置 DashScope API Key"

    from backend.api.ai_service import call_ai_sync
    try:
        return call_ai_sync(prompt, api_key)
    except Exception as e:
        return f"AI 调用出错: {str(e)}"


# ── AI 生成随堂测验 ──

@router.post("/quizzes/ai-generate", summary="AI 自动生成随堂测验")
async def ai_generate_quiz(req: AiGenerateQuiz, request: Request):
    """AI 根据主题自动生成测验题目"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可用")

    type_desc = {
        "single": "单选题",
        "true_false": "判断题",
        "mixed": "混合（单选+判断）",
    }.get(req.question_type, "单选题")

    from backend.prompts.quiz import QUIZ_GENERATE_PROMPT
    prompt = QUIZ_GENERATE_PROMPT.format(
        subject=req.subject,
        topic=req.topic,
        type_desc=type_desc,
        count=req.count,
    )

    result = _call_ai(prompt)
    # 尝试从返回中提取 JSON
    import re
    json_match = re.search(r'\[[\s\S]*\]', result)
    if json_match:
        try:
            questions = json.loads(json_match.group())
            # 按请求数量截取，并统一字段名
            questions = questions[:req.count]
            for q in questions:
                if "svg_code" in q and "svg_content" not in q:
                    q["svg_content"] = q["svg_code"]
                if "has_svg" not in q:
                    q["has_svg"] = 1 if q.get("svg_code") else 0
            return {"questions": questions, "raw": result}
        except json.JSONDecodeError:
            pass
    return {"questions": [], "raw": result, "error": "AI 返回格式异常，请重试或手动输入"}


# ── AI 生成快速投票 ──

@router.post("/polls/ai-generate", summary="AI 自动生成投票")
async def ai_generate_poll(req: AiGeneratePoll, request: Request):
    """AI 根据主题自动生成投票问题和选项"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可用")

    prompt = (
        '请根据主题"' + req.topic + '"生成一个课堂投票，包含' + str(req.option_count) + '个选项。\n\n'
        '请严格按照以下JSON格式输出，不要包含任何其他文字：\n'
        '{\n'
        '  "question": "投票问题",\n'
        '  "options": ["选项1", "选项2", ...]\n'
        '}'
    )

    result = _call_ai(prompt)
    import re
    json_match = re.search(r'\{[\s\S]*\}', result)
    if json_match:
        try:
            data = json.loads(json_match.group())
            return {"poll": data, "raw": result}
        except json.JSONDecodeError:
            pass
    return {"poll": None, "raw": result, "error": "AI 返回格式异常，请重试或手动输入"}


# ── AI 建议回答提问 ──

@router.post("/questions/{question_id}/ai-suggest", summary="AI 建议回答")
async def ai_suggest_answer(question_id: int, request: Request):
    """AI suggests answer based on question content (teacher can modify before submit)"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师可用")

    rows = execute_query(
        "SELECT content FROM interaction_questions WHERE id = ?",
        (question_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="问题不存在")

    content = rows[0][0]
    prompt = (
        '你是一位高中教师。学生在课堂上提出了以下问题，请给出专业、清晰的回答。\n\n'
        '学生问题：' + content + '\n\n'
        '请用中文回答，语气亲切，条理清晰。'
    )

    answer = _call_ai(prompt)
    return {"suggested_answer": answer, "question": content}


class QuizUpdate(BaseModel):
    title: str | None = None
    description: str | None = None
    questions: str | None = None
    status: str | None = None


class PollUpdate(BaseModel):
    question: str | None = None
    options: list[str] | None = None
    poll_type: str | None = None


class QuestionUpdate(BaseModel):
    content: str | None = None


# ── 随堂测验 ──

@router.post("/quizzes", summary="创建随堂测验")
async def create_quiz(req: QuizCreate, request: Request):
    """教师创建随堂测验"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可创建测验")

    # 验证 JSON
    try:
        questions = json.loads(req.questions) if isinstance(req.questions, str) else req.questions
        if not isinstance(questions, list) or len(questions) == 0:
            raise ValueError("试题不能为空")
    except (json.JSONDecodeError, ValueError) as e:
        raise HTTPException(status_code=400, detail=f"试题格式错误: {e}")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    quiz_id = execute_insert_update(
        """INSERT INTO interaction_quizzes
           (creator_username, title, description, questions, status, created_at, updated_at)
           VALUES (?, ?, ?, ?, 'active', ?, ?)""",
        (user["username"], req.title, req.description, json.dumps(questions, ensure_ascii=False), now, now),
    )

    logger.info(f"用户 {user['username']} 创建随堂测验: {req.title} (id={quiz_id})")

    # ── 异步通知学生（不阻塞创建操作） ──
    async def _notify_quiz():
        try:
            from backend.api.notification_router import _notify_users
            from backend.database import execute_query as db_query
            creator = user["username"]
            role_u = user.get("role", 2)
            if role_u == 0:
                students = db_query("SELECT username FROM users WHERE role = 2")
            else:
                grade, cls = _get_user_grade_class(creator)
                if grade:
                    students = db_query(
                        f"SELECT username FROM users WHERE role = 2 AND grade = ?"
                        + (" AND INSTR(',' || ? || ',', ',' || class || ',') > 0" if cls else ""),
                        (grade, cls) if cls else (grade,),
                    )
                else:
                    students = []
            if students:
                _notify_users(
                    [r[0] for r in students], "info",
                    f"新随堂测验「{req.title}」已发布",
                    f"共 {len(questions)} 题，请及时完成",
                    "/interaction",
                )
        except Exception as e:
            logger.warning(f"发送测验通知失败: {e}")
    asyncio.create_task(_notify_quiz())

    return {"message": "测验创建成功", "quiz_id": quiz_id}


def _can_manage_quiz(username: str, role: int, quiz_id: int) -> bool:
    """检查是否有管理测验的权限"""
    if role == 0:
        return True
    rows = execute_query(
        "SELECT creator_username FROM interaction_quizzes WHERE id = ?",
        (quiz_id,),
    )
    return bool(rows and rows[0][0] == username)


@router.put("/quizzes/{quiz_id}", summary="更新测验")
async def update_quiz(quiz_id: int, req: QuizUpdate, request: Request):
    """编辑随堂测验"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="权限不足")
    if not _can_manage_quiz(user["username"], role, quiz_id):
        raise HTTPException(status_code=403, detail="无权修改此测验")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    updates = ["updated_at = ?"]
    params = [now]
    for field in ("title", "description", "status"):
        val = getattr(req, field, None)
        if val is not None:
            updates.append(f"{field} = ?")
            params.append(val)
    if req.questions is not None:
        try:
            json.loads(req.questions) if isinstance(req.questions, str) else req.questions
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="试题格式错误")
        updates.append("questions = ?")
        params.append(req.questions if isinstance(req.questions, str) else json.dumps(req.questions, ensure_ascii=False))

    params.append(quiz_id)
    execute_insert_update(
        f"UPDATE interaction_quizzes SET {', '.join(updates)} WHERE id = ?",
        tuple(params),
    )
    return {"message": "测验已更新"}


@router.delete("/quizzes/{quiz_id}", summary="删除测验")
async def delete_quiz(quiz_id: int, request: Request):
    """删除随堂测验及其答案"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="权限不足")
    if not _can_manage_quiz(user["username"], role, quiz_id):
        raise HTTPException(status_code=403, detail="无权删除此测验")

    execute_insert_update("DELETE FROM interaction_quiz_answers WHERE quiz_id = ?", (quiz_id,))
    execute_insert_update("DELETE FROM interaction_quizzes WHERE id = ?", (quiz_id,))
    return {"message": "测验已删除"}


@router.delete("/quizzes/{quiz_id}/questions/{question_index}", summary="删除测验中的单道题")
async def delete_quiz_question(quiz_id: int, question_index: int, request: Request):
    """删除随堂测验中的单道题，同时清理该题对应的答题记录"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="权限不足")
    if not _can_manage_quiz(user["username"], role, quiz_id):
        raise HTTPException(status_code=403, detail="无权修改此测验")

    # 读出当前测验
    rows = execute_query("SELECT questions FROM interaction_quizzes WHERE id = ?", (quiz_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="测验不存在")
    questions_str = rows[0][0]
    questions = json.loads(questions_str) if isinstance(questions_str, str) else questions_str

    if question_index < 0 or question_index >= len(questions):
        raise HTTPException(status_code=400, detail="题目索引超出范围")

    # 从数组中移除指定题目
    questions.pop(question_index)
    new_questions_str = json.dumps(questions, ensure_ascii=False)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    execute_insert_update(
        "UPDATE interaction_quizzes SET questions = ?, updated_at = ? WHERE id = ?",
        (new_questions_str, now, quiz_id),
    )

    # 清理答题记录：移除该题对应的答案，并重新索引后续题目
    answer_rows = execute_query(
        "SELECT id, answers FROM interaction_quiz_answers WHERE quiz_id = ?",
        (quiz_id,),
    )
    for ans_row in answer_rows:
        ans_id = ans_row[0]
        ans_str = ans_row[1]
        try:
            user_answers = json.loads(ans_str) if isinstance(ans_str, str) else ans_str
            if not isinstance(user_answers, list):
                continue
            # 移除 question_index 对应的答案，并调整后续索引
            new_answers = []
            for a in user_answers:
                qi = a.get("question_index")
                if qi == question_index:
                    continue  # 跳过被删除题的答案
                if qi is not None and qi > question_index:
                    a["question_index"] = qi - 1  # 索引前移
                new_answers.append(a)
            execute_insert_update(
                "UPDATE interaction_quiz_answers SET answers = ? WHERE id = ?",
                (json.dumps(new_answers, ensure_ascii=False), ans_id),
            )
        except (json.JSONDecodeError, TypeError):
            continue

    return {"message": "题目已删除", "questions_remaining": len(questions)}


def _can_manage_poll(username: str, role: int, poll_id: int) -> bool:
    if role == 0:
        return True
    rows = execute_query("SELECT creator_username FROM interaction_polls WHERE id = ?", (poll_id,))
    return bool(rows and rows[0][0] == username)


@router.put("/polls/{poll_id}", summary="更新投票")
async def update_poll(poll_id: int, req: PollUpdate, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="权限不足")
    if not _can_manage_poll(user["username"], role, poll_id):
        raise HTTPException(status_code=403, detail="无权修改")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if req.question is not None:
        execute_insert_update("UPDATE interaction_polls SET question = ?, created_at = ? WHERE id = ?",
                              (req.question, now, poll_id))
    if req.options is not None:
        execute_insert_update("UPDATE interaction_polls SET options = ?, created_at = ? WHERE id = ?",
                              (json.dumps(req.options, ensure_ascii=False), now, poll_id))
    if req.poll_type is not None:
        execute_insert_update("UPDATE interaction_polls SET poll_type = ?, created_at = ? WHERE id = ?",
                              (req.poll_type, now, poll_id))
    return {"message": "投票已更新"}


@router.delete("/polls/{poll_id}", summary="删除投票")
async def delete_poll(poll_id: int, request: Request):
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="权限不足")
    if not _can_manage_poll(user["username"], role, poll_id):
        raise HTTPException(status_code=403, detail="无权删除")

    execute_insert_update("DELETE FROM interaction_poll_votes WHERE poll_id = ?", (poll_id,))
    execute_insert_update("DELETE FROM interaction_polls WHERE id = ?", (poll_id,))
    return {"message": "投票已删除"}


@router.delete("/questions/{question_id}", summary="删除提问")
async def delete_question(question_id: int, request: Request):
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 0:
        # 管理员：删除任何提问
        execute_insert_update("DELETE FROM interaction_questions WHERE id = ?", (question_id,))
    elif role == 1:
        # 教师：只能删除自己班级学生的提问
        grade, cls = _get_user_grade_class(username)
        if cls:
            deleted = execute_insert_update(
                """DELETE FROM interaction_questions WHERE id = ? AND student_username IN (
                    SELECT username FROM users WHERE role = 2 AND grade = ? AND INSTR(',' || ? || ',', ',' || class || ',') > 0
                )""",
                (question_id, grade, cls),
            )
        else:
            deleted = execute_insert_update(
                """DELETE FROM interaction_questions WHERE id = ? AND student_username IN (
                    SELECT username FROM users WHERE role = 2 AND grade = ?
                )""",
                (question_id, grade),
            )
    else:
        # 学生：只能删除自己的提问
        execute_insert_update(
            "DELETE FROM interaction_questions WHERE id = ? AND student_username = ?",
            (question_id, username),
        )
    return {"message": "提问已删除"}


@router.put("/questions/{question_id}", summary="更新提问")
async def update_question(question_id: int, req: QuestionUpdate, request: Request):
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role == 2:
        rows = execute_query(
            "SELECT id FROM interaction_questions WHERE id = ? AND student_username = ?",
            (question_id, username),
        )
        if not rows:
            raise HTTPException(status_code=403, detail="无权修改")
    if req.content is not None:
        execute_insert_update("UPDATE interaction_questions SET content = ? WHERE id = ?",
                              (req.content, question_id))
    return {"message": "提问已更新"}


@router.get("/quizzes", summary="获取随堂测验列表")
async def list_quizzes(
    request: Request,
    status: str = Query("", description="筛选状态"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1),
):
    """获取随堂测验列表"""
    user = get_current_user(request)
    role = user.get("role", 2)
    username = user["username"]

    conditions = []
    params: list = []

    if role == 2:
        # 学生：看自己班级的测验（管理员创建的全体可见，教师创建的需匹配班级）
        grade, cls = _get_user_grade_class(username)
        conditions.append("q.status = 'active'")
        if grade:
            conditions.append("(u.role = 0 OR u.grade = ?)")
            params.append(grade)
        if cls:
            # 教师 class 可能是 "1" 或 "1,2,3,4"，用 INSTR 匹配
            cls_param = f",{cls},"
            conditions.append("(u.role = 0 OR INSTR(',' || u.class || ',', ?) > 0)")
            params.append(cls_param)
    elif role == 1:
        conditions.append("q.creator_username = ?")
        params.append(username)
        if status:
            conditions.append("q.status = ?")
            params.append(status)
    else:
        if status:
            conditions.append("q.status = ?")
            params.append(status)

    where = " AND ".join(conditions) if conditions else "1=1"
    offset = (page - 1) * page_size

    if role == 2:
        # 学生：JOIN users 表筛选同班教师的测验（包括管理员创建的内容）
        rows = execute_query(
            f"""SELECT q.id, q.creator_username, q.title, q.description, q.questions, q.status, q.created_at, q.updated_at,
                        COALESCE(u.name, u.username) AS creator_name
                FROM interaction_quizzes q
                JOIN users u ON q.creator_username = u.username AND u.role IN (0, 1)
                WHERE {where}
                ORDER BY q.created_at DESC LIMIT ? OFFSET ?""",
            tuple(params + [page_size, offset]),
        )
    else:
        rows = execute_query(
            f"""SELECT q.id, q.creator_username, q.title, q.description, q.questions, q.status, q.created_at, q.updated_at,
                        COALESCE(u.name, u.username) AS creator_name
                FROM interaction_quizzes q
                LEFT JOIN users u ON q.creator_username = u.username
                WHERE {where}
                ORDER BY q.created_at DESC LIMIT ? OFFSET ?""",
            tuple(params + [page_size, offset]),
        )

    quizzes = []
    for r in rows:
        q = {
            "id": r[0],
            "creator_username": r[1],
            "title": r[2],
            "description": r[3],
            "questions": json.loads(r[4]) if isinstance(r[4], str) else r[4],
            "status": r[5],
            "created_at": r[6],
            "updated_at": r[7],
            "creator_name": r[8] if len(r) > 8 else r[1],
        }
        # 统一字段名：svg_code → svg_content
        questions = q.get("questions", [])
        if isinstance(questions, list):
            for qs in questions:
                if isinstance(qs, dict) and "svg_code" in qs and "svg_content" not in qs:
                    qs["svg_content"] = qs["svg_code"]
        # 计算答题人数
        count_rows = execute_query(
            "SELECT COUNT(*) FROM interaction_quiz_answers WHERE quiz_id = ?",
            (r[0],),
        )
        q["answer_count"] = count_rows[0][0] if count_rows else 0
        # 学生端标记是否已答题
        if role == 2:
            answered = execute_query(
                "SELECT COUNT(*) FROM interaction_quiz_answers WHERE quiz_id = ? AND student_username = ?",
                (r[0], username),
            )
            q["answered"] = (answered[0][0] if answered else 0) > 0
        quizzes.append(q)

    return {"quizzes": quizzes}


@router.post("/quizzes/{quiz_id}/answer", summary="提交测验答案")
async def submit_quiz_answer(quiz_id: int, req: QuizAnswerSubmit, request: Request):
    """学生提交随堂测验答案"""
    user = get_current_user(request)
    username = user["username"]

    # 验证测验存在且活跃
    role = user.get("role", 2)
    quiz = execute_query(
        "SELECT * FROM interaction_quizzes WHERE id = ? AND status = 'active'",
        (quiz_id,),
    )
    if not quiz:
        raise HTTPException(status_code=404, detail="测验不存在或已结束")

    # 学生只能回答自己班级教师或管理员创建的测验
    if role == 2:
        grade, cls = _get_user_grade_class(username)
        creator = quiz[0][1]
        # 查询创建者是否为管理员（role=0）
        creator_info = execute_query("SELECT role, grade, class FROM users WHERE username=?", (creator,))
        if creator_info:
            creator_role = creator_info[0][0]
            creator_cgrade = creator_info[0][1] or ''
            creator_cclass = creator_info[0][2] or ''
        else:
            creator_role = 2
            creator_cgrade = ''
            creator_cclass = ''

        is_admin_creator = creator_role == 0
        if is_admin_creator:
            teacher_ok = [creator]
        elif cls and creator_cgrade == grade and (not creator_cclass or f",{cls}," in f",{creator_cclass},"):
            teacher_ok = [creator]
        elif not cls and creator_cgrade == grade:
            teacher_ok = [creator]
        else:
            teacher_ok = []
        if not teacher_ok:
            raise HTTPException(status_code=403, detail="无权回答非本班教师的测验")

    # 验证是否已答过
    existing = execute_query(
        "SELECT id FROM interaction_quiz_answers WHERE quiz_id = ? AND student_username = ?",
        (quiz_id, username),
    )
    if existing:
        raise HTTPException(status_code=400, detail="您已经答过此题")

    # 解析答案并评分
    questions = json.loads(quiz[0][4]) if isinstance(quiz[0][4], str) else quiz[0][4]
    user_answers = json.loads(req.answers) if isinstance(req.answers, str) else req.answers

    total_score = 0
    q_score = sum(q.get("score", 1) for q in questions)

    # 分离简答题和其他题
    short_indices = [i for i, q in enumerate(questions) if q.get("type") in ("short", "fill")]
    other_indices = [i for i, q in enumerate(questions) if q.get("type") not in ("short", "fill")]

    # 非简答题：精确匹配
    for idx in other_indices:
        q_type = questions[idx].get("type", "single")
        user_ans = str(next((ua.get("answer", "") for ua in user_answers if ua.get("question_index") == idx), "")).strip()
        correct_ans = str(questions[idx].get("answer", "")).strip()
        if q_type == "multiple":
            user_set = sorted([a.strip().upper() for a in user_ans.split(",") if a.strip()])
            correct_set = sorted([a.strip().upper() for a in correct_ans.split(",") if a.strip()])
            if user_set == correct_set:
                total_score += questions[idx].get("score", 1)
        else:
            if user_ans.upper() == correct_ans.upper():
                total_score += questions[idx].get("score", 1)

    # 简答题：AI 语义批改
    if short_indices:
        try:
            from backend.api.chat_router import get_api_keys
            from backend.api.ai_service import call_ai_async
            api_key, _ = get_api_keys(username)
        except Exception:
            api_key = ""

        if api_key and api_key.strip():
            sem = asyncio.Semaphore(3)

            async def _grade_short(idx):
                q = questions[idx]
                user_ans = str(next((ua.get("answer", "") for ua in user_answers if ua.get("question_index") == idx), "")).strip()
                correct_ans = str(q.get("answer", "")).strip()
                q_score_val = q.get("score", 1)
                async with sem:
                    try:
                        from backend.prompts.teaching import SHORT_ANSWER_GRADING_PROMPT
                        prompt = SHORT_ANSWER_GRADING_PROMPT.format(
                            question_text=str(q.get("question", "")).replace('{', '{{').replace('}', '}}'),
                            correct_answer=correct_ans.replace('{', '{{').replace('}', '}}'),
                            max_score=str(q_score_val),
                            half_score=str(q_score_val * 0.5),
                            near_full=str(q_score_val * 0.8),
                            half_minus=str(q_score_val * 0.4),
                            student_answer=user_ans.replace('{', '{{').replace('}', '}}'),
                        )
                        ai_resp = await call_ai_async(prompt, api_key)
                        jm = re.search(r'\{[^}]+\}', ai_resp)
                        if jm:
                            result = json.loads(jm.group())
                            ai_score = float(result.get("score", 0))
                            ai_score = max(0, min(ai_score, q_score_val))
                            return ai_score if ai_score >= q_score_val * 0.6 else 0.0
                    except Exception:
                        pass
                # AI 失败回退：关键词匹配
                keywords = [k.strip().lower() for k in correct_ans.replace("，", ",").split(",") if k.strip()]
                return q_score_val if keywords and any(kw in user_ans.lower() for kw in keywords) else 0.0

            results = await asyncio.gather(*[_grade_short(idx) for idx in short_indices])
            total_score += sum(results)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    execute_insert_update(
        """INSERT INTO interaction_quiz_answers (quiz_id, student_username, answers, score, submitted_at)
           VALUES (?, ?, ?, ?, ?)""",
        (quiz_id, username, json.dumps(user_answers, ensure_ascii=False), total_score, now),
    )

    # ── 积分奖励 ──
    try:
        from backend.reward_engine import award_participation, award_grade
        quiz_title = quiz[0][2] if len(quiz[0]) > 2 else f"测验#{quiz_id}"
        award_participation(username, "quiz", str(quiz_id), quiz_title)
        award_grade(username, "quiz", str(quiz_id), total_score, q_score, quiz_title)
    except Exception:
        pass

    return {
        "message": "提交成功",
        "score": total_score,
        "total_score": q_score,
        "percentage": round(total_score / max(q_score, 1) * 100, 1),
    }


@router.get("/quizzes/{quiz_id}/my-result", summary="学生查看自己的答题结果")
async def get_my_quiz_result(quiz_id: int, request: Request):
    """学生查看自己的随堂测验答题详情"""
    user = get_current_user(request)
    username = user["username"]
    role = user.get("role", 2)
    if role != 2:
        raise HTTPException(status_code=403, detail="仅学生可查看")

    quiz = execute_query(
        "SELECT * FROM interaction_quizzes WHERE id = ?",
        (quiz_id,),
    )
    if not quiz:
        raise HTTPException(status_code=404, detail="测验不存在")

    # 查询该学生的答题记录
    answers = execute_query(
        "SELECT answers, score, submitted_at FROM interaction_quiz_answers WHERE quiz_id = ? AND student_username = ?",
        (quiz_id, username),
    )
    if not answers:
        raise HTTPException(status_code=404, detail="未找到答题记录")

    questions = json.loads(quiz[0][4]) if isinstance(quiz[0][4], str) else quiz[0][4]
    user_answers = json.loads(answers[0][0]) if isinstance(answers[0][0], str) else answers[0][0]
    total_score = answers[0][1]
    q_score = sum(q.get("score", 1) for q in questions)

    # 每题批改详情
    details = []
    for i, q in enumerate(questions):
        user_ans = ""
        for ua in user_answers:
            if ua.get("question_index") == i:
                user_ans = str(ua.get("answer", "")).strip()
                break
        correct_ans = str(q.get("answer", "")).strip()
        q_type = q.get("type", "single")
        is_correct = False
        if q_type == "multiple":
            user_set = sorted([a.strip().upper() for a in user_ans.split(",") if a.strip()])
            correct_set = sorted([a.strip().upper() for a in correct_ans.split(",") if a.strip()])
            is_correct = user_set == correct_set
        else:
            is_correct = user_ans.upper() == correct_ans.upper()
        details.append({
            "index": i,
            "question": q.get("question", q.get("question_text", "")),
            "options": q.get("options", None),
            "type": q_type,
            "user_answer": user_ans,
            "correct_answer": correct_ans,
            "is_correct": is_correct,
            "score": q.get("score", 1) if is_correct else 0,
            "max_score": q.get("score", 1),
            "explanation": q.get("explanation", ""),
            "svg_content": q.get("svg_content") or q.get("svg_code", ""),
            "has_svg": q.get("has_svg", 1 if q.get("svg_code") or q.get("svg_content") else 0),
            "media_files": q.get("media_files", ""),
        })

    return {
        "quiz_title": quiz[0][2],
        "score": total_score,
        "total_score": q_score,
        "percentage": round(total_score / max(q_score, 1) * 100, 1),
        "details": details,
    }


def _calc_student_correct_count(answer_row: tuple, questions: list) -> int:
    """计算学生的答对题数"""
    try:
        user_answers = json.loads(answer_row[1]) if isinstance(answer_row[1], str) else answer_row[1]
        correct = 0
        for i, q in enumerate(questions):
            q_type = q.get("type", "single")
            correct_ans = str(q.get("answer", "")).strip()
            for ua in user_answers:
                if ua.get("question_index") == i:
                    user_ans = str(ua.get("answer", "")).strip()
                    if q_type == "multiple":
                        us = sorted([x.strip().upper() for x in user_ans.split(",") if x.strip()])
                        cs = sorted([x.strip().upper() for x in correct_ans.split(",") if x.strip()])
                        if us == cs:
                            correct += 1
                    else:
                        if user_ans.upper() == correct_ans.upper():
                            correct += 1
                    break
        return correct
    except Exception:
        return 0


@router.get("/quizzes/{quiz_id}/results", summary="教师查看测验结果统计")
async def get_quiz_results(quiz_id: int, request: Request):
    """教师查看随堂测验结果统计"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="无权查看")

    quiz = execute_query(
        "SELECT * FROM interaction_quizzes WHERE id = ?",
        (quiz_id,),
    )
    if not quiz:
        raise HTTPException(status_code=404, detail="测验不存在")

    questions = json.loads(quiz[0][4]) if isinstance(quiz[0][4], str) else quiz[0][4]
    answers = execute_query(
        "SELECT student_username, answers, score, submitted_at FROM interaction_quiz_answers WHERE quiz_id = ?",
        (quiz_id,),
    )

    # 每题正确率统计
    question_stats = []
    for i, q in enumerate(questions):
        correct_count = 0
        q_type = q.get("type", "single")
        correct_ans = str(q.get("answer", "")).strip()
        for a in answers:
            user_ans = json.loads(a[1]) if isinstance(a[1], str) else a[1]
            for ua in user_ans:
                if ua.get("question_index") == i:
                    user_ans_str = str(ua.get("answer", "")).strip()
                    if q_type == "multiple":
                        user_set = sorted([x.strip().upper() for x in user_ans_str.split(",") if x.strip()])
                        correct_set = sorted([x.strip().upper() for x in correct_ans.split(",") if x.strip()])
                        if user_set == correct_set:
                            correct_count += 1
                    else:
                        if user_ans_str.upper() == correct_ans.upper():
                            correct_count += 1
                    break
        question_stats.append({
            "index": i,
            "question": q.get("question", q.get("question_text", "")),
            "options": q.get("options", []),
            "correct_answer": q.get("answer", ""),
            "type": q.get("type", "single"),
            "svg_content": q.get("svg_content") or q.get("svg_code", ""),
            "has_svg": q.get("has_svg", 1 if q.get("svg_code") or q.get("svg_content") else 0),
            "media_placeholders": q.get("media_placeholders", []),
            "correct_count": correct_count,
            "total_count": len(answers),
            "correct_rate": round(correct_count / max(len(answers), 1) * 100, 1),
        })

    # 建立学生用户名→姓名映射（批量查询）
    usernames = [a[0] for a in answers]
    name_map = {}
    if usernames:
        placeholders = ",".join("?" * len(usernames))
        name_rows = execute_query(
            f"SELECT username, name FROM users WHERE username IN ({placeholders})",
            tuple(usernames),
        )
        name_map = {r[0]: r[1] or r[0] for r in name_rows}

    return {
        "quiz": {
            "id": quiz[0][0],
            "title": quiz[0][2],
            "question_count": len(questions),
        },
        "total_answers": len(answers),
        "question_stats": question_stats,
        "student_answers": [
            {
                "student": a[0],
                "student_name": name_map.get(a[0], a[0]),
                "score": a[2],
                "submitted_at": a[3],
                "correct_count": _calc_student_correct_count(a, questions),
                "total_questions": len(questions),
            }
            for a in answers
        ],
    }


# ── 快速投票 ──

@router.post("/polls", summary="创建快速投票")
async def create_poll(req: PollCreate, request: Request):
    """教师创建快速投票"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role not in (0, 1):
        raise HTTPException(status_code=403, detail="仅教师和管理员可创建投票")

    if len(req.options) < 2:
        raise HTTPException(status_code=400, detail="至少需要2个选项")

    if req.poll_type not in ("single", "multiple"):
        raise HTTPException(status_code=400, detail="poll_type 必须是 single 或 multiple")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    poll_id = execute_insert_update(
        """INSERT INTO interaction_polls (creator_username, question, options, poll_type, status, created_at)
           VALUES (?, ?, ?, ?, 'active', ?)""",
        (user["username"], req.question, json.dumps(req.options, ensure_ascii=False), req.poll_type, now),
    )

    # ── 异步通知学生（不阻塞创建操作） ──
    async def _notify_poll():
        try:
            from backend.api.notification_router import _notify_users
            from backend.database import execute_query as db_query
            creator = user["username"]
            role_u = user.get("role", 2)
            if role_u == 0:
                students = db_query("SELECT username FROM users WHERE role = 2")
            else:
                grade, cls = _get_user_grade_class(creator)
                if grade:
                    students = db_query(
                        f"SELECT username FROM users WHERE role = 2 AND grade = ?"
                        + (" AND INSTR(',' || ? || ',', ',' || class || ',') > 0" if cls else ""),
                        (grade, cls) if cls else (grade,),
                    )
                else:
                    students = []
            if students:
                _notify_users(
                    [r[0] for r in students], "info",
                    f"新投票「{req.question}」已发布",
                    f"共 {len(req.options)} 个选项，请参与投票",
                    "/interaction",
                )
        except Exception as e:
            logger.warning(f"发送投票通知失败: {e}")
    asyncio.create_task(_notify_poll())

    return {"message": "投票创建成功", "poll_id": poll_id}


@router.get("/polls", summary="获取活跃投票")
async def list_polls(request: Request):
    """获取当前活跃的投票列表"""
    user = get_current_user(request)
    role = user.get("role", 2)
    username = user["username"]

    if role == 2:
        # 学生：看自己班级的投票（管理员创建的全体可见，教师创建的需匹配班级）
        grade, cls = _get_user_grade_class(username)
        conditions = ["p.status = 'active'"]
        params: list = []
        if grade:
            conditions.append("(u.role = 0 OR u.grade = ?)")
            params.append(grade)
        if cls:
            cls_param = f",{cls},"
            conditions.append("(u.role = 0 OR INSTR(',' || u.class || ',', ?) > 0)")
            params.append(cls_param)
        where = " AND ".join(conditions)
        rows = execute_query(
            f"""SELECT p.id, p.creator_username, p.question, p.options, p.poll_type, p.status, p.created_at,
                        COALESCE(u.name, u.username) AS creator_name
                FROM interaction_polls p
                JOIN users u ON p.creator_username = u.username AND u.role IN (0, 1)
                WHERE {where}
                ORDER BY p.created_at DESC LIMIT 50""",
            tuple(params),
        )
    elif role == 1:
        # 教师：只看自己创建的投票
        rows = execute_query(
            """SELECT p.id, p.creator_username, p.question, p.options, p.poll_type, p.status, p.created_at,
                        COALESCE(u.name, u.username) AS creator_name
               FROM interaction_polls p
               LEFT JOIN users u ON p.creator_username = u.username
               WHERE p.status = 'active' AND p.creator_username = ?
               ORDER BY p.created_at DESC LIMIT 50""",
            (username,),
        )
    else:
        rows = execute_query(
            """SELECT p.id, p.creator_username, p.question, p.options, p.poll_type, p.status, p.created_at,
                        COALESCE(u.name, u.username) AS creator_name
               FROM interaction_polls p
               LEFT JOIN users u ON p.creator_username = u.username
               WHERE p.status = 'active'
               ORDER BY p.created_at DESC LIMIT 50""",
        )

    polls = []
    for r in rows:
        options = json.loads(r[3]) if isinstance(r[3], str) else r[3]
        poll_type = r[4] if r[4] else "single"
        creator_name = r[7] if len(r) > 7 else r[1]
        vote_counts = []
        for i in range(len(options)):
            cnt = execute_query(
                "SELECT COUNT(*) FROM interaction_poll_votes WHERE poll_id = ? AND selected_option = ?",
                (r[0], i),
            )
            vote_counts.append(cnt[0][0] if cnt else 0)

        # 对于多选投票，计算每人投了几项
        total_votes = sum(vote_counts)
        unique_voters = 0
        if poll_type == "multiple":
            voters = execute_query(
                "SELECT COUNT(DISTINCT student_username) FROM interaction_poll_votes WHERE poll_id = ?",
                (r[0],),
            )
            unique_voters = voters[0][0] if voters else 0

        # 学生端标记是否已投票
        voted = False
        if role == 2:
            voted_row = execute_query(
                "SELECT COUNT(*) FROM interaction_poll_votes WHERE poll_id = ? AND student_username = ?",
                (r[0], username),
            )
            voted = (voted_row[0][0] if voted_row else 0) > 0

        polls.append({
            "id": r[0],
            "creator_username": r[1],
            "creator_name": creator_name,
            "question": r[2],
            "poll_type": poll_type,
            "options": [{"index": i, "text": opt, "votes": vote_counts[i]} for i, opt in enumerate(options)],
            "total_votes": total_votes,
            "unique_voters": unique_voters or total_votes,
            "voted": voted,
            "created_at": r[6],
        })

    return {"polls": polls}


@router.post("/polls/{poll_id}/vote", summary="提交投票")
async def submit_vote(
    poll_id: int,
    request: Request,
    option_index: int = Query(None, description="单选：选项索引"),
    option_indices: str = Query(None, description="多选：逗号分隔的选项索引，如 '0,2,3'"),
):
    """学生参与投票，支持单选和多选"""
    user = get_current_user(request)
    username = user["username"]

    poll = execute_query(
        "SELECT * FROM interaction_polls WHERE id = ? AND status = 'active'",
        (poll_id,),
    )
    if not poll:
        raise HTTPException(status_code=404, detail="投票不存在或已结束")

    # 学生只能参与自己班级教师或管理员创建的投票
    role = user.get("role", 2)
    if role == 2:
        grade, cls = _get_user_grade_class(username)
        creator = poll[0][1]
        # 查询创建者是否为管理员（role=0），管理员创建的所有学生都可参与
        creator_info = execute_query("SELECT role, grade, class FROM users WHERE username=?", (creator,))
        if creator_info:
            creator_role = creator_info[0][0]
            creator_cgrade = creator_info[0][1] or ''
            creator_cclass = creator_info[0][2] or ''
        else:
            creator_role = 2  # 找不到则视为普通用户
            creator_cgrade = ''
            creator_cclass = ''

        is_admin_creator = creator_role == 0
        # 管理员创建 或 创建者的年级班级与学生匹配 → 允许
        if is_admin_creator:
            teacher_ok = [creator]
        elif cls and creator_cgrade == grade and (not creator_cclass or f",{cls}," in f",{creator_cclass},"):
            teacher_ok = [creator]
        elif not cls and creator_cgrade == grade:
            teacher_ok = [creator]
        else:
            teacher_ok = []
        if not teacher_ok:
            raise HTTPException(status_code=403, detail="无权参与非本班教师的投票")

    poll_type = poll[0][6] if len(poll[0]) > 6 and poll[0][6] else "single"
    options = json.loads(poll[0][3]) if isinstance(poll[0][3], str) else poll[0][3]

    # 解析选择的选项
    selected_indices = []
    if poll_type == "single":
        if option_index is None:
            raise HTTPException(status_code=400, detail="请选择一个选项")
        selected_indices = [option_index]
    else:
        if not option_indices:
            raise HTTPException(status_code=400, detail="请至少选择一个选项")
        try:
            selected_indices = [int(i.strip()) for i in option_indices.split(",") if i.strip()]
        except ValueError:
            raise HTTPException(status_code=400, detail="选项索引格式错误")
        if not selected_indices:
            raise HTTPException(status_code=400, detail="请至少选择一个选项")

    # 验证所有选项索引
    for idx in selected_indices:
        if idx < 0 or idx >= len(options):
            raise HTTPException(status_code=400, detail=f"无效的选项索引: {idx}")

    # 检查是否已投票
    existing = execute_query(
        "SELECT id FROM interaction_poll_votes WHERE poll_id = ? AND student_username = ?",
        (poll_id, username),
    )
    if existing:
        if poll_type == "single":
            raise HTTPException(status_code=400, detail="您已经投过票")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if poll_type == "multiple":
        # 多选：先删后插，所有操作在同一个事务中执行，避免数据库锁冲突
        ops = [
            ("DELETE FROM interaction_poll_votes WHERE poll_id = ? AND student_username = ?", (poll_id, username)),
        ]
        for idx in selected_indices:
            ops.append((
                "INSERT INTO interaction_poll_votes (poll_id, student_username, selected_option, created_at) VALUES (?, ?, ?, ?)",
                (poll_id, username, idx, now),
            ))
        execute_batch(ops)
    else:
        # 单选：直接插入一条记录
        execute_insert_update(
            "INSERT INTO interaction_poll_votes (poll_id, student_username, selected_option, created_at) VALUES (?, ?, ?, ?)",
            (poll_id, username, selected_indices[0], now),
        )

    # ── 积分奖励 ──
    try:
        from backend.reward_engine import award_participation
        poll_title = poll[0][2] if len(poll[0]) > 2 else f"投票#{poll_id}"
        award_participation(username, "poll", str(poll_id), poll_title)
    except Exception:
        pass

    return {"message": "投票成功", "selected_count": len(selected_indices)}


@router.get("/polls/{poll_id}/results", summary="查看投票结果")
async def get_poll_results(poll_id: int, request: Request):
    """获取投票实时结果"""
    poll = execute_query(
        "SELECT * FROM interaction_polls WHERE id = ?",
        (poll_id,),
    )
    if not poll:
        raise HTTPException(status_code=404, detail="投票不存在")

    options = json.loads(poll[0][3]) if isinstance(poll[0][3], str) else poll[0][3]
    vote_counts = []
    for i in range(len(options)):
        cnt = execute_query(
            "SELECT COUNT(*) FROM interaction_poll_votes WHERE poll_id = ? AND selected_option = ?",
            (poll_id, i),
        )
        vote_counts.append(cnt[0][0] if cnt else 0)

    total = sum(vote_counts)
    # 计算实际参与人数（去重）
    voters = execute_query(
        "SELECT COUNT(DISTINCT student_username) FROM interaction_poll_votes WHERE poll_id = ?",
        (poll_id,),
    )
    unique_voters = voters[0][0] if voters else 0
    poll_type = poll[0][6] if len(poll[0]) > 6 and poll[0][6] else "single"
    denominator = unique_voters if unique_voters > 0 else total
    return {
        "poll_id": poll_id,
        "question": poll[0][2],
        "poll_type": poll_type,
        "options": [
            {"index": i, "text": opt, "votes": vote_counts[i],
             "percentage": round(vote_counts[i] / max(denominator, 1) * 100, 1)}
            for i, opt in enumerate(options)
        ],
        "total_votes": total,
        "unique_voters": unique_voters,
    }


# ── AI 内容审核 + 重复/频率限制（用于提问和回答的合规检查） ──

import time as _time
import hashlib as _hashlib

# 内容审核跟踪：username -> { rejected_hashes: set, rejection_count: int, window_start: float, blocked_until: float }
_content_review_tracker: dict[str, dict] = {}
_MAX_REJECTIONS = 3           # 时间窗口内最大拒绝次数
_WINDOW_SECONDS = 300         # 时间窗口（5分钟）
_BLOCK_SECONDS = 60           # 触发限制后封禁时长


def _check_review_rate_limit(username: str, content: str) -> tuple[bool, str]:
    """检查是否触发频率限制或重复提交，返回 (是否放行, 提示消息)"""
    now = _time.time()
    tracker = _content_review_tracker.setdefault(username, {
        "rejected_hashes": set(),
        "rejection_count": 0,
        "window_start": now,
        "blocked_until": 0,
    })

    # 1. 检查是否被临时封禁
    if now < tracker["blocked_until"]:
        remain = int(tracker["blocked_until"] - now)
        return False, f"提交过于频繁，请{remain}秒后再试"

    # 2. 检查时间窗口是否过期，过期则重置
    if now - tracker["window_start"] > _WINDOW_SECONDS:
        tracker["rejected_hashes"] = set()
        tracker["rejection_count"] = 0
        tracker["window_start"] = now

    # 3. 检查是否重复提交已被拒绝过的内容
    content_hash = _hashlib.md5(content.encode("utf-8")).hexdigest()
    if content_hash in tracker["rejected_hashes"]:
        return False, "该内容与之前被拒绝的内容相同，请修改后重新提交"

    return True, ""


def _record_rejection(username: str, content: str):
    """记录一次审核拒绝并扣除 2 分"""
    now = _time.time()
    tracker = _content_review_tracker.setdefault(username, {
        "rejected_hashes": set(),
        "rejection_count": 0,
        "window_start": now,
        "blocked_until": 0,
    })
    content_hash = _hashlib.md5(content.encode("utf-8")).hexdigest()
    tracker["rejected_hashes"].add(content_hash)
    tracker["rejection_count"] += 1

    # 扣除 2 分
    try:
        from backend.reward_engine import deduct_points
        deduct_points(username, "内容审核不通过", 2)
    except Exception:
        pass

    # 如果累计拒绝次数达到上限，临时封禁
    if tracker["rejection_count"] >= _MAX_REJECTIONS:
        tracker["blocked_until"] = now + _BLOCK_SECONDS
        logger.warning(f"用户 {username} 触发审核频率限制，封禁 {_BLOCK_SECONDS} 秒")


def _ai_content_review(content: str, username: str, role: str = "question") -> tuple[bool, str]:
    """调用 AI 审核内容是否合规，含重复提交和频率限制，返回 (是否通过, 原因)"""
    # 先检查频率限制和重复提交
    allowed, msg = _check_review_rate_limit(username, content)
    if not allowed:
        return False, msg

    import json
    role_desc = "提问" if role == "question" else "回答"
    prompt = (
        '你是一个课堂内容审核助手。请判断以下学生' + role_desc + '是否包含：\n'
        '1. 违反法律法规的内容\n'
        '2. 违反道德规范、社会公序良俗的内容\n'
        '3. 不文明用语、辱骂、攻击性、歧视性言论\n'
        '4. 色情、暴力、恐怖、血腥等内容\n'
        '5. 广告、垃圾信息\n\n'
        '学生' + role_desc + '：' + content + '\n\n'
        '请严格判断。只返回以下JSON格式（不要包含其他文字）：\n'
        '{"safe": true, "reason": ""} 或 {"safe": false, "reason": "简要说明违规原因（10字以内）"}'
    )
    import os
    api_key = os.environ.get("DASHSCOPE_API_KEY", "")
    if not api_key:
        try:
            from backend.api.config_router import load_config
            cfg = load_config()
            api_key = cfg.get("dashscope_api_key", "")
        except Exception:
            pass
    if not api_key:
        # 无 API Key 时放行（不阻塞功能）
        return True, ""

    try:
        from backend.api.ai_service import call_ai_sync_direct
        # 使用 asyncio.wait_for 添加超时控制，防止 AI 调用挂死
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(call_ai_sync_direct, prompt, api_key)
            try:
                result = future.result(timeout=15)  # AI 审核最多等 15 秒
            except concurrent.futures.TimeoutError:
                logger.warning(f"AI 内容审核超时（15秒），已放行: user={username}")
                return True, ""
        jm = re.search(r'\{[^}]+\}', result)
        if jm:
            data = json.loads(jm.group())
            if not data.get("safe", True):
                # 记录拒绝
                _record_rejection(username, content)
                return False, data.get("reason", "内容不合规")
        return True, ""
    except Exception:
        # AI 调用失败时放行，不阻塞正常使用
        return True, ""


# ── 课堂提问 ──

@router.post("/questions", summary="学生发起提问")
async def ask_question(req: QuestionCreate, request: Request):
    """学生在课堂上提问"""
    user = get_current_user(request)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # ── 每日提问上限检查 ──
    role = user.get("role", 2)
    if role == 2:
        today_start = now[:10] + " 00:00:00"
        today_count = execute_query(
            "SELECT COUNT(*) FROM interaction_questions WHERE student_username = ? AND created_at >= ?",
            (user["username"], today_start),
        )
        if today_count and today_count[0][0] >= 5:
            raise HTTPException(
                status_code=400,
                detail="⚠️ 每日最多提 5 个问题，已达到今日上限。"
            )

    # ── AI 内容审核 ──
    safe, reason = _ai_content_review(req.content, user["username"], "question")
    if not safe:
        logger.info(f"学生 {user['username']} 提问内容不合规已拦截: {reason}")
        raise HTTPException(
            status_code=400,
            detail=f"⚠️ 内容审核未通过：{reason}。请修改后重新提交。"
        )

    qid = execute_insert_update(
        """INSERT INTO interaction_questions (student_username, content, is_anonymous, status, created_at)
           VALUES (?, ?, ?, 'pending', ?)""",
        (user["username"], req.content, 1 if req.is_anonymous else 0, now),
    )

    # ── 积分奖励 ──
    try:
        from backend.reward_engine import award_participation
        award_participation(user["username"], "question", str(qid), req.content[:30])
    except Exception:
        pass

    # ── 异步通知教师 ──
    async def _notify_question():
        try:
            from backend.api.notification_router import _notify_users
            from backend.database import execute_query as db_query
            student_grade, student_cls = _get_user_grade_class(user["username"])
            if student_grade:
                cls_param = f",{student_cls}," if student_cls else ""
                teachers = db_query(
                    f"SELECT username FROM users WHERE role = 1 AND grade = ?"
                    + (" AND INSTR(',' || class || ',', ?) > 0" if student_cls else ""),
                    (student_grade, cls_param) if student_cls else (student_grade,),
                )
                if teachers:
                    _notify_users(
                        [r[0] for r in teachers], "info",
                        f"新课堂提问",
                        f"学生提出了新问题：{req.content[:50]}{'...' if len(req.content) > 50 else ''}",
                        "/interaction",
                    )
        except Exception as e:
            logger.warning(f"发送提问通知失败: {e}")
    asyncio.create_task(_notify_question())

    return {
        "message": "提问成功",
        "question_id": qid,
        "question": {
            "id": qid,
            "content": req.content,
            "is_anonymous": bool(req.is_anonymous),
            "status": "pending",
            "answer": "",
            "created_at": now,
            "answered_at": "",
            "answered_by": "",
            "student_answer_count": 0,
            "approved_answer_count": 0,
            "my_answer_status": None,
            "is_own": True,
        }
    }


@router.get("/questions", summary="获取课堂提问列表")
async def list_questions(
    request: Request,
    status: str = Query("", description="筛选状态"),
    page: int = Query(1, ge=1, description="页码"),
    page_size: int = Query(20, ge=1, le=100, description="每页数量"),
):
    """获取课堂提问（教师看自己班级学生的提问，学生看所有）"""
    user = get_current_user(request)
    role = user.get("role", 2)
    username = user["username"]
    offset = (page - 1) * page_size

    if role == 2:
        # 学生：只看自己班级同学的提问和回答
        grade, cls = _get_user_grade_class(username)
        conditions = []
        params: list = []
        if grade:
            conditions.append("u.grade = ?")
            params.append(grade)
        if cls:
            cls_param = f",{cls},"
            conditions.append("INSTR(',' || u.class || ',', ?) > 0")
            params.append(cls_param)
        where = " AND ".join(conditions) if conditions else "1=1"
        rows = execute_query(
            f"""SELECT q.id, q.student_username, q.content, q.is_anonymous, q.status, q.answer, q.created_at, q.answered_at, q.answered_by
                FROM interaction_questions q
                JOIN users u ON q.student_username = u.username AND u.role = 2
                WHERE {where}
                ORDER BY q.created_at DESC LIMIT ? OFFSET ?""",
            tuple(params + [page_size, offset]),
        )
        # 获取总数
        count_rows = execute_query(
            f"""SELECT COUNT(*) FROM interaction_questions q
                JOIN users u ON q.student_username = u.username AND u.role = 2
                WHERE {where}""",
            tuple(params),
        )
        total = count_rows[0][0] if count_rows else 0
    elif role == 1:
        # 教师：只看自己班级学生的提问
        grade, cls = _get_user_grade_class(username)
        conditions = []
        params: list = []
        if grade:
            conditions.append("u.grade = ?")
            params.append(grade)
        if cls:
            conditions.append("INSTR(',' || ? || ',', ',' || u.class || ',') > 0")
            params.append(cls)  # 原始教师班级字符串，不额外加逗号
        if status:
            conditions.append("q.status = ?")
            params.append(status)
        where = " AND ".join(conditions) if conditions else "1=1"
        rows = execute_query(
            f"""SELECT q.id, q.student_username, q.content, q.is_anonymous, q.status, q.answer, q.created_at, q.answered_at, q.answered_by
                FROM interaction_questions q
                JOIN users u ON q.student_username = u.username AND u.role = 2
                WHERE {where}
                ORDER BY q.created_at DESC LIMIT ? OFFSET ?""",
            tuple(params + [page_size, offset]),
        )
        # 获取总数
        count_rows = execute_query(
            f"""SELECT COUNT(*) FROM interaction_questions q
                JOIN users u ON q.student_username = u.username AND u.role = 2
                WHERE {where}""",
            tuple(params),
        )
        total = count_rows[0][0] if count_rows else 0
    else:
        # 管理员：看全部
        conditions = []
        params: list = []
        if status:
            conditions.append("status = ?")
            params.append(status)
        where = " AND ".join(conditions) if conditions else "1=1"
        rows = execute_query(
            f"""SELECT id, student_username, content, is_anonymous, status, answer, created_at, answered_at, answered_by
                FROM interaction_questions WHERE {where}
                ORDER BY created_at DESC LIMIT ? OFFSET ?""",
            tuple(params + [page_size, offset]),
        )
        # 获取总数
        count_rows = execute_query(
            f"""SELECT COUNT(*) FROM interaction_questions WHERE {where}""",
            tuple(params),
        )
        total = count_rows[0][0] if count_rows else 0

    questions = []
    for r in rows:
        q = {
            "id": r[0],
            "content": r[2],
            "is_anonymous": bool(r[3]),
            "status": r[4],
            "answer": r[5] or "",
            "created_at": r[6],
            "answered_at": r[7] or "",
            "answered_by": r[8] or "",
        }
        if role != 2:
            q["student_username"] = r[1] if not r[3] else "(匿名)"
        else:
            q["is_own"] = (r[1] == username)

        # ── 附加多回答统计 ──
        qid = r[0]
        # 学生回答总数
        cnt = execute_query(
            "SELECT COUNT(*) FROM interaction_question_answers WHERE question_id = ?",
            (qid,),
        )
        q["student_answer_count"] = cnt[0][0] if cnt else 0

        # 已通过的回答数
        approved = execute_query(
            "SELECT COUNT(*) FROM interaction_question_answers WHERE question_id = ? AND status = 'approved'",
            (qid,),
        )
        q["approved_answer_count"] = approved[0][0] if approved else 0

        # 当前登录学生的回答状态
        q["my_answer_status"] = None
        if role == 2:
            my_ans = execute_query(
                "SELECT status FROM interaction_question_answers WHERE question_id = ? AND student_username = ?",
                (qid, username),
            )
            if my_ans:
                q["my_answer_status"] = my_ans[0][0]

        questions.append(q)

    return {"questions": questions, "total": total}


# ── 多学生回答：每个学生限答一次，教师按回答审批 ──

@router.put("/questions/{question_id}/answer", summary="回答提问（教师直接通过，学生写入多回答表）")
async def answer_question(question_id: int, req: QuestionAnswer, request: Request):
    """回答学生提问。教师直接写入主表；学生写入 interaction_question_answers（每人限答一次）。"""
    user = get_current_user(request)
    role = user.get("role", 2)
    username = user["username"]

    if not req.answer or not req.answer.strip():
        raise HTTPException(status_code=400, detail="回答内容不能为空")

    # 查询提问信息
    question_rows = execute_query(
        "SELECT student_username, content, status, answer, answered_by FROM interaction_questions WHERE id = ?",
        (question_id,),
    )
    if not question_rows:
        raise HTTPException(status_code=404, detail="提问不存在")
    asker_username = question_rows[0][0]
    question_content = question_rows[0][1]
    q_status = question_rows[0][2]
    q_answer = question_rows[0][3] or ""
    q_answered_by = question_rows[0][4] or ""

    # 学生：不能回答已标记为已答的问题（教师已答或已审批通过）
    if role == 2 and q_status == "answered":
        raise HTTPException(status_code=400, detail="该问题已被回答")

    # 教师/管理员：允许覆盖/更新自己的回答（与学生回答独立）
    if role != 2 and q_answer and q_answered_by != username:
        # 其他教师/管理员已回答过，询问是否覆盖
        pass  # 允许覆盖

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if role == 2:
        # ── 学生回答：写入多回答表 ──
        if username == asker_username:
            raise HTTPException(status_code=403, detail="不能回答自己的提问")

        # 班级验证
        grade, cls = _get_user_grade_class(username)
        asker_grade, asker_cls = _get_user_grade_class(asker_username)
        if grade != asker_grade:
            raise HTTPException(status_code=403, detail="只能回答同班同学的提问")
        if cls and asker_cls:
            if not (f",{cls}," in f",{asker_cls}," or f",{asker_cls}," in f",{cls},"):
                raise HTTPException(status_code=403, detail="只能回答同班同学的提问")

        # ── AI 内容审核 ──
        safe, reason = _ai_content_review(req.answer, username, "answer")
        if not safe:
            logger.info(f"学生 {username} 回答内容不合规已拦截: {reason}")
            raise HTTPException(
                status_code=400,
                detail=f"⚠️ 回答审核未通过：{reason}。请修改后重新提交。"
            )

        # INSERT OR REPLACE：每人限答一次
        execute_insert_update(
            """INSERT OR REPLACE INTO interaction_question_answers
               (question_id, student_username, answer, status, created_at)
               VALUES (?, ?, ?, 'pending_approval', ?)""",
            (question_id, username, req.answer, now),
        )

        # ── 积分奖励：参与回答 ──
        try:
            from backend.reward_engine import award_participation
            award_participation(username, "question", str(question_id),
                                f"回答：{question_content[:30]}...")
        except Exception:
            pass

        # ── 通知提问学生 ──
        try:
            from backend.api.notification_router import _create_notification
            _create_notification(asker_username, "info",
                "你的提问收到同学回答（待教师审批）",
                f"问题：{question_content[:50]}{'...' if len(question_content) > 50 else ''}",
                "/interaction")
        except Exception as e:
            logger.warning(f"发送回答通知失败: {e}")

        # ── 通知教师审批 ──
        try:
            from backend.api.notification_router import _notify_users
            from backend.database import execute_query as db_query
            g, c = _get_user_grade_class(asker_username)
            if g:
                teachers = db_query(
                    f"SELECT username FROM users WHERE role = 1 AND grade = ?"
                    + (" AND INSTR(',' || ? || ',', ',' || class || ',') > 0" if c else ""),
                    (g, c) if c else (g,),
                )
                if teachers:
                    _notify_users([r[0] for r in teachers], "info",
                        "有学生回答了提问，需要审批",
                        f"问题：{question_content[:50]}{'...' if len(question_content) > 50 else ''}",
                        "/interaction")
        except Exception as e:
            logger.warning(f"发送审批通知失败: {e}")

        return {"message": "回答已提交，等待教师审批"}

    # ── 教师/管理员回答：直接写入主表 ──
    if role == 1:
        grade, cls = _get_user_grade_class(username)
        if cls:
            rows = execute_query(
                """SELECT q.id FROM interaction_questions q
                   JOIN users u ON q.student_username = u.username AND u.role = 2
                   WHERE q.id = ? AND u.grade = ? AND INSTR(',' || ? || ',', ',' || u.class || ',') > 0""",
                (question_id, grade, cls),
            )
        else:
            rows = execute_query(
                """SELECT q.id FROM interaction_questions q
                   JOIN users u ON q.student_username = u.username AND u.role = 2
                   WHERE q.id = ? AND u.grade = ?""",
                (question_id, grade),
            )
        if not rows:
            raise HTTPException(status_code=403, detail="无权回答非本班学生的提问")

    execute_insert_update(
        """UPDATE interaction_questions SET answer = ?, status = 'answered',
           answered_at = ?, answered_by = ? WHERE id = ?""",
        (req.answer, now, username, question_id),
    )

    try:
        from backend.api.notification_router import _create_notification
        _create_notification(asker_username, "info",
            "你的提问已被教师回答",
            f"问题：{question_content[:50]}{'...' if len(question_content) > 50 else ''}",
            "/interaction")
    except Exception as e:
        logger.warning(f"发送回答通知失败: {e}")

    return {"message": "回答成功"}


# ── 获取某个问题的所有学生回答（教师看全部，学生看已通过的） ──

@router.get("/questions/{question_id}/answers", summary="获取问题的学生回答列表")
async def list_question_answers(question_id: int, request: Request):
    """获取某个问题的所有学生回答。教师看全部，学生看已通过的和自己的。"""
    user = get_current_user(request)
    role = user.get("role", 2)
    username = user["username"]

    # 验证问题存在
    q = execute_query("SELECT id FROM interaction_questions WHERE id = ?", (question_id,))
    if not q:
        raise HTTPException(status_code=404, detail="提问不存在")

    if role == 2:
        # 学生：看所有已通过的 + 自己的（含待审批/已拒绝）
        rows = execute_query(
            """SELECT id, student_username, answer, status, created_at
               FROM interaction_question_answers
               WHERE question_id = ? AND (status = 'approved' OR student_username = ?)
               ORDER BY created_at DESC""",
            (question_id, username),
        )
    else:
        # 教师/管理员：看全部
        rows = execute_query(
            """SELECT a.id, a.student_username, a.answer, a.status, a.created_at
               FROM interaction_question_answers a
               JOIN users u ON a.student_username = u.username AND u.role = 2
               WHERE a.question_id = ?
               ORDER BY a.created_at DESC""",
            (question_id,),
        )

    answers = []
    for r in rows:
        ans = {
            "id": r[0],
            "student_username": r[1],
            "answer": r[2],
            "status": r[3],
            "created_at": r[4],
        }
        answers.append(ans)

    return {"answers": answers}


# ── 教师逐条审批学生回答 ──

@router.put("/questions/{question_id}/answers/{answer_id}/approve", summary="教师审批通过某条学生回答")
async def approve_student_answer(question_id: int, answer_id: int, request: Request):
    """教师审批通过某条学生回答"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可审批")

    # 查询该回答
    answer_row = execute_query(
        """SELECT a.id, a.student_username, a.answer, q.content, q.student_username
           FROM interaction_question_answers a
           JOIN interaction_questions q ON a.question_id = q.id
           WHERE a.id = ? AND a.question_id = ?""",
        (answer_id, question_id),
    )
    if not answer_row:
        raise HTTPException(status_code=404, detail="回答不存在")
    answer_username = answer_row[0][1]
    question_content = answer_row[0][3]
    asker_username = answer_row[0][4]

    # 教师：验证该提问来自自己班级
    if role == 1:
        grade, cls = _get_user_grade_class(user["username"])
        rows = execute_query(
            """SELECT q.id FROM interaction_questions q
               JOIN users u ON q.student_username = u.username AND u.role = 2
               WHERE q.id = ? AND u.grade = ?"""
            + (" AND INSTR(',' || ? || ',', ',' || u.class || ',') > 0" if cls else ""),
            (question_id, grade, cls) if cls else (question_id, grade),
        )
        if not rows:
            raise HTTPException(status_code=403, detail="无权审批非本班学生的提问")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    execute_insert_update(
        "UPDATE interaction_question_answers SET status = 'approved' WHERE id = ?",
        (answer_id,),
    )
    # 标记问题已回答（但不覆盖主表 answer 字段，多个学生回答各自独立展示）
    execute_insert_update(
        """UPDATE interaction_questions SET status = 'answered',
           answered_at = COALESCE(answered_at, ?) WHERE id = ?""",
        (now, question_id),
    )

    # ── 积分奖励：回答被审批通过 ──
    try:
        from backend.reward_engine import award_participation
        # 用 question_id + answer_id 作为唯一标识，确保不被重复发放
        award_participation(answer_username, "question",
            f"{question_id}_{answer_id}",
            f"回答被审批通过：{question_content[:30]}...",
            teacher_username=user["username"])
    except Exception:
        pass

    # ── 通知回答者 ──
    try:
        from backend.api.notification_router import _create_notification
        _create_notification(answer_username, "info",
            "你的回答已通过教师审批",
            f"问题：{question_content[:50]}{'...' if len(question_content) > 50 else ''}",
            "/interaction")
    except Exception as e:
        logger.warning(f"发送审批通过通知失败: {e}")

    # ── 通知提问者 ──
    try:
        from backend.api.notification_router import _create_notification
        _create_notification(asker_username, "info",
            "你的提问已有回答（已通过教师审批）",
            f"问题：{question_content[:50]}{'...' if len(question_content) > 50 else ''}",
            "/interaction")
    except Exception as e:
        logger.warning(f"发送审批通知失败: {e}")

    return {"message": "审批通过"}


@router.put("/questions/{question_id}/answers/{answer_id}/reject", summary="教师拒绝某条学生回答")
async def reject_student_answer(question_id: int, answer_id: int, request: Request):
    """教师拒绝某条学生回答，该学生仍可重新回答"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可审批")

    answer_row = execute_query(
        """SELECT a.id, a.student_username, q.content
           FROM interaction_question_answers a
           JOIN interaction_questions q ON a.question_id = q.id
           WHERE a.id = ? AND a.question_id = ?""",
        (answer_id, question_id),
    )
    if not answer_row:
        raise HTTPException(status_code=404, detail="回答不存在")
    answer_username = answer_row[0][1]
    question_content = answer_row[0][2]

    if role == 1:
        grade, cls = _get_user_grade_class(user["username"])
        rows = execute_query(
            """SELECT q.id FROM interaction_questions q
               JOIN users u ON q.student_username = u.username AND u.role = 2
               WHERE q.id = ? AND u.grade = ?"""
            + (" AND INSTR(',' || ? || ',', ',' || u.class || ',') > 0" if cls else ""),
            (question_id, grade, cls) if cls else (question_id, grade),
        )
        if not rows:
            raise HTTPException(status_code=403, detail="无权审批非本班学生的提问")

    execute_insert_update(
        "UPDATE interaction_question_answers SET status = 'rejected' WHERE id = ?",
        (answer_id,),
    )

    # ── 通知回答者 ──
    try:
        from backend.api.notification_router import _create_notification
        _create_notification(answer_username, "info",
            "你的回答未通过教师审批，可重新回答",
            f"问题：{question_content[:50]}{'...' if len(question_content) > 50 else ''}",
            "/interaction")
    except Exception as e:
        logger.warning(f"发送拒绝通知失败: {e}")

    return {"message": "已拒绝"}


# ═══════════════════════════════════════════════════════════
# V3.4 新增：AI 实时答题分析
# ═══════════════════════════════════════════════════════════

@router.get("/quizzes/{quiz_id}/ai-analysis", summary="AI 实时答题分析")
async def ai_quiz_analysis(quiz_id: int, request: Request):
    """AI 分析随堂测验结果，生成教学建议"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可查看")

    quiz = execute_query(
        "SELECT * FROM interaction_quizzes WHERE id = ?",
        (quiz_id,),
    )
    if not quiz:
        raise HTTPException(status_code=404, detail="测验不存在")

    quiz_data = quiz[0]
    questions = json.loads(quiz_data[4]) if isinstance(quiz_data[4], str) else quiz_data[4]

    answers = execute_query(
        "SELECT student_username, answers, score FROM interaction_quiz_answers WHERE quiz_id = ?",
        (quiz_id,),
    )

    participant_count = len(answers)
    total_possible = sum(q.get("score", 1) for q in questions)

    # 每题正确率统计
    question_stats = []
    for i, q in enumerate(questions):
        correct_count = 0
        q_type = q.get("type", "single")
        correct_ans = str(q.get("answer", "")).strip()
        wrong_options = {}  # 统计错误选项分布

        for a in answers:
            user_ans_list = json.loads(a[1]) if isinstance(a[1], str) else a[1]
            user_ans = ""
            for ua in user_ans_list:
                if ua.get("question_index") == i:
                    user_ans = str(ua.get("answer", "")).strip()
                    break

            if q_type == "multiple":
                user_set = sorted([x.strip().upper() for x in user_ans.split(",") if x.strip()])
                correct_set = sorted([x.strip().upper() for x in correct_ans.split(",") if x.strip()])
                if user_set == correct_set:
                    correct_count += 1
                else:
                    # 记录错误选项
                    for ans_part in user_set:
                        wrong_options[ans_part] = wrong_options.get(ans_part, 0) + 1
            else:
                if user_ans.upper() == correct_ans.upper():
                    correct_count += 1
                elif user_ans:
                    wrong_options[user_ans.upper()] = wrong_options.get(user_ans.upper(), 0) + 1

        rate = round(correct_count / max(participant_count, 1) * 100, 1)
        wrong_text = ""
        if wrong_options:
            sorted_wrong = sorted(wrong_options.items(), key=lambda x: -x[1])
            wrong_text = "，错误选项分布：" + "、".join(f"{k}({v}人)" for k, v in sorted_wrong[:3])

        question_stats.append({
            "index": i + 1,
            "text": q.get("question", "")[:60],
            "type": q_type,
            "correct_rate": rate,
            "correct_count": correct_count,
            "total": participant_count,
            "wrong_distribution": wrong_text,
        })

    # 构建统计文本
    stats_text = ""
    for qs in question_stats:
        stats_text += f"第{qs['index']}题 ({qs['type']}): 正确率{qs['correct_rate']}% ({qs['correct_count']}/{qs['total']}人){qs['wrong_distribution']}\n"
        stats_text += f"  题目：{qs['text']}\n\n"

    from backend.api.ai_service import call_ai_async
    from backend.ai_task_manager import task_manager

    async def _do_analysis() -> dict:
        from backend.prompts.interaction import QUIZ_ANALYSIS_PROMPT
        from backend.api.chat_router import get_api_keys

        keys = get_api_keys(user["username"])
        api_key = keys[0] if keys and keys[0] else ""
        if not api_key:
            raise HTTPException(status_code=400, detail="未配置 API Key")

        def _safe(s):
            return str(s).replace('{', '{{').replace('}', '}}')

        prompt = QUIZ_ANALYSIS_PROMPT.format(
            quiz_title=_safe(quiz_data[2] or ""),
            subject=_safe("信息科技"),
            participant_count=_safe(participant_count),
            question_stats=_safe(stats_text),
        )

        analysis = await call_ai_async(prompt, api_key)
        return {
            "analysis": analysis,
            "stats": question_stats,
            "participant_count": participant_count,
            "avg_score": round(sum(a[2] for a in answers) / max(participant_count, 1), 1),
            "total_score": total_possible,
        }

    task_id = await task_manager.create_task(
        description=f"测验 #{quiz_id} AI 答题分析",
        coro_factory=_do_analysis,
    )
    return {"task_id": task_id, "message": "AI 分析已提交，请稍后查询结果"}


# ═══════════════════════════════════════════════════════════
# V3.4 新增：AI 课堂总结
# ═══════════════════════════════════════════════════════════

@router.get("/class-summary", summary="AI 课堂总结")
async def ai_class_summary(
    request: Request,
    grade: str = Query("", description="年级"),
    cls: str = Query("", description="班级"),
    subject: str = Query("", description="学科（由前端传递）"),
    teacher_username: str = Query("", description="教师用户名"),
    time_range: str = Query("本堂课", description="时间范围"),
):
    """AI 综合分析课堂互动数据，生成课堂总结"""
    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可查看")

    query_teacher = teacher_username or user["username"]

    # 构建班级名称
    cls_display = cls
    cls_name = f"{grade}{cls_display}班" if not cls.endswith("班") else f"{grade}{cls}"

    # ── 1. 随堂测验数据 ──
    quizzes = execute_query(
        """SELECT q.id, q.title, q.questions,
                  (SELECT COUNT(*) FROM interaction_quiz_answers WHERE quiz_id = q.id) as answer_count
           FROM interaction_quizzes q
           WHERE q.creator_username = ? AND q.status = 'active'
           ORDER BY q.created_at DESC LIMIT 5""",
        (query_teacher,),
    )

    quiz_data = ""
    quiz_participants = set()
    for q in quizzes:
        questions = json.loads(q[2]) if isinstance(q[2], str) else q[2]
        q_count = len(questions) if isinstance(questions, list) else 0
        quiz_data += f"- 《{q[1]}》: {q_count}题, {q[3]}人参与\n"
        # 获取参与学生
        participants = execute_query(
            "SELECT student_username FROM interaction_quiz_answers WHERE quiz_id = ?",
            (q[0],),
        )
        for p in participants:
            quiz_participants.add(p[0])

    if not quiz_data:
        quiz_data = "暂无随堂测验数据"

    # ── 2. 投票数据 ──
    polls = execute_query(
        """SELECT p.id, p.question, p.options,
                  (SELECT COUNT(*) FROM interaction_poll_votes WHERE poll_id = p.id) as vote_count
           FROM interaction_polls p
           WHERE p.creator_username = ?
           ORDER BY p.created_at DESC LIMIT 5""",
        (query_teacher,),
    )

    poll_data = ""
    poll_participants = set()
    for p in polls:
        options = json.loads(p[2]) if isinstance(p[2], str) else p[2]
        opt_count = len(options) if isinstance(options, list) else 0
        poll_data += f"- 「{p[1]}」: {opt_count}个选项, {p[3]}人投票\n"
        voters = execute_query(
            "SELECT student_username FROM interaction_poll_votes WHERE poll_id = ?",
            (p[0],),
        )
        for v in voters:
            poll_participants.add(v[0])

    if not poll_data:
        poll_data = "暂无投票数据"

    # ── 3. 学生提问 ──
    questions_data = execute_query(
        """SELECT id, content, is_anonymous, status, created_at
           FROM interaction_questions
           ORDER BY created_at DESC LIMIT 10""",
    )

    question_data = ""
    for q in questions_data:
        status = "已回答" if q[3] == "answered" else "待回答"
        question_data += f"- {q[1][:60]} ({status})\n"

    if not question_data:
        question_data = "暂无学生提问"

    # ── 4. 分组讨论 ──
    discussions = execute_query(
        """SELECT d.id, d.title, d.status,
                  (SELECT COUNT(*) FROM discussion_groups g
                   INNER JOIN discussion_members m ON m.group_id = g.id
                   WHERE g.discussion_id = d.id) as member_count,
                  (SELECT COUNT(*) FROM discussion_groups g
                   INNER JOIN discussion_messages m ON m.group_id = g.id
                   WHERE g.discussion_id = d.id) as msg_count
           FROM discussions d
           WHERE d.creator_username = ?
           ORDER BY d.created_at DESC LIMIT 5""",
        (query_teacher,),
    )

    discussion_data = ""
    for d in discussions:
        discussion_data += f"- 《{d[1]}》: {d[3]}人参与, {d[4]}条消息 ({d[2]})\n"

    if not discussion_data:
        discussion_data = "暂无分组讨论数据"

    # 计算总参与人数
    all_participants = quiz_participants | poll_participants
    student_count = len(all_participants)

    from backend.api.ai_service import call_ai_async
    from backend.ai_task_manager import task_manager

    # 将参数复制到局部变量，避免闭包引用外层参数
    _subject = subject
    _time_range = time_range

    async def _do_summary() -> dict:
        from backend.prompts.class_summary import CLASS_SUMMARY_PROMPT
        from backend.api.chat_router import get_api_keys

        keys = get_api_keys(user["username"])
        api_key = keys[0] if keys and keys[0] else ""
        if not api_key:
            raise HTTPException(status_code=400, detail="未配置 API Key")

        def _safe(s):
            return str(s).replace('{', '{{').replace('}', '}}')

        prompt = CLASS_SUMMARY_PROMPT.format(
            subject=_safe(_subject),
            time_range=_safe(_time_range),
            student_count=_safe(student_count),
            quiz_data=_safe(quiz_data),
            poll_data=_safe(poll_data),
            question_data=_safe(question_data),
            discussion_data=_safe(discussion_data),
        )

        summary = await call_ai_async(prompt, api_key)
        return {
            "summary": summary,
            "data": {
                "quiz_count": len(quizzes),
                "poll_count": len(polls),
                "question_count": len(questions_data),
                "discussion_count": len(discussions),
                "student_count": student_count,
            },
        }

    task_id = await task_manager.create_task(
        description="AI 课堂总结",
        coro_factory=_do_summary,
    )
    return {"task_id": task_id, "message": "课堂总结已提交，请稍后查询结果"}


@router.get("/class-summary/export", summary="导出课堂总结为 Word 文档")
async def export_class_summary_docx(
    request: Request,
    grade: str = Query("", description="年级"),
    cls: str = Query("", description="班级"),
    subject: str = Query("", description="学科"),
    teacher_username: str = Query("", description="教师用户名"),
    token: str = Query("", description="认证令牌"),
):
    """导出 AI 课堂总结报告为 Word 文档"""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    if token:
        request.state.user = None
        from backend.auth import decode_jwt_token
        payload = decode_jwt_token(token)
        if payload:
            request.state.user = payload

    user = get_current_user(request)
    role = user.get("role", 2)
    if role == 2:
        raise HTTPException(status_code=403, detail="仅教师和管理员可导出")

    query_teacher = teacher_username or user["username"]
    cls_display = cls
    cls_name = f"{grade}{cls_display}班" if not cls.endswith("班") else f"{grade}{cls}"

    # ── 收集数据（复用 class-summary 逻辑） ──
    quizzes = execute_query(
        """SELECT q.id, q.title, q.questions,
                  (SELECT COUNT(*) FROM interaction_quiz_answers WHERE quiz_id = q.id) as answer_count
           FROM interaction_quizzes q
           WHERE q.creator_username = ? AND q.status = 'active'
           ORDER BY q.created_at DESC LIMIT 5""",
        (query_teacher,),
    )
    quiz_data = ""
    quiz_participants = set()
    for q in quizzes:
        questions = json.loads(q[2]) if isinstance(q[2], str) else q[2]
        q_count = len(questions) if isinstance(questions, list) else 0
        quiz_data += f"- 《{q[1]}》: {q_count}题, {q[3]}人参与\n"
        participants = execute_query(
            "SELECT student_username FROM interaction_quiz_answers WHERE quiz_id = ?", (q[0],),
        )
        for p in participants:
            quiz_participants.add(p[0])
    if not quiz_data:
        quiz_data = "暂无随堂测验数据"

    polls = execute_query(
        """SELECT p.id, p.question, p.options,
                  (SELECT COUNT(*) FROM interaction_poll_votes WHERE poll_id = p.id) as vote_count
           FROM interaction_polls p
           WHERE p.creator_username = ?
           ORDER BY p.created_at DESC LIMIT 5""",
        (query_teacher,),
    )
    poll_data = ""
    poll_participants = set()
    for p in polls:
        options = json.loads(p[2]) if isinstance(p[2], str) else p[2]
        opt_count = len(options) if isinstance(options, list) else 0
        poll_data += f"- 「{p[1]}」: {opt_count}个选项, {p[3]}人投票\n"
        voters = execute_query(
            "SELECT student_username FROM interaction_poll_votes WHERE poll_id = ?", (p[0],),
        )
        for v in voters:
            poll_participants.add(v[0])
    if not poll_data:
        poll_data = "暂无投票数据"

    questions_data = execute_query(
        """SELECT id, content, is_anonymous, status, created_at
           FROM interaction_questions
           ORDER BY created_at DESC LIMIT 10""",
    )
    question_data = ""
    for q in questions_data:
        status = "已回答" if q[3] == "answered" else "待回答"
        question_data += f"- {q[1][:60]} ({status})\n"
    if not question_data:
        question_data = "暂无学生提问"

    discussions = execute_query(
        """SELECT d.id, d.title, d.status,
                  (SELECT COUNT(*) FROM discussion_groups g
                   INNER JOIN discussion_members m ON m.group_id = g.id
                   WHERE g.discussion_id = d.id) as member_count,
                  (SELECT COUNT(*) FROM discussion_groups g
                   INNER JOIN discussion_messages m ON m.group_id = g.id
                   WHERE g.discussion_id = d.id) as msg_count
           FROM discussions d
           WHERE d.creator_username = ?
           ORDER BY d.created_at DESC LIMIT 5""",
        (query_teacher,),
    )
    discussion_data = ""
    for d in discussions:
        discussion_data += f"- 《{d[1]}》: {d[3]}人参与, {d[4]}条消息 ({d[2]})\n"
    if not discussion_data:
        discussion_data = "暂无分组讨论数据"

    all_participants = quiz_participants | poll_participants
    student_count = len(all_participants)

    from backend.api.ai_service import call_ai_async
    from backend.prompts.class_summary import CLASS_SUMMARY_PROMPT
    from backend.api.chat_router import get_api_keys

    keys = get_api_keys(user["username"])
    api_key = keys[0] if keys and keys[0] else ""
    if not api_key:
        raise HTTPException(status_code=400, detail="未配置 API Key")

    def _safe(s):
        return str(s).replace('{', '{{').replace('}', '}}')

    prompt = CLASS_SUMMARY_PROMPT.format(
        subject=_safe(subject or "信息科技"),
        time_range=_safe("本堂课"),
        student_count=_safe(student_count),
        quiz_data=_safe(quiz_data),
        poll_data=_safe(poll_data),
        question_data=_safe(question_data),
        discussion_data=_safe(discussion_data),
    )

    try:
        summary = await call_ai_async(prompt, api_key)
    except Exception as e:
        logger.error(f"AI 课堂总结调用失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 分析出错: {str(e)}")

    # ── 生成 Word 文档 ──
    doc = Document()
    style = doc.styles['Normal']  # type: ignore[union-attr]
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_heading(f"{cls_name} 课堂总结报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"年级：{grade}  班级：{cls_display}  学科：{subject}  参与学生：{student_count}人")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    # 写入统计数据
    doc.add_heading("课堂数据统计", level=2)
    stats_table_data = [
        ("随堂测验", str(len(quizzes))),
        ("投票活动", str(len(polls))),
        ("学生提问", str(len(questions_data))),
        ("分组讨论", str(len(discussions))),
        ("参与学生", str(student_count)),
    ]
    table = doc.add_table(rows=len(stats_table_data), cols=2, style='Light Shading Accent 1')
    for i, (label, value) in enumerate(stats_table_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()
    doc.add_heading("AI 分析总结", level=2)
    # 复用 markdown 转 docx 函数
    _markdown_to_docx(doc, summary)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import urllib.parse
    safe_filename = urllib.parse.quote(f"课堂总结_{cls_name}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"},
    )


# ═══════════════════════════════════════════════════════════
# Word 导出工具
# ═══════════════════════════════════════════════════════════

def _markdown_to_docx(doc, text: str):
    """将 Markdown 文本写入 docx 文档"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- **') and '：' in line:
            content = line.lstrip('- ')
            p = doc.add_paragraph()
            bold_end = content.find('**', 2)
            if bold_end > 0:
                run = p.add_run(content[2:bold_end])
                run.bold = True
                p.add_run(content[bold_end + 2:])
            else:
                p.add_run(content)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f'{i}. ') for i in range(1, 10)):
            doc.add_paragraph(line, style='List Number')
        else:
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            else:
                doc.add_paragraph(line)


# ═══════════════════════════════════════════════════════════
# V3.4 新增：AI 异步任务状态查询
# ═══════════════════════════════════════════════════════════

@router.get("/ai-task/{task_id}", summary="查询 AI 异步任务状态")
async def get_ai_task_status(task_id: str):
    """查询 AI 后台任务的执行状态和结果"""
    from backend.ai_task_manager import task_manager
    task = task_manager.get_task_dict(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在或已过期")
    return task
