"""
Word 试卷生成引擎
使用 python-docx 生成排版规范的试卷文档（打印用）
支持 LaTeX 公式渲染（→图片嵌入）和 SVG/媒体图片嵌入
"""
import io
import json
import re
import os
from typing import Any
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

from backend.logger import logger

# ── 公式渲染（matplotlib 数学文本引擎，无需 LaTeX 系统安装） ──
_HAS_MPL = False
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.mathtext import MathTextParser
    _HAS_MPL = True
except ImportError:
    logger.warning("matplotlib 未安装，LaTeX 公式将以纯文本方式显示")

# ── SVG 转换 ──
_HAS_CAIROSVG = False
try:
    import cairosvg
    _HAS_CAIROSVG = True
except ImportError:
    pass

_HAS_SVGLIB = False
try:
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPM
    _HAS_SVGLIB = True
except ImportError:
    pass

# ═══════════════════════════════════════════════════════════════
# 样式常量
# ═══════════════════════════════════════════════════════════════

FONT_TITLE = "黑体"
FONT_BODY = "宋体"
FONT_HEADING = "黑体"

SIZE_TITLE = Pt(22)
SIZE_SUBTITLE = Pt(14)
SIZE_STUDENT_INFO = Pt(12)
SIZE_SECTION = Pt(13)
SIZE_QUESTION = Pt(12)
SIZE_OPTION = Pt(12)
SIZE_FOOTER = Pt(9)
SIZE_ANSWER_SECTION = Pt(11)

COLOR_LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
COLOR_DARK = RGBColor(0x33, 0x33, 0x33)
COLOR_ANSWER = RGBColor(0xCC, 0x00, 0x00)  # 答案卷用红色标注答案
COLOR_EXPLANATION = RGBColor(0x00, 0x66, 0xCC)  # 解析用蓝色

TYPE_LABELS = {
    "single": "一、单选题",
    "multiple": "二、多选题",
    "true_false": "三、判断题",
    "short": "四、简答题",
}

TYPE_ORDER = ["single", "multiple", "true_false", "short"]


# ═══════════════════════════════════════════════════════════════
# LaTeX 公式渲染（→ PNG 图片）
# ═══════════════════════════════════════════════════════════════

# LaTeX 命令 → Unicode 字符映射（备选方案，当 matplotlib 不可用时使用）
_LATEX_UNICODE_MAP = {
    # 希腊字母
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
    r'\epsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ',
    r'\iota': 'ι', r'\kappa': 'κ', r'\lambda': 'λ', r'\mu': 'μ',
    r'\nu': 'ν', r'\xi': 'ξ', r'\omicron': 'ο', r'\pi': 'π',
    r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ',
    r'\phi': 'φ', r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
    r'\Alpha': 'Α', r'\Beta': 'Β', r'\Gamma': 'Γ', r'\Delta': 'Δ',
    r'\Theta': 'Θ', r'\Lambda': 'Λ', r'\Xi': 'Ξ', r'\Pi': 'Π',
    r'\Sigma': 'Σ', r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
    # 数学符号
    r'\sum': '∑', r'\prod': '∏', r'\int': '∫', r'\oint': '∮',
    r'\infty': '∞', r'\partial': '∂', r'\nabla': '∇',
    r'\rightarrow': '→', r'\leftarrow': '←', r'\Rightarrow': '⇒',
    r'\Leftarrow': '⇐', r'\mapsto': '↦',
    r'\geq': '≥', r'\leq': '≤', r'\neq': '≠', r'\approx': '≈',
    r'\equiv': '≡', r'\times': '×', r'\div': '÷', r'\pm': '±',
    r'\cap': '∩', r'\cup': '∪', r'\subset': '⊂', r'\supset': '⊃',
    r'\subseteq': '⊆', r'\supseteq': '⊇',
    r'\in': '∈', r'\notin': '∉', r'\forall': '∀', r'\exists': '∃',
    r'\neg': '¬', r'\land': '∧', r'\lor': '∨',
    r'\sqrt': '√', r'\angle': '∠', r'\perp': '⊥',
    r'\triangle': '△', r'\square': '□',
    r'\therefore': '∴', r'\because': '∵',
    r'\prime': '′', r'\degree': '°',
    r'\leftrightarrow': '↔',
    r'\circ': '∘', r'\cdot': '·', r'\cdots': '…', r'\vdots': '⋮', r'\ddots': '⋱',
    r'\emptyset': '∅',
    r'\to': '→',
    # 函数名
    r'\sin': 'sin', r'\cos': 'cos', r'\tan': 'tan',
    r'\log': 'log', r'\ln': 'ln', r'\lg': 'lg',
    r'\max': 'max', r'\min': 'min', r'\lim': 'lim',
    r'\sinh': 'sinh', r'\cosh': 'cosh', r'\tanh': 'tanh',
    r'\arcsin': 'arcsin', r'\arccos': 'arccos', r'\arctan': 'arctan',
}

# 下标/上标的 Unicode 映射
_SUPERSCRIPTS = {
    '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
    '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
    '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
    'n': 'ⁿ', 'i': 'ⁱ',
}

_SUBSCRIPTS = {
    '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
    '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
    '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
    'a': 'ₐ', 'e': 'ₑ', 'o': 'ₒ', 'x': 'ₓ',
}


def _latex_to_png(latex_str: str, fontsize: int = 12, dpi: int = 150) -> io.BytesIO | None:
    """将 LaTeX 公式渲染为 PNG 图片（使用 matplotlib 数学文本引擎）

    Args:
        latex_str: LaTeX 公式内容（不含 $）
        fontsize: 字号
        dpi: 图片 DPI

    Returns:
        BytesIO 对象（PNG 数据），渲染失败返回 None
    """
    if not _HAS_MPL:
        return None
    try:
        buf = io.BytesIO()
        fig, ax = plt.subplots(figsize=(0.01, 0.01), dpi=dpi)
        ax.axis('off')
        # 确保内容以 $ 包裹以触发 mathtext 渲染
        display_text = f'${latex_str}$'
        text_obj = ax.text(0, 0, display_text, fontsize=fontsize,
                           verticalalignment='bottom', horizontalalignment='left')
        fig.savefig(buf, format='png', bbox_inches='tight',
                    pad_inches=0.05, transparent=True, dpi=dpi)
        plt.close(fig)
        buf.seek(0)
        if buf.getbuffer().nbytes > 200:  # 确保有实际内容
            return buf
        return None
    except Exception as e:
        logger.warning(f"LaTeX 公式渲染失败: '{latex_str[:50]}' - {e}")
        return None


def _latex_to_unicode(latex_str: str) -> str:
    """将简单的 LaTeX 公式转换为 Unicode 近似文本（备选方案）

    处理: 上标 ^, 下标 _, \frac, \sqrt, 以及常见 LaTeX 命令
    """
    text = latex_str.strip()

    # 去除外层多余的 $（安全处理）
    text = text.strip('$')

    # 处理 \text{...}
    text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)

    # 处理 \frac{a}{b}
    while '\\frac' in text:
        text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text, count=1)

    # 处理 \sqrt[n]{x} 和 \sqrt{x}
    text = re.sub(r'\\sqrt\[([^}]*)\]\{([^}]*)\}', r'√[\1](\2)', text)
    text = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', text)

    # 处理上标 x^y
    def _replace_super(m):
        base = m.group(1)
        exp = m.group(2)
        # 尝试用 Unicode 上标
        if len(exp) == 1 and exp in _SUPERSCRIPTS:
            return f'{base}{_SUPERSCRIPTS[exp]}'
        return f'{base}^{{{exp}}}'

    text = re.sub(r'([a-zA-Z0-9\)\]}])\^\{([^}]*)\}', _replace_super, text)
    text = re.sub(r'([a-zA-Z0-9\)\]}])\^([a-zA-Z0-9])', lambda m: f'{m.group(1)}{_SUPERSCRIPTS.get(m.group(2), "^"+m.group(2))}', text)

    # 处理下标 x_y
    def _replace_sub(m):
        base = m.group(1)
        sub = m.group(2)
        if len(sub) == 1 and sub in _SUBSCRIPTS:
            return f'{base}{_SUBSCRIPTS[sub]}'
        return f'{base}_{{{sub}}}'

    text = re.sub(r'([a-zA-Z0-9\)\]}])\_\{([^}]*)\}', _replace_sub, text)
    text = re.sub(r'([a-zA-Z0-9\)\]}])\_([a-zA-Z0-9])', lambda m: f'{m.group(1)}{_SUBSCRIPTS.get(m.group(2), "_"+m.group(2))}', text)

    # 替换 LaTeX 命令为 Unicode
    for cmd, unicode_char in sorted(_LATEX_UNICODE_MAP.items(), key=lambda x: -len(x[0])):
        text = text.replace(cmd, unicode_char)

    # 清理残余的 {}
    text = text.replace('{', '').replace('}', '')

    # 清理多余空格（保留单个空格）
    text = re.sub(r'  +', ' ', text)

    return text.strip()


def _render_latex_inline(paragraph, text: str, font_name: str, font_size, doc):
    """渲染包含 LaTeX 公式的文本到段落

    支持 $...$（行内公式）和 $$...$$（独立公式）
    会尝试用 matplotlib 渲染为图片；失败时回退到 Unicode 文本
    """
    # 先处理 $$...$$（独立公式，块级）
    parts = re.split(r'(\$\$[^$]+\$\$)', text)
    for i, part in enumerate(parts):
        if part.startswith('$$') and part.endswith('$$'):
            latex_content = part[2:-2].strip()
            if not latex_content:
                continue
            # 尝试渲染为图片
            img_buf = _latex_to_png(latex_content, fontsize=12)
            if img_buf:
                run = paragraph.add_run()
                run.add_picture(img_buf, width=Cm(6))
                # 添加换行
                run2 = paragraph.add_run()
                run2.add_break(WD_BREAK.LINE)
            else:
                # 回退到 Unicode
                unicode_text = _latex_to_unicode(latex_content)
                if unicode_text:
                    run = paragraph.add_run(f" {unicode_text} ")
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = font_size
        else:
            # 处理行内 $...$
            _render_latex_inline_inner(paragraph, part, font_name, font_size)


def _render_latex_inline_inner(paragraph, text: str, font_name: str, font_size):
    """渲染含行内公式 $...$ 的文本"""
    parts = re.split(r'(\$[^$]+\$)', text)
    for part in parts:
        if part.startswith('$') and part.endswith('$'):
            latex_content = part[1:-1].strip()
            if not latex_content:
                continue
            # 尝试渲染为图片
            img_buf = _latex_to_png(latex_content, fontsize=11)
            if img_buf:
                run = paragraph.add_run()
                run.add_picture(img_buf, height=Cm(0.8))
            else:
                # 回退到 Unicode
                unicode_text = _latex_to_unicode(latex_content)
                if unicode_text:
                    run = paragraph.add_run(unicode_text)
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = font_size
        else:
            if part.strip():
                run = paragraph.add_run(part)
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = font_size


# ═══════════════════════════════════════════════════════════════
# SVG / 媒体图片嵌入
# ═══════════════════════════════════════════════════════════════

def _svg_to_png(svg_content: str, width_cm: float = 6) -> io.BytesIO | None:
    """将 SVG 代码转换为 PNG 图片

    优先使用 cairosvg（质量好），备选 svglib，都不行则返回 None
    """
    if not svg_content or not svg_content.strip():
        return None
    try:
        if _HAS_CAIROSVG:
            png_data = cairosvg.svg2png(
                bytestring=svg_content.encode('utf-8'),
                output_width=int(width_cm * 96),
            )
            return io.BytesIO(png_data)
        elif _HAS_SVGLIB:
            buf = io.BytesIO()
            drawing = svg2rlg(io.BytesIO(svg_content.encode('utf-8')))
            renderPM.drawToFile(drawing, buf, fmt='PNG')
            buf.seek(0)
            return buf
    except Exception as e:
        logger.warning(f"SVG 转 PNG 失败: {e}")
    return None


def _embed_media_files(paragraph, media_files: list | str | None, question_id: int):
    """在段落后嵌入媒体图片（从已生成的 media_files 中读取）

    media_files 格式: [{"key": "img1", "type": "image", "url": "..."}]
    """
    if not media_files:
        return

    if isinstance(media_files, str):
        try:
            media_files = json.loads(media_files)
        except (json.JSONDecodeError, TypeError):
            return

    if not isinstance(media_files, list):
        return

    for mf in media_files:
        if not isinstance(mf, dict):
            continue
        mtype = mf.get("type", "")
        if mtype not in ("image",):
            continue
        url = mf.get("url", "")
        if not url:
            continue

        # 尝试定位本地文件
        image_path = _resolve_media_path(url, question_id)
        if image_path and image_path.exists():
            try:
                run = paragraph.add_run()
                # 根据图片类型调整最大宽度
                run.add_picture(str(image_path), width=Cm(8))
                run.add_break(WD_BREAK.LINE)
            except Exception as e:
                logger.warning(f"嵌入媒体图片失败: {url} - {e}")


def _resolve_media_path(url: str, question_id: int) -> Path | None:
    """解析媒体文件 URL 为本地路径"""
    from backend.config import BASE_DIR

    # 如果已经是本地路径
    url_clean = url.split("?")[0].rstrip("/")
    filename = url_clean.split("/")[-1]

    # 可能的位置：question_media/{id}/{filename}
    candidates = [
        BASE_DIR / "question_media" / str(question_id) / filename,
        BASE_DIR / "question_media" / str(question_id) / url_clean,
        Path(url_clean) if os.path.isabs(url_clean) else None,
    ]

    for path in candidates:
        if path and path.exists():
            return path

    return None


def _set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_line_spacing(paragraph, spacing=1.5):
    """设置段落行间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = spacing


def _add_run(paragraph, text, font_name=FONT_BODY, size=SIZE_QUESTION,
             bold=False, color=None, underline=False):
    """添加文本片段"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if underline:
        run.font.underline = True
    return run


def _add_paragraph(doc, text="", font_name=FONT_BODY, size=SIZE_QUESTION,
                   bold=False, alignment=None, space_before=0, space_after=0,
                   color=None, first_line_indent=None):
    """添加段落"""
    p = doc.add_paragraph()
    if text:
        _add_run(p, text, font_name, size, bold, color)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    _set_line_spacing(p, 1.5)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    return p


def _create_student_info_table(doc, school_name, exam_title, subject,
                                grade, semester, exam_desc):
    """创建考生信息栏（使用无边框表格实现对齐）"""
    # 学校名称 + 试卷标题
    p_school = _add_paragraph(doc, school_name, FONT_HEADING, SIZE_SUBTITLE,
                              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_after=2)
    p_title = _add_paragraph(doc, exam_title, FONT_TITLE, SIZE_TITLE,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_after=2)

    # 副标题行
    subtitle_parts = []
    if semester:
        subtitle_parts.append(semester)
    if subject:
        subtitle_parts.append(subject)
    if grade:
        subtitle_parts.append(grade)
    if subtitle_parts:
        _add_paragraph(doc, "  ".join(subtitle_parts), FONT_BODY, SIZE_STUDENT_INFO,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # 考生信息表格（无边框）
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格无边框
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                             start={"val": "none", "sz": "0"},
                             top={"val": "none", "sz": "0"},
                             end={"val": "none", "sz": "0"},
                             bottom={"val": "none", "sz": "0"})

    cells = table.rows[0].cells
    info_items = [
        ("姓名：", "_" * 12),
        ("班级：", "_" * 10),
        ("得分：", "_" * 10),
    ]

    for i, (label, underline) in enumerate(info_items):
        p = cells[i * 2].paragraphs[0]
        _add_run(p, label, FONT_BODY, SIZE_STUDENT_INFO)
        cells[i * 2].width = Cm(2.5)

        p2 = cells[i * 2 + 1].paragraphs[0]
        _add_run(p2, underline, FONT_BODY, SIZE_STUDENT_INFO)
        cells[i * 2 + 1].width = Cm(5)

    # 考试说明
    if exam_desc:
        _add_paragraph(doc, f"说明：{exam_desc}", FONT_BODY, Pt(10),
                       color=COLOR_LIGHT_GRAY, space_before=4, space_after=2)

    # 总分提示行
    _add_paragraph(doc, "", FONT_BODY, Pt(6), space_after=2)

    return doc


def _render_options(paragraph, options_dict: dict[str, str] | None, doc=None):
    """在段落中渲染选项（A. xxx    B. xxx    C. xxx    D. xxx）

    支持选项文本中的 LaTeX 公式 $...$
    """
    if not options_dict:
        return
    keys = list(options_dict.keys())
    # 每行两个选项
    for i in range(0, len(keys), 2):
        line_text = ""
        for j in range(2):
            if i + j < len(keys):
                k = keys[i + j]
                opt_text = options_dict[k]
                # 检查是否含 LaTeX
                if '$' in opt_text and _HAS_MPL:
                    run = paragraph.add_run("\n")
                    run.font.size = Pt(4)  # 微小的换行间隔
                    # 渲染选项标签
                    _add_run(paragraph, f"{k}. ")
                    # 渲染选项文本（含公式）
                    _render_latex_inline_inner(paragraph, opt_text, FONT_BODY, SIZE_OPTION)
                else:
                    line_text += f"{k}. {opt_text}    "
        if line_text:
            run = paragraph.add_run("\n" + line_text.strip())
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            run.font.size = SIZE_OPTION


def _render_true_false_options(paragraph):
    """判断题选项：在题号后加（  ）"""
    run = paragraph.add_run("\n（  ）")
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    run.font.size = SIZE_OPTION


def _render_short_answer_lines(doc, line_count=6):
    """简答题留空行"""
    for _ in range(line_count):
        _add_paragraph(doc, "_______________________________________________",
                       FONT_BODY, Pt(11), color=COLOR_LIGHT_GRAY, space_before=0, space_after=0)


def _add_section_title(doc, title: str, label: str, total_score: float):
    """添加题型标题行"""
    p = _add_paragraph(doc, "", space_before=8, space_after=4)
    _add_run(p, label, FONT_HEADING, SIZE_SECTION, bold=True)
    _add_run(p, f"（共 {total_score:.0f} 分）", FONT_BODY, Pt(11), color=COLOR_LIGHT_GRAY)
    return p


def _sanitize_text(text: str) -> str:
    """清理文本中的特殊字符，防止破坏 XML"""
    if not text:
        return ""
    # 替换 XML 非法字符
    text = text.replace('\x00', '').replace('\r', '')
    return text


def _render_question_text(paragraph, q_text: str, doc=None):
    """渲染题目文本

    支持:
    - LaTeX 公式 $...$ 和 $$...$$（渲染为图片嵌入）
    - Markdown **加粗**（转为普通加粗）
    - 纯文本直接输出
    """
    text = _sanitize_text(q_text)
    if not text:
        return

    # 检查是否包含 LaTeX 公式
    has_latex = ('$' in text)

    if has_latex and _HAS_MPL:
        _render_latex_inline(paragraph, text, FONT_BODY, SIZE_QUESTION, doc)
    else:
        # 无 LaTeX 时，仅处理 **加粗**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                _add_run(paragraph, part[2:-2], bold=True)
            else:
                _add_run(paragraph, part)


# ═══════════════════════════════════════════════════════════════
# 主生成函数
# ═══════════════════════════════════════════════════════════════

def generate_exam_paper(
    exam_info: dict[str, Any],
    questions: list[dict[str, Any]],
    school_name: str = "",
    semester: str = "",
    show_answer_key: bool = False,
) -> io.BytesIO:
    """生成 Word 试卷文档

    Args:
        exam_info: 考试信息（title, subject, grade, total_score, duration, description 等）
        questions: 题目列表（已排序）
        school_name: 学校名称
        semester: 学年学期
        show_answer_key: 是否生成答案卷（True=教师答案卷, False=学生试卷）

    Returns:
        BytesIO 对象（可直接作为 HTTP 响应返回）
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 默认样式 ──
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = SIZE_QUESTION
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    title = _sanitize_text(exam_info.get("title", "试卷"))
    subject = _sanitize_text(exam_info.get("subject", ""))
    grade = _sanitize_text(exam_info.get("target_grade", exam_info.get("grade", "")))
    total_score = exam_info.get("total_score", 100)
    duration = exam_info.get("duration", 45)
    description = _sanitize_text(exam_info.get("description", ""))

    # ── 考生信息栏 ──
    exam_title_text = f"{title}"
    if show_answer_key:
        exam_title_text += "（参考答案及解析）"

    _create_student_info_table(
        doc, school_name, exam_title_text,
        subject, grade, semester,
        description if not show_answer_key else "",
    )

    if show_answer_key:
        _add_paragraph(doc, "【机密·教师专用】", FONT_HEADING, Pt(10),
                       bold=True, color=COLOR_ANSWER,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    else:
        # 考试须知
        notice = f"注意事项：1. 本试卷满分 {total_score:.0f} 分，考试时间 {duration} 分钟。"
        notice += " 2. 选择题请将答案填涂在答题卡上，简答题请在答题纸上作答。"
        _add_paragraph(doc, notice, FONT_BODY, Pt(10), color=COLOR_LIGHT_GRAY,
                       space_before=4, space_after=6)

    # ── 按题型分组 ──
    type_groups: dict[str, list[dict]] = {}
    for q in questions:
        q_type = q.get("type", "single")
        if q_type not in type_groups:
            type_groups[q_type] = []
        type_groups[q_type].append(q)

    # ── 题型全局序号 ──
    question_number = 0

    # ── 遍历题型 ──
    for type_key in TYPE_ORDER:
        if type_key not in type_groups:
            continue

        group = type_groups[type_key]
        if not group:
            continue

        # 计算该题型总分
        type_total = sum(float(q.get("question_score", q.get("score", 5))) for q in group)

        # 题型标题
        label = TYPE_LABELS.get(type_key, type_key)
        _add_section_title(doc, title, label, type_total)

        # 逐题渲染
        for q in group:
            question_number += 1
            q_text = _sanitize_text(q.get("question_text", ""))
            q_score = float(q.get("question_score", q.get("score", 5)))
            q_options = q.get("options")
            if isinstance(q_options, str):
                try:
                    q_options = json.loads(q_options)
                except (json.JSONDecodeError, TypeError):
                    q_options = None

            # 题目行
            p_q = _add_paragraph(doc, "", space_before=4, space_after=1)
            if show_answer_key:
                # 答案卷：题号后标注答案
                correct = _sanitize_text(q.get("correct_answer", ""))
                _add_run(p_q, f"{question_number}. ", FONT_BODY, SIZE_QUESTION, bold=True)
                _render_question_text(p_q, q_text)
                _add_run(p_q, f"  【答案】{correct}", FONT_BODY, Pt(11),
                         bold=True, color=COLOR_ANSWER)
                _add_run(p_q, f"  ({q_score:.0f}分)", FONT_BODY, Pt(10),
                         color=COLOR_LIGHT_GRAY)
            else:
                _add_run(p_q, f"{question_number}. ", FONT_BODY, SIZE_QUESTION, bold=True)
                _render_question_text(p_q, q_text)
                _add_run(p_q, f"  ({q_score:.0f}分)", FONT_BODY, Pt(10),
                         color=COLOR_LIGHT_GRAY)

            # 选项渲染
            if type_key == "single" or type_key == "multiple":
                _render_options(p_q, q_options, doc)
            elif type_key == "true_false":
                _render_true_false_options(p_q)

            # ── SVG 配图嵌入 ──
            svg_content = q.get("svg_content")
            has_svg = q.get("has_svg", 0)
            if svg_content and (has_svg or str(has_svg) == "1"):
                p_svg = _add_paragraph(doc, "", space_before=2, space_after=2)
                p_svg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                svg_png = _svg_to_png(svg_content, width_cm=10)
                if svg_png:
                    run = p_svg.add_run()
                    run.add_picture(svg_png, width=Cm(10))

            # ── 媒体文件嵌入（已上传的图片） ──
            media_files = q.get("media_files")
            if media_files:
                _embed_media_files(p_q, media_files, q.get("id", 0))

            # 简答题留空（仅在学生试卷中）
            if type_key == "short" and not show_answer_key:
                line_count = max(3, int(q_score / 2))
                _render_short_answer_lines(doc, line_count)

            # 答案卷：显示解析
            if show_answer_key:
                explanation = _sanitize_text(q.get("explanation", ""))
                if explanation:
                    p_exp = _add_paragraph(doc, "", space_before=1, space_after=2)
                    _add_run(p_exp, "【解析】", FONT_BODY, Pt(10),
                             bold=True, color=COLOR_EXPLANATION)
                    _add_run(p_exp, explanation, FONT_BODY, Pt(10),
                             color=COLOR_EXPLANATION)
                # 知识点
                kp = _sanitize_text(q.get("knowledge_points", ""))
                if kp:
                    p_kp = _add_paragraph(doc, "", space_before=0, space_after=2)
                    _add_run(p_kp, "【知识点】", FONT_BODY, Pt(9),
                             bold=True, color=COLOR_LIGHT_GRAY)
                    _add_run(p_kp, kp, FONT_BODY, Pt(9),
                             color=COLOR_LIGHT_GRAY)

    # ── 页脚：页码 ──
    _add_page_number(doc)

    # ── 生成输出 ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_answer_key(
    exam_info: dict[str, Any],
    questions: list[dict[str, Any]],
    school_name: str = "",
    semester: str = "",
) -> io.BytesIO:
    """生成教师答案卷（含答案和解析）"""
    return generate_exam_paper(
        exam_info=exam_info,
        questions=questions,
        school_name=school_name,
        semester=semester,
        show_answer_key=True,
    )


def _add_page_number(doc):
    """为文档添加页码（页脚居中）"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 使用 Word 域代码实现页码
        run = p.add_run()
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_char1)

        run2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run2._element.append(instr)

        run3 = p.add_run()
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fld_char2)

        # 居中
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_answer_sheet(
    exam_info: dict[str, Any],
    questions: list[dict[str, Any]],
) -> io.BytesIO:
    """生成答题卡（选择题填涂区域 + 简答题作答区）"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)

    title = _sanitize_text(exam_info.get("title", "试卷"))

    _add_paragraph(doc, f"{title} — 答题卡", FONT_TITLE, SIZE_SUBTITLE,
                   bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # 考生信息
    table_info = doc.add_table(rows=1, cols=4)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_cells = table_info.rows[0].cells
    labels = ["姓名：________", "班级：________", "考号：________", "得分：________"]
    for i, lbl in enumerate(labels):
        info_cells[i].text = lbl
        for paragraph in info_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_paragraph(doc, "", space_after=6)

    # 按题型生成填涂区域
    type_groups: dict[str, list[dict]] = {}
    for q in questions:
        q_type = q.get("type", "single")
        if q_type not in type_groups:
            type_groups[q_type] = []
        type_groups[q_type].append(q)

    question_number = 0

    for type_key in TYPE_ORDER:
        if type_key not in type_groups:
            continue
        group = type_groups[type_key]
        if not group:
            continue

        label = TYPE_LABELS.get(type_key, type_key)

        if type_key in ("single", "multiple"):
            _add_paragraph(doc, label, FONT_HEADING, SIZE_SECTION, bold=True, space_before=6, space_after=4)

            # 创建填涂表格（每行 5 题）
            cols = 5
            rows_count = (len(group) + cols - 1) // cols
            table = doc.add_table(rows=rows_count, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            idx = 0
            for r in range(rows_count):
                for c in range(cols):
                    if idx < len(group):
                        q = group[idx]
                        question_number += 1
                        num = question_number
                        cell = table.rows[r].cells[c]
                        if type_key == "single":
                            cell.text = f"{num}. [A] [B] [C] [D]"
                        else:
                            cell.text = f"{num}. [A] [B] [C] [D] [E]"
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                        idx += 1

        elif type_key == "true_false":
            _add_paragraph(doc, label, FONT_HEADING, SIZE_SECTION, bold=True, space_before=6, space_after=4)

            # 判断题每行 8 题
            p_tf = _add_paragraph(doc, "", space_after=4)
            for q in group:
                question_number += 1
                _add_run(p_tf, f"{question_number}. [对] [错]    ")

        elif type_key == "short":
            _add_paragraph(doc, label, FONT_HEADING, SIZE_SECTION, bold=True, space_before=6, space_after=4)
            for q in group:
                question_number += 1
                q_text = _sanitize_text(q.get("question_text", ""))[:40]
                _add_paragraph(doc, f"{question_number}. {q_text}", FONT_BODY, Pt(10),
                               space_before=2, space_after=1)
                _render_short_answer_lines(doc, 4)

    _add_page_number(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
